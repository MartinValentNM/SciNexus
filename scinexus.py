# SciNexus v24.37
# Opravy v24.37:
#  • Banner "Nalezena automatická záloha": info zaujme celou šířku, tlačítka pod ním
#    → žádné překrývání záložek; CSS rozšířen pro sloupce i kontejnery
#  • Quick Presets → přejmenováno na "Rychlé nastavení" v celém CZ rozhraní
#  • Správa Rychlých nastavení: přidány vestavěné scénáře (Rychlá extrakce+validace,
#    Hyolitha velký PDF, Historický text překlad, Plný Workflow)
#    → každý scénář je viditelný a aplikovatelný přímo v záložce Nastavení
#    → smazání uživatelských nastavení opraveno (tlačítko vždy viditelné)
#  • chunk → blok: všechna CZ UI slova "chunk" nahrazena srozumitelnějším "blok"
#    (Velikost bloku, Paralelní extrakce bloků, Nastavení bloků překladu, atd.)
#  • Uložená přihlašovací jména: known_users.json ukládá max 30 jmen
#    → rozbalovací nabídka v sidebaru před přihlášením ("Přihlásit se jako…")
#    → nový uživatel se automaticky uloží po prvním přihlášení
#  • Per-user adresáře: users/<jméno>/ obsahuje vlastní nastavení každého uživatele
#    → šablony, slovníky, presets, session autosave, history, extraction snapshots
#    → activate_user_profile() přepne všechny globální cesty po přihlášení
#    → ostatní uživatelé nevidí vaše šablony ani presets
# Opravy v24.36:
#  • _temp/ soubory se hromadily: _temp_prune(keep_last=10) přidán přímo do
#    _temp_save_txt(), _temp_save_docx(), _temp_save_json() — po každém uložení
#    se automaticky zachová jen posledních 10 souborů daného prefixu
#    (platí pro překlad i extrakci — sdílejí stejné save funkce)
#  • use_container_width=True/False → width="stretch"/"content" (Streamlit deprecation fix)
#    → odstraněny runtime warnings "Please replace use_container_width with width"
#  • Překlad přepracován na "1 chunk = 1 Streamlit rerun"
#    → UI reaguje na tlačítka Stop/Pause/Resume mezi každým chunkem
#    → dříve byl překlad synchronní smyčka: UI bylo zablokované, tlačítka nefungovala
#  • Pauza: nastaví translate_paused=True + rerun; překlad se zastaví před dalším chunkem
#    → záložka zůstane na tab 1 (query_params["tab"]="1" + _gsb_jump_to_tab=1)
#    → tlačítko se změní na "▶️ Pokračovat", zobrazí počet přeložených bloků
#  • Resume: nastaví translate_running=True + rerun; pokračuje od dalšího chunku ve frontě
#    → tr_chunk_queue obsahuje zbývající chunky, tr_live_partial dosavadní přeložený text
#  • Živé okno: zobrazuje tr_live_partial vždy (viditelné i po pauze/rerunu)
#  • gsb_done_stay_on_tab: přidáno gsb_tab → tlačítko "Přejít na záložku" v gsb baneru
#  • OFF tlačítko: time.sleep(1.5) před os._exit(0)
# Opravy v24.34:
#  • Banner "Nalezena automatická záloha": CSS margin-top na stTabs + margin-bottom na alerty
#    → banner už nepřekrývá záložky bez ohledu na výšku banneru
#  • OFF tlačítko: st.markdown(unsafe_allow_html) → st.html() + window.parent.close()
#    → st.markdown sanitizuje <script> tagy (CSP), st.html() je renderuje v iframe bez sanitizace
#    → window.parent.close() zavírá skutečnou záložku (ne iframe)
#  • BS4: opraven bug "try: BS4_AVAILABLE = True" bez importu → přidán "from bs4 import..."
#    → dříve BS4_AVAILABLE bylo vždy True i bez nainstalované knihovny → pád při použití
#  • Živé okno překladu: přežije pauzu/rerun
#    → translate_text_live ukládá current_text do session_state["tr_live_partial"] po každém bloku
#    → Stop/Pause blok zobrazuje tr_live_partial pod tlačítky (popisek "pozastaveno"/"běží")
#    → při zastavení se tr_live_partial vymaže; při novém překladu se resetuje
# Opravy v24.32:
#  • JS tab switching: window.parent.document místo document — záložky jsou v parent frame,
#    st.html() běží v sandboxed iframe → document.querySelectorAll() nenašel záložky
#    → po překladu/extrakci/validaci se záložka nyní spolehlivě zachová
#  • JS clickTab: retry loop (max 25× po 60ms) + ověření aria-selected → spolehlivější přepnutí
#  • JS klávesové zkratky (Alt+1–9): přesunuty do parent dokumentu (stejný fix jako výše)
#  • JS watchTabs: history.replaceState + sessionStorage voláno přes getWin() (parent window)
#  • MutationObserver: sleduje parent document.body s try/catch fallbackem
# Opravy v24.31:
#  • Přihlašovací stránka: bannery (session info, _temp soubory, autosave) se zobrazí AŽ PO přihlášení
#    → presence.login_dialog() přesunuto PŘED bannery (dříve bylo za nimi)
#  • Pozastavení překladu: tlačítka Stop/Pozastavit přesunuta VNE z if do_translate_btn bloku
#    → tlačítka přežijí rerun a jsou viditelná i po pozastavení (neskáče na záložku extrakce)
#    → zachovává se již přeložený text a tlačítko "Pokračovat"
#    → po Stop/Pause vždy nastaven _gsb_jump_to_tab=1 (záložka Překlad)
#  • Stop/Pause extrakce: podobná oprava — tlačítka přesunuta před if run_btn blok
#    → vnitřní duplikáty odstraněny (zamezení duplicate key chyby Streamlitu)
#  • Stop/Pause validace: stejná oprava — tlačítka přesunuta před Spustit validaci
#  • Obnova session: po obnovení synchronizace aliasů (translation_ready, last_extracted_text/last_extraction_text)
#    → detailní hláška co bylo obnoveno (překlad, extrakce, validace, text k překladu)
#    → stejná oprava pro tlačítko v sidebaru
#  • Tlačítko OFF: nyní ukončuje CMD/terminál nadřazený procesu (psutil, pokud dostupný)
#    → záložka se zavře přes window.close() po 600 ms
# Opravy v24.19 — Živá okna, Stop/Pause všude, zachování záložek, klávesové zkratky
#  • Živé okno překladu: HTML div + JS scroll na konec po každém bloku (místo text_area)
#  • Živé okno extrakce: HTML div green style, aktualizace po každém bloku smyčky
#  • Stop/Pause překlad: tlačítka viditelná PŘED a BĚHEM překladu (ne až po dokončení)
#  • Stop/Pause extrakce: nová tlačítka ⏹️/⏸️ nad smyčkou souborů
#  • Stop/Pause validace: nová tlačítka ⏹️/⏸️ před zahájením validace
#  • Zachování záložky: JS klikne na tab POUZE při explicitním jump_target (ne při každém rerunu)
#  • Klávesové zkratky: Alt+1–9 (nebo Ctrl+Shift+1–9) pro přepínání záložek (vlastní listener, nativní accesskey nefunguje v Chrome)
#  • Status bar "Přejít": opravena logika — funguje správně s novým JS
# Opravy v24.4 — Validace kódu (pyflakes)
#  • Kritické: HYOLITHA_FIELD_MAP, HYOLITHA_*_COLS — definice chyběly (NameError při exportu)
#  • Kritické: _TEMP_DIR použit před definicí v banner bloku → _TEMP_DIR_EARLY forward alias
#  • Kritické: ext_sys_prompt použit před definicí → session_state fallback
#  • Odstraněny nepoužité proměnné: _td, bg_tab_i, dark, text_lc, taxon_lower_w,
#    HIER, start (gsb_render), _sve
#  • Odstraněn duplicitní global _BACKUP_DONE_TODAY v _do_backup (nested fn)
#  • Odstraněny nepoužité importy: hashlib, BeautifulSoup (top-level),
#    defaultdict (top-level), qn, OxmlElement (lazy import uvnitř funkce)
#  • Odstraněn zbytečný f-prefix na LANGUAGE_AUTO_LABELED
# Opravy v24.3 — Dokončení ochrany dat
#  • Chat: atomický zápis konverzací (.tmp → os.replace), autosave po každé odpovědi
#  • Style polish: výsledek uložen do _temp/ ihned + autosave + v SESSION_AUTOSAVE_KEYS
#  • Style polish přijetí: autosave před zobrazením success zprávy
#  • Načtení snapshotu extrakce: sync obou klíčů (last_extracted_text + last_extraction_text)
#    + autosave před rerun
#  • Workflow validace: _temp save ihned po dokončení validačního kroku
#  • Hyolitha export records: _temp save ihned (extrakce i workflow) + v SESSION_AUTOSAVE_KEYS
#  • SESSION_AUTOSAVE_KEYS rozšířeny: style_polish_result, hyolitha_export_records, ab_results
# Opravy v24.2 — Ochrana rozdělané práce
#  • Workflow: průběžné uložení do _temp/ po každém kroku (current_text + log JSON)
#  • Workflow: finální výsledek uložen do _temp/ + autosave + sync last_extracted_text
#  • Workflow extrakce: opravena synchronizace last_extracted_text (bylo jen last_extraction_text)
#  • Re-validace nenalezených: autosave před st.rerun()
#  • Bulk re-validace upravených řádků: autosave před st.rerun()
#  • Styl překladu: výsledek uložen do _temp/ ihned po dokončení (ne až po kliknutí "Přijmout")
#  • Terminologie překladu: výsledek uložen do _temp/ ihned po dokončení
#  • Zpětný překlad: výsledek uložen do _temp/ ihned po dokončení
#  • Přijetí stylu/terminologie: autosave před st.rerun()
#  • A/B test: autosave po dokončení + ab_results přidán do SESSION_AUTOSAVE_KEYS
#  • Workflow: _temp_prune před spuštěním (zabrání hromadění starých kroků)
# Opravy v24.1 — Ochrana proti ztrátě dat
#  • Kritická oprava: "last_extraction_text" ≠ "last_extracted_text" — typo bug způsoboval,
#    že výsledky extrakce nikdy nebyly zahrnuty do autosave ani do UI snapshotů
#    → synchronizace obou klíčů + přidání last_extraction_text do SESSION_AUTOSAVE_KEYS
#  • Automatický snapshot extrakce po dokončení (extraction_history/ + timestampový label)
#  • save_extraction_version: atomický zápis (.tmp → os.replace) + cleanup chyb
#  • save_session_to_disk: rotující .bak záloha předchozí verze (ochrana proti korrupci)
#  • load_session_from_disk: fallback na .bak pokud je primary soubor poškozený
#  • Okamžitý autosave po dokončení extrakce, překladu i validace (nečeká na N rerunů)
#  • _autosave_if_due: interval 5→3 reruns + hash-check (ukládá jen při změně dat)
#  • _temp_prune(): zachová jen posledních N sad temp souborů (zabrání plnění disku)
#  • _temp_prune volán na začátku extrakce, překladu i validace
#  • _temp_cleanup pro extrakci: smaže jen při 100% úspěchu (při chybách zachová temp)
#  • Iterační temp soubory překladu: smažou se jen pokud finál není prázdný
#  • UI banner při startu: upozorní na soubory z přerušené operace v _temp/
#    s tlačítkem pro smazání nebo zavření notifikace
# Opravy v24.0 — Stabilita, průběžné ukládání a paralelní zpracování
#  • Stabilita: _LLMConcurrencyManager.set_max() bezpečná (přidá/odebere tokeny semaforu,
#    nekonstruuje nový objekt — čekající thready nejsou ztraceny)
#  • Stabilita: chat_completion — sdílená _LLM_SESSION (connection pool),
#    broken pipe recovery (_reset_llm_session), thread-safe session_state přístup
#  • Stabilita: _get_http_session / _reset_http_session — double-checked locking,
#    broken pipe recovery při ConnectionError/ChunkedEncodingError
#  • Stabilita: SQLite disk cache — double-checked locking, mmap 64 MB,
#    _reset_cache_conn() + retry při DatabaseError
#  • Paralelismus: do_translate — modul-level _THREAD_OFFLINE/_THREAD_GSP propagace
#    (worker thready bez ScriptRunContext čtou fallback hodnoty)
#  • Paralelismus: do_translate — count_lock oddělený od cache_lock (eliminuje False contention),
#    _safe_progress() obaluje progress_cb (ScriptRunContext chyby ignorovány)
#  • Paralelismus: extrakce paralelní — context overflow halving (stejný vzor jako překlad),
#    _ext_count_lock + safe progress update, propagace session_state před thready
#  • Paralelismus: validace paralelní — max 4 souběžné taxony (vnitřně 8 DB threadů = 32 spojení),
#    _val_result_lock pro thread-safe append, error handling per-taxon
#  • Průběžné ukládání: _temp_save_docx fallback na .txt při chybě python-docx
#  • Průběžné ukládání: do_translate_batched propaguje session_state před dávkami
# Opravy v23.19 — Přesné URL vzory GNverifier + PaleoDB
#
# GNVERIFIER — URL vzory:
#  • Webový odkaz: ?capitalize=on&format=html&names=Gompholites
#    pro víceslovná jména:  ?capitalize=on&format=html&names=alfaites+romeo  (+ místo %20)
#  • _gnv_web_name = taxon.replace(' ', '+')  → správný formát dle dokumentace
#
# PALEODB — přesné URL vzory dle dokumentace (dev.paleobiodb.org):
#  • Jeden taxon (found URL):
#    /data1.2/taxa/single.json?name=Gompholites&show=full&datainfo=true
#  • Všechny druhy rodu (wildcard found URL):
#    /data1.2/taxa/list.json?base_name=Gompholites&rank=species&show=class&rowcount=true&datainfo=true
#  • Subtaxony (non-wildcard list):
#    /data1.2/taxa/list.json?base_name=Gompholites&show=class&rowcount=true&datainfo=true
#  • Prefix search (wildcard fallback/default):
#    /data1.2/taxa/list.json?match_name=Gompholi%&show=class&rowcount=true&datainfo=true
#  • Wildcards: * → % (sekvence), ? → _ (jeden znak), . = zkratka (PBDB nativní)
#  • Konverze: _pb_wc_name = taxon.replace("*","%").replace("?","_")
#  • Parametr datainfo=true v zobrazovacích URL (dle doporučení docs)
# Opravy v23.18 — GNverifier URL syntaxe + PaleoDB/Fossilworks API přepis
#
# GNVERIFIER:
#  • Oprava URL pro webový odkaz: ?capitalize=on&format=html&names=X
#    (dříve /names/X — nefungovalo)
#  • POST payload: přidán "withCapitalization": true (ekvivalent capitalize=on)
#    → GNverifier nyní najde "Gompholites", "alnus" apod. díky kapitalizaci
#  • Match typy přeloženy do češtiny v details i gnv_preferred tabulce
#  • Oprava bug: result["summary"]["found"] += 1 bylo mimo if block → počítalo
#    každý GNverifier dotaz jako found bez ohledu na výsledek
#
# PALEODB + FOSSILWORKS — přepis na správné API endpointy:
#  • Primárně taxa/single.json?name=X&show=full,classext&vocab=pbdb
#    (dříve taxa/list.json s nesprávným parametrem show=attr,classext,parent)
#  • Fallback: taxa/list.json?base_name=X pro subtaxony nebo
#    taxa/list.json?match_name=X% pro wildcard dotazy (dle dokumentace PBDB)
#  • Parametr vocab=pbdb → plné názvy polí místo zkratek (nam→taxon_name atd.)
#  • Lidsky čitelné hodnoty ve details:
#    - rank: "rod (genus)" místo "5", "druh (species)" místo "2"
#    - status: "platné jméno" místo "belongs to", "synonymum" místo kódů
#    - extinct: "ano (fosilní)" / "ne (recentní)" místo "0"/"1"
#    - klíče polí: "název (name)", "kmen (phylum)", "třída (class)" atd.
#  • Wildcard hits využívají match_name pro abecední shodu (dle PBDB docs)
#  • URL odkaz vede přímo na Fossilworks taxonInfo s taxon_no
#
# Opravy v23.17 — Cache flush + GNverifier integrace
#
# AUTOMATICKÉ MAZÁNÍ CACHE:
#  • Cache (memory + SQLite) se nyní vymaže pro dotazované taxony vždy při spuštění validace
#    → zabrání zobrazení starých záporných výsledků po opravách kódu
#    → mazáno selektivně (jen aktuální taxony × DB, ne celá cache)
#
# INTEGRACE GLOBAL NAMES VERIFIER (GNverifier):
#  • Nová DB "GNverifier" — POST /api/v1/verifications
#  • Verifikuje jméno najednou proti 100+ biodiverzitním databázím
#  • Vrátí: bestResult (nejvyšší skóre), preferredResults (PaleoDB, CoL, GBIF, WoRMS, IRMNG, ITIS)
#  • Match typy: Exact, Fuzzy, FuzzyRelaxed, Partial (pro jednoslovné rody)
#  • Preferované zdroje: 1=CoL, 11=GBIF, 172=PaleoDB/Fossilworks, 169=WoRMS, 167=IRMNG, 3=ITIS
#  • V detailu výsledku: tabulka preferredResults s přímými odkazy
#  • Přidán do presetů: Hyolitha, Paleontologie, Mořské organismy, Botanika, Zoologie
#  • API: https://verifier.globalnames.org/api/v1/verifications (veřejné, bez API klíče)
#
# Opravy v23.16 — Validace: přepínání záložek + Alnus/Betula not found
#
# PŘEPÍNÁNÍ NA ZÁLOŽKU EXTRAKCE PO VALIDACI:
#  • Oprava: gsb_done() nyní vždy nastaví _gsb_jump_to_tab na záložku operace (gsb_tab)
#    Dříve: po kliknutí ✕ nebo při rerunu po validaci se mohl zobrazit starý banner
#    extrakce a přepnout na záložku 0 (Extrakce)
#    Nyní: gsb_done() automaticky nastaví jump na správnou záložku (2 = Validace)
#
# ALNUS/BETULA NOT FOUND (cache a BioLib):
#  ⚠️ DŮLEŽITÉ: Pokud Alnus/Betula stále nenalezeno, je pravděpodobně
#     uložen starý záporný výsledek v SQLite cache.
#     ŘEŠENÍ: Sidebar → 🗑️ Vymazat cache validace → spustit znovu
#
#  • Oprava BioLib: přidán _name_matches filter pro tx_name (bral první hit bez ověření)
#    + podpora alternativního API formátu (status:"ok" místo "found")
#  • Oprava CoL: HTTP 400 fallback — zkusí dotaz bez datasetKey=COL pokud první selže
#    (datasetKey=COL způsoboval 400 pro některé znakové sady)
#
# Opravy v23.15 — Genus expansion: wildcard_hits ve všech podporovaných DB
#
# GENUS EXPANSION (tlačítko "🔍 Hledat druhy rodu X*" v detailu validace):
#  • Nové: Fossilworks — wildcard_hits ze všech záznamů při X% dotazu
#  • Nové: GBIF — /species/{key}/children + /species/search?genus=X pro wildcard_hits
#  • Nové: CoL — /nameusage/{id}/children pro rod/čeleď/řád úroveň
#  • Nové: ITIS — wildcard_hits ze searchByScientificNameWildcard výsledků
#  • Nové: IPNI — prefix match při wildcard dotazu, wildcard_hits (botanické druhy)
#  • Nové: Tropicos — prefix match při wildcard dotazu, wildcard_hits (botanika)
#  • Nové: IFPNI — prefix match při wildcard dotazu, wildcard_hits (fosilní rostliny)
#  • _WC_SUPPORTED rozšířena: WoRMS, IRMNG, PaleoDB, Fossilworks, ITIS, CoL,
#    GBIF, IPNI, Tropicos, IFPNI (dříve jen 6 DB, nyní 10)
#  • UI: aktualizován hint o wildcard podpoře v záložce Validace
#
# PRAKTICKÉ POUŽITÍ:
#  - Zadej "Hyolithes*" nebo "Betula*" do pole taxonů
#  - Validuj normálně — DB označené ✅ vrátí nalezené záznamy
#  - V detailu každého výsledku se zobrazí tlačítko "🔍 Hledat druhy rodu Hyolithes"
#  - Po kliknutí se zobrazí tabulka všech druhů s přímými odkazy
#  - Alternativně: zadej rovnou "Hyolithes" (jednoslovně) a klikni na tlačítko
#    které se zobrazí automaticky pod výsledky validace rodu
# Opravy v23.14 — ScriptRunContext + nestabilní výsledky validace
#
# SCRIPTRUNCONTEXT WARNING (ThreadPoolExecutor):
#  • Oprava: validate_taxon_name čte st.session_state["val_timeout"] JEDNOU
#    v hlavním threadu a předává jako parametr timeout do worker threadů
#    (dříve každý worker thread volal st.session_state.get() = ScriptRunContext warning)
#  • Oprava: _validate_taxon_single_db přijímá timeout: int = VAL_TIMEOUT parametr,
#    TO = timeout (ne TO = st.session_state.get(...))
#  • Oprava: _safe_get má try/except kolem st.session_state přístupu,
#    fallback na VAL_TIMEOUT konstantu když ScriptRunContext není dostupný
#
# NESTABILNÍ VÝSLEDKY VALIDACE (Alnus někdy found/někdy not found, Betula nenalezena):
#  • Přidána funkce _name_matches(api_name, query) — správné porovnání jmen:
#    - "Betula L." vs "Betula" → match ✅ (autor suffix nevadí)
#    - "Betula pendula" vs "Betula" → no match ✅ (druh ≠ rod)
#    - "BETULA" vs "Betula" → match ✅ (case-insensitive)
#  • _name_matches použita v: CoL, GBIF fallback, IPNI, IRMNG, WoRMS
#    (dříve striktní lower().strip() == lower().strip() odmítal záznamy s autorem)
#
# MEDIAFILESTORAGERROR (Bad filename *.bin):
#  • Není chyba aplikace — Streamlit interní in-memory storage expiruje při rerunu.
#    Neškodná chyba v logu, uživatel ji nevidí. Oprava je na straně Streamlit.
#
# Opravy v23.13 — Validace: mazání pole + opravy databází
#
# MAZÁNÍ POLE TAXONŮ:
#  • Oprava: st.session_state["val_taxa_manual"] = "" před renderem widgetu (správný vzor)
#    Dříve: pop() + value= parametr nefungoval spolehlivě při Streamlit rerunu
#    Nyní: přímé nastavení session_state[key] = "" zajistí prázdné pole po smazání
#  • Odstraněn konfliktní value= parametr z text_area — widget čte z session_state[key]
#
# OPRAVY DATABÁZÍ:
#  • Tropicos: odstraněn fallback na první výsledek bez přesné shody
#    (candidates_t = exact_t if exact_t else data → found=False pokud není přesná shoda)
#    → Gompholites nyní správně vrátí not found (není rostlina)
#    → Alnus nyní správně vrátí found jen pokud jméno odpovídá přesně
#  • IPNI: přidán filtr přesné shody (bral první výsledek bez ověření jména)
#    → Gompholites nyní správně not found (není v IPNI)
#    → Alnus Mill. nyní správně found (přesná shoda "Alnus")
#  • Mikrotax: zpřísnění detekce — jméno musí být v <title> nebo <h1>/<h2>,
#    ne jen kdekoliv v HTML těle stránky (eliminuje false positive pro menu/navigation)
#  • Poznámky k jednotlivým DB pro Alnus / Gompholites:
#    - BioLib:      Alnus ✅ (strom, česká flora) | Gompholites ❌ (fosilní, není v BioLib)
#    - CoL:         Alnus ✅ | Gompholites ❌
#    - IFPNI:       Alnus ❌ (fosilní rostliny) | Gompholites ❌
#    - Fossilworks: Alnus ❌ (není fosilní rod) | Gompholites ✅ (hyolit)
#    - GBIF:        Alnus ✅ | Gompholites ✅ (záznamy výskytu)
#    - IPNI:        Alnus ✅ (rostlina) | Gompholites ❌ (není rostlina)
#    - IRMNG:       Alnus ✅ | Gompholites ✅ (rod dle IRMNG)
#    - ITIS:        Alnus ✅ | Gompholites ❌ (není v ITIS)
#    - Mikrotax:    Alnus ❌ | Gompholites ❌ (pouze mikroorg.)
#    - PaleoDB:     Alnus ❌ nebo ✅ záleží | Gompholites ✅
#    - Plazi:       Alnus ✅/❌ záleží na tratmentech | Gompholites ✅ (pokud jsou tratmenty)
#    - WoRMS:       Alnus ❌ (sladkovodní) | Gompholites ❌/✅ záleží
#    - ZooBank:     vyžaduje CAPTCHA — zobrazí "Neověřeno" s odkazem
#    - Tropicos:    Alnus ✅ (botanická DB) | Gompholites ❌ (po opravě)
# Opravy v23.12 — Výkon: zbývající systémové optimalizace
#  • Výkon: normalize_stratigraphy_local — pre-kompilované _ICS_PATTERNS (54 termínů,
#    žádná re.compile za runtime; dříve 54 kompilací per každé volání funkce)
#  • Výkon: find_duplicate_taxa — prefix indexing (3znakový prefix dict)
#    eliminuje ~95 % Levenshtein porovnání pro databázi 1500 taxonů
#  • Výkon: records_to_hyolitha_dfs — vektorizovaná DataFrame konstrukce,
#    opravena logika filtru (OR místo AND pro prázdné Rod/Druh)
#  • Výkon: lazy import PyPDF2 a python-docx — cold start aplikace rychlejší
#    o ~0.5–1 s (import se provede jen při prvním použití PDF/DOCX)
#  • Výkon: normalize_record_keys — _NORM_KEY_CACHE cachuje lower().strip() per klíč
#    (pro batch 1500 záznamů × 20 polí = 30k operací → O(1) lookup)
#  • Výkon: validation_results_to_df — cache výsledku do session state
#    (DataFrame se nesestavuje znovu při každém Streamlit rerunu záložky Validace)
#  • Výkon: workflow validace — ThreadPoolExecutor (stejný vzor jako tab Validace)
#    (bylo sekvenční; nyní paralelní s max(4, lms_max_concurrent) workery)
#  • Výkon: _GLOSSARY_RE_CACHE — LRU eviction při max 500 položkách
#    (zabraňuje neomezenému růstu paměti pro long-running instance)
#  • Výkon: fts_index_results — odstraněna zbytečná vnitřní alokace fts_rows_with_id
#  • Oprava: fts_index_results — zbytečný `fts_rows` list (shromažďoval data které
#    se nikdy nepoužila — výsledek nedokončené refaktorizace)
# Opravy v23.11 — Výkon: systémové optimalizace
#  • Výkon: validate_taxon_name — paralelní dotazy na DB (ThreadPoolExecutor, max 8 workerů)
#    každá DB dostane vlastní thread; čas validace klesá z ~15s na ~2–3s pro 10 DB
#  • Výkon: chat_completion_stream — O(n²) string concat → O(n) list append + join
#    (při 500+ tokenech výstup se každý += tvořil nový string objekt)
#  • Výkon: suggest_chunk_size — @st.cache_data → @functools.lru_cache (bez Streamlit overhead)
#  • Výkon: _enc(), _plus() — přesun z lokálních definic uvnitř validate_taxon_name
#    na modul-level (eliminuje new object creation při každém volání)
#  • Výkon: _auto_nightly_backup — in-memory flag _BACKUP_DONE_TODAY + daemon thread
#    (dříve Path.exists() syscall při každém save_to_history)
#  • Výkon: dedup hash — md5 → Python built-in hash() (5–10× rychlejší pro dedup)
#  • Výkon: HTTP connection pool — pool_maxsize 10 → 20 (pro 14 DB × 8 paralel. workerů)
#  • Výkon: chat_completion_stream — přidán parametr stop (konzistence s chat_completion)
#  • Oprava: read_uploaded_file cache klíč normalizuje page_spec přes parse_page_ranges
#    ("1-3" a "1,2,3" nyní dají stejný cache hit)
# Opravy v23.10 — Výkon: obecné optimalizace po auditu
#  • Výkon: _JSON_FENCE_RE, _MULTI_NEWLINE_RE, _PAGE_BATCH_RE — pre-kompilované regex konstanty
#    (nahrazuje 19× inline re.sub(r"```json|```",...) kompilovaných za runtime)
#  • Výkon: offline_db_* — sdílené SQLite spojení přes _get_offline_conn() (WAL mode)
#    (dříve open()/close() na každý dotaz = zbytečný OS overhead)
#  • Výkon: fts_index_results — batch zpracování + jediný commit() (dříve commit per záznam)
#  • Výkon: fts_conn() — přidány PRAGMA synchronous=NORMAL + cache_size=2000
#  • Výkon: preprocess_text_for_llm — token_reduction_pct přes len()//4 místo split()
#    (split() na 100k znaků = ~50ms; len()//4 = <1µs)
#  • Výkon: _GLOSSARY_RE_CACHE cachuje kompilované regexpy pro slovník
#    (pro dávkový překlad N textů se stejným slovníkem regex nesestavuje N×)
#  • Oprava: _cached_sidebar_stats měla trojitý @st.cache_data dekorátor → opraveno na jeden
#  • Oprava: src_lang chyběl v do_translate volání pro terminologii a zpětný překlad
#  • Oprava: workflow extrakce/překlad/čištění používaly hardcoded max_tokens=8000/temp=0.15
#    → nahrazeno konstantami _MAX_TOKENS_* a voláním do_translate()
#  • Oprava: levenshtein — early-exit pro přesnou shodu a délkový rozdíl > 2
# Opravy v23.9 — Výkon: překlad & extrakce
#  • Výkon: chat_completion — stop sekvence pro JSON extrakci (zastaví hned po ']')
#  • Výkon: chat_completion — volitelný parametr stop_at pro překlad ("</translation>")
#  • Výkon: backoff s jitterem (random 0–30 % nad base) pro všechna retry místa
#  • Výkon: DLQ (Dead Letter Queue) pro selželvší chunky — sbírá chyby, nepřeruší batch
#  • Výkon: do_translate paralelní — semaphore-guard bez busy-wait (nahrazuje ThreadPoolExecutor)
#  • Výkon: _adaptive_max_tokens — zpřesnění: jazykový multiplikátor (CJK ×1.5, slovanské ×1.2)
#  • Výkon: suggest_chunk_size — opravena přepočet kontext window (byl ×3, teď ×1 / chars per token)
#  • Výkon: chunk_text_smart — minimální chunk size 2 000 znaků (zamezuje malým fragmentům)
#  • Výkon: extrakce sekvenční — stop sekvence "```\n" pro JSON výstupy
#  • Výkon: extrakce — deduplikace bloků i v sekvenčním režimu (cache přes hash)
#  • Výkon: do_translate_batched — přidán DLQ report po dávkách
#  • Doporučení: výchozí nastavení pro batch ("Quick Start") v záložce Nápověda
# Opravy v23.8 — Persistentní ukládání:
#  • Průběžný checkpoint každých N chunků v sekvenční extrakci (ext_progress_*)
#  • Checkpoint po každém dokončeném souboru (partial=True pokud ještě zbývají)
#  • Úklid ext_progress_* checkpointů po úspěšném dokončení extrakce
#  • Konstanta _EXT_CHUNK_CHECKPOINT_INTERVAL (default 5 chunků)
# Opravy v23.7:
#  • Oprava: Dismiss ✕ na status baru zůstane na aktuální záložce (ne tab 0)
#  • Oprava: CoL — pouze přesná shoda jména (žádný fallback na první náhodný výsledek)
#  • Oprava: IFPNI — přidán filtr přesné shody jména; nenastavovat found=True z bool(text)
#  • Oprava: GBIF fallback search — filtr přesné shody canonicalName/scientificName
#  • Oprava: Plazi XML fallback — ověření přítomnosti jména taxonu v odpovědi
#  • Oprava: Fossilworks — found=True jen pokud tid existuje A jméno odpovídá
#  • Oprava: PaleoDB — found=True jen pokud taxno existuje A jméno odpovídá
# Opravy v23.6:
#  • Oprava: StreamlitAPIException val_taxa_manual — clear button používá příznak _val_clear_pending,
#    widget key se maže přes pop() NA ZAČÁTKU záložky před renderem widgetu (ne po jeho vytvoření)
#  • Oprava: Status bar "probíhá validace" zůstával — gsb_done() se nyní volá PŘED st.rerun(),
#    ne až za ním kde se nikdy nedostal (rerun přeruší zbytek kódu)
#  • Oprava: NameError idx — for r in results_v → for idx, r in enumerate(results_v)
#  • Oprava: fallback gsb_done na konci záložky pokud gsb_active=True po obnovení session
# Opravy v23.5:
#  • Oprava: NameError _has_wildcard v genus expansion (přejmenováno na _taxon_has_wc)
#  • Status bar: odstraněn časomír (⏱), přidána popisná zpráva CZ/EN ("Probíhá validace…" atd.)
#  • Status bar: spinner CSS animace místo statické ikonky
#  • Status bar: po dokončení zůstává (✕ tlačítko pro zavření) — nezmizí automaticky
#  • Status bar: sidebar také bez časomíru
#  • Oprava: PaleoDB fallback URL = Fossilworks search (ne prázdný browse_url)
#  • Oprava: Fossilworks URL fallback = taxonSearch?taxon_name= (ne chybný taxonInfo bez ID)
#  • Oprava: IFPNI — víc ID polí pro přímý odkaz, lepší fallback search URL
#  • Oprava: PaleoDB wildcard hits — bezpečná int konverze rnk (None→0)
# Opravy v23.4:
#  • Nové: Genus expansion — tlačítko "Hledat druhy rodu X*" v detailu výsledku validace
#    vrátí druhy z DB (PaleoDB/IRMNG/WoRMS/GBIF/CoL/ITIS/Fossilworks) s klikacími odkazy
#  • Oprava: Recovery banner se nezobrazuje pokud běží operace (gsb_active) → dismiss nepřerušuje validaci
#  • Oprava: PaleoDB browse_url = taxon_no= (ne taxon_name=)
#  • Oprava: ZooBank CAPTCHA — zobrazí varování místo false positive, search link místo prázdného URL
#  • Oprava: Tropicos — přímý odkaz /name/NameId + filtr přesné shody jména
#  • Oprava: IFPNI — přímý odkaz /name/{id}
#  • Oprava: CoL — datasetKey=COL pro alfanumerická usage ID
#  • Oprava: IRMNG fallback URL = browse URL (ne API endpoint)
# Opravy v23.3:
#  • Oprava: VŠECHNY st.rerun() v záložkách Překlad (tab 1) a Validace (tab 2) mají jump target
#    — smazat překlad, přijmout styl/terminologii, uložit/smazat slovník, re-validovat, offline DB
#  • Oprava: CoL API endpoint + správné usage.id pro URL (alfanumerické, ne numerické)
#  • Oprava: PaleoDB URL = taxon_no z oid (ne jméno v URL)
#  • Oprava: Fossilworks URL = taxon_no z oid
#  • Oprava: IRMNG + WoRMS anti-false-positive filtr (přesná shoda jména)
#  • Oprava: Smazat pole taxonů — resetuje i widget key val_taxa_manual
# Opravy v23.2:
#  • Oprava: všechna rerun uvnitř záložky Validace nyní zachovávají tab 2 (neskáče na Extrakci)
#  • Oprava: tlačítko Smazat ve validaci maže i widget key (val_taxa_manual) — obsah se skutečně vymaže
#  • Oprava: CoL — správný API endpoint (api.catalogueoflife.org) + správné URL ID (usage.id, ne numerické)
#  • Oprava: CoL — filtr přesné shody jména (zamezuje false positive)
#  • Oprava: PaleoDB — URL nyní používá taxon_no z oid pole (ne název v URL)
#  • Oprava: Fossilworks — URL nyní používá taxon_no z oid pole
#  • Oprava: IRMNG — filtr přesné shody pro like=false (zamezuje false positive)
#  • Oprava: WoRMS — filtr přesné shody pro exact dotaz (zamezuje false positive)
# Opravy v23.1:
#  • Oprava: validace po dokončení zůstává na záložce Validace (ne tab 0)
#  • Oprava: swap ⇄ – flag-before-widgets vzor (odstraněn key= konflikt se selectboxy)
#  • Oprava: swap ⇄ – přeskok na tab 1 funguje spolehlivě (retry JS smyčka)
#  • Oprava: "Přejít na záložku" – JS retry smyčka (400ms start, max 3s)
#  • Wildcards: správná per-DB podpora (WoRMS/IRMNG/PaleoDB/Fossilworks/ITIS = ano; GBIF/ZooBank/CoL = ne)
#  • Wildcards: všechny nalezené záznamy uloženy do wildcard_hits (ne jen první)
#  • Wildcards: klikatelná tabulka všech wildcard hitů v detailu výsledků (max 50 záznamů)
#  • Wildcards: DB bez podpory dostávají čistý název (wildcards odstraněny)
# Opravy v23.0:
#  • Oprava: NameError '_L is not defined' při načtení session z disku (recovery banner)
#  • UI: gsb_render — klikatelný status bar po dokončení → tlačítko "Přejít na záložku"
#  • UI: gsb_render — tlačítko ✕ pro zavření notifikace bez přesunu
#  • Překlad: swap ⇄ prohazuje i text (výsledek překladu → vstupní pole)
#  • Výkon: do_translate — adaptive max_tokens (kratší chunk → méně tokenů → rychlejší)
#  • Výkon: do_translate paralelní — deduplikace chunků (opakující se záhlaví/bloky)
#  • Výkon: do_translate paralelní — per-chunk retry 3× s exponential backoff (0.5s, 1s)
# Opravy v20.0:
#  • Výkon: load_history() s in-memory cache (čte disk jen při změně souboru)
#  • Výkon: read_pdf OCR respektuje page_spec (nečte celý dokument)
#  • Výkon: preprocess_text Counter místo dict (rychlejší repeat-header detekce)
#  • Výkon: val_timeout vytažen mimo DB smyčku (1 session_state lookup místo N)
#  • Výkon: HTTP session na úrovni modulu (přežije reruns) + connection pool
#  • Stabilita: save_session_to_disk + checkpoint_save atomický zápis (os.replace)
#  • Stabilita: _safe_get retry na ConnectionError + backoff 500ms
#  • Stabilita: shutil inline import odstraněn
#  • UI: tmavý/světlý režim toggle v sidebaru (🌙/☀️)
#  • UI: CSS animace tab-panel fadeIn, focus ring, table striping, toast pozice
#  • UI: sidebar stats jako kompaktní grid (méně šumu)
#  • UI: status-pill CSS třída, skeleton loader CSS
# Opravy v19.0:
#  • Extrakce: two-pass přes chat_completion_queued (fronta + concurrency)
#  • Extrakce: opraven bug _eff_prompt — confidence + JSON schema se nevylučují
#  • Extrakce: auto-checkpoint při chybě souboru (disk, ne jen session state)
#  • Extrakce: chunk_text s overlap parametrem (200 znaků default)
#  • Extrakce: per-chunk retry v paralelním režimu (3× s exponential backoff)
#  • Extrakce: overlap slider v PDF nastavení (0–800 znaků)
# Opravy v18.0:
#  • Výkon: SQLite connection pool (WAL mode) — 60–80 % rychlejší disk cache
#  • Výkon: chat_completion GSP early-exit, lokální cache session state v hot path
#  • Výkon: workflow tab používá chat_completion_queued (fronta + concurrency)
#  • Výkon: Counter import přesunut na top-level
#  • GUI: sidebar přesunut Model manager + per-model prompty do záložky Nastavení
#  • GUI: status bar pro dlouhé úlohy
#  • GUI: desktop notifikace po dokončení extrakce/překladu
#  • Nové: Checkpoint systém — pozastavení a obnovení dávkové extrakce
#  • Nové: FTS full-text search v historii výsledků extrakce
#  • Nové: A/B prompt testing (formalizovaný, s metrikami)
# Opravy v17.0:
#  • Presence systém: přihlašování uživatelů, kdo je online, aktivní záložka
#  • Fronta LLM: pořadí úloh při souběžném použití více uživateli
# Opravy v16.3:
#  • Validace: persistentní HTTP Session (connection pooling, 30–50 % rychlejší)
#  • Validace: disk cache SQLite – výsledky přežijí restart aplikace
#  • Validace: detekce konfliktů mezi DB (CoL/PaleoDB status mismatch)
#  • Validace: tlačítko „Znovu validovat jen nenalezené"
#  • Validace: PBDB occurrence data (počet nálezů + stratigrafický rozsah)
#  • Validace: export filtr (vše / jen nalezené / jen nenalezené)
#  • Validace: vylepšený detail – klikatelné hypertextové odkazy per DB
# Opravy v16.2:
#  • Čištění dat: prompt vždy viditelný (mimo expander), uložení/načtení promptů
#  • Překlad: text se ukládá do session state – tlačítko Přeložit funguje po opuštění pole
#  • Validace: odstraněno Přednastavení, default = všechny DB zaškrtnuty
#  • Plazi: přepnuto na tb.plazi.org/GgServer/search, aktivní hypertextové odkazy
#  • BioLib: anti-robot bypass přes Session + cookies
#  • Záložka Kódování odstraněna

import io
import json
import logging
import os
import re
import sqlite3
import subprocess
import shutil
import warnings as _warnings_early
# Potlač pynvml FutureWarning PŘED importem torch/transformers (přichází z torch.cuda.__init__)
_warnings_early.filterwarnings("ignore", message=".*pynvml.*deprecated.*", category=FutureWarning)
_warnings_early.filterwarnings("ignore", message=".*nvidia-ml-py.*", category=FutureWarning)
import time as _time
import zipfile
import unicodedata
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import contextmanager
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
from collections import Counter

import requests
import streamlit as st
# ── SILENCE ANNOYING STARTUP WARNINGS ─────────────────────────────────────
# ───────────────────────────────────────────────────────────────────────────
import presence
# PyPDF2 a python-docx jsou lazy-importovány v příslušných funkcích (read_pdf, to_docx_bytes)
# — nezpomalují cold start aplikace pokud uživatel v dané session PDF/DOCX nepoužije
import pandas as pd

# ── SILENCE ALL ANNOYING STARTUP WARNINGS (v24.30) ────────────────────────
import warnings
warnings.filterwarnings("ignore", message=".*use_container_width.*")
warnings.filterwarnings("ignore", message=".*Please replace.*use_container_width.*")

# Velmi agresivní filtr na všechny transformers __path__ warnings
warnings.filterwarnings(
    "ignore",
    message=r"Accessing `__path__` from `\.models\..*`",
    category=FutureWarning
)
warnings.filterwarnings(
    "ignore",
    message=r"Accessing `__path__` from",
    category=FutureWarning
)
# Potlač FutureWarning z deprecated pynvml balíčku (správný je nvidia-ml-py)
warnings.filterwarnings("ignore", message=".*pynvml.*deprecated.*", category=FutureWarning)
warnings.filterwarnings("ignore", message=".*nvidia-ml-py.*", category=FutureWarning)
# ───────────────────────────────────────────────────────────────────────────

# ── Thread-safe přístup k session_state ──────────────────────────────────────
session_lock = threading.Lock()

def safe_get(key: str, default=None):
    """Thread-safe čtení ze session_state."""
    with session_lock:
        return st.session_state.get(key, default)

def safe_set(key: str, value):
    """Thread-safe zápis do session_state."""
    with session_lock:
        st.session_state[key] = value

# ── Safe LLM call s retry a recovery ─────────────────────────────────────────
@contextmanager
def safe_llm_call(base_url, model, messages, temp=0.1, max_tokens=4000, **kwargs):
    """Context manager pro bezpečné volání LLM s retry (5×) a exponential backoff.

    Použití:
        with safe_llm_call(base_url, model, messages) as result:
            print(result)

    Při 5 neúspěšných pokusech zobrazí st.error a vyvolá výjimku.
    """
    for attempt in range(5):
        try:
            yield chat_completion(base_url, model, messages,
                                  temp=temp, max_tokens=max_tokens, **kwargs)
            return
        except Exception as e:
            if attempt == 4:
                st.error(f"LLM selhalo po 5 pokusech: {e}")
                raise
            wait = (2 ** attempt) + 0.1
            st.warning(f"LLM chyba (pokus {attempt+1}/5) – čekám {wait:.1f}s...")
            _time.sleep(wait)

# ── Deduplikace chunků (opakující se záhlaví/patičky) ────────────────────────
def deduplicate_chunks(chunks: List[str]) -> List[str]:
    """Odstraní opakující se záhlaví a patičky mezi chunky.

    Typicky vznikají při OCR PDF, kde každá stránka opakuje název časopisu,
    záhlaví tabulky nebo patičku s číslem strany.
    Detekuje shodu prvních 3 řádků (záhlaví) s předchozím chunkem a odstraní je.
    """
    if not chunks:
        return chunks
    cleaned: List[str] = []
    last_header = ""
    for chunk in chunks:
        lines = chunk.splitlines()
        if len(lines) > 5:
            header = "\n".join(lines[:3]).strip()
            if header and header == last_header:
                lines = lines[3:]  # odstranit opakující se záhlaví
            last_header = header
        cleaned.append("\n".join(lines).strip())
    return [c for c in cleaned if c]

# ── Unified Chunking Config (v24.17) ──────────────────────────────────────────
class ChunkConfig:
    """Sjednocená konfigurace chunkování pro všechny záložky aplikace."""
    DEFAULT_TARGET    = 2200
    DEFAULT_OVERLAP   = 350
    MIN_CHUNK         = 1800
    ULTRA_LARGE_BATCH = 40   # stránek najednou pro dokumenty 1000+ stran

    @staticmethod

    @classmethod
    def smart_chunk(cls, text: str,
                    target_chars: int = None,
                    overlap: int = None) -> List[str]:
        """Deleguje na smart_chunk_text s výchozími hodnotami z ChunkConfig."""
        return smart_chunk_text(
            text,
            target_chars=target_chars if target_chars is not None else cls.DEFAULT_TARGET,
            overlap=overlap       if overlap       is not None else cls.DEFAULT_OVERLAP,
        )

# ── Checkpoint & Resume (v24.16) ──────────────────────────────────────────────
def save_checkpoint(operation: str, data: Dict, step: int = 0):
    """Atomicky uloží checkpoint do _temp/checkpoints/ (.tmp → replace)."""
    checkpoint_dir = Path("_temp/checkpoints")
    checkpoint_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = checkpoint_dir / f"checkpoint_{operation}_{step}_{ts}.json"
    tmp = path.with_suffix(".tmp")
    payload = {
        "operation": operation,
        "timestamp": datetime.now().isoformat(),
        "step": step,
        "data": data,
    }
    tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(path)

def list_checkpoints(operation: str = None):
    """Vrátí seznam checkpointů, volitelně filtrovaných dle operace (seřazeno desc)."""
    checkpoint_dir = Path("_temp/checkpoints")
    if not checkpoint_dir.exists():
        return []
    cps = []
    for f in checkpoint_dir.glob("*.json"):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            if operation is None or data.get("operation") == operation:
                cps.append(data)
        except Exception:
            continue
    return sorted(cps, key=lambda x: x.get("timestamp", ""), reverse=True)

# ══════════════════════════════════════════════════════
# v24.20 – Auto-Resume + Advanced Chunking + Live Progress + Safe Threading
# ══════════════════════════════════════════════════════

def find_interrupted_operations() -> List[Dict]:
    """Najde všechny přerušené operace v checkpoints a _temp/."""
    interrupted = []
    # Checkpoints
    for cp in list_checkpoints():
        interrupted.append({
            "type":      "checkpoint",
            "operation": cp.get("operation", "unknown"),
            "timestamp": cp.get("timestamp", ""),
            "step":      cp.get("step", 0),
            "data":      cp.get("data", {}),
            "source":    "checkpoint",
        })
    # Temp soubory
    temp_dir = Path("_temp")
    if temp_dir.exists():
        for f in temp_dir.glob("*.json"):
            if any(x in f.name.lower()
                   for x in ["extraction", "wf_", "translate", "partial"]):
                try:
                    data = json.loads(f.read_text(encoding="utf-8"))
                    interrupted.append({
                        "type":      "temp",
                        "operation": data.get("operation", f.name),
                        "timestamp": data.get("saved_at", str(f.stat().st_mtime)),
                        "data":      data,
                        "source":    str(f),
                    })
                except Exception:
                    continue
    return sorted(interrupted, key=lambda x: x.get("timestamp", ""), reverse=True)


def render_auto_resume_banner():
    """Zobrazí banner při startu aplikace, pokud existuje přerušená práce
    z checkpointů (nikoli z _temp/ — to má vlastní banner výše)."""
    interrupted = find_interrupted_operations()
    if not interrupted:
        return
    if st.session_state.get("_auto_resume_banner_dismissed"):
        return

    latest = interrupted[0]
    _L_loc = st.session_state.get("lang", "cz")
    with st.container(border=True):
        st.warning(
            tt(f"🔄 **Nalezena přerušená operace:** {latest['operation']} "
               f"({latest.get('timestamp', 'nedávno')})",
               f"🔄 **Interrupted operation found:** {latest['operation']} "
               f"({latest.get('timestamp', 'recently')})", _L_loc)
        )
        col1, col2, col3 = st.columns([3, 2, 2])
        with col1:
            if st.button(
                tt("▶️ Pokračovat v poslední operaci",
                   "▶️ Resume last operation", _L_loc),
                type="primary", width="stretch", key="resume_latest"
            ):
                st.session_state["resume_data"]      = latest.get("data")
                st.session_state["resume_operation"] = latest["operation"]
                op = latest["operation"].lower()
                if "workflow" in op:
                    keep_current_tab(6)
                elif any(x in op for x in ["extrakce", "extraction"]):
                    keep_current_tab(0)
                elif any(x in op for x in ["překlad", "translate"]):
                    keep_current_tab(1)
                st.session_state["_gsb_jump_to_tab"] = st.session_state.get("current_tab_index", 0)
                st.rerun()
        with col2:
            if st.button(
                tt("🗑️ Smazat checkpointy", "🗑️ Clear checkpoints", _L_loc),
                width="stretch", key="clear_checkpoints_banner"
            ):
                cp_dir = Path("_temp/checkpoints")
                if cp_dir.exists():
                    for _cf in cp_dir.glob("*.json"):
                        try: _cf.unlink()
                        except Exception: pass
                st.session_state["_auto_resume_banner_dismissed"] = True
                st.success(tt("✅ Checkpointy smazány", "✅ Checkpoints cleared", _L_loc))
                st.rerun()
        with col3:
            if st.button(tt("✕ Zavřít", "✕ Dismiss", _L_loc),
                         key="dismiss_auto_resume_banner"):
                st.session_state["_auto_resume_banner_dismissed"] = True
                st.rerun()


def safe_threaded_operation(max_workers: int = 8, name: str = "worker"):
    """Decorator pro bezpečné paralelní operace s automatickým checkpointem při chybě.

    Použití:
        @safe_threaded_operation(max_workers=4, name="extraction")
        def process_item(item):
            ...

        results = process_item([item1, item2, item3])
    """
    def decorator(func):
        def wrapper(*args, **kwargs):
            try:
                items = args[0] if args and isinstance(args[0], list) else []
                with ThreadPoolExecutor(
                    max_workers=max_workers,
                    thread_name_prefix=name
                ) as pool:
                    futures = {pool.submit(func, item): item for item in items}
                    results = []
                    for fut in as_completed(futures):
                        try:
                            results.append(fut.result())
                        except Exception as _fe:
                            safe_set(f"{name}_last_error", str(_fe))
                            save_checkpoint(name, {
                                "error":   str(_fe),
                                "partial": str(futures[fut])[:200],
                            })
                    return results
            except Exception as e:
                safe_set(f"{name}_last_error", str(e))
                save_checkpoint(name, {"error": str(e), "partial": str(kwargs)[:200]})
                try:
                    st.error(f"Chyba v {name}: {e}")
                except Exception:
                    pass
                raise
        return wrapper
    return decorator


# ══════════════════════════════════════════════════════
# v24.24 – FULL STABILITY & PRODUCTIVITY PACK
# ══════════════════════════════════════════════════════

# ── 5. Persistentní nastavení per-model ───────────────────────────────────────
PER_MODEL_SETTINGS_FILE = "per_model_settings.json"

def load_per_model_settings() -> Dict:
    if Path(PER_MODEL_SETTINGS_FILE).exists():
        try:
            return json.loads(Path(PER_MODEL_SETTINGS_FILE).read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}

def save_per_model_settings(settings: Dict):
    tmp = Path(PER_MODEL_SETTINGS_FILE + ".tmp")
    tmp.write_text(json.dumps(settings, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(Path(PER_MODEL_SETTINGS_FILE))


# ── 2. Granulární per-file resume ─────────────────────────────────────────────
PARTIAL_RESULTS_DIR = Path("_temp/partial")
PARTIAL_RESULTS_DIR.mkdir(parents=True, exist_ok=True)

def list_partial_results() -> List[Dict]:
    """Vrátí seznam všech partial results (seřazeno desc)."""
    results = []
    for f in PARTIAL_RESULTS_DIR.glob("*.json"):
        try:
            results.append(json.loads(f.read_text(encoding="utf-8")))
        except Exception:
            continue
    return sorted(results, key=lambda x: x.get("timestamp", ""), reverse=True)

def render_partial_resume_section():
    """UI sekce pro obnovu rozpracovaných souborů (volá se ze záložky Extrakce/Překlad)."""
    partials = list_partial_results()
    if not partials:
        return
    _L_loc = st.session_state.get("lang", "cz")
    st.subheader(tt("🔄 Rozpracované soubory", "🔄 Partial results", _L_loc))
    for p in partials[:8]:
        col1, col2, col3 = st.columns([4, 2, 2])
        with col1:
            st.write(f"**{p['filename']}** — {p['operation']} "
                     f"({p['progress']*100:.0f} %)")
        with col2:
            if st.button(tt("▶️ Pokračovat", "▶️ Resume", _L_loc),
                         key=f"res_{p['filename'][:30]}_{p['operation']}"):
                st.session_state["resume_partial"] = p
                st.rerun()
        with col3:
            if st.button(tt("🗑️ Smazat", "🗑️ Delete", _L_loc),
                         key=f"del_{p['filename'][:30]}_{p['operation']}"):
                safe_name = re.sub(r'[^\w.-]', '_', Path(p['filename']).stem)[:80]
                _pf = PARTIAL_RESULTS_DIR / f"{p['operation']}_{safe_name}.json"
                try:
                    _pf.unlink(missing_ok=True)
                except Exception:
                    pass
                st.rerun()


# ── 3. Centralizovaný logger + diagnostics ────────────────────────────────────
LOG_FILE = Path("app_errors.log")

def _setup_logger() -> logging.Logger:
    logger = logging.getLogger("lm_utility_pro")
    if not logger.handlers:      # přidej handler jen jednou (Streamlit reruns)
        logger.setLevel(logging.INFO)
        handler = logging.FileHandler(LOG_FILE, encoding="utf-8")
        handler.setFormatter(
            logging.Formatter('%(asctime)s | %(levelname)s | %(message)s')
        )
        logger.addHandler(handler)
    return logger

APP_LOGGER = _setup_logger()

def log_error(operation: str, message: str,
              exc: Exception = None, context: Dict = None):
    """Zapíše chybu do souboru, session_state a zobrazí st.error."""
    full_msg = f"[{operation}] {message}"
    if exc:
        APP_LOGGER.error(full_msg, exc_info=True)
    else:
        APP_LOGGER.error(full_msg)
    if "error_log" not in st.session_state:
        st.session_state["error_log"] = []
    st.session_state["error_log"].append({
        "timestamp": datetime.now().isoformat(),
        "operation": operation,
        "message":   message,
        "context":   context or {},
    })
    try:
        st.error(f"❌ {full_msg}")
    except Exception:
        pass

def log_info(operation: str, message: str):
    """Zapíše info zprávu do souboru (bez st.info)."""
    APP_LOGGER.info(f"[{operation}] {message}")

def show_diagnostics():
    """UI komponenta pro zobrazení error logu a diagnostiky (volá se ze záložky Nastavení)."""
    _L_loc = st.session_state.get("lang", "cz")
    if st.button(tt("📋 Zobrazit Error Log + Diagnostika",
                    "📋 Show Error Log + Diagnostics", _L_loc),
                 key="btn_show_diag"):
        err_log = st.session_state.get("error_log", [])
        if err_log:
            st.dataframe(pd.DataFrame(err_log), width="stretch")
            if st.button(tt("🗑️ Vymazat log", "🗑️ Clear log", _L_loc),
                         key="btn_clear_errlog"):
                st.session_state["error_log"] = []
                st.rerun()
        else:
            st.success(tt("✅ Žádné chyby v logu", "✅ No errors in log", _L_loc))
        # Systémové info
        with st.expander(tt("ℹ️ Systémové info", "ℹ️ System info", _L_loc)):
            st.code(
                f"Python:    {__import__('sys').version.split()[0]}\n"
                f"Log file:  {LOG_FILE.resolve()}\n"
                f"Temp dir:  {Path('_temp').resolve()}\n"
                f"Checkpoints: {len(list_checkpoints())}\n"
                f"Partials:  {len(list_partial_results())}\n"
            )


# ── 4. Background Task Queue (sidebar widget) ─────────────────────────────────
def add_task(task_id: str, name: str, operation: str, progress: float = 0.0):
    """Přidá nebo aktualizuje úlohu ve frontě."""
    queue = st.session_state.get("task_queue", [])
    queue = [t for t in queue if t["id"] != task_id]
    queue.append({
        "id":        task_id,
        "name":      name,
        "operation": operation,
        "progress":  progress,
        "started":   datetime.now().isoformat(),
    })
    st.session_state["task_queue"] = queue

def update_task(task_id: str, progress: float):
    """Aktualizuje progress úlohy."""
    for t in st.session_state.get("task_queue", []):
        if t["id"] == task_id:
            t["progress"] = min(1.0, max(0.0, progress))
            break

def remove_task(task_id: str):
    """Odstraní dokončenou/zastavenou úlohu."""
    st.session_state["task_queue"] = [
        t for t in st.session_state.get("task_queue", [])
        if t["id"] != task_id
    ]

def render_task_queue():
    """Zobrazí aktivní úlohy v sidebaru (volá se v bloku `with st.sidebar`)."""
    queue = st.session_state.get("task_queue", [])
    if not queue:
        return
    _L_loc = st.session_state.get("lang", "cz")
    with st.sidebar.expander(
        tt("🔄 **Aktivní úlohy**", "🔄 **Active tasks**", _L_loc), expanded=True
    ):
        for t in queue:
            st.progress(t["progress"],
                        text=f"{t['name']} ({t['operation']})")


# ══════════════════════════════════════════════════════
# v24.25 – FINAL POLISH & COMPLETION
# ══════════════════════════════════════════════════════

# ── 1. Unified Settings Dashboard ─────────────────────────────────────
def render_unified_settings():
    """Centrální přehled všech nastavení aplikace (per-model, chunk, logger, backup)."""
    _L_loc = st.session_state.get("lang", "cz")
    st.markdown(tt("#### ⚙️ Unified Settings Dashboard",
                   "#### ⚙️ Unified Settings Dashboard", _L_loc))
    search = st.text_input(
        tt("🔎 Hledat v nastaveních", "🔎 Search settings", _L_loc),
        placeholder=tt("teplota, blok, model…", "temperature, chunk, model…", _L_loc),
        key="unified_settings_search",
    )
    # Sestav přehled klíčových nastavení ze session_state
    _setting_keys = [
        ("lms_base_url",          tt("LM Studio URL",            "LM Studio URL",           _L_loc)),
        ("lms_max_concurrent",    tt("Max souběžných pož.",      "Max concurrent req.",      _L_loc)),
        ("ext_temp",              tt("Teplota extrakce",         "Extraction temperature",   _L_loc)),
        ("tr_chunk_target",       tt("Velikost bloku (překlad)", "Chunk size (translation)", _L_loc)),
        ("ext_overlap",           tt("Překryv bloků",            "Chunk overlap",            _L_loc)),
        ("taxa_per_chunk",        tt("Počet taxonů v bloku",     "Taxa per chunk",           _L_loc)),
        ("val_timeout",           tt("Časový limit validace (s)","Validation timeout (s)",   _L_loc)),
        ("lang",                  tt("Jazyk UI",                 "UI language",              _L_loc)),
    ]
    rows = []
    for key, label in _setting_keys:
        val = st.session_state.get(key, "—")
        if search and search.lower() not in label.lower() and search.lower() not in key.lower():
            continue
        rows.append({"Nastavení": label, "Klíč": key, "Hodnota": str(val)})
    if rows:
        st.dataframe(rows, width="stretch", hide_index=True)
    else:
        st.info(tt("Žádná nastavení neodpovídají hledání.", "No settings match your search.", _L_loc))


# ── 2. Operation Timeline ─────────────────────────────────────────────
def log_operation(operation: str, status: str = "ok", detail: str = ""):
    """Zaznamená operaci do timeline v session_state. Volej po každé dokončené operaci."""
    if "operation_history" not in st.session_state:
        st.session_state["operation_history"] = []
    st.session_state["operation_history"].append({
        "time":      datetime.now().strftime("%H:%M:%S"),
        "operation": operation,
        "status":    status,
        "detail":    detail,
    })
    # Limit: uchovej max 50 záznamů
    st.session_state["operation_history"] = st.session_state["operation_history"][-50:]

def render_operation_timeline():
    """Zobrazí posledních 12 operací jako timeline."""
    _L_loc = st.session_state.get("lang", "cz")
    hist = st.session_state.get("operation_history", [])
    if not hist:
        st.caption(tt("Zatím žádné operace.", "No operations yet.", _L_loc))
        return
    st.markdown(tt("#### 📜 Čas operace", "#### 📜 Operation Timeline", _L_loc))
    _STATUS_ICON = {"ok": "✅", "error": "❌", "warning": "⚠️", "info": "ℹ️"}
    for entry in reversed(hist[-12:]):
        icon = _STATUS_ICON.get(entry.get("status", "ok"), "•")
        detail = f" — {entry['detail']}" if entry.get("detail") else ""
        st.caption(f"{icon} `{entry['time']}` **{entry['operation']}**{detail}")


# ── 3. Smart Auto-Cleanup ─────────────────────────────────────────────
def smart_auto_cleanup():
    """Při startu automaticky uklidí temp soubory starší než 7 dní."""
    try:
        _temp_prune("all", keep_last=5)
        _cutoff = datetime.now().timestamp() - 7 * 86400  # 7 dní
        _removed = 0
        for _f in Path("_temp").rglob("*"):
            try:
                if _f.is_file() and _f.stat().st_mtime < _cutoff:
                    _f.unlink(missing_ok=True)
                    _removed += 1
            except Exception:
                pass
        if _removed:
            APP_LOGGER.info(f"Chytré čištění: odstraněno {_removed} starých souborů")
    except Exception as _e:
        APP_LOGGER.warning(f"Chytré čištění selhalo: {_e}")


# ── 4. Keyboard Shortcuts Cheat Sheet ────────────────────────────────
def show_keyboard_cheatsheet():
    """Zobrazí přehled klávesových zkratek v expanderu."""
    _L_loc = st.session_state.get("lang", "cz")
    with st.expander(tt("⌨️ Klávesové zkratky",
                        "⌨️ Keyboard shortcuts", _L_loc)):
        st.markdown(tt("""
**Globální zkratky:**
| Zkratka | Akce |
|---|---|
| **Alt+1–9** | Přepnutí záložek 1–9 |
| **Ctrl+Enter** | Odeslat zprávu v chatu |
| **Alt+1** | Záložka Extrakce |
| **Alt+2** | Záložka Překlad |
| **Alt+3** | Záložka Validace |
| **Alt+4** | Záložka Chat |
| **Alt+7** | Záložka Workflow |
""", """
**Global shortcuts:**
| Shortcut | Action |
|---|---|
| **Alt+1–9** | Switch to tab 1–9 |
| **Ctrl+Enter** | Send chat message |
| **Alt+1** | Extraction tab |
| **Alt+2** | Translation tab |
| **Alt+3** | Validation tab |
| **Alt+4** | Chat tab |
| **Alt+7** | Workflow tab |
""", _L_loc))


# ── 5. Factory Reset ─────────────────────────────────────────────────
def render_factory_reset():
    """UI pro factory reset — vymaže všechna lokální data aplikace."""
    _L_loc = st.session_state.get("lang", "cz")
    st.markdown(tt("#### 🗑️ Tovární nastavení", "#### 🗑️ Factory Reset", _L_loc))
    st.warning(tt("⚠️ Tovární nastavení smaže VŠECHNA lokální data (session, cache, checkpointy, slovníky). "
                  "Tato akce je nevratná.",
                  "⚠️ Factory reset will delete ALL local data (session, cache, checkpoints, glossaries). "
                  "This action cannot be undone.", _L_loc))
    _fr_confirm = st.checkbox(
        tt("Rozumím — chci smazat všechna data", "I understand — delete all data", _L_loc),
        key="factory_reset_confirm",
    )
    if st.button(
        tt("🗑️ Provést Tovární nastavení", "🗑️ Perform Factory Reset", _L_loc),
        type="secondary", key="factory_reset_btn",
        disabled=not _fr_confirm,
    ):
        _deleted = 0
        # JSON soubory v kořeni (session, per-model settings, templates, history...)
        for _p in Path(".").glob("*.json"):
            try: _p.unlink(missing_ok=True); _deleted += 1
            except Exception: pass
        # Temp soubory
        for _p in Path("_temp").rglob("*"):
            try:
                if _p.is_file(): _p.unlink(missing_ok=True); _deleted += 1
            except Exception: pass
        # Vymaž session state (kromě základní inicializace)
        for _k in list(st.session_state.keys()):
            if not _k.startswith("_st_"):
                del st.session_state[_k]
        st.success(tt(f"✅ Tovární nastavení dokončeno — smazáno {_deleted} souborů. Aplikace se restartuje.",
                      f"✅ Factory reset complete — deleted {_deleted} files. Restarting.",
                      _L_loc))
        _time.sleep(1)
        st.rerun()


# ── 5. Přidej záznamy do timeline při startu (volitelně) ──────────────
def _log_startup():
    log_operation("startup", "info", f"v24.25 · {datetime.now().strftime('%Y-%m-%d')}")


# ══════════════════════════════════════════════════════
# v24.26 – LLM-as-Judge + Semantic RAG Cache + Agentic Workflow
# ══════════════════════════════════════════════════════

# ── Semantic RAG Cache ────────────────────────────────────────────────────────
try:
    from sentence_transformers import SentenceTransformer as _SentenceTransformer
    import chromadb as _chromadb
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False

def _init_rag():
    """Lazy inicializace RAG klienta (jen pokud je dostupná knihovna)."""
    if not RAG_AVAILABLE:
        return
    if "rag_client" not in st.session_state:
        try:
            st.session_state["rag_client"]     = _chromadb.PersistentClient(path="./rag_cache")
            st.session_state["rag_collection"] = st.session_state["rag_client"].get_or_create_collection("extraction_cache")
            st.session_state["rag_model"]      = _SentenceTransformer("all-MiniLM-L6-v2")
        except Exception as _e:
            APP_LOGGER.warning(f"RAG init selhal: {_e}")

def rag_cache_store(key: str, text: str, result: Any):
    """Uloží výsledek extrakce do semantic cache (ChromaDB + sentence embeddings)."""
    if not RAG_AVAILABLE:
        return
    _init_rag()
    try:
        embedding = st.session_state["rag_model"].encode(text[:2000]).tolist()
        st.session_state["rag_collection"].upsert(
            documents=[text[:2000]],
            embeddings=[embedding],
            metadatas=[{"result": json.dumps(result, ensure_ascii=False, default=str), "key": key}],
            ids=[key],
        )
    except Exception as _e:
        APP_LOGGER.warning(f"rag_cache_store selhal: {_e}")

def rag_cache_search(query: str, top_k: int = 3) -> Optional[Any]:
    """Najde sémanticky podobné předchozí extrakce a vrátí výsledek s nejvyšším skóre."""
    if not RAG_AVAILABLE:
        return None
    _init_rag()
    try:
        embedding = st.session_state["rag_model"].encode(query[:2000]).tolist()
        results = st.session_state["rag_collection"].query(
            query_embeddings=[embedding], n_results=top_k
        )
        metas = (results.get("metadatas") or [[]])[0]
        if metas:
            return json.loads(metas[0]["result"])
    except Exception as _e:
        APP_LOGGER.warning(f"rag_cache_search selhal: {_e}")
    return None


# ── LLM-as-Judge + Auto Quality Scoring ──────────────────────────────────────
def judge_extraction_quality(records: list, original_text: str = "",
                              base_url: str = "", model: str = "") -> Dict:
    """Hodnotí kvalitu extrakce druhým LLM voláním.
    Vrátí dict: {'score': int, 'comment': str, 'issues': list}.
    Pokud je LLM Judge vypnutý nebo selže, vrátí výchozí skóre.
    """
    if not st.session_state.get("enable_llm_judge", False):
        return {"score": 75, "comment": "LLM Judge je vypnutý", "issues": []}
    if not base_url or not model:
        base_url = st.session_state.get("lms_base_url", "http://localhost:1234")
        model    = st.session_state.get("selected_model", "")
    if not model:
        return {"score": 60, "comment": "Žádný model není načten", "issues": []}

    judge_prompt = (
        f"Ohodnoť kvalitu extrakce taxonomických záznamů (0–100 bodů).\n"
        f"Počet záznamů: {len(records)}\n"
        f"První 3 záznamy: {json.dumps(records[:3], ensure_ascii=False)[:700]}\n\n"
        "Vrať POUZE JSON (bez markdown):\n"
        '{ "score": číslo, "comment": "krátké hodnocení", "issues": ["seznam problémů"] }'
    )
    try:
        resp = chat_completion(
            base_url, model,
            [{"role": "user", "content": judge_prompt}],
            temp=0.1, max_tokens=600,
        )
        # Odstraň případné ```json fences
        resp_clean = re.sub(r"```[a-z]*\n?|```", "", resp).strip()
        parsed = json.loads(resp_clean)
        return {
            "score":   int(parsed.get("score", 60)),
            "comment": str(parsed.get("comment", "")),
            "issues":  list(parsed.get("issues", [])),
        }
    except Exception as _e:
        APP_LOGGER.warning(f"judge_extraction_quality selhal: {_e}")
        return {"score": 60, "comment": f"Hodnocení selhalo: {_e}", "issues": []}


# ── Agentic / Self-correcting Workflow ───────────────────────────────────────
def run_agentic_workflow(text: str, goal: str = "Zpracuj pro Hyolitha databázi",
                          base_url: str = "", model: str = "",
                          extraction_prompt: str = "") -> Dict:
    """Jednoduchý agent: extrahuje → ohodnotí → případně opraví.

    Kroky:
      1. Extrakce JSON ze vstupního textu
      2. LLM Judge hodnocení kvality
      3. Self-correction pokud score < 70
    Vrátí: {'result': Any, 'log': List[str], 'quality': Dict}
    """
    if not base_url:
        base_url = st.session_state.get("lms_base_url", "http://localhost:1234")
    if not model:
        model = st.session_state.get("selected_model", "")
    if not extraction_prompt:
        extraction_prompt = st.session_state.get(
            "ext_prompt",
            "Extrahuj taxonomické záznamy jako JSON seznam objektů s poli: "
            "rod, druh, autor, rok, lokalita, stratigrafie, popis."
        )

    log: List[str] = [f"🤖 Agent spuštěn — cíl: {goal}"]

    if not model:
        log.append("❌ Žádný model není načten — agent přerušen")
        return {"result": [], "log": log, "quality": {"score": 0, "comment": "Chybí model", "issues": []}}

    # ── Krok 1: Extrakce ─────────────────────────────────────────────
    log.append("→ Krok 1: Extrakce dat z textu")
    try:
        ext_resp = chat_completion(
            base_url, model,
            [{"role": "system", "content": extraction_prompt},
             {"role": "user",   "content": text[:8000]}],
            temp=0.05, max_tokens=4000,
            stop=["```\n"],
        )
        ext_clean = re.sub(r"```[a-z]*\n?|```", "", ext_resp).strip()
        # Pokus o parsování JSON
        raw_result = json.loads(ext_clean) if ext_clean.startswith("[") else ext_clean
        records = raw_result if isinstance(raw_result, list) else []
        log.append(f"   ✅ Extrahováno {len(records)} záznamů")
    except Exception as _e:
        log.append(f"   ❌ Extrakce selhala: {_e}")
        return {"result": [], "log": log, "quality": {"score": 0, "comment": str(_e), "issues": []}}

    # ── Krok 2: LLM Judge ────────────────────────────────────────────
    log.append("→ Krok 2: Hodnocení kvality (LLM Judge)")
    quality = judge_extraction_quality(records, text, base_url, model)
    log.append(f"   📊 Skóre: {quality['score']}/100 — {quality['comment']}")
    if quality["issues"]:
        for _iss in quality["issues"]:
            log.append(f"   ⚠️ {_iss}")

    # ── Krok 3: Self-correction (skóre < 70) ─────────────────────────
    if quality.get("score", 100) < 70:
        log.append("→ Krok 3: Self-correction (nízká kvalita → oprava)")
        try:
            corr_resp = chat_completion(
                base_url, model,
                [{"role": "system", "content":
                    "Jsi expert na taxonomická data. Oprav chyby v JSON extrakci. "
                    "Vrať POUZE opravený JSON seznam, bez komentářů."},
                 {"role": "user", "content":
                    f"Problémy: {json.dumps(quality['issues'], ensure_ascii=False)}\n\n"
                    f"Původní extrakce:\n{json.dumps(records, ensure_ascii=False)[:3000]}"}],
                temp=0.05, max_tokens=4000,
            )
            corr_clean = re.sub(r"```[a-z]*\n?|```", "", corr_resp).strip()
            corrected = json.loads(corr_clean) if corr_clean.startswith("[") else records
            log.append(f"   ✅ Opraveno {len(corrected)} záznamů")
            raw_result = corrected
        except Exception as _e:
            log.append(f"   ⚠️ Oprava selhala ({_e}) — ponechávám původní výsledek")

    log.append("✅ Agent dokončen")
    return {"result": raw_result, "log": log, "quality": quality}


# ══════════════════════════════════════════════════════
# v24.27 – QUICK PRESETS (persistentní, uživatelsky definované)
# ══════════════════════════════════════════════════════

PRESETS_FILE = "quick_presets.json"

def load_quick_presets() -> Dict:
    """Načte uloženě presets ze souboru (per-user po aktivaci profilu)."""
    if Path(PRESETS_FILE).exists():
        try:
            return json.loads(Path(PRESETS_FILE).read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}

def save_quick_presets(presets: Dict):
    """Atomicky uloží presets na disk (per-user po aktivaci profilu)."""
    tmp = Path(PRESETS_FILE + ".tmp")
    tmp.parent.mkdir(parents=True, exist_ok=True)
    tmp.write_text(json.dumps(presets, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(Path(PRESETS_FILE))


def defer_apply_settings(settings: Dict):
    """
    Bezpečná alternativa k přímému zápisu do session_state pro klíče,
    které jsou zároveň klíče widgetů (temp_t, tr_parallel, wf_chunk_target atd.).
    Uloží hodnoty do staging klíče _pending_settings; ty se aplikují
    na začátku PŘÍŠTÍHO rerunu před inicializací widgetů — bez Streamlit chyby.
    """
    pending = st.session_state.get("_pending_settings", {})
    pending.update(settings)
    st.session_state["_pending_settings"] = pending


# Inicializace při startu
if "quick_presets" not in st.session_state:
    st.session_state["quick_presets"] = load_quick_presets()


# ── 6. Self-test / Health Check ───────────────────────────────────────────────
def run_health_check(base_url: str = "", model: str = ""):
    """Spustí sérii diagnostických testů a zobrazí výsledky."""
    _L_loc = st.session_state.get("lang", "cz")
    st.subheader(tt("🧪 Self-test / Health Check", "🧪 Self-test / Health Check", _L_loc))
    results: Dict[str, str] = {}
    with st.spinner(tt("Probíhá diagnostika…", "Running diagnostics…", _L_loc)):
        # Disk
        try:
            _tp = Path("_temp"); _tp.mkdir(exist_ok=True)
            (_tp / "_hc_test.txt").write_text("ok"); (_tp / "_hc_test.txt").unlink()
            results[tt("Disk (_temp/)", "Disk (_temp/)", _L_loc)] = "✅ OK"
        except Exception as e:
            results[tt("Disk (_temp/)", "Disk (_temp/)", _L_loc)] = f"❌ {e}"
        # SQLite cache
        try:
            _conn = sqlite3.connect(":memory:")
            _conn.execute("CREATE TABLE t (x TEXT)"); _conn.close()
            results["SQLite"] = "✅ OK"
        except Exception as e:
            results["SQLite"] = f"❌ {e}"
        # LM Studio ping
        if base_url:
            try:
                _r = requests.get(base_url.rstrip("/") + "/v1/models", timeout=3)
                results[tt("LM Studio API", "LM Studio API", _L_loc)] = (
                    f"✅ HTTP {_r.status_code}" if _r.status_code < 400 else f"⚠️ HTTP {_r.status_code}")
            except Exception as e:
                results[tt("LM Studio API", "LM Studio API", _L_loc)] = f"❌ {e}"
        else:
            results[tt("LM Studio API", "LM Studio API", _L_loc)] = tt(
                "⚠️ URL nezadána", "⚠️ URL not set", _L_loc)
        # Session state keys
        _key_count = len(st.session_state)
        results[tt("Klíče relace", "Session state keys", _L_loc)] = f"✅ {_key_count}"
        # Checkpoints
        results[tt("Kontrolní bod", "Checkpoints", _L_loc)] = str(len(list_checkpoints()))
        # Partial results
        results[tt("Částečný výsledek", "Partial results", _L_loc)] = str(len(list_partial_results()))

    st.dataframe(
        pd.DataFrame([{tt("Test", "Test", _L_loc): k,
                       tt("Výsledek", "Result", _L_loc): v}
                      for k, v in results.items()]),
        width="stretch",
    )
    if all("✅" in v for v in results.values()):
        st.success(tt("✅ Všechny testy prošly", "✅ All tests passed", _L_loc))
    else:
        st.warning(tt("⚠️ Některé testy selhaly", "⚠️ Some tests failed", _L_loc))


# ── 1. AI Prompt Optimizer ────────────────────────────────────────────────────
def optimize_prompt_with_llm(original_prompt: str, base_url: str,
                             model: str, task_type: str = "extraction") -> str:
    """Vylepší prompt pomocí LLM pro maximální přesnost a min. halucinací.
    Vrátí vylepšený prompt nebo původní při chybě.
    """
    if not model or not base_url:
        return original_prompt
    _L_loc = st.session_state.get("lang", "cz")
    system = tt(
        "Jsi expert na tvorbu promptů pro LLM. Vylepši následující prompt pro "
        "maximální přesnost a minimalizaci halucinací. Zachej jazyk promptu. "
        "Výstup: POUZE vylepšený prompt, bez komentářů.",
        "You are an expert prompt engineer. Improve the following prompt for "
        "maximum accuracy and minimum hallucinations. Keep the prompt language. "
        "Output: ONLY the improved prompt, no comments.", _L_loc
    )
    try:
        resp = chat_completion(
            base_url, model,
            [{"role": "system", "content": system},
             {"role": "user", "content":
              f"{tt('Úloha', 'Task', _L_loc)}: {task_type}\n\n"
              f"{tt('Původní prompt', 'Original prompt', _L_loc)}:\n{original_prompt}"}],
            temp=0.3, max_tokens=1200,
        )
        return (resp or "").strip() or original_prompt
    except Exception as e:
        log_error("optimize_prompt", str(e))
        return original_prompt


# ── Nightly Backup + Session Versioning ───────────────────────────────────────
BACKUP_DIR = Path("backups")
BACKUP_DIR.mkdir(exist_ok=True)

# ── Globální flag pro nightly backup (jen jednou za session) ─────────────────
if "_NIGHTLY_BACKUP_DONE_TODAY" not in st.session_state:
    st.session_state["_NIGHTLY_BACKUP_DONE_TODAY"] = False

def do_nightly_backup():
    """Noční záloha — spustí se maximálně jednou za session/den.
    Používá session flag (_NIGHTLY_BACKUP_DONE_TODAY) místo file.exists() kontroly."""
    if st.session_state["_NIGHTLY_BACKUP_DONE_TODAY"]:
        return   # v této session již proběhlo
    today = datetime.now().strftime("%Y-%m-%d")
    # Pokud ZIP ze dne existuje, označíme jako hotovo a skončíme
    if (BACKUP_DIR / f"session_backup_{today}.zip").exists():
        st.session_state["_NIGHTLY_BACKUP_DONE_TODAY"] = True
        return
    try:
        # Primární cesta: plný ZIP export
        zip_bytes = export_session_zip()
        (BACKUP_DIR / f"session_backup_{today}.zip").write_bytes(zip_bytes)
        log_info("nightly_backup", "ZIP záloha uložena")
        try:
            st.toast("💾 Noční záloha vytvořena", icon="✅")
        except Exception:
            pass
    except Exception as _e_zip:
        # Fallback: jednoduchý JSON
        try:
            _json_target = BACKUP_DIR / f"session_backup_{today}.json"
            _json_target.write_text(
                json.dumps({
                    "timestamp":               datetime.now().isoformat(),
                    "version":                 VERSION,
                    "last_extraction_text":    st.session_state.get("last_extraction_text", ""),
                    "last_validation_results": st.session_state.get("last_validation_results", []),
                    "hyolitha_export_records": st.session_state.get("hyolitha_export_records", []),
                    "translation_result":      st.session_state.get("translation_result", ""),
                }, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            log_info("nightly_backup", "JSON fallback záloha uložena")
            try:
                st.toast("💾 Noční záloha vytvořena (JSON)", icon="✅")
            except Exception:
                pass
        except Exception as _e_json:
            log_error("nightly_backup", f"Záloha selhala: {_e_zip} / {_e_json}")
            if "error_log" in st.session_state:
                st.session_state["error_log"].append({
                    "timestamp": datetime.now().isoformat(),
                    "operation": "nightly_backup",
                    "message":   str(_e_json),
                })
    finally:
        # Označit jako hotovo bez ohledu na výsledek — neopakovat v této session
        st.session_state["_NIGHTLY_BACKUP_DONE_TODAY"] = True


def save_session_version():
    """Uloží snapshot session_state s timestampem (pro verzování)."""
    try:
        ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = Path("_temp/session_versions") / f"session_v{ts}.json"
        path.parent.mkdir(exist_ok=True)
        data = {k: v for k, v in st.session_state.items()
                if not k.startswith("_") and isinstance(v, (str, int, float, bool, list, dict, type(None)))}
        path.write_text(json.dumps(data, ensure_ascii=False, default=str), encoding="utf-8")
    except Exception as e:
        log_error("save_session_version", str(e))


# ── Remote Access Protection (volitelné, výchozí vypnuto) ────────────────────
REMOTE_PASSWORD   = os.environ.get("LMU_PASSWORD", "")   # prázdné = ochrana vypnuta
REMOTE_AUTH_ENABLED = bool(REMOTE_PASSWORD)               # aktivuje se jen pokud je heslo nastaveno

def export_to_darwin_core(records: list) -> bytes:
    """Exportuje záznamy do Darwin Core CSV (UTF-8).
    Mapuje české sloupce na standardní DwC termíny.
    """
    if not records:
        return b""
    df = pd.DataFrame(records)
    mapping = {
        "rod":          "genus",
        "druh":         "specificEpithet",
        "autor":        "scientificNameAuthorship",
        "rok":          "year",
        "lokalita":     "locality",
        "stratigrafie": "verbatimEventDate",
    }
    df = df.rename(columns={k: v for k, v in mapping.items() if k in df.columns})
    return df.to_csv(index=False).encode("utf-8")

try:
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

# ── Zachování aktivní záložky (v24.18) ────────────────────────────────────────
def keep_current_tab(tab_index: int = None):
    """Uloží index aktuální záložky do session_state.
    v24.18: také udržuje 'active_tab_index' alias pro obnovení session.
    """
    if "current_tab_index" not in st.session_state:
        st.session_state["current_tab_index"] = 0
    if "active_tab_index" not in st.session_state:
        st.session_state["active_tab_index"] = 0
    if tab_index is not None:
        st.session_state["current_tab_index"] = tab_index
        st.session_state["active_tab_index"]  = tab_index

# ── OFF tlačítko — zavře záložku prohlížeče + ukončí CMD okno (v24.31) ────────
def shutdown_app():
    """Spolehlivě ukončí Streamlit proces + záložku prohlížeče + CMD okno."""
    st.session_state["shutdown_mode"] = True
    st.rerun()

if st.session_state.get("shutdown_mode"):
    st.html("""
    <div style="text-align:center; margin-top:80px; padding:40px;">
        <h2 style="color:#22c55e;">&#x2705; Aplikace byla úspěšně ukončena</h2>
        <p style="font-size:1.3em;">Záložka bude automaticky zavřena.</p>
    </div>
    """)
    # st.html() běží v sandboxed iframe → zavírám přes window.parent (skutečná záložka)
    st.html("""
    <script>
    setTimeout(function() {
        try { window.parent.open('', '_self', ''); window.parent.close(); } catch(e) {}
        try { window.top.open('', '_self', ''); window.top.close(); } catch(e) {}
        try { window.parent.location.replace('about:blank'); } catch(e) {}
        try { window.location.replace('about:blank'); } catch(e) {}
    }, 600);
    </script>
    """)
    import os, signal, sys, time as _time
    _time.sleep(1.5)
    try:
        import psutil
        _parent = psutil.Process(os.getpid()).parent()
        if _parent and _parent.name().lower() in ("cmd.exe", "powershell.exe", "pwsh.exe", "bash", "sh"):
            _parent.terminate()
    except Exception:
        pass
    try:
        os.kill(os.getpid(), signal.SIGTERM)
    except Exception:
        pass
    os._exit(0)
    st.stop()
try:
    from bs4 import BeautifulSoup as _BS4Check  # noqa: F401
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from langdetect import detect
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False
    detect = None

try:
    from pydantic import BaseModel, field_validator, ValidationError as PydanticValidationError
    from typing import Optional as PydOpt
    PYDANTIC_AVAILABLE = True

    class ExtractionRecord(BaseModel):
        """Pydantic model pro validaci extrahovaného taxonomického záznamu."""
        druh:         PydOpt[str] = None
        rod:          PydOpt[str] = None
        autor:        PydOpt[str] = None
        rok:          PydOpt[int] = None
        lokalita:     PydOpt[str] = None
        stratigrafie: PydOpt[str] = None
        popis:        PydOpt[str] = None

        @field_validator("rok", mode="before")
        @classmethod
        def validate_rok(cls, v):
            if v is None or v == "": return None
            try:
                yr = int(str(v).strip()[:4])
                if 1758 <= yr <= 2030:
                    return yr
            except (ValueError, TypeError):
                pass
            return None

        @field_validator("druh", "rod", mode="before")
        @classmethod
        def clean_taxon_name(cls, v):
            if v is None: return None
            s = str(v).strip()
            # Odstraň LLM artefakty
            if s.lower() in ("null","none","n/a","unknown","?","-",""):
                return None
            return s

    def validate_extraction_records(records: list) -> tuple:
        """Validuje seznam extrahovaných záznamů přes Pydantic.
        Vrátí (valid_records, errors_list).
        """
        valid, errors = [], []
        for i, rec in enumerate(records):
            if not isinstance(rec, dict):
                errors.append(f"Záznam {i}: není dict")
                continue
            try:
                validated = ExtractionRecord(**{k: v for k, v in rec.items()
                                                if k in ExtractionRecord.model_fields})
                d = validated.model_dump(exclude_none=True)
                # Zachovej extra pole (lokalita, popis atd.) i po validaci
                d.update({k: v for k, v in rec.items()
                           if k not in ExtractionRecord.model_fields})
                valid.append(d)
            except PydanticValidationError as e:
                errors.append(f"Záznam {i} ({rec.get('druh','?')}): {str(e)[:120]}")
                valid.append(rec)  # zachovej i nevalidní záznamy
        return valid, errors

except ImportError:
    PYDANTIC_AVAILABLE = False
    ExtractionRecord = None

    def validate_extraction_records(records: list) -> tuple:
        """Fallback bez Pydantic — jednoduchá sanitizace."""
        valid, errors = [], []
        NULL_VALS = {"null","none","n/a","unknown","?","-",""}
        for rec in records:
            if not isinstance(rec, dict):
                continue
            cleaned = {}
            for k, v in rec.items():
                if isinstance(v, str) and v.strip().lower() in NULL_VALS:
                    cleaned[k] = None
                else:
                    cleaned[k] = v
            # Normalizace roku
            if "rok" in cleaned and cleaned["rok"] is not None:
                try:
                    yr = int(str(cleaned["rok"])[:4])
                    cleaned["rok"] = yr if 1758 <= yr <= 2030 else None
                except (ValueError, TypeError):
                    cleaned["rok"] = None
            valid.append(cleaned)
        return valid, errors


# ══════════════════════════════════════════════════════
# LOKALIZACE (CZ / EN)
# ══════════════════════════════════════════════════════
# Lokalizační slovník CZ / EN
# Každý klíč = tuple (CZ, EN)

T = {
    # ── App ──────────────────────────────────────────────
    "app_title":          ("SciNexus", "SciNexus"),
    "page_icon":          ("🧠", "🧠"),

    # ── Sidebar ──────────────────────────────────────────
    "lang_toggle":        ("🇬🇧 English", "🇨🇿 Čeština"),
    "api_url":            ("LM Studio API URL", "LM Studio API URL"),
    "load_models":        ("🔄 Načíst modely", "🔄 Load models"),
    "models_ok":          ("✅ {n} modelů", "✅ {n} models"),
    "active_model":       ("Aktivní model", "Active model"),
    "no_model":           ("⚠️ Načti modely", "⚠️ Load models first"),
    "offline_mode":       ("🔌 Offline / demo", "🔌 Offline / demo"),
    "offline_help":       ("Simuluje LLM bez LM Studio", "Simulates LLM without LM Studio"),
    "val_timeout":        ("Časový limit validace (s)", "Validation timeout (s)"),
    "global_sys_prompt":  ("🌐 Globální systémový prompt", "🌐 Global system prompt"),
    "gsp_help":           ("Přidán ke každému LLM volání", "Prepended to every LLM call"),
    "prompt_profiles":    ("💾 Profily promptů", "💾 Prompt profiles"),
    "profile_name":       ("Název profilu", "Profile name"),
    "profile_placeholder":("např. Hyolitha_CZ_v1", "e.g. Hyolitha_EN_v1"),
    "save_profile":       ("💾 Uložit profil", "💾 Save profile"),
    "profile_saved":      ("Profil '{n}' uložen", "Profile '{n}' saved"),
    "load_profile":       ("Načíst profil", "Load profile"),
    "load_btn":           ("📂 Načíst", "📂 Load"),
    "profile_loaded":     ("Profil '{n}' načten", "Profile '{n}' loaded"),
    "cache_taxa":         ("🧬 Cache: {n} taxonů", "🧬 Cache: {n} taxa"),
    "glossaries_count":   ("📚 Slovníku: {n}", "📚 Glossaries: {n}"),
    "templates_count":    ("📝 Šablony: {n}", "📝 Templates: {n}"),
    "clear_cache":        ("🗑️ Vymazat cache validace", "🗑️ Clear validation cache"),
    "cache_cleared":      ("Cache vymazána", "Cache cleared"),
    "profiles_help":      ("Uloží GSP + šablony + slovníky", "Saves GSP + templates + glossaries"),

    # ── Tabs ─────────────────────────────────────────────
    "tab_extract":        ("🔍 Extrakce", "🔍 Extraction"),
    "tab_translate":      ("🌐 Překlad", "🌐 Translation"),
    "tab_validate":       ("🧬 Validace", "🧬 Validation"),
    "tab_chat":           ("🤖 DeepSeek", "🤖 DeepSeek"),
    "tab_clean":          ("🧹 Čištění dat", "🧹 Data Cleaning"),
    "tab_style":          ("✍️ Stylistika", "✍️ Style Polish"),
    "tab_search":         ("🔎 Vyhledávání", "🔎 Web Search"),
    "tab_code":           ("💻 Kódování", "💻 Coding"),
    "tab_workflow":       ("⚙️ Workflow", "⚙️ Workflow"),
    "tab_stats":          ("📊 Statistiky", "📊 Statistics"),
    "tab_history":        ("📜 Historie", "📜 History"),

    # ── Extraction ───────────────────────────────────────
    "ext_title":          ("### 🔍 Extrakce dat z dokumentů", "### 🔍 Data Extraction from Documents"),
    "ext_upload":         ("Nahrát soubory (PDF, DOCX, TXT)", "Upload files (PDF, DOCX, TXT)"),
    "ext_mode":           ("Režim extrakce", "Extraction mode"),
    "ext_temp":           ("Teplota", "Temperature"),
    "ext_maxtok":         ("Max tokenů výstupu", "Max output tokens"),
    "ext_prompt_label":   ("Extrakční prompt", "Extraction prompt"),
    "ext_template":       ("Šablona", "Template"),
    "ext_own_prompt":     ("— vlastní prompt —", "— custom prompt —"),
    "ext_pdf_settings":   ("⚙️ Rozšířené nastavení PDF + bloky", "⚙️ Advanced PDF + chunking settings"),
    "ext_pages":          ("Rozsah stran (např. 1-5,8)", "Page range (e.g. 1-5,8)"),
    "ext_ocr":            ("Použít OCR", "Use OCR"),
    "ext_chunk_size":     ("Velikost bloku (znaky)", "Block size (chars)"),
    "ext_sys_prompt":     ("🌐 Systémový prompt extrakce", "🌐 Extraction system prompt"),
    "ext_run":            ("🚀 Spustit extrakci", "🚀 Run extraction"),
    "ext_done":           ("Zpracováno {n} souborů / {b} bloků. Taxonů: {t}", "Processed {n} files / {b} blocks. Taxa: {t}"),
    "ext_export":         ("#### Export výsledků", "#### Export results"),
    "ext_table":          ("#### 📊 Extrahovaná tabulka", "#### 📊 Extracted table"),
    "ext_no_json":        ("Výstup není JSON – CSV export nedostupný", "Output is not JSON – CSV export unavailable"),
    "ext_hyolitha":       ("#### 🦪 Export pro databázi Hyolitha (NM Praha)", "#### 🦪 Export for Hyolitha DB (NM Prague)"),
    "ext_hyolitha_cap":   ("Mapuje pole na schéma Access databáze.", "Maps fields to Access database schema."),
    "ext_hyolitha_xlsx":  ("⬇️ Hyolitha XLSX (4 listy)", "⬇️ Hyolitha XLSX (4 sheets)"),
    "ext_hyolitha_csv":   ("⬇️ Hyolitha CSV ZIP (4 soubory)", "⬇️ Hyolitha CSV ZIP (4 files)"),
    "ext_hyolitha_none":  ("Hyolitha export dostupný po JSON extrakci.", "Hyolitha export available after JSON extraction."),
    "ext_preview_taxa":   ("Náhled: Taxa", "Preview: Taxa"),
    "ext_tpl_manage":     ("🗂️ Správa prompt šablon", "🗂️ Manage prompt templates"),
    "ext_tpl_note":       ("Šablony jsou trvalé.", "Templates are persistent."),
    "ext_tpl_name":       ("Název nové šablony", "New template name"),
    "ext_tpl_desc":       ("Popis", "Description"),
    "ext_tpl_lang":       ("Doporučený zdrojový jazyk", "Recommended source language"),
    "ext_tpl_prompt":     ("Text promptu", "Prompt text"),
    "ext_tpl_save":       ("💾 Uložit šablonu", "💾 Save template"),
    "ext_tpl_saved":      ("Šablona '{n}' uložena", "Template '{n}' saved"),
    "ext_tpl_delete":     ("Smazat šablonu", "Delete template"),
    "ext_tpl_del_btn":    ("🗑️ Smazat", "🗑️ Delete"),
    "ext_tpl_deleted":    ("Šablona '{n}' smazána", "Template '{n}' deleted"),

    # ── Translation ──────────────────────────────────────
    "tr_title":           ("### 🌐 Překlad textu", "### 🌐 Text Translation"),
    "tr_mode_single":     ("📄 Jeden text", "📄 Single text"),
    "tr_mode_batch":      ("📦 Dávkový (více souborů)", "📦 Batch (multiple files)"),
    "tr_mode_compare":    ("⚖️ Porovnat dva překlady", "⚖️ Compare two translations"),
    "tr_src_lang":        ("Zdrojový jazyk", "Source language"),
    "tr_tgt_lang":        ("Cílový jazyk", "Target language"),
    "tr_temp":            ("Teplota", "Temperature"),
    "tr_preserve":        ("Zachovat latinské termíny", "Preserve Latin terms"),
    "tr_glossary":        ("Aktivní slovník (volitelně)", "Active glossary (optional)"),
    "tr_no_gloss":        ("— bez slovníku —", "— no glossary —"),
    "tr_src_manual":      ("✏️ Ručně", "✏️ Manual"),
    "tr_src_file":        ("📄 Ze souboru", "📄 From file"),
    "tr_src_extract":     ("🔗 Z extrakce", "🔗 From extraction"),
    "tr_input_label":     ("Text k překladu", "Text to translate"),
    "tr_upload":          ("Soubor", "File"),
    "tr_extract_preview": ("Náhled extrakce", "Extraction preview"),
    "tr_no_extract":      ("Nejprve proveď extrakci.", "Run extraction first."),
    "tr_run":             ("🌐 Přeložit", "🌐 Translate"),
    "tr_detected":        ("Detekovaný jazyk: **{l}**", "Detected language: **{l}**"),
    "tr_view_result":     ("📄 Překlad", "📄 Translation"),
    "tr_view_diff":       ("↔️ Srovnání", "↔️ Comparison"),
    "tr_result_label":    ("Přeložený text (editovatelný)", "Translated text (editable)"),
    "tr_refine":          ("#### ✨ Doladění překladu", "#### ✨ Translation Refinement"),
    "tr_ref_term":        ("🔬 Terminologie", "🔬 Terminology"),
    "tr_ref_style":       ("✍️ Styl", "✍️ Style"),
    "tr_ref_back":        ("🔄 Zpětný překlad", "🔄 Back-translation"),
    "tr_accept":          ("✅ Přijmout jako aktivní překlad", "✅ Accept as active translation"),
    "tr_diff_btn":        ("↔️ Diff", "↔️ Diff"),
    "tr_batch_upload":    ("Nahrát soubory k překladu", "Upload files for translation"),
    "tr_batch_fmt":       ("Formát výstupních souborů", "Output format"),
    "tr_batch_addorig":   ("Přidat originál vedle překladu", "Include original alongside translation"),
    "tr_batch_run":       ("📦 Spustit dávkový překlad", "📦 Run batch translation"),
    "tr_batch_done":      ("Přeloženo {ok}/{total} souborů", "Translated {ok}/{total} files"),
    "tr_batch_dl":        ("⬇️ Stáhnout ZIP ({n} souborů)", "⬇️ Download ZIP ({n} files)"),
    "tr_cmp_orig":        ("Originální text", "Original text"),
    "tr_cmp_a":           ("Překlad A", "Translation A"),
    "tr_cmp_b":           ("Překlad B", "Translation B"),
    "tr_cmp_label_a":     ("Popis A (volitelně)", "Label A (optional)"),
    "tr_cmp_label_b":     ("Popis B (volitelně)", "Label B (optional)"),
    "tr_cmp_criteria":    ("Hodnotící kritéria", "Evaluation criteria"),
    "tr_cmp_run":         ("⚖️ Porovnat překlady", "⚖️ Compare translations"),
    "tr_cmp_result":      ("### ⚖️ Výsledek porovnání", "### ⚖️ Comparison result"),
    "tr_gloss_section":   ("#### 📚 Správa slovníků", "#### 📚 Glossary management"),
    "tr_gloss_view":      ("Zobrazit / editovat", "View / edit"),
    "tr_gloss_new":       ("Nový slovník", "New glossary"),
    "tr_gloss_content":   ("Obsah (originál → překlad)", "Content (original → translation)"),
    "tr_gloss_save":      ("💾 Uložit změny", "💾 Save changes"),
    "tr_gloss_saved":     ("Uloženo", "Saved"),
    "tr_gloss_delete":    ("🗑️ Smazat slovník", "🗑️ Delete glossary"),
    "tr_gloss_deleted":   ("'{n}' smazán", "'{n}' deleted"),
    "tr_gloss_export":    ("⬇️ Exportovat JSON", "⬇️ Export JSON"),
    "tr_gloss_new_name":  ("Název slovníku", "Glossary name"),
    "tr_gloss_src":       ("Zdrojový jazyk slovníku", "Glossary source language"),
    "tr_gloss_tgt":       ("Cílový jazyk slovníku", "Glossary target language"),
    "tr_gloss_terms":     ("Termíny (originál → překlad)", "Terms (original → translation)"),
    "tr_gloss_import":    ("Nebo importovat z JSON", "Or import from JSON"),
    "tr_gloss_imported":  ("Importováno {n} termínů", "Imported {n} terms"),
    "tr_gloss_gen":       ("🤖 Generovat slovník (LLM)", "🤖 Generate glossary (LLM)"),
    "tr_gloss_no_ext":    ("Nejprve proveď extrakci.", "Run extraction first."),
    "tr_gloss_generated": ("Vygenerováno {n} termínů", "Generated {n} terms"),
    "tr_gloss_raw":       ("Surový výstup LLM", "Raw LLM output"),
    "tr_gloss_save_btn":  ("💾 Uložit nový slovník", "💾 Save new glossary"),
    "tr_no_gloss_yet":    ("Žádné slovníku. Vytvoř v záložce 'Nový slovník'.", "No glossaries. Create one in 'New glossary' tab."),
    "tr_active_gloss":    ("📚 Aktivní slovník **{n}**: {c} termínů – {t}", "📚 Active glossary **{n}**: {c} terms – {t}"),

    # ── Validation ───────────────────────────────────────
    "val_title":          ("### 🧬 Validace taxonomických jmen", "### 🧬 Taxonomic Name Validation"),
    "val_src_manual":     ("📝 Ruční", "📝 Manual"),
    "val_src_file":       ("📄 Ze souboru", "📄 From file"),
    "val_src_extract":    ("🔗 Z extrakce", "🔗 From extraction"),
    "val_input_label":    ("Taxony (jeden na řádek)", "Taxa (one per line)"),
    "val_col_select":     ("Sloupec s taxony", "Column with taxa"),
    "val_loaded":         ("Načteno {n} taxonů", "Loaded {n} taxa"),
    "val_no_extract":     ("Nejprve proveď extrakci.", "Run extraction first."),
    "val_edit_label":     ("Taxony z extrakce ({n} nalezeno) – uprav", "Taxa from extraction ({n} found) – edit"),
    "val_db_header":      ("#### 🗄️ Databáze", "#### 🗄️ Databases"),
    "val_preset":         ("Přednastavení", "Preset"),
    "val_max":            ("Max. taxonů", "Max. taxa"),
    "val_parallel":       ("⚡ Paralelní validace", "⚡ Parallel validation"),
    "val_run":            ("🚀 Spustit validaci", "🚀 Run validation"),
    "val_progress":       ("Validuji…", "Validating…"),
    "val_done":           ("✅ Hotovo", "✅ Done"),
    "val_high":           ("✅ Vysoká", "✅ High"),
    "val_medium":         ("⚠️ Střední", "⚠️ Medium"),
    "val_low":            ("❓ Nízká", "❓ Low"),
    "val_none":           ("❌ Nenalezeno", "❌ Not found"),
    "val_filter":         ("Zobrazit confidence", "Show confidence"),
    "val_showing":        ("Zobrazeno {s}/{t} taxonů", "Showing {s}/{t} taxa"),
    "val_detail":         ("📋 Podrobné výsledky po taxonech", "📋 Detailed results by taxon"),
    "val_dl_csv":         ("⬇️ CSV", "⬇️ CSV"),
    "val_dl_xlsx":        ("⬇️ Excel", "⬇️ Excel"),
    "val_dl_json":        ("⬇️ JSON", "⬇️ JSON"),
    "val_dup_found":      ("⚠️ Možné duplikáty / překlepy ({n} skupin)", "⚠️ Possible duplicates / typos ({n} groups)"),
    "val_dedup":          ("Automaticky deduplikovat", "Auto-deduplicate"),
    "val_deduped":        ("Deduplikováno: {n} unikátních taxonů", "Deduplicated: {n} unique taxa"),
    "val_no_dup":         ("✅ Žádné podezřelé duplikáty.", "✅ No suspicious duplicates."),

    # ── Chat ─────────────────────────────────────────────
    "chat_title":         ("### 💬 Chat s LLM", "### 💬 Chat with LLM"),
    "chat_sys_prompt":    ("Systémový prompt", "System prompt"),
    "chat_temp":          ("Teplota", "Temperature"),
    "chat_clear":         ("🗑️ Smazat historii", "🗑️ Clear history"),
    "chat_export":        ("⬇️ Exportovat konverzaci", "⬇️ Export conversation"),
    "chat_input":         ("Zpráva…", "Message…"),
    "chat_no_model":      ("Nejprve načti model.", "Load a model first."),

    # ── Data cleaning ────────────────────────────────────
    "clean_title":        ("### 🧹 Čištění a analýza dat", "### 🧹 Data Cleaning & Analysis"),
    "clean_upload":       ("Excel nebo CSV", "Excel or CSV"),
    "clean_loaded":       ("Načteno {r} řádků × {c} sloupců", "Loaded {r} rows × {c} columns"),
    "clean_dedup":        ("🔁 Odstranit duplicity", "🔁 Remove duplicates"),
    "clean_empty":        ("🧽 Prázdné řádky", "🧽 Empty rows"),
    "clean_trim":         ("✂️ Oříznout bílé znaky", "✂️ Trim whitespace"),
    "clean_dedup_done":   ("Odstraněno {n} duplicit", "Removed {n} duplicates"),
    "clean_empty_done":   ("Odstraněno {n} prázdných řádků", "Removed {n} empty rows"),
    "clean_trim_done":    ("Hotovo", "Done"),
    "clean_llm_norm":     ("🤖 LLM normalizace sloupce", "🤖 LLM column normalisation"),
    "clean_col":          ("Sloupec", "Column"),
    "clean_norm_inst":    ("Instrukce", "Instructions"),
    "clean_norm_run":     ("Normalizovat", "Normalise"),
    "clean_stats":        ("📊 Statistiky", "📊 Statistics"),
    "clean_dl":           ("⬇️ Stáhnout upravenou tabulku (CSV)", "⬇️ Download cleaned table (CSV)"),

    # ── Coding ───────────────────────────────────────────
    "code_title":         ("### 💻 Kódovací asistent", "### 💻 Coding Assistant"),
    "code_task":          ("Úkol", "Task"),
    "code_gen":           ("⚡ Generovat", "⚡ Generate"),
    "code_analyze":       ("🔍 Analyzovat", "🔍 Analyse"),
    "code_debug":         ("🐛 Debug", "🐛 Debug"),
    "code_explain":       ("💬 Vysvětlit", "💬 Explain"),
    "code_lang":          ("Jazyk", "Language"),
    "code_input":         ("Popis nebo kód", "Description or code"),
    "code_run":           ("▶️ Spustit", "▶️ Run"),

    # ── Workflow ─────────────────────────────────────────
    "wf_title":           ("### ⚙️ Workflow – řetězení úkolů", "### ⚙️ Workflow – task chaining"),
    "wf_step1":           ("1. Vstupní soubor", "1. Input file"),
    "wf_upload":          ("Zdrojový dokument", "Source document"),
    "wf_pages":           ("Rozsah stran PDF", "PDF page range"),
    "wf_step2":           ("2. Kroky procesu", "2. Pipeline steps"),
    "wf_add_step":        ("Přidat krok", "Add step"),
    "wf_add_btn":         ("➕ Přidat", "➕ Add"),
    "wf_remove":          ("🗑️ Odebrat: {s}", "🗑️ Remove: {s}"),
    "wf_clear":           ("🗑️ Smazat celý proces", "🗑️ Clear entire pipeline"),
    "wf_no_steps":        ("Přidej alespoň jeden krok.", "Add at least one step."),
    "wf_step3":           ("3. Parametry", "3. Parameters"),
    "wf_ext_prompt":      ("Prompt pro Extrakci", "Extraction prompt"),
    "wf_tgt_lang":        ("Cílový jazyk Překladu", "Translation target language"),
    "wf_valid_dbs":       ("Databáze pro Validaci", "Validation databases"),
    "wf_strat_llm":       ("Stratigrafická normalizace: LLM (pomalejší, přesnější)", "Stratigraphic normalisation: LLM (slower, more accurate)"),
    "wf_step4":           ("4. Spuštění", "4. Run"),
    "wf_run":             ("▶️ Spustit celý workflow", "▶️ Run entire workflow"),
    "wf_result":          ("### 📊 Výsledek workflow", "### 📊 Workflow result"),
    "wf_output":          ("Výstup", "Output"),
    "wf_log":             ("📋 Log", "📋 Log"),
    "wf_val_section":     ("#### Validace taxonů", "#### Taxon validation"),
    "wf_hyolitha":        ("#### 🦪 Hyolitha export ({n} záznamů)", "#### 🦪 Hyolitha export ({n} records)"),
    "wf_xlsx":            ("⬇️ XLSX (4 listy)", "⬇️ XLSX (4 sheets)"),
    "wf_csv":             ("⬇️ CSV ZIP", "⬇️ CSV ZIP"),

    # ── Statistics ───────────────────────────────────────
    "stats_title":        ("### 📊 Statistiky a přehled relace", "### 📊 Statistics & Session Overview"),
    "stats_cache":        ("Taxonů v cache", "Taxa in cache"),
    "stats_chat":         ("Zpráv v chatu", "Chat messages"),
    "stats_glossaries":   ("Slovníku", "Glossaries"),
    "stats_templates":    ("Prompt šablony", "Prompt templates"),
    "stats_last_val":     ("#### 🧬 Poslední validace", "#### 🧬 Last validation"),
    "stats_conf_dist":    ("**Rozložení jistoty**", "**Confidence distribution**"),
    "stats_db_cov":       ("**Pokrytí databázemi**", "**Database coverage**"),
    "stats_not_found":    ("❌ Nenalezené ({n})", "❌ Not found ({n})"),
    "stats_last_tr":      ("#### 🌐 Poslední překlad", "#### 🌐 Last translation"),
    "stats_orig_chars":   ("Znaky originálu", "Original chars"),
    "stats_tr_chars":     ("Znaky překladu", "Translated chars"),
    "stats_ratio":        ("Poměr délky", "Length ratio"),
    "stats_tokens":       ("Odhad tokenů", "Estimated tokens"),
    "stats_ops":          ("#### 📜 Operace", "#### 📜 Operations"),
    "stats_glossaries_h": ("#### 📚 Slovníku", "#### 📚 Glossaries"),

    # ── History ──────────────────────────────────────────
    "hist_title":         ("### 📜 Historie operací", "### 📜 Operation History"),
    "hist_empty":         ("Historie je prázdná.", "History is empty."),
    "hist_time":          ("Čas", "Time"),
    "hist_op":            ("Operace", "Operation"),
    "hist_detail":        ("Detail", "Detail"),
    "hist_clear":         ("🗑️ Smazat historii", "🗑️ Clear history"),
    "hist_cleared":       ("Smazáno", "Cleared"),

    # ── Misc ─────────────────────────────────────────────
    "txt_dl":             ("📄 TXT", "📄 TXT"),
    "json_dl":            ("📋 JSON", "📋 JSON"),
    "docx_dl":            ("📝 DOCX", "📝 DOCX"),
    "csv_dl":             ("⬇️ CSV", "⬇️ CSV"),
    "xlsx_dl":            ("⬇️ Excel (.xlsx)", "⬇️ Excel (.xlsx)"),
    "footer":             ("Extrakce · Překlad · Validace · Workflow · Chat",
                           "Extraction · Translation · Validation · Workflow · Chat"),

    # ── Další záložky ──────────────────────────────────────
    "tab_abtest":         ("🧪 A/B Test", "🧪 A/B Test"),
    "tab_settings":       ("⚙️ Nastavení", "⚙️ Settings"),
    "tab_deepseek":       ("💬 LMS Chat", "💬 LMS Chat"),
    "tab_help":           ("❓ Nápověda", "❓ Help"),

    # ── Sidebar: model connection ─────────────────────────
    "sb_connect":            ("🔄 Připojit", "🔄 Connect"),
    "sb_connect_help":       ("Načte aktuálně aktivní model z LM Studia",
                              "Loads the active model from LM Studio API"),
    "sb_connecting":         ("Připojuji…", "Connecting…"),
    "sb_list_models":        ("📋 Načíst modely", "📋 List models"),
    "sb_list_models_on":     ("Načte seznam všech stažených modelů",
                              "Lists all downloaded models via lms ls"),
    "sb_list_models_off":    ("lms CLI není dostupný — nainstaluj LM Studio CLI",
                              "lms CLI not available — install LM Studio CLI"),
    "sb_loading_list":       ("Načítám seznam modelů…", "Loading model list…"),
    "sb_loaded_n_models":    ("✅ Načteno {n} modelů", "✅ Loaded {n} models"),
    "sb_active_model":       ("▶️ Aktivní model", "▶️ Active model"),
    "sb_switch_model":       ("🔀 Přepnout model", "🔀 Switch model"),
    "sb_lms_missing":        ("⚠️ lms CLI není dostupný. Stáhni na: lmstudio.ai/docs/cli",
                              "⚠️ lms CLI not available. Download at: lmstudio.ai/docs/cli"),
    "sb_click_list":         ("Klikni **📋 Načíst modely** pro zobrazení stažených modelů",
                              "Click **📋 List models** to show downloaded models"),
    "sb_debug_lms":          ("🔍 Debug — výstup lms ls", "🔍 Debug — lms ls output"),
    "sb_debug_hint":         ("Pokud vidíš modely, ale seznam je prázdný, napiš mi výstup a opravím parsování.",
                              "If you see models but list is empty, share this output and I'll fix parsing."),
    "sb_select_load":        ("Vyber model ke spuštění", "Select model to load"),
    "sb_model_active":       ("✅ Tento model je právě aktivní",
                              "✅ This model is currently active"),
    "sb_loading":            ("⏳ Načítám…", "⏳ Loading…"),
    "sb_load_model":         ("🚀 Načíst model", "🚀 Load model"),
    "sb_load_model_help":    ("Uvolní aktuální model a načte vybraný model",
                              "Unloads current model and loads selected via lms load"),
    "sb_unload":             ("🗑️ Uvolnit", "🗑️ Unload"),
    "sb_unload_help":        ("Uvolní aktuálně načtený model z paměti",
                              "Frees the currently loaded model from RAM"),
    "sb_unloading":          ("Uvolňuji…", "Unloading…"),
    "sb_unloaded":           ("✅ Model uvolněn", "✅ Model unloaded"),
    "sb_loading_status":     ("🚀 Načítám {m}…", "🚀 Loading {m}…"),
    "sb_unloading_current":  ("Uvolňuji aktuální model…", "Unloading current model…"),
    "sb_loading_new":        ("Načítám {m}…", "Loading {m}…"),
    "sb_loading_takes":      ("_(může trvat 1–3 minuty podle velikosti)_",
                              "_(may take 1–3 minutes depending on size)_"),
    "sb_loaded_ok":          ("✅ Model načten!", "✅ Model loaded!"),
    "sb_load_error":         ("❌ Chyba načítání", "❌ Load error"),
    "sb_model_active_toast": ("✅ {m} je aktivní", "✅ {m} is active"),

    # ── Sidebar: parallel requests ────────────────────────
    "sb_parallel":           ("⚡ Paralelní požadavky", "⚡ Parallel requests"),
    "sb_max_concurrent":     ("Počet souběžných LLM požadavků", "Max concurrent LLM requests"),
    "sb_par_help":           ("LM Studio 0.4+ continuous batching.\n"
                              "Nastav stejnou hodnotu v LM Studio:\n"
                              "  Developer → Server → Max Concurrent Predictions\n\n"
                              "1 = sériově (bezpečné, pomalé)\n"
                              "4 = doporučený paralelní limit pro LM Studio\n"
                              "Využito u: překlad, extrakce, validace",
                              "LM Studio 0.4+ continuous batching.\n"
                              "Set the same value in LM Studio:\n"
                              "  Developer → Server → Max Concurrent Predictions\n\n"
                              "1 = serialized (safe, slow)\n"
                              "4 = default LM Studio limit — recommended\n"
                              "Used in: translation, extraction, validation"),
    "sb_par_serial":         ("⚪ Sériově", "⚪ Serial"),
    "sb_par_recommended":    ("🟢 4× (doporučeno)", "🟢 4× (recommended)"),

    # ── Sidebar: session persistence ──────────────────────
    "sb_session_persist":    ("💾 Zachování relace", "💾 Session persistence"),
    "sb_last_backup":        ("Poslední záloha: {d}", "Last saved: {d}"),
    "sb_save":               ("💾 Uložit", "💾 Save"),
    "sb_save_help":          ("Uloží výsledky validace, extrakce a překladu na disk",
                              "Saves validation, extraction and translation results to disk"),
    "sb_load":               ("📂 Načíst", "📂 Load"),
    "sb_load_help":          ("Načte poslední uloženou relaci z disku",
                              "Loads the last saved session from disk"),
    "sb_session_saved":      ("✅ Relace uložena", "✅ Session saved"),
    "sb_save_error":         ("❌ Chyba ukládání", "❌ Save error"),
    "sb_loaded_keys":        ("✅ Načteno {n} klíčů ({d})", "✅ Loaded {n} keys ({d})"),
    "sb_no_session":         ("⚠️ Žádná uložená relace", "⚠️ No saved session"),
    "sb_autosave_ago":       ("🔄 Autosave: před {n} min", "🔄 Autosave: {n} min ago"),
    "sb_autosave_time":      ("🔄 Autosave: {t}", "🔄 Autosave: {t}"),
    "sb_autosave_none":      ("🔄 Autosave: zatím nic", "🔄 Autosave: nothing yet"),

    # ── Sidebar: stats grid ───────────────────────────────
    "sb_stat_valcache":      ("🧬 Cache validace", "🧬 Validation cache"),
    "sb_stat_glossaries":    ("📚 Slovníky", "📚 Glossaries"),
    "sb_stat_templates":     ("📝 Šablony", "📝 Templates"),
    "sb_stat_fts":           ("🔎 FTS index", "🔎 FTS index"),
    "sb_stat_checkpoints":   ("💾 Kontrolní body", "💾 Checkpoints"),

    # ── Sidebar: hint + operations ────────────────────────
    "sb_hint":               ("⚙️ Model manager, prompty → záložka <b>Nastavení</b>",
                              "⚙️ Model manager, prompts → <b>Settings</b> tab"),
    "sb_op_running":         ("⚙️ {l}", "⚙️ {l}"),
}

LANG_INDEX = {"cz": 0, "en": 1}

def t(key: str, lang: str = "cz", **kwargs) -> str:
    """Vrátí přeložený řetězec pro daný klíč a jazyk."""
    pair = T.get(key, (key, key))
    idx  = LANG_INDEX.get(lang, 0)
    s    = pair[idx] if idx < len(pair) else pair[0]
    if kwargs:
        for k, v in kwargs.items():
            s = s.replace("{" + k + "}", str(v))
    return s


def tt(cz: str, en: str, lang: str = None, **kwargs) -> str:
    """Inline překlad bez nutnosti rozšiřovat slovník T.

    Použití místo ternárů:
        # místo:
        label = "Spustit extrakci" if _L == "cz" else "Run extraction"
        # napiš:
        label = tt("Spustit extrakci", "Run extraction", _L)

    Když se nezadá lang, vezme se ze session_state["lang"].
    Podporuje {placeholder}y přes kwargs, stejně jako t().
    """
    if lang is None:
        try:
            lang = st.session_state.get("lang", "cz")
        except Exception:
            lang = "cz"
    s = en if lang == "en" else cz
    if kwargs:
        for k, v in kwargs.items():
            s = s.replace("{" + k + "}", str(v))
    return s


# ══════════════════════════════════════════════════════
# CSS TEMA (dark / light – full sidebar override)
# ══════════════════════════════════════════════════════
# Kompletní dark/light theme CSS
# Přepisuje Streamlit defaulty včetně sidebaru, inputů, tlačítek, tabů

@st.cache_data(show_spinner=False)
def get_theme_css() -> str:
    """Vrátí CSS pro světlý (light) režim. Tmavý režim byl odstraněn.
    Cachováno — CSS se generuje jen jednou per session, ne při každém rerun."""
    bg_main   = "#f8fafc"
    bg_side   = "#f1f5f9"
    bg_card   = "#ffffff"
    bg_input  = "#ffffff"
    bg_hover  = "#e2e8f0"
    bg_btn    = "#2563b0"
    bg_btn_h  = "#1d4ed8"
    bg_tab_a  = "#2563b0"
    border    = "#cbd5e1"
    txt_main  = "#1e293b"
    txt_muted = "#64748b"
    txt_head  = "#0f172a"
    txt_link  = "#2563b0"
    accent1   = "#059669"
    accent2   = "#d97706"
    accent3   = "#4f46e5"
    accent4   = "#dc2626"
    accent5   = "#ea580c"
    scrollbar = "#cbd5e1"
    return f"""
<style>
/* ══════════════════════════════════════════════════════
   RESET & ROOT
══════════════════════════════════════════════════════ */
:root {{
    --bg-main:   {bg_main};
    --bg-side:   {bg_side};
    --bg-card:   {bg_card};
    --bg-input:  {bg_input};
    --border:    {border};
    --txt-main:  {txt_main};
    --txt-muted: {txt_muted};
    --txt-head:  {txt_head};
    --txt-link:  {txt_link};
    --accent1:   {accent1};
    --accent2:   {accent2};
    --accent3:   {accent3};
    --accent4:   {accent4};
    --accent5:   {accent5};
    --radius:    12px;
    --shadow:    0 4px 20px rgba(15,23,42,{'0.08'});
}}

/* ── App background ─────────────────────────────────── */
.stApp, .stApp > div {{
    background-color: {bg_main} !important;
    color: {txt_main} !important;
}}

/* ── TOP BAR / HEADER / TOOLBAR (světlý pruh nahoře) ── */
header[data-testid="stHeader"],
header[data-testid="stHeader"] > div,
header[data-testid="stHeader"] > div > div {{
    background-color: {bg_main} !important;
    border-bottom: 1px solid {border} !important;
}}
/* Streamlit toolbar tlačítka nahoře (Deploy, Settings, menu) */
header[data-testid="stHeader"] button,
header[data-testid="stHeader"] a,
[data-testid="stToolbar"],
[data-testid="stToolbar"] > div,
[data-testid="stDeployButton"],
[data-testid="stDecoration"] {{
    background-color: {bg_main} !important;
    color: {txt_muted} !important;
    border-color: {border} !important;
}}
[data-testid="stDecoration"] {{
    background: {'linear-gradient(90deg, #2563b0 0%, #f1f5f9 100%)'} !important;
    height: 3px !important;
}}
/* Streamlit running man / status */
[data-testid="stStatusWidget"] {{
    background-color: {bg_main} !important;
    color: {txt_muted} !important;
}}
/* Hamburger / main menu */
#MainMenu, #MainMenu button {{
    background-color: {bg_main} !important;
    color: {txt_main} !important;
}}
/* Footer */
footer, footer * {{
    background-color: {bg_main} !important;
    color: {txt_muted} !important;
    border-color: {border} !important;
}}
/* Block container - hlavní oblast pod headerem */
.stMainBlockContainer,
[data-testid="stAppViewBlockContainer"],
.stAppViewBlockContainer > div,
[data-testid="stMain"],
[data-testid="stVerticalBlock"] {{
    background-color: {bg_main} !important;
}}
/* Formuláře */
[data-testid="stForm"] {{
    background-color: {bg_card} !important;
    border: 1px solid {border} !important;
    border-radius: var(--radius) !important;
}}
/* Popover / tooltip */
[data-testid="stPopover"],
[data-baseweb="popover"] > div,
[data-baseweb="tooltip"] > div {{
    background-color: {bg_card} !important;
    border: 1px solid {border} !important;
    color: {txt_main} !important;
}}
/* Modal dialogy */
[data-testid="stModal"] > div,
[data-baseweb="dialog"] {{
    background-color: {bg_card} !important;
    border: 1px solid {border} !important;
}}
/* Notifications / stAlert */
[data-testid="stNotification"] {{
    background-color: {bg_card} !important;
    border: 1px solid {border} !important;
}}
/* Number input tlačítka +/- */
[data-baseweb="input"] button {{
    background-color: {bg_input} !important;
    color: {txt_main} !important;
    border-color: {border} !important;
}}
/* Multiselect tagy */
[data-baseweb="tag"] {{
    background-color: {'#dbeafe'} !important;
    color: {'#1d4ed8'} !important;
}}
[data-baseweb="tag"] span {{
    color: {'#1d4ed8'} !important;
}}
/* Slider track */
[data-baseweb="slider"] [role="slider"] {{
    background-color: {bg_btn} !important;
    border-color: {bg_btn} !important;
}}
/* Radio buttons */
[data-testid="stRadio"] label,
[data-testid="stRadio"] p {{
    color: {txt_main} !important;
}}
/* Checkbox */
[data-testid="stCheckbox"] label,
[data-testid="stCheckbox"] p {{
    color: {txt_main} !important;
}}
/* stSelectbox dropdown background */
[data-baseweb="select"] [data-baseweb="popover"] li {{
    background-color: {bg_card} !important;
    color: {txt_main} !important;
}}
[data-baseweb="select"] [data-baseweb="popover"] li:hover,
[data-baseweb="menu"] li:hover {{
    background-color: {bg_hover} !important;
}}
/* File uploader */
[data-testid="stFileUploader"] {{
    background-color: {bg_card} !important;
    border: 2px dashed {border} !important;
    border-radius: var(--radius) !important;
}}
[data-testid="stFileUploader"] section,
[data-testid="stFileUploader"] > div {{
    background-color: {bg_card} !important;
}}
[data-testid="stFileUploader"] label,
[data-testid="stFileUploader"] span,
[data-testid="stFileUploader"] p,
[data-testid="stFileUploader"] button {{
    color: {txt_muted} !important;
    background-color: transparent !important;
}}
/* stSpinner */
[data-testid="stSpinner"] {{
    color: {txt_muted} !important;
}}
/* Vega/Plotly charts background */
[data-testid="stVegaLiteChart"],
[data-testid="stPlotlyChart"] {{
    background-color: {bg_card} !important;
    border-radius: var(--radius) !important;
}}
/* Column containers */
[data-testid="column"] {{
    background-color: transparent !important;
}}

/* ── SIDEBAR – full override ────────────────────────── */
section[data-testid="stSidebar"],
section[data-testid="stSidebar"] > div,
section[data-testid="stSidebar"] > div > div {{
    background-color: {bg_side} !important;
    border-right: 1px solid {border} !important;
}}
/* Symetrické mezery zleva i zprava — obsah sidebaru
   se nelepí na pravý okraj */
section[data-testid="stSidebar"] [data-testid="stSidebarUserContent"],
section[data-testid="stSidebar"] .block-container {{
    padding-left: 1.1rem !important;
    padding-right: 1.1rem !important;
}}
/* Sidebar collapse button */
section[data-testid="stSidebar"] button[kind="headerNoPadding"] {{
    background: {bg_side} !important;
    color: {txt_main} !important;
    border: 1px solid {border} !important;
}}
/* Sidebar scrollbar */
section[data-testid="stSidebar"] ::-webkit-scrollbar {{
    width: 4px;
}}
section[data-testid="stSidebar"] ::-webkit-scrollbar-thumb {{
    background: {scrollbar}; border-radius: 4px;
}}

/* ── ALL text in sidebar ────────────────────────────── */
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] div,
section[data-testid="stSidebar"] .stMarkdown {{
    color: {txt_main} !important;
}}
section[data-testid="stSidebar"] .stCaption,
section[data-testid="stSidebar"] small {{
    color: {txt_muted} !important;
}}

/* ── Sidebar inputs ─────────────────────────────────── */
section[data-testid="stSidebar"] input,
section[data-testid="stSidebar"] textarea,
section[data-testid="stSidebar"] select {{
    background-color: {bg_input} !important;
    color: {txt_main} !important;
    border: 1px solid {border} !important;
    border-radius: var(--radius) !important;
}}
section[data-testid="stSidebar"] [data-baseweb="select"] > div,
section[data-testid="stSidebar"] [data-baseweb="input"] > div {{
    background-color: {bg_input} !important;
    border-color: {border} !important;
}}
section[data-testid="stSidebar"] [data-baseweb="select"] span,
section[data-testid="stSidebar"] [data-baseweb="select"] div {{
    color: {txt_main} !important;
}}

/* ── Sidebar buttons ────────────────────────────────── */
section[data-testid="stSidebar"] .stButton > button {{
    background-color: {bg_btn} !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: var(--radius) !important;
    font-weight: 500 !important;
    transition: background .15s ease !important;
}}
section[data-testid="stSidebar"] .stButton > button:hover {{
    background-color: {bg_btn_h} !important;
}}
section[data-testid="stSidebar"] .stButton > button[kind="secondary"] {{
    background-color: {bg_card} !important;
    color: {txt_main} !important;
    border: 1px solid {border} !important;
}}

/* ── Sidebar expander ───────────────────────────────── */
section[data-testid="stSidebar"] details,
section[data-testid="stSidebar"] details > div {{
    background-color: {bg_card} !important;
    border: 1px solid {border} !important;
    border-radius: var(--radius) !important;
    color: {txt_main} !important;
}}
section[data-testid="stSidebar"] details summary,
section[data-testid="stSidebar"] details summary span {{
    color: {txt_main} !important;
    font-weight: 600 !important;
}}

/* ── Sidebar divider ────────────────────────────────── */
section[data-testid="stSidebar"] hr {{
    border-color: {border} !important;
    margin: .6rem 0 !important;
}}

/* ── Sidebar checkbox ───────────────────────────────── */
section[data-testid="stSidebar"] [data-testid="stCheckbox"] label {{
    color: {txt_main} !important;
}}
section[data-testid="stSidebar"] [data-testid="stSlider"] label,
section[data-testid="stSidebar"] [data-testid="stSlider"] p {{
    color: {txt_main} !important;
}}

/* ── Main area inputs ───────────────────────────────── */
input, textarea {{
    background-color: {bg_input} !important;
    color: {txt_main} !important;
    border: 1px solid {border} !important;
    border-radius: var(--radius) !important;
}}
[data-baseweb="select"] > div {{
    background-color: {bg_input} !important;
    border-color: {border} !important;
    color: {txt_main} !important;
    border-radius: var(--radius) !important;
}}
[data-baseweb="select"] li {{
    background-color: {bg_card} !important;
    color: {txt_main} !important;
}}
[data-baseweb="select"] li:hover {{
    background-color: {bg_hover} !important;
}}

/* ── Buttons ────────────────────────────────────────── */
/* Hlavní primární tlačítka */
.stButton > button,
.stButton > button[kind="primary"],
[data-testid="baseButton-primary"],
button[data-testid="baseButton-primary"] {{
    background-color: {bg_btn} !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: var(--radius) !important;
    font-weight: 500 !important;
    padding: .4rem .9rem !important;
    transition: background .15s, box-shadow .15s !important;
}}
.stButton > button:hover,
.stButton > button[kind="primary"]:hover {{
    background-color: {bg_btn_h} !important;
    box-shadow: var(--shadow) !important;
}}
/* Sekundární tlačítka */
.stButton > button[kind="secondary"],
[data-testid="baseButton-secondary"],
button[data-testid="baseButton-secondary"] {{
    background-color: {bg_card} !important;
    color: {txt_main} !important;
    border: 1px solid {border} !important;
}}
.stButton > button[kind="secondary"]:hover {{
    background-color: {bg_hover} !important;
}}
/* Tertiary / minimal tlačítka */
.stButton > button[kind="tertiary"],
[data-testid="baseButton-tertiary"] {{
    background-color: transparent !important;
    color: {txt_main} !important;
    border: 1px solid {border} !important;
}}
/* Icon-only tlačítka (expander, copy) */
button[kind="icon"],
[data-testid="baseButton-icon"] {{
    background-color: transparent !important;
    color: {txt_muted} !important;
    border: none !important;
}}
button[kind="icon"]:hover {{
    background-color: {bg_hover} !important;
    color: {txt_main} !important;
}}
/* Download button */
.stDownloadButton > button,
[data-testid="baseButton-download"] {{
    background-color: {'#e6f4ea'} !important;
    color: {accent1} !important;
    border: 1px solid {accent1} !important;
    border-radius: var(--radius) !important;
}}
.stDownloadButton > button:hover {{
    background-color: {'#c8e8cc'} !important;
}}
/* Form submit */
button[kind="primaryFormSubmit"] {{
    background-color: {bg_btn} !important;
    color: #ffffff !important;
}}
/* Veškeré zbývající tlačítka která Streamlit generuje dynamicky */
[data-testid^="baseButton-"]:not([data-testid="baseButton-icon"]) {{
    background-color: {bg_btn} !important;
    color: #ffffff !important;
    border-radius: var(--radius) !important;
}}

/* ── Tabs ───────────────────────────────────────────── */
.stTabs [data-baseweb="tab-list"] {{
    background-color: {bg_side} !important;
    border-radius: var(--radius) !important;
    padding: 4px 6px !important;
    gap: 4px !important;
    border: 1px solid {border} !important;
}}
.stTabs [data-baseweb="tab"] {{
    background-color: transparent !important;
    color: {txt_muted} !important;
    border-radius: 6px !important;
    padding: .35rem .7rem !important;
    font-weight: 500 !important;
    font-size: .85rem !important;
    border: none !important;
    transition: background .15s, color .15s !important;
}}
.stTabs [aria-selected="true"] {{
    background-color: {bg_tab_a} !important;
    color: #ffffff !important;
    font-weight: 700 !important;
}}
.stTabs [data-baseweb="tab"]:hover {{
    background-color: {bg_hover} !important;
    color: {txt_main} !important;
}}
.stTabs [data-baseweb="tab-panel"] {{
    background-color: {bg_main} !important;
    padding-top: 1rem !important;
}}
/* Inner sub-tabs (translation refinement etc.) */
.stTabs .stTabs [data-baseweb="tab-list"] {{
    background-color: {bg_card} !important;
}}
.stTabs .stTabs [aria-selected="true"] {{
    background-color: {'#3b82f6'} !important;
}}

/* ── Expanders ──────────────────────────────────────── */
details {{
    background-color: {bg_card} !important;
    border: 1px solid {border} !important;
    border-radius: var(--radius) !important;
}}
details summary {{
    color: {txt_main} !important;
    font-weight: 600 !important;
    padding: .5rem .75rem !important;
}}
details summary:hover {{
    background-color: {bg_hover} !important;
    border-radius: var(--radius) !important;
}}
details > div {{
    background-color: {bg_card} !important;
    border-top: 1px solid {border} !important;
    padding: .75rem !important;
}}

/* ── Cards / metrics ────────────────────────────────── */
[data-testid="metric-container"] {{
    background-color: {bg_card} !important;
    border: 1px solid {border} !important;
    border-radius: var(--radius) !important;
    padding: .75rem 1rem !important;
}}
[data-testid="metric-container"] label {{
    color: {txt_muted} !important;
    font-size: .8rem !important;
}}
[data-testid="metric-container"] [data-testid="stMetricValue"] {{
    color: {txt_head} !important;
    font-weight: 700 !important;
}}

/* ── Dataframe ──────────────────────────────────────── */
[data-testid="stDataFrame"], iframe {{
    background-color: {bg_card} !important;
    border: 1px solid {border} !important;
    border-radius: var(--radius) !important;
}}

/* ── Info / Success / Warning / Error boxes ─────────── */
[data-testid="stAlert"] {{
    border-radius: var(--radius) !important;
    border-left-width: 4px !important;
}}
.stAlert[data-baseweb="notification"][kind="info"] {{
    background: {'#eff6ff'} !important;
    border-color: {accent3} !important;
}}
.stAlert[data-baseweb="notification"][kind="success"] {{
    background: {'#f0fdf4'} !important;
    border-color: {accent1} !important;
}}
.stAlert[data-baseweb="notification"][kind="warning"] {{
    background: {'#fffbeb'} !important;
    border-color: {accent2} !important;
}}
.stAlert[data-baseweb="notification"][kind="error"] {{
    background: {'#fef2f2'} !important;
    border-color: {accent4} !important;
}}

/* ── Divider ────────────────────────────────────────── */
hr {{
    border-color: {border} !important;
    margin: .75rem 0 !important;
}}

/* ── Typography ─────────────────────────────────────── */
h1, h2, h3, h4 {{ color: {txt_head} !important; }}
p, li, label, span {{ color: {txt_main} !important; }}
.stCaption, small {{ color: {txt_muted} !important; }}
a {{ color: {txt_link} !important; }}
code {{
    background: {bg_card} !important;
    color: {'#1d4ed8'} !important;
    padding: 1px 5px !important;
    border-radius: 4px !important;
    font-size: .88em !important;
}}

/* ── Progress bar ───────────────────────────────────── */
[data-testid="stProgressBar"] > div {{
    background-color: {border} !important;
    border-radius: 4px !important;
}}
[data-testid="stProgressBar"] > div > div {{
    background: linear-gradient(90deg, {'#2563b0, #0ea5e9'}) !important;
    border-radius: 4px !important;
}}

/* ── Scrollbar (main) ───────────────────────────────── */
::-webkit-scrollbar {{ width: 5px; height: 5px; }}
::-webkit-scrollbar-track {{ background: {bg_main}; }}
::-webkit-scrollbar-thumb {{ background: {scrollbar}; border-radius: 4px; }}
::-webkit-scrollbar-thumb:hover {{ background: {border}; }}

/* ── Checkbox & Radio ───────────────────────────────── */
[data-testid="stCheckbox"] label,
[data-testid="stRadio"] label {{
    color: {txt_main} !important;
}}

/* ── Slider ─────────────────────────────────────────── */
[data-testid="stSlider"] label,
[data-testid="stSlider"] p {{
    color: {txt_main} !important;
}}
[data-testid="stSlider"] [data-baseweb="slider"] [role="slider"] {{
    background-color: {bg_btn} !important;
    border-color: {bg_btn} !important;
}}

/* ── File uploader ──────────────────────────────────── */
[data-testid="stFileUploader"] {{
    background: {bg_card} !important;
    border: 2px dashed {border} !important;
    border-radius: var(--radius) !important;
}}
[data-testid="stFileUploader"] section {{
    background: {bg_card} !important;
}}
[data-testid="stFileUploader"] label,
[data-testid="stFileUploader"] span,
[data-testid="stFileUploader"] p {{
    color: {txt_muted} !important;
}}

/* ══════════════════════════════════════════════════════
   CUSTOM COMPONENTS
══════════════════════════════════════════════════════ */

/* ── Validation result cards ────────────────────────── */
.val-high {{
    background: {'#f0fdf4'};
    border-left: 4px solid {accent1};
    padding: .6rem 1rem; margin: .3rem 0; border-radius: var(--radius);
    color: {txt_main} !important;
}}
.val-medium {{
    background: {'#fffbeb'};
    border-left: 4px solid {accent2};
    padding: .6rem 1rem; margin: .3rem 0; border-radius: var(--radius);
    color: {txt_main} !important;
}}
.val-low {{
    background: {'#eef2ff'};
    border-left: 4px solid {accent3};
    padding: .6rem 1rem; margin: .3rem 0; border-radius: var(--radius);
    color: {txt_main} !important;
}}
.val-none {{
    background: {'#fef2f2'};
    border-left: 4px solid {accent4};
    padding: .6rem 1rem; margin: .3rem 0; border-radius: var(--radius);
    color: {txt_main} !important;
}}
.dup-warn {{
    background: {'#fff7ed'};
    border-left: 4px solid {accent5};
    padding: .5rem 1rem; margin: .2rem 0; border-radius: var(--radius);
    font-size: .85rem; color: {txt_main} !important;
}}

/* ── Stream output box ──────────────────────────────── */
.stream-box {{
    background: {bg_card};
    border: 1px solid {border};
    padding: 1rem; border-radius: var(--radius);
    font-family: 'JetBrains Mono', 'Fira Code', monospace;
    font-size: .85rem;
    white-space: pre-wrap;
    color: {txt_main};
    max-height: 400px;
    overflow-y: auto;
}}

/* ── Pipeline step badges ───────────────────────────── */
.step-badge {{
    display: inline-block;
    background: {bg_btn};
    color: #ffffff;
    border-radius: 20px;
    padding: 3px 12px;
    font-size: .78rem;
    font-weight: 600;
    margin: 2px;
    letter-spacing: .02em;
}}

/* ── Glossary tag ───────────────────────────────────── */
.gloss-tag {{
    display: inline-block;
    background: {'#dbeafe'};
    color: {'#1d4ed8'};
    border-radius: 6px;
    padding: 1px 8px;
    font-size: .78rem;
    margin: 2px;
}}

/* ── Section header accent ──────────────────────────── */
.section-header {{
    border-left: 4px solid {bg_btn};
    padding-left: .6rem;
    margin-bottom: .5rem;
    color: {txt_head};
    font-weight: 700;
}}

/* ── Tooltip / caption boxes ────────────────────────── */
.info-box {{
    background: {'#eef2ff'};
    border: 1px solid {'#c7d2fe'};
    border-radius: var(--radius);
    padding: .5rem .75rem;
    font-size: .85rem;
    color: {txt_main};
    margin: .4rem 0;
}}

/* ── Status pill ────────────────────────────────────── */
.status-pill {{
    display: inline-flex; align-items: center; gap: 5px;
    padding: 3px 10px; border-radius: 20px; font-size: .78rem;
    font-weight: 600; letter-spacing: .02em;
}}
.status-pill.ok   {{ background:{'#dcfce7'}; color:{accent1}; }}
.status-pill.warn {{ background:{'#fef3c7'}; color:{accent2}; }}
.status-pill.err  {{ background:{'#fee2e2'};  color:{accent4}; }}

/* ── Toast positioning fix ──────────────────────────── */
[data-testid="stToast"] {{
    bottom: 2rem !important;
    right: 1.5rem !important;
    border-radius: var(--radius) !important;
    box-shadow: 0 4px 16px rgba(0,0,0,.18) !important;
}}

/* ── Skeleton loader ────────────────────────────────── */
@keyframes shimmer {{
    0%   {{ background-position: -600px 0; }}
    100% {{ background-position: 600px 0; }}
}}
.skeleton {{
    background: linear-gradient(90deg,{border} 25%,{bg_hover} 50%,{border} 75%);
    background-size: 600px 100%;
    animation: shimmer 1.4s infinite;
    border-radius: 4px; height: 16px; margin: 4px 0;
}}

/* ── Table striping ─────────────────────────────────── */
[data-testid="stDataFrame"] tr:nth-child(even) > td {{
    background-color: {'rgba(0,0,0,.02)'} !important;
}}

/* ── Sidebar nav hint ───────────────────────────────── */
.sb-hint {{
    font-size: .73rem; color: {txt_muted}; padding: 2px 0;
    letter-spacing: .01em;
}}

/* ── Focus ring for accessibility ───────────────────── */
button:focus-visible, input:focus-visible, textarea:focus-visible {{
    outline: 2px solid {bg_btn} !important;
    outline-offset: 2px !important;
}}

/* ── Smooth tab transition ──────────────────────────── */
.stTabs [data-baseweb="tab-panel"] {{
    animation: fadeIn .18s ease;
}}
@keyframes fadeIn {{
    from {{ opacity: 0; transform: translateY(4px); }}
    to   {{ opacity: 1; transform: translateY(0); }}
}}

/* ── Compact expander in sidebar ───────────────────── */
section[data-testid="stSidebar"] details summary {{
    padding: .35rem .6rem !important;
    font-size: .83rem !important;
}}

/* ── Responzivní layout: 4K + wide monitory (v24.25) ─ */
@media (min-width: 1920px) {{
    .stMainBlockContainer {{
        max-width: 1680px !important;
        margin: 0 auto !important;
        padding: 2rem 3.5rem !important;
    }}
    .stApp {{ font-size: 16px !important; }}
    [data-testid="stDataFrame"] {{ font-size: .9rem !important; }}
}}
@media (min-width: 2560px) {{
    /* 2.5K / 4K monitory */
    .stMainBlockContainer {{
        max-width: 2200px !important;
        padding: 2.5rem 5rem !important;
    }}
    .stApp {{ font-size: 17px !important; }}
    .stButton > button {{
        font-size: 1rem !important;
        padding: .55rem 1.4rem !important;
    }}
    [data-baseweb="tab"] {{ font-size: .95rem !important; padding: .7rem 1.2rem !important; }}
}}

/* ── Compact layout: laptopy a menší obrazovky ───────── */
@media (max-width: 1280px) {{
    .stMainBlockContainer {{
        padding: 1rem 1.2rem !important;
    }}
    .stApp {{ font-size: 14px !important; }}
}}
@media (max-width: 900px) {{
    .stMainBlockContainer {{
        padding: .75rem .8rem !important;
    }}
    section[data-testid="stSidebar"] {{
        min-width: 220px !important;
        max-width: 280px !important;
    }}
}}
</style>
"""


# ══════════════════════════════════════════════════════
# KONSTANTY
# ══════════════════════════════════════════════════════
APP_TITLE = "SciNexus"
VERSION   = "1.0 SciNexus"

HISTORY_FILE        = "processing_history.json"
GLOSSARY_FILE       = "glossaries.json"
TEMPLATES_FILE      = "prompt_templates.json"
PROMPT_PROFILE_FILE = "prompt_profiles.json"
CLEAN_PROMPT_FILE   = "clean_prompts.json"

# ── Per-user adresáře: každý uživatel má vlastní složku se svými nastaveními ──
KNOWN_USERS_FILE    = "known_users.json"   # seznam uložených přihlašovacích jmen
USERS_BASE_DIR      = Path("users")        # users/<jméno>/

def _user_dir(username: str) -> Path:
    """Vrátí cestu k adresáři uživatele, vytvoří ho pokud neexistuje."""
    if not username or username.strip() == "":
        return Path(".")
    d = USERS_BASE_DIR / username.strip()
    d.mkdir(parents=True, exist_ok=True)
    return d

def load_known_users() -> list:
    """Načte seznam uložených přihlašovacích jmen."""
    try:
        if Path(KNOWN_USERS_FILE).exists():
            return json.loads(Path(KNOWN_USERS_FILE).read_text(encoding="utf-8"))
    except Exception:
        pass
    return []

def save_known_users(users: list):
    """Přidá jméno do seznamu známých uživatelů (bez duplikátů, max 30)."""
    try:
        tmp = Path(KNOWN_USERS_FILE + ".tmp")
        unique = list(dict.fromkeys(users))[:30]
        tmp.write_text(json.dumps(unique, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(Path(KNOWN_USERS_FILE))
    except Exception:
        pass

def activate_user_profile(username: str):
    """Přepne FILE konstanty na per-user cesty pro přihlášeného uživatele.
    Volá se jednou po přihlášení — přesměruje všechny globální souborové cesty."""
    global HISTORY_FILE, GLOSSARY_FILE, TEMPLATES_FILE
    global PROMPT_PROFILE_FILE, CLEAN_PROMPT_FILE
    global PRESETS_FILE, SESSION_AUTOSAVE_FILE, SESSION_PERSIST_FILE, EXTRACTION_HISTORY_DIR
    if not username or username.strip() == "":
        return
    _ud = _user_dir(username)
    HISTORY_FILE           = str(_ud / "processing_history.json")
    GLOSSARY_FILE          = str(_ud / "glossaries.json")
    TEMPLATES_FILE         = str(_ud / "prompt_templates.json")
    PROMPT_PROFILE_FILE    = str(_ud / "prompt_profiles.json")
    CLEAN_PROMPT_FILE      = str(_ud / "clean_prompts.json")
    PRESETS_FILE           = str(_ud / "quick_presets.json")
    SESSION_AUTOSAVE_FILE  = str(_ud / "session_autosave.json")
    SESSION_PERSIST_FILE   = str(_ud / "session_state.json")
    EXTRACTION_HISTORY_DIR = _ud / "extraction_history"

DEFAULT_BASE_URL  = os.getenv("LMSTUDIO_BASE_URL", "http://127.0.0.1:1234/v1")
MAX_PREVIEW_CHARS = 4000
DEFAULT_TIMEOUT   = 600
VAL_TIMEOUT       = 12   # timeout pro validační HTTP požadavky
MAX_HISTORY_ITEMS = 50

LANGUAGE_OPTIONS = [
    "čeština", "angličtina", "čínština", "francouzština",
    "japonština", "korejština", "latina", "němčina",
    "polština", "ruština", "slovenština", "španělština", "švédština",
    "arabština", "italština", "nizozemština", "portugalština", "turečtina",
    "ukrajinština",
]
LANGUAGE_TO_CODE = dict(zip(LANGUAGE_OPTIONS,
    ["cs","en","zh","fr","ja","ko","la","de","pl","ru","sk","es","sv",
     "ar","it","nl","pt","tr","uk"]))
CODE_TO_LANGUAGE = {v: k for k, v in LANGUAGE_TO_CODE.items()}
AUTO_LANGUAGE_LABEL = "🔍 Automaticky rozpoznat"

# Vlajky pro jazyky (emoji)
LANGUAGE_FLAGS: Dict[str, str] = {
    "čeština":       "🇨🇿",
    "angličtina":    "🇬🇧",
    "čínština":      "🇨🇳",
    "francouzština": "🇫🇷",
    "japonština":    "🇯🇵",
    "korejština":    "🇰🇷",
    "latina":        "🏛️",
    "němčina":       "🇩🇪",
    "polština":      "🇵🇱",
    "ruština":       "🇷🇺",
    "slovenština":   "🇸🇰",
    "španělština":   "🇪🇸",
    "švédština":     "🇸🇪",
    "arabština":     "🇸🇦",
    "italština":     "🇮🇹",
    "nizozemština":  "🇳🇱",
    "portugalština": "🇵🇹",
    "turečtina":     "🇹🇷",
    "ukrajinština":  "🇺🇦",
}

def lang_label(lang: str) -> str:
    """Vrátí jazyk s vlajkou: 🇨🇿 čeština"""
    flag = LANGUAGE_FLAGS.get(lang, "🌐")
    return f"{flag} {lang}"

LANGUAGE_OPTIONS_LABELED = [lang_label(l) for l in LANGUAGE_OPTIONS]
LANGUAGE_AUTO_LABELED    = "🔍 Automaticky rozpoznat"

def lang_from_label(label: str) -> str:
    """Ze '🇨🇿 čeština' vrátí 'čeština'"""
    parts = label.split(" ", 1)
    return parts[1] if len(parts) > 1 else label


EXTRACTION_MODES = {
    "Doslovně":  {"desc": "Přesná extrakce – žádná interpretace",
                  "system": "Extrahuj data PŘESNĚ, jak jsou uvedena. Neinterpretuj, nedoplňuj."},
    "Standard":  {"desc": "Standardní extrakce s normalizací",
                  "system": "Extrahuj data a normalizuj je (sjednoť formáty, oprav překlepy)."},
    "Obohacená":  {"desc": "Obohacená extrakce s kontextem",
                  "system": "Extrahuj data, normalizuj a doplň taxonomický/stratigrafický kontext."},
}

# Výchozí prompt šablony (budou obohaceny z disku při načtení)
DEFAULT_TEMPLATES = {
    "Hyolitha – obecná extrakce": {
        "desc": "Základní extrakce pro hyolitovou literaturu",
        "lang_hint": "angličtina",
        "prompt": "Extrahuj taxonomické záznamy jako JSON seznam objektů se sloupci: "
                  "rod, druh, autor, rok, lokalita, stratigrafie, formace, paleoregio. "
                  "Jedno occurrence = jeden objekt.",
    },
    "Čínské prameny (zh→en)": {
        "desc": "Extrakce z čínsky psaných paleontologických prací",
        "lang_hint": "čínština",
        "prompt": "Text je v čínštině. Přelož taxonomická jména do latinky (rodové a druhové jméno), "
                  "autora romanizuj (Pinyin). Extrahuj: род (rod), 种 (druh), 产地 (lokalita), "
                  "地层 (stratigrafie). Výstup jako JSON.",
    },
    "Ruské prameny (ru→en)": {
        "desc": "Extrakce z rusky psaných paleontologických prací",
        "lang_hint": "ruština",
        "prompt": "Text je v ruštině. Transliteruj jména autorů do latinky (BGN/PCGN). "
                  "Extrahuj: род, вид, автор, год, местонахождение, стратиграфия. "
                  "Výstup jako JSON seznam.",
    },
    "Francouzské prameny (fr→en)": {
        "desc": "Extrakce z francouzsky psaných paleontologických prací (Barrande a kol.)",
        "lang_hint": "francouzština",
        "prompt": "Texte en français. Extrais les données taxonomiques: genre, espèce, auteur, "
                  "année, localité, stratigraphie. Normalise les noms de localités. "
                  "Output as JSON list.",
    },
    "Bibliografická extrakce": {
        "desc": "Extrakce bibliografických citací z literatury",
        "lang_hint": "angličtina",
        "prompt": "Extrahuj bibliografické citace jako JSON: autor, rok, název_článku, "
                  "časopis, ročník, strany, DOI (pokud existuje). Jedna citace = jeden objekt.",
    },
}

TAXONOMIC_DATABASES = {
    # ── abecedně ──────────────────────────────────────────────────────────────
    "BioLib": {
        "name": "BioLib.cz",
        "search_url":  "https://www.biolib.cz/en/search/?q=",
        "taxon_url":   "https://www.biolib.cz/en/taxon/?action=search&string=",
        "browse_base": "https://www.biolib.cz",
        "type": "web",
        "default": False,
        "icon": "🌿",
        "desc": "Česká biologická databáze",
    },
    "CoL": {
        "name": "Catalogue of Life",
        "api_base":   "https://api.catalogueoflife.org/nameusage/search",
        "browse_url": "https://www.catalogueoflife.org/data/taxon/",
        "search_url": "https://www.catalogueoflife.org/data/search?q=",
        "type": "api",
        "default": False,
        "icon": "📖",
        "desc": "Globální katalog druhů",
    },
    "Fossilworks": {
        "name": "Fossilworks",
        "api_url":      "https://paleobiodb.org/data1.2/taxa/list.json?name=",
        "browse_url":   "https://fossilworks.org/?a=taxonInfo&taxon_no=",
        "fallback_url": "https://fossilworks.org/?a=taxonSearch&taxon_name=",
        "type": "api",
        "default": False,
        "icon": "🦴",
        "desc": "Veřejné rozhraní PaleoBioDB",
    },
    "GBIF": {
        "name": "GBIF",
        "api_match":  "https://api.gbif.org/v1/species/match",
        "api_search": "https://api.gbif.org/v1/species/search",
        "browse_url": "https://www.gbif.org/species/",
        "search_url": "https://www.gbif.org/species/search?q=",
        "type": "api",
        "default": True,
        "icon": "🌍",
        "desc": "Globální biodiverzitní databáze",
    },
    "IFPNI": {
        "name": "IFPNI",
        "api_search": "https://ifpni.org/api/v1/name/search?q=",
        "api_record": "https://ifpni.org/api/v1/name/",
        "browse_url": "https://ifpni.org/name/",
        "search_page": "https://ifpni.org/search?q=",
        "type": "api",
        "default": False,
        "icon": "🌱",
        "desc": "Mezinárodní index názvů fosilních rostlin",
    },
    "IPNI": {
        "name": "IPNI",
        "api_base":   "https://www.ipni.org/api/1/search",
        "browse_url": "https://www.ipni.org/n/",
        "type": "api",
        "default": False,
        "icon": "🌸",
        "desc": "Mezinárodní index názvů rostlin",
    },
    "IRMNG": {
        "name": "IRMNG",
        "search_url": "https://www.irmng.org/rest/AphiaRecordsByName/",
        "browse_url": "https://www.irmng.org/aphia.php?p=taxdetails&id=",
        "type": "api",
        "default": False,
        "icon": "🔬",
        "desc": "Registr mořských a nemořských rodů",
    },
    "ITIS": {
        "name": "ITIS",
        "api_base":  "https://www.itis.gov/ITISWebService/jsonservice/searchByScientificName",
        "api_any":   "https://www.itis.gov/ITISWebService/jsonservice/searchByScientificNameWildcard",
        "api_hier":  "https://www.itis.gov/ITISWebService/jsonservice/getFullHierarchyFromTSN",
        "browse_url":"https://www.itis.gov/servlet/SingleRpt/SingleRpt?search_topic=TSN&search_value=",
        "type": "api",
        "default": False,
        "icon": "🇺🇸",
        "desc": "Integrovaný taxonomický systém",
    },
    "Mikrotax": {
        "name": "Mikrotax",
        "pforams_url":  "https://www.mikrotax.org/pforams/index.php?taxon=",
        "nannotax_url": "https://www.mikrotax.org/Nannotax3/index.php?taxon=",
        "ostracod_url": "https://www.mikrotax.org/Ostracoda/index.php?taxon=",
        "browse_url":   "https://www.mikrotax.org/",
        "type": "web",
        "default": False,
        "icon": "🔭",
        "desc": "Mikrofosílie (nanofosílie, foraminifera, ostrakódi)",
    },
    "PaleoDB": {
        "name": "Paleobiology Database",
        "search_url": "https://paleobiodb.org/data1.2/taxa/list.json?name=",
        "single_url": "https://paleobiodb.org/data1.2/taxa/single.json?name=",
        "browse_url": "https://paleobiodb.org/classic/displayTaxonInfo?taxon_no=",
        "type": "api",
        "default": True,
        "icon": "🦕",
        "desc": "Paleontologická databáze fosilních taxonů",
    },
    "Plazi": {
        "name": "Plazi TreatmentBank",
        # GgServer: taxonomicName=; OpenBiodiv SPARQL fallback; browse přes treatmentbank.org
        "gg_url":     "https://tb.plazi.org/GgServer/search?taxonomicName=",
        "rest_url":   "https://tb.plazi.org/GgServer/search?q=",
        "browse_url": "https://treatmentbank.org/treatment?q=",
        "type": "api",
        "default": False,
        "icon": "📄",
        "desc": "Taxonomické popisy z literatury",
    },
    "Tropicos": {
        "name": "Tropicos (MBG)",
        # Tropicos REST API – Missouri Botanical Garden
        "api_search": "https://services.tropicos.org/Name/Search?name=",
        "api_suffix": "&type=wildcard&commonnames=false&format=json",
        "browse_url": "https://www.tropicos.org/name/",
        "search_page": "https://www.tropicos.org/Search?name=",
        "type": "api",
        "default": False,
        "icon": "🌺",
        "desc": "Botanická nomenklatura",
    },
    "WoRMS": {
        "name": "WoRMS",
        # AphiaRecordsByName: jméno musí být URL-encoded, marine_only=false
        # AphiaRecordsByVernacular: alternativa
        # AphiaSearch: fulltext, vrátí list i pro neexaktní shody
        "aphia_name":   "https://www.marinespecies.org/rest/AphiaRecordsByName/",
        "aphia_search": "https://www.marinespecies.org/rest/AphiaRecordsByNames",
        "browse_url":   "https://www.marinespecies.org/aphia.php?p=taxdetails&id=",
        "search_page":  "https://www.marinespecies.org/aphia.php?p=search&action=search&taxon=",
        "type": "api",
        "default": False,
        "icon": "🌊",
        "desc": "Světový rejstřík mořských druhů",
    },
    "ZooBank": {
        "name": "ZooBank",
        "api_search":  "https://zoobank.org/Search.json?search_term=",
        "api_names":   "https://zoobank.org/NomenclaturalActs.json?search_term=",
        "browse_url":  "https://zoobank.org/Names/",
        "type": "api",
        "default": False,
        "icon": "🐛",
        "desc": "Registr zoologické nomenklatury (ICZN)",
    },
    "GNverifier": {
        "name": "Global Names Verifier",
        # POST /api/v1/verifications — batch až 5000 jmen, 100+ zdrojů najednou
        # Preferované zdroje pro paleontologii: 1=CoL, 11=GBIF, 172=PaleoDB, 169=WoRMS,
        # 4=NCBI, 167=IRMNG, 3=ITIS, 8=Tropicos, 165=IPNI, 185=Plazi
        "api_url":    "https://verifier.globalnames.org/api/v1/verifications",
        # browse_url NESMÍ být použit jako fallback — generujeme URL ručně s ?capitalize=on
        "browse_url": None,
        "type": "api",
        "default": True,
        "icon": "🌐",
        "desc": "Global Names Verifier — 100+ databází najednou, fuzzy matching, best match score",
        # Preferované zdroje ID pro paleontologii (vrátí se i v preferredResults)
        "preferred_sources": [1, 11, 172, 169, 167, 3],
    },
}

DB_PRESETS = {
    "Hyolitha (doporučeno)": ["PaleoDB", "GBIF", "IRMNG", "Plazi", "GNverifier"],
    "Paleontologie":         ["PaleoDB", "Fossilworks", "GBIF", "IRMNG", "Plazi", "GNverifier"],
    "Mořské organismy":      ["WoRMS", "GBIF", "IRMNG", "GNverifier"],
    "Botanika":              ["IPNI", "IFPNI", "Tropicos", "CoL", "GBIF", "GNverifier"],
    "Zoologie":              ["GBIF", "ZooBank", "ITIS", "CoL", "WoRMS", "GNverifier"],
    "Všechny databáze":      list(TAXONOMIC_DATABASES.keys()),
}

# ICS stratigrafická nomenklatura – výběr nejčastějších synonym
ICS_SYNONYMS: Dict[str, str] = {
    # Kambrium
    "lower cambrian": "Cambrian Stage 2–3", "early cambrian": "Terreneuvian–Series 2",
    "middle cambrian": "Miaolingian", "upper cambrian": "Furongian",
    "late cambrian": "Furongian",
    # Ordovik
    "lower ordovician": "Early Ordovician", "middle ordovician": "Middle Ordovician",
    "upper ordovician": "Late Ordovician",
    # Silur
    "lower silurian": "Llandovery–Wenlock", "upper silurian": "Ludlow–Přídolí",
    # Devon
    "lower devonian": "Early Devonian", "middle devonian": "Middle Devonian",
    "upper devonian": "Late Devonian",
    # Formální jednotky – české synonymy
    "spodní kambrium": "Terreneuvian–Series 2",
    "střední kambrium": "Miaolingian",
    "svrchní kambrium": "Furongian",
    "spodní ordovik": "Early Ordovician",
    "střední ordovik": "Middle Ordovician",
    "svrchní ordovik": "Late Ordovician",
    "spodní silur": "Llandovery–Wenlock",
    "svrchní silur": "Ludlow–Přídolí",
    "spodní devon": "Early Devonian",
    "střední devon": "Middle Devonian",
    "svrchní devon": "Late Devonian",
}

# Pre-kompilované regexpy pro ICS_SYNONYMS — jednou při startu, ne při každém volání
# normalize_stratigraphy_local (54 termínů × volání per soubor = zbytečná re-kompilace)
_ICS_PATTERNS: List[Tuple[re.Pattern, str]] = [
    (re.compile(re.escape(orig), re.IGNORECASE), ics)
    for orig, ics in ICS_SYNONYMS.items()
]

# ══════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════
_defaults: Dict[str, Any] = {
    "chat_messages":          [],
    "validation_cache":       {},
    "workflow_steps":         [],
    "available_models":       [],
    "lms_model_list":         [],    # seznam modelů z `lms ls`
    "lms_active_model":       None,  # ID aktuálně načteného modelu
    "lms_loading":            False, # příznak probíhajícího načítání
    "last_extraction_text":   "",
    "last_extraction_taxa":   [],
    "last_validation_results":[],
    "translation_original":   "",
    "translation_result":     "",
    "translation_src_lang":   "",
    "translation_tgt_lang":   "",
    "tr_manual_text":         "",
    "tr_src_text":            "",
    "translation_iterations": [],
    "translation_ready":      False,
    "glossaries":             {},
    "prompt_templates":       {},
    "offline_mode":           False,
    "model_status":           "unknown",
    "val_timeout":            12,
    "prompt_profiles":        {},
    "global_system_prompt":   "Jsi vědecký asistent specializovaný na paleontologii, taxonomii a stratigrafii. Odpovídej přesně a odborně.",
    "hyolitha_export_records": [],
    "lang": "cz",
    "clean_prompts": {},
    "style_polish_result": "",
    "tr_src_lang_val": "🔍 Automaticky rozpoznat",
    "tr_tgt_lang_val": "angličtina",
    "search_history":      [],
    "search_results_raw":  [],
    "search_synthesis":    "",
    "search_query_last":   "",
    "val_selected_dbs":    [],
    # v18.0 — nové klíče
    "ab_results":          {},    # výsledky A/B testu
    "ext_pause_requested": False, # příznak pro checkpoint/pause
    "ext_overlap":         200,   # překryv mezi bloky (znaky)
    # v24.17 — nové klíče
    "translate_running":   False, # příznak aktivního překladu (Stop/Pause)
    "translate_paused":    False, # příznak pozastaveného překladu
    "tr_live_partial":     "",    # průběžný text živého překladu (přežije pauzu)
    "tr_chunk_queue":      None,  # v24.35: fronta chunků čekajících na překlad (list nebo None)
    "tr_chunk_total":      0,     # v24.35: celkový počet chunků aktuálního překladu
    "tr_sys_msg":          "",    # v24.35: sys_msg pro resume (uložen při startu překladu)
    "tr_actual_src":       "",    # v24.35: detekovaný zdrojový jazyk pro resume
    "tr_show_compare":     False, # zobrazit side-by-side porovnání
    "current_tab_index":   0,     # zachování aktivní záložky po operaci
    # v24.18 — nové klíče
    "active_tab_index":    0,     # alias pro obnovení relace (v24.18)
    "extraction_running":  False, # příznak aktivní extrakce (live okno)
    "taxa_per_chunk":      3,     # počet taxonů na blok (slider)
    # v24.19 — nové klíče
    "validation_running":  False, # příznak aktivní validace (Stop/Pause)
    "validation_paused":   False, # příznak pozastavené validace
    # v24.20 — nové klíče
    "resume_data":         None,  # data pro auto-resume operace
    "resume_operation":    "",    # název obnovované operace
    "_auto_resume_banner_dismissed": False,  # zamezuje opakovanému zobrazení
    # v24.26 — nové klíče
    "enable_llm_judge":    False,  # LLM-as-Judge hodnocení kvality extrakce
    # v24.24 — nové klíče
    "per_model_settings":  {},    # persistentní nastavení per-model (načte se níže)
    "task_queue":          [],    # fronta aktivních úloh (sidebar widget)
    "error_log":           [],    # in-memory error log
    "resume_partial":      None,  # data pro per-file resume
    "remote_authenticated": not REMOTE_AUTH_ENABLED,  # True pokud ochrana vypnuta
}
for _k, _v in _defaults.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ── v24.24: Načtení per_model_settings z disku při startu ────────────────────
if not st.session_state.get("per_model_settings"):
    st.session_state["per_model_settings"] = load_per_model_settings()

st.set_page_config(page_title="SciNexus v1.0", layout="wide", page_icon="🧠")

# ── Aplikuj téma (vždy světlý režim) ─────────────────────────
st.markdown(get_theme_css(), unsafe_allow_html=True)

# ══════════════════════════════════════════════════════
# POMOCNÉ FUNKCE – SOUBORY
# ══════════════════════════════════════════════════════
def preview_text(text: str, limit: int = MAX_PREVIEW_CHARS) -> str:
    return text[:limit] + ("\n[…zkráceno…]" if len(text) > limit else "")

def estimate_tokens(text: str) -> int:
    return max(1, len(text) // 4)

def read_txt(file) -> str:
    return file.read().decode("utf-8", errors="ignore")

def read_docx(file) -> str:
    from docx import Document  # lazy import — docx se načte jen pokud uživatel nahraje .docx
    return "\n".join(p.text for p in Document(file).paragraphs)

def parse_page_ranges(spec: str, total: int) -> List[int]:
    selected: set = set()
    for part in spec.split(','):
        part = part.strip()
        if '-' in part:
            a, b = map(int, part.split('-'))
            selected.update(range(max(1,a)-1, min(total,b)))
        elif part.isdigit():
            p = int(part)
            if 1 <= p <= total:
                selected.add(p-1)
    return sorted(selected)

def read_pdf(file, page_spec: str = None, use_ocr: bool = False, ocr_dpi: int = 300) -> str:
    if use_ocr and OCR_AVAILABLE:
        raw = file.read()
        images = convert_from_bytes(raw, dpi=ocr_dpi)
        total = len(images)
        pages = parse_page_ranges(page_spec, total) if page_spec else list(range(total))
        return "\n".join(
            f"--- Strana {p+1} ---\n{pytesseract.image_to_string(images[p], lang='ces+eng')}"
            for p in pages
        )
    from PyPDF2 import PdfReader  # lazy import — PyPDF2 se načte jen pokud uživatel nahraje PDF
    reader = PdfReader(file)
    total  = len(reader.pages)
    pages  = parse_page_ranges(page_spec, total) if page_spec else list(range(total))
    return "\n".join(
        f"--- Strana {p+1} ---\n{reader.pages[p].extract_text() or ''}"
        for p in pages
    )

def read_uploaded_file(file, page_spec: str = None, use_ocr: bool = False,
                        ocr_dpi: int = 300) -> str:
    """Přečte nahraný soubor s cache podle jméno+velikost (nečte znovu při rerunu).
    Cache klíč normalizuje page_spec přes parse_page_ranges — "1-3" a "1,2,3"
    dají stejný klíč a nebudou způsobovat zbytečné cache miss.
    """
    # Normalizace page_spec pro stabilní cache klíč
    _norm_spec = ""
    if page_spec:
        try:
            from PyPDF2 import PdfReader as _PdfReader  # lazy import
            _pr = _PdfReader(file) if Path(file.name).suffix.lower() == ".pdf" else None
            _total = len(_pr.pages) if _pr else 9999
            file.seek(0)
            _norm_spec = ",".join(str(p) for p in parse_page_ranges(page_spec, _total))
        except Exception:
            _norm_spec = page_spec.strip()
    cache_key = f"_file_cache_{file.name}_{file.size}_{_norm_spec}_{use_ocr}_{ocr_dpi}"
    if cache_key in st.session_state:
        file.seek(0)
        return st.session_state[cache_key]
    file.seek(0)
    ext = Path(file.name).suffix.lower()
    if ext == ".txt":   result = read_txt(file)
    elif ext == ".docx": result = read_docx(file)
    elif ext == ".pdf":  result = read_pdf(file, page_spec, use_ocr, ocr_dpi)
    else: raise ValueError(f"Nepodporovaný formát: {ext}")
    if len(result) < 10_000_000:
        st.session_state[cache_key] = result
    return result


# ══════════════════════════════════════════════════════
# OCR PREPROCESSING – čistění textu před LLM
# ══════════════════════════════════════════════════════
_PAGE_HEADER_RE = re.compile(
    r'(?m)^[ \t]*---\s*Strana\s*\d+\s*---[ \t]*$\n?'
)
_PAGE_NUM_RE    = re.compile(r'(?m)^\s*\d{1,4}\s*$\n?')
_BIBREF_RE      = re.compile(
    r'\b(?:doi|DOI):\s*\S+|\b\d{4}\.\s+[A-Z][^.]{5,80}\.\s+(?:In:|pp?\.|vol\.|ed\.|eds\.)',
    re.IGNORECASE
)
_REPEAT_HEADER_MIN = 3   # minimální počet opakování záhlaví pro detekci

# Pre-kompilované regexpy používané opakovaně v hot-path funkcích
# (kompilace 19× za runtime na každé volání je zbytečná)
_JSON_FENCE_RE   = re.compile(r"```json|```")          # strip markdown code fences
_MULTI_NEWLINE_RE = re.compile(r'\n{3,}')              # normalizace prázdných řádků
_PAGE_BATCH_RE   = re.compile(r'(--- Strana \d+ ---\n?)')  # dělení stránek v batched překladu

def preprocess_text_for_llm(text: str,
                              remove_page_markers: bool = True,
                              remove_page_numbers: bool = True,
                              remove_bibref: bool = False) -> Tuple[str, Dict]:
    """Lokálně vyčistí text před odesláním do LLM.
    Odstraní záhlaví stran, čísla stránek, opakující se záhlaví tabulek.
    Vrátí (vyčištěný_text, statistiky_oprav).
    """
    stats: Dict[str, int] = {}
    result = text

    if remove_page_markers:
        cleaned, n = _PAGE_HEADER_RE.subn('', result)
        if n:
            stats["page_markers"] = n
            result = cleaned

    if remove_page_numbers:
        cleaned, n = _PAGE_NUM_RE.subn('', result)
        if n:
            stats["page_numbers"] = n
            result = cleaned

    # Detekce opakujících se záhlaví tabulek pomocí Counter (rychlejší)
    lines = result.split('\n')
    line_counts = Counter(ln.strip() for ln in lines if len(ln.strip()) > 10)
    repeat_headers = {ln for ln, cnt in line_counts.items() if cnt >= _REPEAT_HEADER_MIN}
    if repeat_headers:
        kept = [ln for ln in lines if ln.strip() not in repeat_headers]
        stats["repeat_headers"] = len(lines) - len(kept)
        result = '\n'.join(kept)

    if remove_bibref:
        cleaned, n = _BIBREF_RE.subn('', result)
        if n:
            stats["bibref"] = n
            result = cleaned

    # Normalizace vícenásobných prázdných řádků
    result = _MULTI_NEWLINE_RE.sub('\n\n', result)
    # Rychlý odhad tokenů: len//4 místo split() — o ~10× rychlejší pro velké texty
    orig_tokens  = max(1, len(text) // 4)
    clean_tokens = max(1, len(result) // 4)
    stats["token_reduction_pct"] = round(100 * (1 - clean_tokens / orig_tokens), 1)

    return result, stats


# ══════════════════════════════════════════════════════
# AUTOMATICKÝ VÝBĚR CHUNK SIZE
# ══════════════════════════════════════════════════════
# Odhadované context windows pro běžné modely (znaky ≈ tokeny * 4)
_MODEL_CONTEXT = {
    # Hodnoty jsou v znacích (chars ≈ tokens × 3 pro latinku, × 1.5 pro CJK)
    # Context windows v tokenech → chars = tokens × 3
    "qwen2.5":    131072 * 3,   # 128k ctx → ~393k chars
    "qwen3":       32768 * 3,   # 32k ctx  → ~98k chars
    "mistral":     32768 * 3,   # 32k ctx  → ~98k chars
    "deepseek":    65536 * 3,   # 64k ctx  → ~196k chars
    "llama":        8192 * 3,   # 8k ctx   → ~24k chars
    "default":     16384 * 3,   # 16k ctx  → ~49k chars
}

# Cílové max_tokens pro výstup podle typu úlohy
_MAX_TOKENS_EXTRACT = 4096   # extrakce JSON — přísný strop
_MAX_TOKENS_TRANSLATE = 8000  # překlad — větší variabilita délky

# Stop sekvence pro různé typy úloh
_STOP_JSON   = ["\n```", "```\n"]   # zastavit za koncem JSON bloku
_STOP_TRANSL = []                    # překlad — bez stop sekvencí (volný text)

import functools as _functools

@_functools.lru_cache(maxsize=256)
def suggest_chunk_size(text_len: int, model_name: str = "") -> int:
    """Navrhne bezpečný chunk_size pro překlad.

    Vzorec:
      dostupné tokeny = context_window × 0.75   (25 % safety margin)
      dostupné tokeny -= SYS_PROMPT_OVERHEAD     (odečti system prompt)
      dostupné tokeny -= _MAX_TOKENS_TRANSLATE   (odečti max výstupní tokeny)
      max_chunk_chars  = dostupné tokeny × 3     (chars per token pro latinku)

    Zabrání "Context size has been exceeded" při překladu.
    """
    _SYS_OVERHEAD_TOKENS = 600   # odhad: system prompt + user wrapper + safety margin
    ctx_tokens = _MODEL_CONTEXT["default"] // 3   # přepočet zpět na tokeny
    mn = model_name.lower()
    for key, val in _MODEL_CONTEXT.items():
        if key in mn:
            ctx_tokens = val // 3
            break
    # Odečti overhead: safety margin 25 % + system prompt + výstupní tokeny
    available = int(ctx_tokens * 0.75) - _SYS_OVERHEAD_TOKENS - _MAX_TOKENS_TRANSLATE
    available = max(available, 1000)   # absolutní minimum
    max_chunk = available * 3          # tokeny → chars (latinková aproximace)
    min_chunk = 3_000
    max_chunk = max(min_chunk, max_chunk)
    if text_len <= max_chunk:
        return max_chunk
    ideal = max(min_chunk, text_len // 4)
    return max(min_chunk, min(ideal, max_chunk))


# ══════════════════════════════════════════════════════
# DOI RESOLVER
# ══════════════════════════════════════════════════════
def resolve_doi(doi: str, timeout: int = 6) -> Optional[Dict]:
    """Ověří DOI přes doi.org API a vrátí metadata.
    Vrátí dict s: valid, title, author, year, journal nebo None při chybě.
    """
    doi = doi.strip().lstrip("https://doi.org/").lstrip("doi:")
    try:
        url = f"https://doi.org/{doi}"
        r = _safe_get(url, timeout=timeout,
                      headers={"Accept": "application/vnd.citationstyles.csl+json"},
                      allow_redirects=True)
        if r.status_code == 200:
            try:
                data = r.json()
                authors = data.get("author", [])
                auth_str = "; ".join(
                    f"{a.get('family','')}, {a.get('given','')}" for a in authors[:3]
                ).strip("; ")
                issued = data.get("issued", {}).get("date-parts", [[None]])[0]
                year   = issued[0] if issued else None
                return {
                    "valid":   True,
                    "title":   data.get("title",""),
                    "author":  auth_str,
                    "year":    year,
                    "journal": data.get("container-title",""),
                    "doi":     doi,
                }
            except Exception:
                return {"valid": True, "doi": doi}
        elif r.status_code == 404:
            return {"valid": False, "doi": doi, "error": "DOI nenalezeno"}
    except Exception as e:
        return {"valid": None, "doi": doi, "error": str(e)[:60]}
    return None


# ══════════════════════════════════════════════════════
# ZOOBANK NAME REGISTRATION CHECK
# ══════════════════════════════════════════════════════
def check_zoobank_registration(taxon: str, timeout: int = 8) -> Optional[Dict]:
    """Ověří, zda je taxon registrován v ZooBank.
    Vrátí dict: {found, lsid, name, pub_date} nebo None.
    """
    try:
        url = (f"http://zoobank.org/Search.json"
               f"?q={requests.utils.quote(taxon)}&matchAny=false")
        r = _safe_get(url, timeout=timeout)
        if r.status_code != 200:
            return {"found": False, "error": f"HTTP {r.status_code}"}
        data = r.json()
        if isinstance(data, list) and data:
            hit = data[0]
            return {
                "found":    True,
                "lsid":     hit.get("lsid",""),
                "name":     hit.get("label",""),
                "pub_date": hit.get("pub_date",""),
                "url":      f"https://zoobank.org/{hit.get('lsid','')}",
            }
        return {"found": False}
    except Exception as e:
        return {"found": None, "error": str(e)[:60]}


# ══════════════════════════════════════════════════════
# REQUEST QUEUE – serializace LLM dotazů
# ══════════════════════════════════════════════════════
import threading as _threading

class _LLMConcurrencyManager:
    """Řídí souběžnost LLM požadavků pro LM Studio continuous batching.

    LM Studio 0.4+ podporuje parallel requests (Max Concurrent Predictions, default 4).
    Více souběžných požadavků = vyšší throughput díky continuous batching.

    max_concurrent=1  → původní serializace (bezpečné, pomalé)
    max_concurrent=4  → default LM Studio limit, optimální pro překlad chunků
    max_concurrent=8+ → pokud jsi v LM Studio zvýšil limit

    Stabilita v24:
    • set_max() bezpečně přidá/odebere tokeny semaforu bez ztráty čekajících threadů
    • _lock chrání konzistenci _max a semaforových operací
    """
    def __init__(self, max_concurrent: int = 1):
        self._lock = _threading.Lock()
        self._semaphore = _threading.Semaphore(max_concurrent)
        self._max = max_concurrent

    def set_max(self, n: int):
        """Bezpečně změní max souběžnost — přidá nebo odebere tokeny semaforu."""
        n = max(1, n)
        with self._lock:
            diff = n - self._max
            if diff > 0:
                for _ in range(diff):
                    self._semaphore.release()   # přidej tokeny
            elif diff < 0:
                # Blokující acquire() by mohlo deadlocknout — použijeme non-blocking
                # a snížíme _max; tokeny se dočerpají až dokončí běžící tasky
                removed = 0
                for _ in range(-diff):
                    if self._semaphore.acquire(blocking=False):
                        removed += 1
                # Pokud se nepodařilo odebrat všechny tokeny, odebereme je postupně
                # při příštím uvolnění (přes _pending_reduce)
                self._pending_reduce = (-diff) - removed
            else:
                self._pending_reduce = 0
            self._max = n

    @property
    def max_concurrent(self) -> int:
        return self._max

    def submit(self, fn, *args, **kwargs):
        """Čeká na volný slot, spustí fn, pak uvolní slot."""
        self._semaphore.acquire()
        try:
            return fn(*args, **kwargs)
        finally:
            # Pokud čekáme na snížení kapacity, pohltíme token místo uvolnění
            with self._lock:
                pending = getattr(self, "_pending_reduce", 0)
                if pending > 0:
                    self._pending_reduce = pending - 1
                    # token se „pohltí" — nesaháme na semafor
                else:
                    self._semaphore.release()


_llm_queue = _LLMConcurrencyManager(max_concurrent=1)  # start serializovaně, UI to změní


def chat_completion_queued(base_url, model, messages, **kwargs):
    """Verze chat_completion přes concurrency manager + presence fronta."""
    _job = presence.llm_start()
    try:
        result = _llm_queue.submit(chat_completion, base_url, model, messages, **kwargs)
    finally:
        presence.llm_done(_job)
    return result


def chunk_text(text: str, max_chars: int = 8000, overlap: int = 0) -> List[str]:
    """Dělí text na bloky po větách, každý ≤ max_chars znaků.
    Překryv: počet znaků z konce předchozího bloku přidaných na začátek dalšího
             (lepší kontext na hranicích; doporučeno 200–500 pro extrakci).
    """
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks, cur, cur_len = [], [], 0
    for s in sentences:
        if cur_len + len(s) > max_chars and cur:
            chunks.append(" ".join(cur))
            # Overlap: vezmi poslední věty do nového chunku jako kontext
            if overlap > 0:
                overlap_sents = []
                overlap_len   = 0
                for prev_s in reversed(cur):
                    if overlap_len + len(prev_s) > overlap:
                        break
                    overlap_sents.insert(0, prev_s)
                    overlap_len += len(prev_s)
                cur     = overlap_sents + [s]
                cur_len = overlap_len + len(s)
            else:
                cur, cur_len = [s], len(s)
        else:
            cur.append(s); cur_len += len(s)
    if cur:
        chunks.append(" ".join(cur))
    return chunks or [text]


def chunk_text_smart(text: str, max_chars: int = 8000) -> List[str]:
    """Chytré dělení textu: preferuje celé odstavce, pak celé věty.
    Ideální pro překlad – zachovává kontext odstavce co nejdéle.
    Minimální chunk: 2 000 znaků (zamezuje zbytečně malým fragmentům).
    """
    _min_chunk = 2_000
    # Rozděl na odstavce (prázdný řádek = odstavec)
    paragraphs = re.split(r'\n\s*\n', text)
    chunks: List[str] = []
    current_parts: List[str] = []
    current_len = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        # Odstavec se vejde do aktuálního chunku
        if current_len + len(para) + 2 <= max_chars:
            current_parts.append(para)
            current_len += len(para) + 2
        else:
            # Uložíme aktuální chunk (pokud něco je)
            if current_parts:
                chunks.append("\n\n".join(current_parts))
                current_parts, current_len = [], 0
            # Odstavec sám se vejde do nového chunku
            if len(para) <= max_chars:
                current_parts = [para]
                current_len = len(para)
            else:
                # Odstavec je příliš dlouhý – rozsekat po větách
                sentences = re.split(r'(?<=[.!?])\s+', para)
                sent_buf: List[str] = []
                sent_len = 0
                for s in sentences:
                    if sent_len + len(s) + 1 > max_chars and sent_buf:
                        chunks.append(" ".join(sent_buf))
                        sent_buf, sent_len = [s], len(s)
                    else:
                        sent_buf.append(s)
                        sent_len += len(s) + 1
                if sent_buf:
                    # Zbytek vět jako nový chunk-začátek
                    current_parts = [" ".join(sent_buf)]
                    current_len = sent_len

    if current_parts:
        chunks.append("\n\n".join(current_parts))

    # Sloučení příliš malých chunků: fragment < _min_chunk připoj k předchozímu
    # (pokud by výsledek nepřekročil max_chars), jinak k násled.
    merged: List[str] = []
    for ch in (chunks or [text]):
        if merged and len(ch) < _min_chunk and len(merged[-1]) + len(ch) + 2 <= max_chars:
            merged[-1] = merged[-1] + "\n\n" + ch
        else:
            merged.append(ch)
    return merged or [text]


def smart_chunk_text(text: str, target_chars: int = 2200, overlap: int = 350) -> List[str]:
    """Chytré rozdělení po blocích pro překlad / workflow.
    Preferuje konec odstavce nebo věty a přidává překryv mezi bloky.
    """
    text = (text or "").strip()
    if not text:
        return [""]
    if len(text) <= int(target_chars * 1.3):
        return [text]

    paragraphs = re.split(r'\n\s*\n', text)
    chunks: List[str] = []
    current = ""
    sentence_split_re = re.compile(r'(?<=[.!?])\s+')

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        candidate = f"{current}\n\n{para}".strip() if current else para
        if len(candidate) <= target_chars:
            current = candidate
            continue

        if current:
            chunks.append(current.strip())
            prev_overlap = current[-overlap:] if overlap > 0 else ""
            current = f"{prev_overlap}\n\n{para}".strip() if prev_overlap else para
            if len(current) <= int(target_chars * 1.15):
                continue

        sentences = sentence_split_re.split(para)
        buf = ""
        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue
            sent_candidate = f"{buf} {sent}".strip() if buf else sent
            if len(sent_candidate) <= target_chars:
                buf = sent_candidate
            else:
                if buf:
                    chunks.append(buf.strip())
                    prev_overlap = buf[-overlap:] if overlap > 0 else ""
                    buf = f"{prev_overlap} {sent}".strip() if prev_overlap else sent
                else:
                    chunks.append(sent[:target_chars].strip())
                    buf = sent[max(0, target_chars - overlap):].strip() if overlap > 0 else ""
        current = buf.strip()

    if current:
        chunks.append(current.strip())

    return [c for c in chunks if c.strip()] or [text]


def translate_text_live(base_url: str,
                        model: str,
                        text: str,
                        sys_msg: str,
                        temp: float = 0.15,
                        target_chars: int = 2200,
                        overlap: int = 350,
                        use_smart: bool = True,
                        live_placeholder=None,
                        progress_cb=None,
                        src_lang: str = "",
                        temp_prefix: str = "preklad_live") -> str:
    """Překládá text po blocích s živým náhledem a průběžným ukládáním do _temp/."""
    if use_smart:
        chunks = smart_chunk_text(text, target_chars=target_chars, overlap=overlap)
    else:
        chunks = chunk_text_smart(text, max_chars=target_chars)

    total = max(1, len(chunks))
    live_placeholder = live_placeholder or st.empty()
    translated_chunks: List[str] = []

    for i, chunk in enumerate(chunks, 1):
        # Zkontrolovat, zda nebyl překlad zastaven
        if not st.session_state.get("translate_running", True):
            # Uložit částečný výsledek
            partial = "\n\n".join([p for p in translated_chunks if p]).strip()
            st.session_state["translation_result"] = partial
            st.session_state["translation_ready"] = True
            # Ukončit funkci, vrátit částečný překlad
            return partial
        translated = do_translate(
            base_url, model, chunk, sys_msg,
            temp=temp,
            chunk_size=max(800, min(target_chars, 8000)),
            progress_cb=None,
            parallel=False,
            max_workers=1,
            src_lang=src_lang,
        )
        translated = (translated or "").strip()
        translated_chunks.append(translated)
        current_text = "\n\n".join([p for p in translated_chunks if p]).strip()

        # ── v24.34: průběžně ukládej do session_state → přežije pauzu/rerun ──
        try:
            st.session_state["tr_live_partial"] = current_text
        except Exception:
            pass

        try:
            _preview_html = current_text[-6000:].replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace("\n","<br>")
            live_placeholder.markdown(
                f'''<div id="lmu-live-box" style="background:#f8fafc;border:1px solid #cbd5e1;
border-radius:6px;padding:.6rem 1rem;font-size:.82em;line-height:1.55;
max-height:320px;overflow-y:auto;color:#1e293b;white-space:pre-wrap;font-family:inherit">
{_preview_html}
</div>
<script>
(function(){{var b=document.getElementById("lmu-live-box");if(b)b.scrollTop=b.scrollHeight;}})();
</script>''',
                unsafe_allow_html=True,
            )
        except Exception:
            try:
                live_placeholder.caption(f"⚙️ Blok {i}/{total}…")
            except Exception:
                pass

        try:
            _temp_save_txt(temp_prefix, i, current_text, total)
            _temp_save_docx(temp_prefix, i, current_text, total)
        except Exception:
            pass

        if progress_cb:
            try:
                progress_cb(i, total, chunk[:80])
            except Exception:
                pass

    return "\n\n".join([p for p in translated_chunks if p]).strip()


def fix_spaced_taxon_names(text: str) -> str:
    """Opraví chybné mezery uvnitř taxonomických jmen způsobené OCR/přepisem.
    Např. 'H exitheca' → 'Hexitheca', 'P aleozoic' → 'Paleozoic'.
    Logika: pokud je osamocené velké písmeno těsně před slovem, a dohromady
    tvoří slovo s velkým prvním písmenem, spojíme je.
    Zachová legitimní zkratky jako 'H. exitheca' (s tečkou).
    """
    # Pattern: osamocené velké písmeno (bez tečky za ním) + mezera + slovo začínající malým písmenem
    # → spojit do jednoho slova s velkým prvním písmenem
    fixed = re.sub(
        r'\b([A-Z])\s+([a-z]{2,})\b',
        lambda m: m.group(1) + m.group(2),
        text
    )
    return fixed

# ══════════════════════════════════════════════════════
# POMOCNÉ FUNKCE – LLM
# ══════════════════════════════════════════════════════
import random as _random

def _backoff_jitter(attempt: int, base: float = 0.5, jitter: float = 0.3) -> float:
    """Exponential backoff s náhodným jitterem (zabrání synchronizovaným retry bouřím).
    attempt=0 → ~0.5s, attempt=1 → ~1s, attempt=2 → ~2s (+0–30 % náhoda)
    """
    delay = base * (2 ** attempt)
    return delay * (1.0 + _random.uniform(0, jitter))


# Sdílená LLM session — stejný vzor jako _HTTP_SESSION pro validaci
_LLM_SESSION: Optional[requests.Session] = None
_LLM_SESSION_LOCK = _threading.Lock()

def _get_llm_session() -> requests.Session:
    """Vrátí sdílenou requests.Session pro LM Studio (thread-safe, connection pool)."""
    global _LLM_SESSION
    if _LLM_SESSION is None:
        with _LLM_SESSION_LOCK:
            if _LLM_SESSION is None:
                s = requests.Session()
                adapter = requests.adapters.HTTPAdapter(
                    pool_connections=4, pool_maxsize=8, max_retries=0)
                s.mount("http://", adapter)
                s.mount("https://", adapter)
                _LLM_SESSION = s
    return _LLM_SESSION

def _reset_llm_session():
    """Zahodí a znovu vytvoří LLM session (při connection chybách)."""
    global _LLM_SESSION
    with _LLM_SESSION_LOCK:
        old = _LLM_SESSION
        _LLM_SESSION = None
        if old:
            try:
                old.close()
            except Exception:
                pass

# Thread-local cache pro session_state hodnoty čtené z worker threadů
_THREAD_OFFLINE = False   # modul-level fallback; UI vlákno nastaví před spuštěním threadů
_THREAD_GSP     = ""      # global system prompt fallback


def chat_completion(base_url: str, model: str, messages: list,
                    temp: float = 0.1, max_tokens: int = 4000,
                    timeout: int = DEFAULT_TIMEOUT, stream: bool = False,
                    stop: Optional[List[str]] = None):
    # Thread-safe čtení session_state — worker thready nemají ScriptRunContext
    try:
        _offline = st.session_state.get("offline_mode", _THREAD_OFFLINE)
        _gsp     = st.session_state.get("global_system_prompt", _THREAD_GSP).strip()
    except Exception:
        _offline = _THREAD_OFFLINE
        _gsp     = _THREAD_GSP

    # Offline / demo mode – vrátí simulovanou odpověď
    if _offline:
        last_user = next((m["content"] for m in reversed(messages) if m["role"]=="user"), "")
        demo = (f"[DEMO VÝSTUP]\nModel: {model}\n"
                f"Vstup ({len(last_user)} znaků): {last_user[:120]}...\n\n"
                f"[{'JSON' if 'JSON' in str(messages) or 'json' in str(messages) else 'TEXT'}] "
                f"Demo odpověď – zapni LM Studio pro reálné výsledky.")
        return demo

    # Injektuj globální systémový prompt — pouze pokud je neprázdný
    if _gsp:
        if messages and messages[0].get("role") == "system":
            if _gsp not in messages[0]["content"]:
                messages = [{"role": "system", "content": _gsp + "\n\n" + messages[0]["content"]}] + messages[1:]
        else:
            messages = [{"role": "system", "content": _gsp}] + messages

    # Bezpečnostní zkrácení kontextu — dynamický limit podle modelu
    _mn_cap = model.lower()
    _ctx_cap = _MODEL_CONTEXT.get("default")
    for _k, _v in _MODEL_CONTEXT.items():
        if _k in _mn_cap:
            _ctx_cap = _v
            break
    # Limit: 70 % context window v znacích, minus výstupní tokeny (×3 chars)
    _hard_limit = max(8_000, int(_ctx_cap * 0.70) - max_tokens * 3)
    total_chars = sum(len(m.get("content", "")) for m in messages)
    if total_chars > _hard_limit:
        messages = [
            {**m, "content": m["content"][:_hard_limit] + "\n\n[... zkráceno z důvodu limitu kontextu ...]"}
            if m.get("role") == "user" and len(m.get("content", "")) > _hard_limit
            else m
            for m in messages
        ]

    url     = f"{base_url.rstrip('/')}/chat/completions"
    payload = {"model": model, "messages": messages,
               "temperature": temp, "max_tokens": max_tokens, "stream": stream}
    if stop:
        payload["stop"] = stop

    session = _get_llm_session()
    try:
        if stream:
            r = session.post(url, json=payload, timeout=timeout, stream=True)
            r.raise_for_status()
            return r
        r = session.post(url, json=payload, timeout=timeout)
    except (requests.exceptions.ConnectionError, requests.exceptions.ChunkedEncodingError):
        # Broken pipe / stale connection — zahodíme session a zkusíme znovu
        _reset_llm_session()
        session = _get_llm_session()
        if stream:
            r = session.post(url, json=payload, timeout=timeout, stream=True)
            r.raise_for_status()
            return r
        r = session.post(url, json=payload, timeout=timeout)

    if r.status_code == 400:
        try:
            err_detail = r.json()
        except Exception:
            err_detail = r.text[:300]
        raise requests.exceptions.HTTPError(
            f"400 Bad Request – LM Studio odmítlo požadavek.\n"
            f"Nejčastější příčiny:\n"
            f"  • Celková délka kontextu překračuje limit modelu (zkus kratší text nebo méně výsledků)\n"
            f"  • max_tokens ({max_tokens}) + vstup přesahuje context window\n"
            f"Detail: {err_detail}",
            response=r)
    r.raise_for_status()
    # Nepiš do session_state z worker threadů — jen z UI vlákna
    try:
        st.session_state["model_status"] = "ok"
    except Exception:
        pass
    return r.json()["choices"][0]["message"]["content"].strip()

def chat_completion_stream(base_url: str, model: str, messages: list,
                           temp: float = 0.7, max_tokens: int = 4000,
                           timeout: int = DEFAULT_TIMEOUT, placeholder=None,
                           stop: Optional[List[str]] = None) -> str:
    resp = chat_completion(base_url, model, messages, temp, max_tokens, timeout,
                           stream=True, stop=stop)
    parts: List[str] = []   # O(n) append místo O(n²) += concatenation
    for line in resp.iter_lines():
        if not line: continue
        line = line.decode("utf-8")
        if line.startswith("data: "):
            data = line[6:]
            if data == "[DONE]": break
            try:
                delta = json.loads(data)["choices"][0]["delta"]
                if "content" in delta:
                    parts.append(delta["content"])
                    if placeholder:
                        placeholder.markdown(
                            f'<div class="stream-box">{"".join(parts)}</div>',
                            unsafe_allow_html=True)
            except Exception:
                pass
    return "".join(parts)

def list_models(base_url: str) -> List[str]:
    r = requests.get(f"{base_url.rstrip('/')}/models", timeout=20)
    r.raise_for_status()
    return [m["id"] for m in r.json().get("data", [])]


# ── LM Studio CLI (lms) funkce ────────────────────────
def _lms_available() -> bool:
    """Zkontroluje, zda je lms CLI dostupný v PATH."""
    return shutil.which("lms") is not None


def lms_ls() -> Tuple[List[Dict], str]:
    """Načte seznam všech stažených modelů přes `lms ls`.
    Vrátí (list_modelů, raw_výstup_pro_debug).
    Každý model: {'id', 'size_gb', 'architecture', 'context', '_raw'}
    """
    if not _lms_available():
        return [], "lms CLI není dostupné (není v PATH)"
    try:
        result = subprocess.run(
            ["lms", "ls"],
            capture_output=True, text=True, timeout=20,
            encoding="utf-8", errors="replace")
        raw = result.stdout + (("\nSTDERR: " + result.stderr) if result.stderr.strip() else "")

        if result.returncode != 0:
            return [], f"lms ls selhalo (kód {result.returncode}):\n{raw}"

        # Odstraň ANSI escape kódy
        import re as _re
        clean = _re.sub(r'\x1b\[[0-9;]*[mGKHF]', '', raw)

        models = []
        lines = [l.rstrip() for l in clean.splitlines()]

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            # Přeskoč záhlaví tabulky (obsahuje "Model", "Size", "Architecture")
            low = stripped.lower()
            if any(h in low for h in ["model", "size", "arch", "─", "━", "─", "=", "---"]):
                continue
            # Přeskoč prázdné oddělovače
            if all(c in " │|─━─=─" for c in stripped):
                continue

            # Parsuj různé formáty:
            # 1) "model-name-q4_k_m   4.2 GB   llama   32768"
            # 2) "│ model-name │ 4.2 GB │ llama │ 32768 │"  (tabulka s │)
            # 3) "model-name" (pouze název)

            if "│" in stripped or "|" in stripped:
                # Tabulkový formát s oddělovači
                sep = "│" if "│" in stripped else "|"
                parts = [p.strip() for p in stripped.split(sep) if p.strip()]
            else:
                # Whitespace-separated
                parts = stripped.split()

            if not parts:
                continue

            model_id = parts[0].strip()
            # Přeskoč prázdná nebo nesmyslná jména
            if not model_id or len(model_id) < 3:
                continue
            # Model ID musí vypadat jako název souboru/modelu
            if not _re.match(r'^[\w\.\-\/]+$', model_id):
                continue

            models.append({
                "id":           model_id,
                "size_gb":      parts[1] + (" " + parts[2] if len(parts) > 2 and parts[2].upper() in ("GB","MB","KB","TB") else "") if len(parts) > 1 else "",
                "architecture": parts[3] if len(parts) > 3 else (parts[2] if len(parts) > 2 else ""),
                "context":      parts[4] if len(parts) > 4 else "",
                "_raw":         stripped,
            })

        return models, raw

    except Exception as e:
        return [], f"Výjimka: {e}"


def lms_ps() -> Optional[str]:
    """Vrátí název aktuálně načteného modelu přes `lms ps`."""
    if not _lms_available():
        return None
    try:
        result = subprocess.run(
            ["lms", "ps"],
            capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            for line in result.stdout.strip().splitlines():
                line = line.strip()
                if line and not line.lower().startswith("model") and not line.startswith("#"):
                    return line.split()[0]  # první sloupec = model ID
    except Exception:
        pass
    return None


def lms_load(model_id: str) -> Tuple[bool, str]:
    """Načte model přes `lms load`. Vrátí (úspěch, zpráva)."""
    if not _lms_available():
        return False, "lms CLI není dostupný. Nainstaluj ho: lmstudio.ai/docs/cli"
    try:
        result = subprocess.run(
            ["lms", "load", model_id],
            capture_output=True, text=True, timeout=120)
        if result.returncode == 0:
            return True, result.stdout.strip() or f"Model {model_id} načten."
        else:
            err = result.stderr.strip() or result.stdout.strip()
            return False, err or f"Chyba při načítání modelu (kód {result.returncode})"
    except subprocess.TimeoutExpired:
        return False, "Timeout — načítání modelu trvá příliš dlouho (>120 s)"
    except Exception as e:
        return False, str(e)


def lms_unload() -> Tuple[bool, str]:
    """Uvolní aktuálně načtený model přes `lms unload`."""
    if not _lms_available():
        return False, "lms CLI není dostupný"
    try:
        result = subprocess.run(
            ["lms", "unload", "--all"],
            capture_output=True, text=True, timeout=30)
        return result.returncode == 0, result.stdout.strip() or "Modely uvolněny."
    except Exception as e:
        return False, str(e)


def detect_language(text: str) -> Optional[str]:
    """Detekuje jazyk textu pomocí langdetect. Vrátí český název jazyka nebo None."""
    if not LANGDETECT_AVAILABLE or not detect:
        return None
    try:
        return CODE_TO_LANGUAGE.get(detect(text[:1000]))
    except Exception:
        return None

# ══════════════════════════════════════════════════════
# POMOCNÉ FUNKCE – EXPORT
# ══════════════════════════════════════════════════════
def to_txt_bytes(text: str) -> bytes:
    return text.encode("utf-8")

def to_json_bytes(data) -> bytes:
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")

def to_docx_bytes(text: str) -> bytes:
    from docx import Document  # lazy import
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def to_docx_track_changes(original: str, translated: str) -> bytes:
    """Exportuje překlad jako DOCX s vizuálním track-changes zobrazením.
    Odstraněný text (originál) je červeně přeškrtnutý, nový (překlad) zeleně podtržený.
    """
    from docx import Document  # lazy import
    from docx.shared import RGBColor, Pt

    doc = Document()
    doc.add_heading("Překlad s vyznačením změn", level=1)

    orig_pars  = [p.strip() for p in original.split("\n")   if p.strip()]
    trans_pars = [p.strip() for p in translated.split("\n") if p.strip()]

    for i in range(max(len(orig_pars), len(trans_pars))):
        o = orig_pars[i]  if i < len(orig_pars)  else ""
        t = trans_pars[i] if i < len(trans_pars) else ""

        para = doc.add_paragraph()
        # Originál — červeně přeškrtnutý
        if o:
            run_del = para.add_run(o + " ")
            run_del.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run_del.font.strike = True
            run_del.font.size = Pt(10)
        # Překlad — zeleně podtržený
        if t:
            run_ins = para.add_run(t)
            run_ins.font.color.rgb = RGBColor(0x00, 0x88, 0x00)
            run_ins.font.underline = True
            run_ins.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def segment_diff_html(original: str, translated: str) -> str:
    """Segment-level diff: odstavce 1:1 vedle sebe s barevným zvýrazněním rozdílů."""
    import difflib
    orig_pars  = [p.strip() for p in original.split("\n")   if p.strip()]
    trans_pars = [p.strip() for p in translated.split("\n") if p.strip()]
    rows = ""
    for i in range(max(len(orig_pars), len(trans_pars))):
        o = orig_pars[i]  if i < len(orig_pars)  else ""
        t = trans_pars[i] if i < len(trans_pars) else ""
        matcher = difflib.SequenceMatcher(None, o.split(), t.split())
        orig_hl = []
        tran_hl = []
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            orig_words = " ".join(o.split()[i1:i2])
            tran_words = " ".join(t.split()[j1:j2])
            if tag == "equal":
                orig_hl.append(orig_words)
                tran_hl.append(tran_words)
            elif tag == "replace":
                orig_hl.append(f'<span style="background:#fee2e2;color:#991b1b;'
                                f'border-radius:2px;padding:0 2px">{orig_words}</span>')
                tran_hl.append(f'<span style="background:#dcfce7;color:#166534;'
                                f'border-radius:2px;padding:0 2px">{tran_words}</span>')
            elif tag == "delete":
                orig_hl.append(f'<span style="background:#fee2e2;color:#991b1b;'
                                f'text-decoration:line-through;border-radius:2px;'
                                f'padding:0 2px">{orig_words}</span>')
            elif tag == "insert":
                tran_hl.append(f'<span style="background:#dcfce7;color:#166534;'
                                f'border-radius:2px;padding:0 2px">{tran_words}</span>')
        row_bg = "#ffffff" if i % 2 == 0 else "#f9fafb"
        rows += (f'<tr style="background:{row_bg}">'
                 f'<td style="padding:7px 12px;vertical-align:top;border-bottom:1px solid #e5e7eb;'
                 f'width:50%;font-size:.85rem;color:#111;line-height:1.5">'
                 f'{" ".join(orig_hl)}</td>'
                 f'<td style="padding:7px 12px;vertical-align:top;border-bottom:1px solid #e5e7eb;'
                 f'width:50%;font-size:.85rem;color:#111;line-height:1.5">'
                 f'{" ".join(tran_hl)}</td></tr>')
    return ('<table style="width:100%;border-collapse:collapse;border:1px solid #e5e7eb;'
            'border-radius:6px;overflow:hidden">'
            '<thead><tr>'
            '<th style="padding:8px 12px;text-align:left;background:#f3f4f6;'
            'color:#374151;font-weight:600;border-bottom:2px solid #d1d5db">Originál</th>'
            '<th style="padding:8px 12px;text-align:left;background:#f0fdf4;'
            'color:#374151;font-weight:600;border-bottom:2px solid #d1d5db">Překlad</th>'
            '</tr></thead><tbody>' + rows + '</tbody></table>')


def to_xlsx_bytes(df: pd.DataFrame) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        df.to_excel(w, index=False)
    return buf.getvalue()

def make_zip(files: Dict[str, bytes]) -> bytes:
    """files = {filename: content_bytes}"""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in files.items():
            zf.writestr(name, data)
    return buf.getvalue()

def simple_diff_html(original: str, translated: str) -> str:
    orig_pars  = [p.strip() for p in original.split("\n")   if p.strip()]
    trans_pars = [p.strip() for p in translated.split("\n") if p.strip()]
    rows = ""
    for i in range(max(len(orig_pars), len(trans_pars))):
        o = orig_pars[i]  if i < len(orig_pars)  else ""
        t = trans_pars[i] if i < len(trans_pars) else ""
        rows += (f'<tr><td style="padding:4px 8px;vertical-align:top;border-bottom:1px solid #333;'
                 f'width:50%;color:#ccc">{o}</td>'
                 f'<td style="padding:4px 8px;vertical-align:top;border-bottom:1px solid #333;'
                 f'width:50%">{t}</td></tr>')
    return ('<table style="width:100%;border-collapse:collapse;font-size:.85rem"><thead><tr>'
            '<th style="padding:6px 8px;text-align:left;background:#222">Originál</th>'
            '<th style="padding:6px 8px;text-align:left;background:#1a3a2a">Překlad</th>'
            '</tr></thead><tbody>' + rows + '</tbody></table>')

def render_export_buttons(text: str, base_name: str, metadata: dict = None):
    c1, c2, c3, c4 = st.columns(4)
    c1.download_button("📄 TXT",  to_txt_bytes(text),  f"{base_name}.txt",  width='stretch')
    c2.download_button("📋 JSON", to_json_bytes({"content": text, "metadata": metadata or {}}),
                       f"{base_name}.json", width='stretch')
    c3.download_button("📝 DOCX", to_docx_bytes(text), f"{base_name}.docx", width='stretch')
    c4.download_button("📄 PDF",  to_pdf_bytes(text),  f"{base_name}.pdf",  width='stretch')

def to_pdf_bytes(text: str) -> bytes:
    """Jednoduchý PDF export bez externích závislostí – UTF-8 přes reportlab, fallback na latin-1."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                 leftMargin=2.5*cm, rightMargin=2.5*cm,
                                 topMargin=2.5*cm, bottomMargin=2.5*cm)
        styles = getSampleStyleSheet()
        story = []
        for line in text.split("\n"):
            if line.strip():
                story.append(Paragraph(line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"),
                                        styles["Normal"]))
            else:
                story.append(Spacer(1, 6))
        doc.build(story)
        return buf.getvalue()
    except ImportError:
        # Fallback: minimální PDF bez závislostí (Latin-1, základní ASCII)
        safe = text.encode("latin-1", errors="replace").decode("latin-1")
        lines = safe.split("\n")
        pdf_lines = []
        y = 800
        pdf_lines.append("%PDF-1.4")
        stream_content = "BT\n/F1 10 Tf\n"
        for line in lines:
            escaped = line.replace("\\","\\\\").replace("(","\\(").replace(")","\\)")
            stream_content += f"10 {y} Td ({escaped}) Tj\n0 -14 Td\n"
            y -= 14
        stream_content += "ET"
        stream_bytes = stream_content.encode("latin-1", errors="replace")
        pdf_lines.append("1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj")
        pdf_lines.append("2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj")
        pdf_lines.append("3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
                         "/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj")
        pdf_lines.append(f"4 0 obj\n<< /Length {len(stream_bytes)} >>\nstream\n"
                         + stream_content + "\nendstream\nendobj")
        pdf_lines.append("5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj")
        body = "\n".join(pdf_lines).encode("latin-1", errors="replace")
        xref_pos = len(body)
        xref = ("xref\n0 6\n0000000000 65535 f \n"
                + "".join(f"{len(p):010d} 00000 n \n" for p in pdf_lines)
                + f"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF")
        return body + xref.encode("latin-1", errors="replace")

# ══════════════════════════════════════════════════════
# POMOCNÉ FUNKCE – SLOVNÍKY (persistentní)
# ══════════════════════════════════════════════════════
def load_glossaries() -> Dict[str, Dict[str, str]]:
    if os.path.exists(GLOSSARY_FILE):
        try:
            with open(GLOSSARY_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def save_glossaries(data: Dict[str, Dict[str, str]]):
    Path(GLOSSARY_FILE).parent.mkdir(parents=True, exist_ok=True)
    with open(GLOSSARY_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def glossary_to_text(g: Dict[str, str]) -> str:
    return "\n".join(f"{k} → {v}" for k, v in g.items())

def text_to_glossary(text: str) -> Dict[str, str]:
    result = {}
    for line in text.strip().split("\n"):
        if "→" in line:
            parts = line.split("→", 1)
            k, v = parts[0].strip(), parts[1].strip()
            if k: result[k] = v
    return result

_GLOSSARY_RE_CACHE: Dict[str, re.Pattern] = {}
_GLOSSARY_RE_CACHE_MAX = 500   # max položek — LRU eviction při překročení

# ══════════════════════════════════════════════════════
# POMOCNÉ FUNKCE – PROMPT ŠABLONY (persistentní)
# ══════════════════════════════════════════════════════
def load_templates() -> Dict[str, Dict]:
    if os.path.exists(TEMPLATES_FILE):
        try:
            with open(TEMPLATES_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return dict(DEFAULT_TEMPLATES)

def save_templates(data: Dict[str, Dict]):
    Path(TEMPLATES_FILE).parent.mkdir(parents=True, exist_ok=True)
    with open(TEMPLATES_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# Inicializace šablon do session state (jednou za session)
if not st.session_state["prompt_templates"]:
    st.session_state["prompt_templates"] = load_templates()

# Inicializace slovníků do session state
if not st.session_state["glossaries"]:
    st.session_state["glossaries"] = load_glossaries()


# ══════════════════════════════════════════════════════
# POMOCNÉ FUNKCE – HISTORIE
# ══════════════════════════════════════════════════════
_HISTORY_CACHE: Optional[list] = None
_HISTORY_MTIME: float = 0.0

def load_history() -> list:
    """Načte historii ze souboru s in-memory cache — čte disk jen pokud se soubor změnil."""
    global _HISTORY_CACHE, _HISTORY_MTIME
    try:
        mtime = os.path.getmtime(HISTORY_FILE) if os.path.exists(HISTORY_FILE) else 0.0
        if _HISTORY_CACHE is not None and mtime == _HISTORY_MTIME:
            return _HISTORY_CACHE
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                _HISTORY_CACHE = json.load(f)
            _HISTORY_MTIME = mtime
            return _HISTORY_CACHE
    except Exception:
        pass
    _HISTORY_CACHE = []
    return []

def save_to_history(operation: str, metadata: dict):
    global _HISTORY_CACHE, _HISTORY_MTIME, HISTORY_FILE
    hist = load_history()
    hist.insert(0, {"operation": operation,
                    "timestamp": datetime.now().isoformat(),
                    "metadata": metadata})
    hist = hist[:MAX_HISTORY_ITEMS]
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(hist, f, ensure_ascii=False, indent=2)
    # Invalidate cache po zápisu
    _HISTORY_CACHE = hist
    _HISTORY_MTIME = os.path.getmtime(HISTORY_FILE)
    _auto_nightly_backup()


def _auto_nightly_backup():
    """Uloží zálohu processing_history.json jednou za den do backups/.
    Používá in-memory příznak _BACKUP_DONE_TODAY — zabraňuje opakovaným
    Path.exists() syscallům při každém save_to_history().
    Běží v daemon threadu — neblokuje main thread.
    """
    if _BACKUP_DONE_TODAY:
        return

    def _do_backup():
        global _BACKUP_DONE_TODAY
        try:
            today_str = datetime.now().strftime("%Y-%m-%d")
            backup_dir = Path("backups")
            backup_file = backup_dir / f"history_{today_str}.json"
            if not backup_file.exists() and os.path.exists(HISTORY_FILE):
                backup_dir.mkdir(exist_ok=True)
                shutil.copy2(HISTORY_FILE, backup_file)
                existing = sorted(backup_dir.glob("history_*.json"))
                for old in existing[:-30]:
                    old.unlink()
            _BACKUP_DONE_TODAY = True
        except Exception:
            pass

    _threading.Thread(target=_do_backup, daemon=True).start()


# ── Versioning výsledků extrakce ──────────────────────
EXTRACTION_HISTORY_DIR = Path("extraction_history")

def save_extraction_version(data: Any, label: str = ""):
    """Uloží timestampovaný snapshot výsledků extrakce (atomický zápis)."""
    try:
        EXTRACTION_HISTORY_DIR.mkdir(parents=True, exist_ok=True)
        ts    = datetime.now().strftime("%Y%m%d_%H%M%S")
        fname = EXTRACTION_HISTORY_DIR / f"extraction_{ts}.json"
        tmp   = str(fname) + ".tmp"
        payload = {
            "timestamp": datetime.now().isoformat(),
            "label":     label,
            "data":      data if isinstance(data, (dict, list)) else str(data),
        }
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        os.replace(tmp, fname)           # atomický přesun — žádná partial write
        # Ponecháme jen posledních 50 snapshotů
        existing = sorted(EXTRACTION_HISTORY_DIR.glob("extraction_*.json"))
        for old in existing[:-50]:
            try:
                old.unlink()
            except Exception:
                pass
        return str(fname)
    except Exception:
        try:
            os.remove(str(fname) + ".tmp")
        except Exception:
            pass
        return None

def list_extraction_versions() -> List[Dict]:
    """Vrátí seznam uložených snapshotů extrakce."""
    if not EXTRACTION_HISTORY_DIR.exists():
        return []
    result = []
    for f in sorted(EXTRACTION_HISTORY_DIR.glob("extraction_*.json"), reverse=True):
        try:
            with open(f, "r", encoding="utf-8") as fh:
                meta = json.load(fh)
            result.append({
                "file": str(f),
                "timestamp": meta.get("timestamp","")[:16].replace("T"," "),
                "label": meta.get("label",""),
                "size": f.stat().st_size
            })
        except Exception:
            pass
    return result[:20]

# ══════════════════════════════════════════════════════
# SESSION PERSISTENCE (uložení/načtení stavu na disk)
# ══════════════════════════════════════════════════════
SESSION_PERSIST_FILE  = "session_state.json"
SESSION_AUTOSAVE_FILE = "session_autosave.json"   # automatická záloha každých N rerunů

SESSION_PERSIST_KEYS = [
    # Extrakce
    "last_extracted_text", "last_extracted_df_json", "last_extraction_taxa",
    "last_extraction_text",
    # Překlad — výsledky
    "translation_original", "translation_result",
    "translation_src_lang", "translation_tgt_lang", "translation_ready",
    "translation_iterations",
    # Překlad — vstup a nastavení
    "tr_src_text", "tr_src_lang_val", "tr_tgt_lang_val",
    "trans_mode_sel", "temp_t", "tr_iter", "tr_parallel",
    # Validace
    "last_validation_results", "val_selected_dbs", "val_taxa_input",
    # Chat
    "chat_messages", "chat_injected_context",
    # Nastavení
    "glossaries", "prompt_templates", "clean_prompts",
    "lms_max_concurrent", "lang",
    # Ostatní výsledky
    "ab_results", "style_polish_result", "hyolitha_export_records",
]

# Klíče pro autosave — jen těžce získatelná data (překlad, extrakce, validace)
SESSION_AUTOSAVE_KEYS = [
    # Extrakce
    "last_extracted_text", "last_extracted_df_json", "last_extraction_taxa",
    "last_extraction_text",      # alias: nastavuje extraction loop (sync s last_extracted_text)
    # Překlad — výsledky
    "translation_original", "translation_result", "translation_ready",
    "translation_src_lang", "translation_tgt_lang",
    "translation_iterations",    # iterace překladu pro stažení
    # Překlad — vstup a nastavení
    "tr_src_text",               # text v ručním textovém poli
    "tr_src_lang_val",           # zdrojový jazyk
    "tr_tgt_lang_val",           # cílový jazyk
    "trans_mode_sel",            # režim (Jeden text / Dávkový / Porovnat)
    "temp_t",                    # teplota překladu
    "tr_iter",                   # počet iterací
    "tr_parallel",               # paralelní překlad
    # Validace, chat, ostatní
    "last_validation_results", "chat_messages",
    "ab_results",                # A/B test výsledky
    "style_polish_result",       # výsledek stylistiky
    "hyolitha_export_records",   # Hyolitha záznamy připravené k exportu
]

_AUTOSAVE_INTERVAL = 3   # uloží každých N rerunů Streamlitu (dříve 5)
_EXT_CHUNK_CHECKPOINT_INTERVAL = 5   # průběžný checkpoint každých N chunků (sekvenční režim)


def save_session_to_disk(filepath: str = None,
                         keys: list = None) -> bool:
    """Uloží klíčové části session state do JSON souboru na disk.
    Atomický zápis přes .tmp → os.replace — odolné vůči výpadku napájení.
    Zároveň udržuje .bak kopii předchozí verze pro případ korrupce.
    """
    global SESSION_PERSIST_FILE, SESSION_AUTOSAVE_FILE
    if filepath is None:
        filepath = SESSION_PERSIST_FILE
    keys = keys or SESSION_PERSIST_KEYS
    data = {}
    for k in keys:
        v = st.session_state.get(k)
        if v is None:
            continue
        try:
            json.dumps(v)
            data[k] = v
        except (TypeError, ValueError):
            pass
    if not data:
        return False
    tmp_path = filepath + ".tmp"
    bak_path = filepath + ".bak"
    try:
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump({"saved_at": datetime.now().isoformat(), "data": data},
                      f, ensure_ascii=False, indent=2)
        # Rotuj .bak před přepsáním
        if os.path.exists(filepath):
            try:
                os.replace(filepath, bak_path)
            except Exception:
                pass
        os.replace(tmp_path, filepath)
        return True
    except Exception:
        try:
            os.remove(tmp_path)
        except Exception:
            pass
        return False


def load_session_from_disk(filepath: str = None,
                           keys: list = None) -> dict:
    """Načte uložený stav session ze souboru.
    Při korrupci primary souboru zkusí .bak zálohu.
    Vrátí slovník načtených klíčů.
    """
    if filepath is None:
        filepath = SESSION_PERSIST_FILE
        keys = keys or SESSION_PERSIST_KEYS

    def _try_load(path: str) -> dict:
        with open(path, "r", encoding="utf-8") as f:
            payload = json.load(f)
        data     = payload.get("data", {})
        saved_at = payload.get("saved_at", "")
        loaded   = {}
        for k, v in data.items():
            if k in keys:
                st.session_state[k] = v
                loaded[k] = True
        loaded["_saved_at"] = saved_at
        return loaded

    if os.path.exists(filepath):
        try:
            return _try_load(filepath)
        except Exception:
            pass   # primary poškozená — zkus .bak

    bak_path = filepath + ".bak"
    if os.path.exists(bak_path):
        try:
            result = _try_load(bak_path)
            result["_from_backup"] = True
            return result
        except Exception:
            pass

    return {}


def _autosave_if_due():
    """Automaticky uloží session každých N rerunů pokud jsou data k uložení.
    Ukládá jen tehdy, kdy se data od poslední zálohy skutečně změnila
    (hash kontrola — zabrání zbytečným I/O zápisům při read-only rerunech).
    """
    counter = st.session_state.get("_autosave_counter", 0) + 1
    st.session_state["_autosave_counter"] = counter
    if counter % _AUTOSAVE_INTERVAL == 0:
        has_data = any(st.session_state.get(k) for k in SESSION_AUTOSAVE_KEYS)
        if has_data:
            # Rychlý hash klíčových hodnot — porovnej s předchozím hashem
            try:
                _sig = hash(tuple(
                    len(str(st.session_state.get(k, "")))
                    for k in SESSION_AUTOSAVE_KEYS
                ))
            except Exception:
                _sig = counter
            if _sig != st.session_state.get("_autosave_sig"):
                st.session_state["_autosave_sig"] = _sig
                save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)


def session_persist_info(filepath: str = SESSION_PERSIST_FILE) -> Optional[str]:
    """Vrátí čas poslední zálohy session nebo None."""
    if not os.path.exists(filepath):
        return None
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            payload = json.load(f)
        return payload.get("saved_at","")
    except Exception:
        return None


# ══════════════════════════════════════════════════════
# FUZZY PRE-SEARCH – oprava OCR chyb před validací
# ══════════════════════════════════════════════════════
# Časté záměny v OCR taxonomických jmen
_OCR_SUBSTITUTIONS = [
    (r'\b0([a-z])',  r'O\1'),   # 0 → O na začátku slova
    (r'([A-Z])0',   r'\g<1>O'), # 0 → O za verzálkou
    (r'([a-z])l([A-Z])', r'\g<1>I\2'),  # l → I před verzálkou
    (r'\bI([a-z]{3,})', r'l\1'),  # I → l na začátku malého slova
    (r'([a-z])rn([a-z])', r'\g<1>m\2'),  # rn → m
    (r'([a-z])ii([a-z])', r'\g<1>ii\2'), # necháme — ii je validní (genitivní přípona)
    (r'  +', ' '),              # dvojité mezery
    (r'(\w)-\n(\w)', r'\1\2'),  # zalomení přes pomlčku
]

def fuzzy_fix_taxon_name(name: str) -> Tuple[str, List[str]]:
    """Aplikuje OCR korekce na taxonomické jméno.
    Vrátí (opravené_jméno, seznam_provedených_oprav).
    """
    original = name
    fixes = []
    result = name.strip()
    # Odstraň čísla stránek přilepená k jménu (např. "Hyolithes234")
    cleaned = re.sub(r'\d{2,}$', '', result).strip()
    if cleaned != result:
        fixes.append(f"odstraněna čísla: '{result}' → '{cleaned}'")
        result = cleaned
    for pattern, repl in _OCR_SUBSTITUTIONS:
        new = re.sub(pattern, repl, result)
        if new != result:
            fixes.append(f"OCR fix: '{result}' → '{new}'")
            result = new
    # Oprav první písmeno na verzálku (pokud bylo přepsáno)
    if result and result[0].islower():
        result = result[0].upper() + result[1:]
        fixes.append(f"verzálka: '{original[0]}' → '{result[0]}'")
    return result, fixes


# ══════════════════════════════════════════════════════
# TAXONOMICKÁ VALIDACE
# ══════════════════════════════════════════════════════
def _normalize_for_dup(name: str) -> str:
    """Normalizuje jméno pro porovnávání duplikátů: lower, strip diakritiky, single space."""
    nfkd = unicodedata.normalize("NFKD", name.lower().strip())
    ascii_str = "".join(c for c in nfkd if not unicodedata.combining(c))
    return re.sub(r'\s+', ' ', ascii_str)


def _name_matches(api_name: str, query: str) -> bool:
    """Ověří zda jméno z API odpovídá hledanému taxonu.

    Řeší typické problémy API odpovědí:
    - "Betula L." vs "Betula" (author suffix)
    - "Betula pendula" vs "Betula" (species vs genus — neakceptovat, příliš volné)
    - "BETULA" vs "Betula" (case)
    - Přesná shoda má prioritu; pokud api_name začíná query + mezera/konec → genus match
    """
    if not api_name or not query:
        return False
    a = api_name.strip()
    q = query.strip()
    # 1) Přesná shoda (case-insensitive)
    if a.lower() == q.lower():
        return True
    # 2) API vrátí "Rod Autor" (jednoslovný taxon + autor) — rod odpovídá
    # Pouze pokud query je jednoslovné (rod) a api_name začíná query + mezera
    if ' ' not in q and a.lower().startswith(q.lower() + ' '):
        # Ověř že zbytek po prvním slovu vypadá jako autor (velké písmeno nebo závorka)
        rest = a[len(q):].strip()
        if rest and (rest[0].isupper() or rest[0] in '(,'):
            return True
    return False

def find_duplicate_taxa(taxa: List[str]) -> List[Dict]:
    """
    Vrátí skupiny taxonů, které jsou si podezřele podobné
    (stejný normalizovaný tvar nebo Levenshtein vzdálenost ≤ 2).

    Optimalizace:
    - Prefix indexing: kandidáti sdílející 3znakový prefix jsou testováni Levenshtein;
      ostatní páry přeskočeny (typicky eliminuje ~95 % porovnání pro databázi 1500 taxonů)
    - Early-exit v levenshtein pro délkový rozdíl > 2
    """
    def levenshtein(a: str, b: str) -> int:
        if a == b: return 0
        if len(a) < len(b): a, b = b, a
        if not b: return len(a)
        if len(a) - len(b) > 2: return len(a) - len(b)
        prev = list(range(len(b)+1))
        for i, ca in enumerate(a):
            curr = [i+1]
            for j, cb in enumerate(b):
                curr.append(min(prev[j+1]+1, curr[j]+1,
                                prev[j] + (0 if ca == cb else 1)))
            prev = curr
        return prev[-1]

    normalized = [(_normalize_for_dup(t), t) for t in taxa]

    # Prefix index: 3-znakový prefix normalizovaného jména → seznam indexů
    # Taxony sdílející prefix jsou kandidáti na duplicitu; ostatní páry přeskočíme
    from collections import defaultdict as _defaultdict
    prefix_idx: Dict[str, List[int]] = _defaultdict(list)
    for i, (ni, _) in enumerate(normalized):
        prefix_idx[ni[:3]].append(i)
        # Také indexujeme přilehlé prefixy pro zachycení 1znakové substituce na začátku
        if len(ni) >= 3:
            prefix_idx[ni[1:4]].append(i)

    groups = []
    used = set()
    for i, (ni, ti) in enumerate(normalized):
        if i in used:
            continue
        group = [ti]
        # Kandidáti: pouze taxony se sdíleným prefixem
        candidates_idx = set()
        for pfx in (ni[:3], ni[1:4] if len(ni) >= 4 else ""):
            if pfx:
                candidates_idx.update(prefix_idx.get(pfx, []))
        candidates_idx.discard(i)

        for j in sorted(candidates_idx):
            if j <= i or j in used:
                continue
            nj, tj = normalized[j]
            if ni == nj or levenshtein(ni, nj) <= 2:
                group.append(tj)
                used.add(j)
        if len(group) > 1:
            groups.append({"canonical": ti, "duplicates": group})
            used.add(i)
    return groups

# Module-level URL encoding helpers — dříve definovány uvnitř validate_taxon_name
# (nová definice při každém volání = zbytečný overhead)
def _enc(t: str) -> str:
    """URL encode mezer jako %20."""
    return t.replace(" ", "%20")

def _plus(t: str) -> str:
    """URL encode mezer jako + (form encoding)."""
    return t.replace(" ", "+")

# In-memory příznak pro nightly backup — zabraňuje opakovaným Path.exists() syscallům
_BACKUP_DONE_TODAY: bool = False
_HTTP_SESSION: Optional[requests.Session] = None

_HTTP_SESSION_LOCK = _threading.Lock()

def _get_http_session() -> requests.Session:
    global _HTTP_SESSION
    if _HTTP_SESSION is None:
        with _HTTP_SESSION_LOCK:
            if _HTTP_SESSION is None:
                s = requests.Session()
                s.headers.update({"User-Agent": "HyolithaTaxonValidator/1.0"})
                adapter = requests.adapters.HTTPAdapter(
                    pool_connections=20, pool_maxsize=20, max_retries=0)
                s.mount("http://", adapter)
                s.mount("https://", adapter)
                _HTTP_SESSION = s
    return _HTTP_SESSION

def _reset_http_session():
    """Zahodí a znovu vytvoří HTTP session (při connection chybách)."""
    global _HTTP_SESSION
    with _HTTP_SESSION_LOCK:
        old = _HTTP_SESSION
        _HTTP_SESSION = None
        if old:
            try:
                old.close()
            except Exception:
                pass

def _safe_get(url: str, timeout: int = None, **kwargs) -> requests.Response:
    """HTTP GET s retry (max 3 pokusy) a exponential backoff s jitterem.
    Sdílená requests.Session na úrovni modulu — přežije Streamlit reruns.
    timeout=None → použije VAL_TIMEOUT konstantu (ne st.session_state — bezpečné z threadů).
    """
    if timeout is None:
        # Bezpečné i z worker threadů — žádný ScriptRunContext přístup
        try:
            timeout = st.session_state.get("val_timeout", VAL_TIMEOUT)
        except Exception:
            timeout = VAL_TIMEOUT
    session = _get_http_session()
    last_exc = None
    for attempt in range(3):
        try:
            return session.get(url, timeout=timeout, **kwargs)
        except requests.exceptions.Timeout as e:
            last_exc = e
            if attempt < 2:
                _time.sleep(_backoff_jitter(attempt, base=0.3))
        except (requests.exceptions.ConnectionError,
                requests.exceptions.ChunkedEncodingError) as e:
            last_exc = e
            # Obnov session — stale connection pool po dlouhé nečinnosti
            _reset_http_session()
            session = _get_http_session()
            if attempt < 2:
                _time.sleep(_backoff_jitter(attempt, base=0.3))
    raise last_exc


# ── Disk cache (SQLite) s connection poolem ──────────────────────────────────
_DISK_CACHE_PATH = Path("validation_cache.db")
_DISK_CACHE_CONN: Optional[sqlite3.Connection] = None

_DISK_CACHE_LOCK = _threading.Lock()

def _get_cache_conn() -> sqlite3.Connection:
    """Vrátí sdílené SQLite připojení (WAL mode, thread-safe, double-checked locking)."""
    global _DISK_CACHE_CONN
    if _DISK_CACHE_CONN is None:
        with _DISK_CACHE_LOCK:
            if _DISK_CACHE_CONN is None:
                conn = sqlite3.connect(
                    str(_DISK_CACHE_PATH), check_same_thread=False, timeout=15)
                conn.execute("PRAGMA journal_mode=WAL")
                conn.execute("PRAGMA synchronous=NORMAL")
                conn.execute("PRAGMA cache_size=4000")
                conn.execute("PRAGMA temp_store=MEMORY")
                conn.execute("PRAGMA mmap_size=67108864")  # 64 MB mmap
                _DISK_CACHE_CONN = conn
    return _DISK_CACHE_CONN

def _reset_cache_conn():
    """Zahodí SQLite spojení (při chybách — znovu se otevře při příštím přístupu)."""
    global _DISK_CACHE_CONN
    with _DISK_CACHE_LOCK:
        old = _DISK_CACHE_CONN
        _DISK_CACHE_CONN = None
        if old:
            try:
                old.close()
            except Exception:
                pass

def _disk_cache_init():
    """Inicializuje SQLite disk cache pro výsledky validace."""
    conn = _get_cache_conn()
    conn.execute("""CREATE TABLE IF NOT EXISTS val_cache (
        cache_key TEXT PRIMARY KEY,
        result_json TEXT,
        created_at TEXT DEFAULT (datetime('now'))
    )""")
    conn.commit()

def _disk_cache_get(cache_key: str) -> Optional[Dict]:
    """Načte výsledek z disk cache. Vrátí None pokud neexistuje."""
    try:
        row = _get_cache_conn().execute(
            "SELECT result_json FROM val_cache WHERE cache_key=?", (cache_key,)
        ).fetchone()
        return json.loads(row[0]) if row else None
    except Exception:
        return None

def _disk_cache_set(cache_key: str, result: Dict, _commit: bool = True):
    """Uloží výsledek do disk cache.
    _commit=False umožňuje hromadné vkládání bez commit po každém záznamu —
    volající pak zavolá conn.commit() sám (viz _disk_cache_set_batch).
    """
    try:
        conn = _get_cache_conn()
        conn.execute(
            "INSERT OR REPLACE INTO val_cache (cache_key, result_json) VALUES (?,?)",
            (cache_key, json.dumps(result, ensure_ascii=False))
        )
        if _commit:
            conn.commit()
    except Exception:
        pass


def _disk_cache_set_batch(items: List[Tuple[str, Dict]]):
    """Hromadné uložení výsledků do disk cache v jedné transakci.
    Výrazně rychlejší než N × _disk_cache_set pro dávkovou validaci.
    Při poškozené DB resetuje spojení a zkusí znovu.
    """
    if not items:
        return
    for _retry in range(2):
        try:
            conn = _get_cache_conn()
            conn.executemany(
                "INSERT OR REPLACE INTO val_cache (cache_key, result_json) VALUES (?,?)",
                [(k, json.dumps(v, ensure_ascii=False)) for k, v in items]
            )
            conn.commit()
            return
        except sqlite3.DatabaseError:
            _reset_cache_conn()   # poškozená DB — znovu otevře při příštím přístupu
            if _retry == 0:
                continue
        except Exception:
            return

def _disk_cache_clear():
    """Vymaže celou disk cache."""
    try:
        conn = _get_cache_conn()
        conn.execute("DELETE FROM val_cache")
        conn.commit()
    except Exception:
        pass

def _disk_cache_count() -> int:
    """Vrátí počet záznamů v disk cache."""
    try:
        return _get_cache_conn().execute(
            "SELECT COUNT(*) FROM val_cache").fetchone()[0]
    except Exception:
        return 0

# Inicializace disk cache při startu
try:
    _disk_cache_init()
except Exception:
    pass

# Inicializace presence systému
presence.init()

# === Per-user profile activation ===
if st.session_state.get("presence_name"):
    if "activate_user_profile" in globals():
        activate_user_profile(st.session_state["presence_name"])


def validate_taxon_name(taxon: str, databases: List[str], cache: dict = None,
                         use_offline: bool = False) -> Dict:
    """Validuje jméno taxonu proti vybraným databázím.
    Timeout se čte ze session_state JEDNOU v hlavním threadu a předává jako parametr —
    zabraňuje ScriptRunContext warning při přístupu ze session_state z worker threadů.
    """
    # Čtení session_state POUZE v hlavním threadu (ne v workerech)
    _timeout = VAL_TIMEOUT
    try:
        _timeout = st.session_state.get("val_timeout", VAL_TIMEOUT)
    except Exception:
        pass  # Jsme v threadu bez ScriptRunContext — použij default
    cache_key = taxon + '::' + ','.join(sorted(databases))
    # 1) Memory cache
    if cache and cache_key in cache:
        return cache[cache_key]
    # 2) Disk cache (SQLite)
    disk_hit = _disk_cache_get(cache_key)
    if disk_hit is not None:
        if cache is not None:
            cache[cache_key] = disk_hit
        return disk_hit

    # 3) Paralelní dotazy — každá DB dostane vlastní podvolání se [db]
    #    (bezpečné: každé podvolání je nezávislé, výsledky se sloučí níže)
    if len(databases) > 1:
        _per_db_results: Dict[str, Dict] = {}
        _per_db_lock = _threading.Lock()

        def _query_db_single(db):
            r = _validate_taxon_single_db(taxon, [db], use_offline=use_offline,
                                          timeout=_timeout)
            with _per_db_lock:
                _per_db_results[db] = r.get("results", {}).get(db, {"found": False})

        with ThreadPoolExecutor(max_workers=min(8, len(databases))) as _pool:
            list(_pool.map(_query_db_single, databases))

        # Sloučení výsledků do standardního formátu
        merged: Dict[str, Any] = {
            "taxon": taxon,
            "results": _per_db_results,
            "summary": {"total": len(databases), "found": 0,
                        "confidence": "not_found", "details": {}}
        }
        PREF = ("PaleoDB", "GBIF", "IRMNG", "WoRMS", "CoL", "ITIS",
                "ZooBank", "Fossilworks", "Plazi", "BioLib")
        for db in databases:
            db_r = _per_db_results.get(db, {})
            if db_r.get("found"):
                merged["summary"]["found"] += 1
                det = db_r.get("details", {})
                if det:
                    cur = merged["summary"]["details"]
                    if not cur or db in PREF[:4]:
                        merged["summary"]["details"] = det
        found = merged["summary"]["found"]
        ratio = found / max(1, len(databases))
        merged["summary"]["confidence"] = (
            "high"      if ratio >= 0.45 else
            "medium"    if ratio >= 0.20 else
            "low"       if ratio  > 0   else
            "not_found"
        )
        if cache is not None:
            cache[cache_key] = merged
        _disk_cache_set(cache_key, merged)
        return merged

    # Single DB — přímé volání (bez overhead ThreadPoolExecutor)
    return _validate_taxon_single_db(taxon, databases, cache=cache,
                                     use_offline=use_offline, timeout=_timeout)


def _validate_taxon_single_db(taxon: str, databases: List[str], cache: dict = None,
                                use_offline: bool = False,
                                timeout: int = VAL_TIMEOUT) -> Dict:
    """Interní implementace validace — sekvenční smyčka přes DB.
    Voláno z validate_taxon_name buď přímo (single DB) nebo jako worker
    pro jednotlivé DB v paralelním režimu (každé volání dostane [db]).
    timeout: předán z hlavního threadu, nevyžaduje ScriptRunContext.
    """
    cache_key = taxon + '::' + ','.join(sorted(databases))
    # 1) Memory cache
    if cache and cache_key in cache:
        return cache[cache_key]
    # 2) Disk cache (SQLite)
    disk_hit = _disk_cache_get(cache_key)
    if disk_hit is not None:
        if cache is not None:
            cache[cache_key] = disk_hit
        return disk_hit
    # 3) Offline SQLite fallback
    if use_offline:
        offline_hit = offline_db_lookup(taxon)
        if offline_hit:
            result = {
                "taxon": taxon,
                "results": {"Offline DB": {
                    "found": True, "url": "",
                    "details": offline_hit,
                    "error": None
                }},
                "summary": {
                    "total": 1, "found": 1,
                    "confidence": "medium",
                    "details": offline_hit
                }
            }
            if cache is not None:
                cache[cache_key] = result
            return result


    result: Dict[str, Any] = {
        "taxon": taxon,
        "results": {},
        "summary": {"total": len(databases), "found": 0, "confidence": "not_found", "details": {}}
    }

    # _enc() a _plus() jsou definovány na úrovni modulu (výše)

    # Wildcard konverze pro DB které wildcards podporují
    # ? → _ (jeden znak), * → % (libovolný počet znaků)
    # WoRMS/IRMNG: like=true + % v URL; PaleoDB: % v name; GBIF: % v query
    _has_wildcard = ("?" in taxon or "*" in taxon)

    # Která DB podporují wildcard dotazy a vracejí wildcard_hits (genus expansion):
    # WoRMS/IRMNG:    like=true + % v URL → vrací list hits
    # PaleoDB/Fossilworks: % v name parametru → vrací records[]
    # ITIS:           searchByScientificNameWildcard → wildcard nativně + wildcard_hits
    # CoL:            q parametr + /children endpoint → wildcard_hits
    # GBIF:           /species/{key}/children nebo /species/search?genus= → wildcard_hits
    # IPNI:           q parametr, prefix match → wildcard_hits
    # Tropicos:       type=wildcard v API → wildcard_hits
    # IFPNI:          q parametr, prefix match → wildcard_hits
    # ZooBank/Plazi/BioLib/Mikrotax: wildcards nemají efekt nebo API nepodporuje
    _WC_SUPPORTED = frozenset({"WoRMS", "IRMNG", "PaleoDB", "Fossilworks", "ITIS",
                                "CoL", "GBIF", "IPNI", "Tropicos", "IFPNI"})

    def _wc_for_db(name: str, db_name: str) -> str:
        """Převede uživatelské wildcards na formát pro konkrétní DB.
        DB mimo _WC_SUPPORTED dostanou čisté jméno bez wildcards.
        """
        if not _has_wildcard:
            return name
        if db_name not in _WC_SUPPORTED:
            # Pro DB bez wildcard podpory odstraníme wildcards a použijeme čistý název
            return name.replace("*", "").replace("?", "").strip()
        if db_name in ("WoRMS", "IRMNG"):
            # AphiaRecordsByName: % v path + like=true; _ = jeden znak
            return name.replace("*", "%25").replace("?", "_")
        if db_name in ("PaleoDB", "Fossilworks"):
            # data1.2: % jako wildcard v name=
            return name.replace("*", "%25").replace("?", "_")
        if db_name == "ITIS":
            # searchByScientificNameWildcard: % jako wildcard
            return name.replace("*", "%25").replace("?", "_")
        if db_name == "CoL":
            # ChecklistBank: partial/fuzzy match přes q=
            return name.replace("*", "").replace("?", "").strip()
        return name

    # Taxonomické klíče pro hierarchii (v pořadí pro zobrazení)
    # Timeout předán jako parametr z hlavního threadu — žádný st.session_state přístup z workeru
    TO = timeout

    for db in databases:
        cfg = TAXONOMIC_DATABASES.get(db)
        if not cfg:
            continue
        db_result: Dict[str, Any] = {"found": False}
        # Jméno taxonu upravené pro tuto konkrétní DB (wildcards)
        taxon_for_db = _wc_for_db(taxon, db)

        try:
            # ═══════════════════════════════════════════════════════
            # BioLib.cz  –  HTML scraping, 2 fallback URL
            # ═══════════════════════════════════════════════════════
            if db == "BioLib":
                # Nové BioLib REST API — přímý odkaz přes tx_id
                taxon_plus = _plus(taxon_for_db)
                api_url = (f"https://www.biolib.cz/rest/findname/"
                           f"?name={taxon_plus}"
                           f"&client=nmapp1&codeword=ZVlaleNM48%23")
                db_result["url"] = f"https://www.biolib.cz/en/search/?q={taxon_plus}"
                try:
                    r = _safe_get(api_url, timeout=TO,
                                  headers={"Accept": "application/json",
                                           "User-Agent": "HyolithaTaxonValidator/1.0"})
                    if r.status_code == 200:
                        data = r.json()
                        # BioLib REST API vrací success:true + status:"found"/"not found"
                        # nebo alternativně status:"ok" s results listem
                        _bl_names = (data.get("names") or
                                     data.get("results") or
                                     (data.get("data", []) if isinstance(data.get("data"), list) else []))
                        _bl_found_status = (data.get("status","") in ("found","ok") or
                                            bool(_bl_names))
                        if _bl_found_status and _bl_names:
                            names = _bl_names
                            # Filtr: přijmi jen záznamy kde jméno odpovídá
                            exact_bl = [h for h in names
                                        if _name_matches(h.get("tx_name","") or h.get("name",""), taxon)]
                            hit = (exact_bl[0] if exact_bl else None)
                            if hit:
                                tx_id  = hit.get("tx_id", "")
                                tx_name = hit.get("tx_name", taxon_for_db)
                                tx_auth = hit.get("tx_authority", "")
                                tx_rank_num = str(hit.get("tx_rank",""))
                                RANK_MAP = {"10":"kingdom","20":"phylum","30":"class",
                                            "40":"order","50":"family","60":"genus",
                                            "100":"species","120":"subspecies","150":"genus",
                                            "160":"species"}
                                rank_str = RANK_MAP.get(tx_rank_num, f"rank_{tx_rank_num}")
                                if tx_id:
                                    db_result["found"] = True
                                    db_result["url"]   = f"https://www.biolib.cz/cz/taxon/id{tx_id}/"
                                    db_result["details"] = {
                                        "name":      tx_name,
                                        "authority": tx_auth,
                                        "rank":      rank_str,
                                        "tx_id":     tx_id,
                                        "tn_id":     hit.get("tn_id",""),
                                        "match":     str(hit.get("match","")),
                                    }
                        elif data.get("status") == "not found":
                            db_result["found"] = False
                    elif r.status_code == 403:
                        db_result["error"] = "403 – IP není na whitelistu BioLib API"
                        db_result["url"]   = f"https://www.biolib.cz/en/search/?q={taxon_plus}"
                    else:
                        db_result["error"] = f"HTTP {r.status_code}"
                except Exception as e_bl:
                    db_result["error"] = str(e_bl)[:80]

            # ═══════════════════════════════════════════════════════
            # Catalogue of Life  –  ChecklistBank REST API
            # ═══════════════════════════════════════════════════════
            elif db == "CoL":
                _col_search_url = cfg.get("search_url", "https://www.catalogueoflife.org/data/search?q=")
                db_result["url"] = f"{_col_search_url}{_enc(taxon)}"
                # Zkus nejprve s datasetKey=COL (omezí na CoL dataset s alfanum. ID)
                # Fallback bez datasetKey pokud vrátí 400
                _col_data = None
                for _col_params in (
                    f"?q={_enc(taxon)}&limit=20&datasetKey=COL",
                    f"?q={_enc(taxon)}&limit=20",
                ):
                    r = _safe_get(f"{cfg['api_base']}{_col_params}",
                                  headers={"Accept":"application/json"}, timeout=TO)
                    if r.status_code == 200 and r.content:
                        _col_data = r.json()
                        break
                    elif r.status_code == 400:
                        continue  # zkus bez datasetKey
                    elif r.status_code == 403:
                        db_result["error"] = "API nedostupné (403)"
                        break
                    else:
                        db_result["error"] = f"HTTP {r.status_code}"
                        break

                if _col_data is not None:
                    items = _col_data.get("result",[])
                    exact = [it for it in items
                             if _name_matches(it.get("name","") or "", taxon)]
                    if exact:
                        it  = exact[0]
                        # Usage ID je v usage.id (alfanum klíč pro URL), ne v item["id"]
                        usage = it.get("usage") or {}
                        col_usage_id = (usage.get("id","") or
                                        it.get("usageId","") or
                                        it.get("id",""))
                        # Klasifikace
                        cl = {c["rank"].lower(): c["name"]
                              for c in (it.get("classification") or
                                        usage.get("classification",[]))
                              if c.get("rank")}
                        name_str = (it.get("name","") or
                                    usage.get("label","") or
                                    usage.get("name",""))
                        rank_str = (it.get("rank","") or
                                    usage.get("rank","") or "").lower()
                        status_str = (it.get("status","") or
                                      usage.get("status","") or "")
                        if col_usage_id and not col_usage_id.isdigit():
                            col_url = f"{cfg['browse_url']}{col_usage_id}"
                        else:
                            col_url = f"{_col_search_url}{_enc(taxon)}"
                        db_result["found"] = True
                        db_result["url"]   = col_url
                        db_result["details"] = {
                            "name":     name_str,
                            "rank":     rank_str,
                            "status":   status_str,
                            "kingdom":  cl.get("kingdom",""),
                            "phylum":   cl.get("phylum",""),
                            "class":    cl.get("class",""),
                            "order":    cl.get("order",""),
                            "family":   cl.get("family",""),
                            "genus":    cl.get("genus",""),
                            "col_id":   col_usage_id,
                        }
                        # Genus expansion: /nameusage/{id}/children pro rod-level záznamy
                        if _has_wildcard and col_usage_id and rank_str in ("genus","family","order"):
                            try:
                                rc = _safe_get(
                                    f"https://api.catalogueoflife.org/nameusage/{col_usage_id}/children?limit=100",
                                    headers={"Accept":"application/json"}, timeout=TO)
                                if rc.status_code == 200 and rc.content:
                                    kids = rc.json().get("result", [])
                                    if kids:
                                        db_result["wildcard_hits"] = [
                                            {"name":   (k.get("name","") or
                                                        (k.get("usage") or {}).get("label","") or
                                                        (k.get("usage") or {}).get("name","")),
                                             "rank":   (k.get("rank","") or
                                                        (k.get("usage") or {}).get("rank","")).lower(),
                                             "status": ((k.get("usage") or {}).get("status","") or
                                                        k.get("status","")),
                                             "col_id": ((k.get("usage") or {}).get("id","") or k.get("id",""))}
                                            for k in kids
                                            if k.get("name") or (k.get("usage") or {}).get("label")
                                        ]
                            except Exception:
                                pass
                elif r.status_code == 403:
                    db_result["error"] = "API nedostupné (403)"
                    db_result["url"] = f"{_col_search_url}{_enc(taxon)}"
                else:
                    db_result["error"] = f"HTTP {r.status_code}"

            # ═══════════════════════════════════════════════════════
            # IFPNI  –  International Fossil Plant Names Index
            # ═══════════════════════════════════════════════════════
            elif db == "IFPNI":
                # IFPNI API: /api/v1/name/search?q=<n>
                # Wildcard: limit zvýšen na 100, akceptuje prefix shodu
                query_name = taxon_for_db
                _ifpni_limit = 100 if _has_wildcard else 10
                r = _safe_get(
                    f"{cfg['api_search']}{_enc(query_name)}&limit={_ifpni_limit}",
                    headers={"Accept": "application/json"}, timeout=TO)
                if r.status_code == 200:
                    try:
                        data  = r.json()
                        items = (data if isinstance(data, list)
                                 else data.get("results", data.get("data", [])))
                        if _has_wildcard:
                            _ifpni_base = query_name.rstrip("*%?_").lower()
                            matched_ifpni = [x for x in items
                                             if (x.get("fullName","") or x.get("name","") or
                                                 x.get("scientificName","") or "").lower().startswith(_ifpni_base)]
                            if matched_ifpni:
                                it = matched_ifpni[0]
                                nid = str(it.get("id","") or it.get("nameId","") or
                                          it.get("ipniId","") or it.get("recordId","") or "")
                                name_str = (it.get("fullName","") or it.get("name","")
                                            or it.get("scientificName","") or it.get("taxonName",""))
                                db_result["found"] = True
                                db_result["url"] = (f"https://ifpni.org/name/{nid}"
                                                    if nid and nid != "None" and len(nid) > 2
                                                    else f"https://ifpni.org/search?q={_enc(name_str or taxon)}")
                                db_result["details"] = {
                                    "name":     name_str,
                                    "rank":     (it.get("rank","") or it.get("taxonRank","")).lower(),
                                    "status":   it.get("nomenclaturalStatus","") or it.get("status",""),
                                    "family":   it.get("family",""),
                                    "authors":  it.get("authors","") or it.get("authorsString",""),
                                    "year":     str(it.get("year","") or it.get("publishYear","")),
                                    "ifpni_id": nid,
                                }
                                if len(matched_ifpni) > 1:
                                    db_result["wildcard_hits"] = [
                                        {"name":    (x.get("fullName","") or x.get("name","") or
                                                     x.get("scientificName","")),
                                         "rank":    (x.get("rank","") or x.get("taxonRank","")).lower(),
                                         "status":  x.get("nomenclaturalStatus","") or x.get("status",""),
                                         "authors": x.get("authors","") or x.get("authorsString",""),
                                         "ifpni_id": str(x.get("id","") or x.get("nameId",""))}
                                        for x in matched_ifpni
                                    ]
                        else:
                            taxon_lower_ifpni = taxon.lower().strip()
                            items_exact = [x for x in items
                                           if (x.get("fullName","") or x.get("name","") or
                                               x.get("scientificName","") or "").lower().strip() == taxon_lower_ifpni]
                            if items_exact:
                                it  = items_exact[0]
                                nid = str(it.get("id","") or it.get("nameId","") or
                                          it.get("ipniId","") or it.get("recordId","") or "")
                                name_str = (it.get("fullName","") or it.get("name","")
                                            or it.get("scientificName","") or it.get("taxonName",""))
                                ifpni_record_url = (f"https://ifpni.org/name/{nid}"
                                                    if nid and nid != "None" and len(nid) > 2
                                                    else f"https://ifpni.org/search?q={_enc(name_str or taxon)}")
                                db_result["found"] = True
                                db_result["url"]   = ifpni_record_url
                                db_result["details"] = {
                                    "name":     name_str,
                                    "rank":     (it.get("rank","") or it.get("taxonRank","")).lower(),
                                    "status":   it.get("nomenclaturalStatus","") or it.get("status",""),
                                    "family":   it.get("family",""),
                                    "authors":  it.get("authors","") or it.get("authorsString",""),
                                    "year":     str(it.get("year","") or it.get("publishYear","")),
                                    "ifpni_id": nid,
                                }
                    except Exception:
                        db_result["found"] = False
                elif r.status_code == 404:
                    db_result["found"] = False
                else:
                    db_result["error"] = f"HTTP {r.status_code}"
                db_result.setdefault("url", cfg["search_page"] + _enc(taxon_for_db))

            # ═══════════════════════════════════════════════════════
            # Fossilworks  –  PaleoBioDB data1.2 API → Fossilworks browse URL
            # ═══════════════════════════════════════════════════════
            elif db == "Fossilworks":
                # Fossilworks = veřejné rozhraní PaleoDB → používáme stejné data1.2 API
                # taxa/single pro přesnou shodu, taxa/list pro wildcard
                _FW_RANK = {2:"druh (species)", 3:"poddruh (subspecies)", 5:"rod (genus)",
                            9:"čeleď (family)", 13:"řád (order)", 17:"třída (class)",
                            20:"kmen (phylum)", 23:"říše (kingdom)"}
                _FW_STATUS = {"belongs to": "platné jméno", "subjective synonym of": "subjektivní synonymum",
                              "objective synonym of": "objektivní synonymum", "nomen dubium": "nomen dubium",
                              "nomen nudum": "nomen nudum", "replaced by": "nahrazeno"}
                db_result["url"] = f"{cfg['fallback_url']}{_enc(taxon)}"
                _fw_found = False

                # 1) taxa/single — přesná shoda
                try:
                    _fs_url = (f"https://paleobiodb.org/data1.2/taxa/single.json"
                               f"?name={_enc(taxon)}&show=full,classext&vocab=pbdb")
                    r = _safe_get(_fs_url, timeout=TO)
                    if r.status_code == 200 and r.content:
                        _fs_recs = r.json().get("records", [])
                        if _fs_recs:
                            rec  = _fs_recs[0]
                            pname = (rec.get("taxon_name","") or rec.get("name","") or rec.get("nam",""))
                            taxno = str(rec.get("taxon_no","") or re.sub(r"\D","", rec.get("oid","") or ""))
                            if _name_matches(pname, taxon) and taxno:
                                _fw_found = True
                                rnk = rec.get("taxon_rank","") or rec.get("rank","") or rec.get("rnk","")
                                sta = rec.get("taxon_status","") or rec.get("status","") or rec.get("sta","")
                                ext = rec.get("is_extant","") or rec.get("ext","")
                                db_result["found"] = True
                                db_result["url"]   = f"https://fossilworks.org/?a=taxonInfo&taxon_no={taxno}"
                                db_result["details"] = {
                                    "název (name)":       pname,
                                    "rank":               rnk if isinstance(rnk,str) else _FW_RANK.get(int(rnk),str(rnk)),
                                    "status":             _FW_STATUS.get(sta, sta),
                                    "vyhynulý (extinct)": "ano (fosilní)" if str(ext) in ("0","false","False") else ("ne (recentní)" if str(ext) in ("1","true","True") else ""),
                                    "kmen (phylum)":      rec.get("phylum","") or rec.get("phm",""),
                                    "třída (class)":      rec.get("class","") or rec.get("cll",""),
                                    "řád (order)":        rec.get("order","") or rec.get("ord",""),
                                    "čeleď (family)":     rec.get("family","") or rec.get("fml",""),
                                    "rod (genus)":        rec.get("genus","") or rec.get("gnl",""),
                                    "pbdb_id":            taxno,
                                }
                except Exception:
                    pass

                # 2) taxa/list pro wildcard dotaz
                if _has_wildcard and not _fw_found or (not _fw_found):
                    try:
                        _fw_list_url = (f"https://paleobiodb.org/data1.2/taxa/list.json"
                                        f"?{'match_name' if _has_wildcard else 'base_name'}={_enc(taxon_for_db)}"
                                        f"&show=class&rowcount=true&vocab=pbdb&limit=200")
                        r2 = _safe_get(_fw_list_url, timeout=TO)
                        if r2.status_code == 200 and r2.content:
                            _fw_recs = r2.json().get("records", [])
                            if _fw_recs and not _fw_found:
                                rec2  = next((x for x in _fw_recs if _name_matches(
                                    x.get("taxon_name","") or x.get("nam",""), taxon)), _fw_recs[0])
                                txn2  = str(rec2.get("taxon_no","") or re.sub(r"\D","", rec2.get("oid","") or ""))
                                pnm2  = rec2.get("taxon_name","") or rec2.get("nam","")
                                if txn2:
                                    db_result["found"] = True
                                    db_result["url"]   = f"https://fossilworks.org/?a=taxonInfo&taxon_no={txn2}"
                                    rnk2 = rec2.get("taxon_rank","") or rec2.get("rnk","")
                                    sta2 = rec2.get("taxon_status","") or rec2.get("sta","")
                                    db_result["details"] = {
                                        "název (name)":   pnm2,
                                        "rank":           rnk2 if isinstance(rnk2,str) else _FW_RANK.get(int(rnk2),str(rnk2)),
                                        "status":         _FW_STATUS.get(sta2, sta2),
                                        "kmen (phylum)":  rec2.get("phylum","") or rec2.get("phm",""),
                                        "třída (class)":  rec2.get("class","") or rec2.get("cll",""),
                                        "řád (order)":    rec2.get("order","") or rec2.get("ord",""),
                                        "čeleď (family)": rec2.get("family","") or rec2.get("fml",""),
                                        "pbdb_id":        txn2,
                                    }
                            if _has_wildcard and _fw_recs and len(_fw_recs) > 1:
                                db_result["wildcard_hits"] = [
                                    {"name":   (x.get("taxon_name","") or x.get("nam","")),
                                     "rank":   (x.get("taxon_rank","") or x.get("rank","") or
                                                _FW_RANK.get(int(x.get("rnk",0) or 0), str(x.get("rnk","")))),
                                     "status": _FW_STATUS.get(
                                         x.get("taxon_status","") or x.get("sta",""),
                                         x.get("taxon_status","") or x.get("sta","")),
                                     "pbdb_id": (x.get("taxon_no","") or re.sub(r"\D","", x.get("oid","") or "")),
                                     "taxno":   (x.get("taxon_no","") or re.sub(r"\D","", x.get("oid","") or ""))}
                                    for x in _fw_recs
                                ]
                    except Exception:
                        pass

            # ═══════════════════════════════════════════════════════
            # GBIF  –  /species/match s fallback na /species/search
            # Pro wildcard/genus: /species/{key}/children + /species/search?genus=X
            # ═══════════════════════════════════════════════════════
            elif db == "GBIF":
                # 1) Přesný match (nebo pro wildcard: match na čisté jméno rodu)
                _gbif_clean = taxon_for_db  # bez wildcard znaků pro /match
                r    = _safe_get(f"{cfg['api_match']}?name={_enc(_gbif_clean)}&verbose=false", timeout=TO)
                data = r.json() if r.status_code == 200 and r.content else {}
                if data.get("matchType") not in ("NONE", None) and data.get("usageKey"):
                    key = data.get("usageKey","")
                    db_result["found"] = True
                    db_result["url"]   = f"{cfg['browse_url']}{key}"
                    db_result["details"] = {
                        "name":       data.get("scientificName",""),
                        "rank":       data.get("rank","").lower(),
                        "status":     data.get("status",""),
                        "confidence": str(data.get("confidence","")),
                        "kingdom":    data.get("kingdom",""),
                        "phylum":     data.get("phylum",""),
                        "class":      data.get("class_",data.get("class","")),
                        "order":      data.get("order",""),
                        "family":     data.get("family",""),
                        "genus":      data.get("genus",""),
                        "species":    data.get("species",""),
                        "gbif_key":   str(key),
                    }
                    # Genus expansion: stáhni dětské taxony přes /species/{key}/children
                    if _has_wildcard and key and data.get("rank","").upper() in ("GENUS", "FAMILY", "ORDER"):
                        try:
                            rc = _safe_get(
                                f"https://api.gbif.org/v1/species/{key}/children?limit=100",
                                timeout=TO)
                            if rc.status_code == 200 and rc.content:
                                kids = rc.json().get("results", [])
                                if kids:
                                    db_result["wildcard_hits"] = [
                                        {"name":     k.get("canonicalName","") or k.get("scientificName",""),
                                         "rank":     k.get("rank","").lower(),
                                         "status":   k.get("taxonomicStatus",""),
                                         "gbif_key": str(k.get("key","") or k.get("nubKey",""))}
                                        for k in kids if k.get("canonicalName") or k.get("scientificName")
                                    ]
                        except Exception:
                            pass
                else:
                    # 2) Fallback: full-text search
                    r2 = _safe_get(f"{cfg['api_search']}?q={_enc(_gbif_clean)}&limit=10&status=ACCEPTED", timeout=TO)
                    if r2.status_code == 200 and r2.content:
                        res2  = r2.json().get("results",[])
                        res2_exact = [x for x in res2
                                      if _name_matches(x.get("canonicalName","") or x.get("scientificName","") or "", taxon)]
                        res2 = res2_exact if res2_exact else []
                        if res2:
                            it2 = res2[0]
                            key2 = it2.get("key","") or it2.get("nubKey","")
                            db_result["found"] = True
                            db_result["url"]   = f"{cfg['browse_url']}{key2}" if key2 else f"{cfg['search_url']}{_enc(taxon)}"
                            db_result["details"] = {
                                "name":    it2.get("scientificName",""),
                                "rank":    it2.get("rank","").lower(),
                                "status":  it2.get("taxonomicStatus",""),
                                "kingdom": it2.get("kingdom",""),
                                "phylum":  it2.get("phylum",""),
                                "class":   it2.get("class",""),
                                "order":   it2.get("order",""),
                                "family":  it2.get("family",""),
                                "genus":   it2.get("genus",""),
                                "gbif_key":str(key2),
                            }
                            # Genus expansion fallback: /species/search?genus=X
                            if _has_wildcard and key2:
                                try:
                                    genus_clean = _gbif_clean.rstrip("%*?")
                                    rg = _safe_get(
                                        f"https://api.gbif.org/v1/species/search?"
                                        f"genus={_enc(genus_clean)}&rank=SPECIES&status=ACCEPTED&limit=100",
                                        timeout=TO)
                                    if rg.status_code == 200 and rg.content:
                                        spp = rg.json().get("results", [])
                                        if spp:
                                            db_result["wildcard_hits"] = [
                                                {"name":     s.get("canonicalName","") or s.get("scientificName",""),
                                                 "rank":     s.get("rank","").lower(),
                                                 "status":   s.get("taxonomicStatus",""),
                                                 "gbif_key": str(s.get("key","") or s.get("nubKey",""))}
                                                for s in spp if s.get("canonicalName") or s.get("scientificName")
                                            ]
                                except Exception:
                                    pass
                        else:
                            db_result["url"] = f"{cfg['search_url']}{_enc(taxon)}"
                    else:
                        db_result["url"] = f"{cfg['search_url']}{_enc(taxon)}"

            # ═══════════════════════════════════════════════════════
            # IPNI  –  REST API
            # ═══════════════════════════════════════════════════════
            elif db == "IPNI":
                # Fallback search URL v novém formátu
                db_result["url"] = f"https://www.ipni.org/search?q={_enc(taxon_for_db)}"
                # Pro wildcard: zvyš limit a použij wildcard query; jinak přesná shoda
                _ipni_q   = taxon_for_db if _has_wildcard else taxon
                _ipni_lim = 100 if _has_wildcard else 10
                r = _safe_get(f"{cfg['api_base']}?q={_enc(_ipni_q)}&page=1&perPage={_ipni_lim}",
                              headers={"Accept":"application/json"}, timeout=TO)
                if r.status_code == 200:
                    items = r.json().get("results",[])
                    if _has_wildcard:
                        # Wildcard: akceptuj vše co začíná základem dotazu
                        _ipni_base = _ipni_q.rstrip("*%?").lower()
                        matched_ipni = [x for x in items
                                        if (x.get("name","") or "").lower().startswith(_ipni_base)]
                        if matched_ipni:
                            it = matched_ipni[0]
                            nid = it.get("id","")
                            db_result["found"] = True
                            db_result["url"] = (f"{cfg['browse_url']}{nid}" if nid
                                                else f"https://www.ipni.org/search?q={_enc(taxon_for_db)}")
                            db_result["details"] = {
                                "name":    it.get("name",""),
                                "rank":    it.get("rank","").lower(),
                                "family":  it.get("family",""),
                                "authors": it.get("authors",""),
                                "year":    str(it.get("publicationYear","")),
                                "ipni_id": nid,
                            }
                            if len(matched_ipni) > 1:
                                db_result["wildcard_hits"] = [
                                    {"name":     x.get("name",""),
                                     "rank":     x.get("rank","").lower(),
                                     "status":   "",
                                     "ipni_id":  x.get("id",""),
                                     "authors":  x.get("authors",""),
                                     "family":   x.get("family","")}
                                    for x in matched_ipni
                                ]
                    else:
                        exact_ipni = [x for x in items
                                      if _name_matches(x.get("name","") or "", taxon)]
                        if exact_ipni:
                            it  = exact_ipni[0]
                            nid = it.get("id","")
                            db_result["found"] = True
                            db_result["url"]   = (f"{cfg['browse_url']}{nid}"
                                                  if nid
                                                  else f"https://www.ipni.org/search?q={_enc(taxon_for_db)}")
                            db_result["details"] = {
                                "name":    it.get("name",""),
                                "rank":    it.get("rank","").lower(),
                                "family":  it.get("family",""),
                                "authors": it.get("authors",""),
                                "year":    str(it.get("publicationYear","")),
                                "ipni_id": nid,
                            }
                else:
                    db_result["error"] = f"HTTP {r.status_code}"

            # ═══════════════════════════════════════════════════════
            # IRMNG  –  AphiaRecordsByName REST, fallback like=true
            # ═══════════════════════════════════════════════════════
            elif db == "IRMNG":
                for like in ("false","true"):
                    r = _safe_get(
                        f"{cfg['search_url']}{_enc(taxon)}?like={like}&marine_only=false",
                        timeout=TO)
                    if r.status_code in (200,206):
                        try:
                            data = r.json()
                        except Exception:
                            data = []
                        if isinstance(data,list) and data:
                            # Filtr: přesná shoda jména (case-insensitive) pro like=false
                            # Pro like=true akceptujeme i částečné shody
                            taxon_lower = taxon.lower().strip()
                            if like == "false":
                                exact = [r for r in data
                                         if _name_matches(r.get("scientificname","") or "", taxon)]
                                filtered = exact if exact else []
                            else:
                                # like=true: akceptuj záznamy kde jméno začíná dotazem
                                filtered = [r for r in data
                                            if r.get("scientificname","").lower().strip().startswith(
                                               taxon_lower.rstrip("%").lower())]
                                if not filtered:
                                    filtered = data  # fallback
                            if filtered:
                                rec = filtered[0]
                                aid = rec.get("AphiaID","")
                                db_result["found"] = True
                                db_result["url"]   = (f"{cfg['browse_url']}{aid}" if aid
                                      else f"https://www.irmng.org/aphia.php?p=search&searchtype=taxdetails&AphiaID={_enc(taxon)}")
                                db_result["details"] = {
                                    "name":     rec.get("scientificname",""),
                                    "rank":     rec.get("rank","").lower(),
                                    "status":   rec.get("status",""),
                                    "kingdom":  rec.get("kingdom",""),
                                    "phylum":   rec.get("phylum",""),
                                    "class":    rec.get("class",""),
                                    "order":    rec.get("order",""),
                                    "family":   rec.get("family",""),
                                    "genus":    rec.get("genus",""),
                                    "irmng_id": str(aid),
                                }
                                # Wildcard: ulož všechny nalezené záznamy
                                if _has_wildcard and len(filtered) > 1:
                                    db_result["wildcard_hits"] = [
                                        {"name": r.get("scientificname",""), "rank": r.get("rank","").lower(),
                                         "status": r.get("status",""), "irmng_id": str(r.get("AphiaID",""))}
                                        for r in filtered
                                    ]
                                break
                    elif r.status_code == 204:
                        db_result["error"] = "Nenalezeno (204)"
                        break

            # ═══════════════════════════════════════════════════════
            # ITIS  –  JSON service + hierarchy lookup
            # ═══════════════════════════════════════════════════════
            elif db == "ITIS":
                r = _safe_get(f"{cfg['api_base']}?srchKey={_enc(taxon_for_db)}",
                              headers={"Accept":"application/json"}, timeout=TO)
                if r.status_code == 200:
                    try:
                        snames = [x for x in (r.json().get("scientificNames",[]) or []) if x]
                        if snames:
                            it  = snames[0]
                            tsn = it.get("tsn","")
                            db_result["found"] = True
                            db_result["url"]   = f"{cfg['browse_url']}{tsn}" if tsn else "https://www.itis.gov/"
                            hier = {}
                            if tsn:
                                try:
                                    r2 = _safe_get(f"{cfg['api_hier']}?tsn={tsn}",
                                                   headers={"Accept":"application/json"}, timeout=TO)
                                    if r2.status_code == 200:
                                        hier = {item.get("rankName","").lower(): item.get("taxonName","")
                                                for item in (r2.json().get("hierarchyList") or [])
                                                if item and item.get("rankName")}
                                except Exception:
                                    pass
                            db_result["details"] = {
                                "name":    it.get("combinedName",""),
                                "tsn":     str(tsn),
                                "kingdom": hier.get("kingdom",""),
                                "phylum":  hier.get("phylum","") or hier.get("division",""),
                                "class":   hier.get("class",""),
                                "order":   hier.get("order",""),
                                "family":  hier.get("family",""),
                                "genus":   hier.get("genus",""),
                            }
                            # Wildcard: sbírej všechny výsledky jako wildcard_hits
                            if _has_wildcard and len(snames) > 1:
                                db_result["wildcard_hits"] = [
                                    {"name":   s.get("combinedName",""),
                                     "rank":   "",
                                     "status": "",
                                     "tsn":    str(s.get("tsn",""))}
                                    for s in snames
                                ]
                        else:
                            # Fallback: wildcard search (pro genus expansion nebo nenalezená jména)
                            _itis_wc_key = _enc(taxon_for_db) if _has_wildcard else f"{_enc(taxon)}%25"
                            rw = _safe_get(f"{cfg['api_any']}?srchKey={_itis_wc_key}",
                                           headers={"Accept":"application/json"}, timeout=TO)
                            if rw.status_code == 200 and rw.content:
                                sw = [x for x in (rw.json().get("scientificNames",[]) or []) if x]
                                if sw:
                                    tsn2 = sw[0].get("tsn","")
                                    db_result["found"] = True
                                    db_result["url"]   = f"{cfg['browse_url']}{tsn2}" if tsn2 else "https://www.itis.gov/"
                                    db_result["details"] = {"name": sw[0].get("combinedName",""), "tsn": str(tsn2)}
                                    # Wildcard: sbírej všechny výsledky
                                    if _has_wildcard and len(sw) > 1:
                                        db_result["wildcard_hits"] = [
                                            {"name":   s.get("combinedName",""),
                                             "rank":   "",
                                             "status": "",
                                             "tsn":    str(s.get("tsn",""))}
                                            for s in sw
                                        ]
                    except Exception as e:
                        db_result["error"] = str(e)[:60]
                else:
                    db_result["error"] = f"HTTP {r.status_code}"

            # ═══════════════════════════════════════════════════════
            # Mikrotax  –  3 sub-databáze: pforams, nannotax, ostrakódi
            # ═══════════════════════════════════════════════════════
            elif db == "Mikrotax":
                genus = taxon.split()[0].lower()
                sub_urls = [
                    ("foraminifera",  cfg["pforams_url"]  + _enc(taxon_for_db)),
                    ("nanofossils",   cfg["nannotax_url"] + _enc(taxon_for_db)),
                    ("ostracods",     cfg["ostracod_url"] + _enc(taxon_for_db)),
                ]
                db_result["url"] = cfg["browse_url"]
                for sub_name, sub_url in sub_urls:
                    try:
                        r = _safe_get(sub_url, timeout=TO,
                                      headers={"User-Agent":"Mozilla/5.0"})
                        if r.status_code == 200:
                            # Zpřísněný match: jméno musí být v <title> nebo <h1>/<h2>,
                            # ne jen kdekoliv v HTML (to by bylo false positive pro run-of-site menu)
                            import re as _re2
                            title_m = _re2.search(r'<title[^>]*>(.*?)</title>', r.text, _re2.I | _re2.S)
                            h1_m    = _re2.search(r'<h[12][^>]*>(.*?)</h[12]>', r.text, _re2.I | _re2.S)
                            title_text = (title_m.group(1) if title_m else "").lower()
                            h1_text    = (h1_m.group(1)    if h1_m    else "").lower()
                            # found=True jen pokud jméno odpovídá titulku/nadpisu stránky
                            if genus in title_text or genus in h1_text:
                                db_result["found"] = True
                                db_result["url"]   = sub_url
                                db_result["details"] = {"group": sub_name}
                                break
                    except Exception:
                        pass

            # ═══════════════════════════════════════════════════════
            # PaleoDB  –  data1.2 REST API, show=attr,classext,parent
            # ═══════════════════════════════════════════════════════
            elif db == "PaleoDB":
                # PaleoDB data1.2 API — dev.paleobiodb.org
                # URL vzory dle dokumentace:
                # Jeden taxon:  /taxa/single.json?name=X&show=full&datainfo=true
                # Druhy rodu:   /taxa/list.json?base_name=X&rank=species&show=class&rowcount=true&datainfo=true
                # Subtaxony:    /taxa/list.json?base_name=X&show=class&rowcount=true&datainfo=true
                # Prefix:       /taxa/list.json?match_name=X%&show=class&rowcount=true&datainfo=true
                # Wildcards v jménech: % = sekvence znaků, _ = jeden znak, . = zkratka
                _RANK_NUM = {2:"druh (species)", 3:"poddruh (subspecies)", 5:"rod (genus)",
                             9:"čeleď (family)", 13:"řád (order)", 17:"třída (class)",
                             20:"kmen (phylum)", 23:"říše (kingdom)"}
                _STATUS_MAP = {"belongs to": "platné jméno",
                               "subjective synonym of": "subjektivní synonymum",
                               "objective synonym of": "objektivní synonymum",
                               "replaced by": "nahrazeno", "nomen dubium": "nomen dubium",
                               "nomen nudum": "nomen nudum",
                               "invalid subgroup of": "neplatný podtaxon",
                               "corrected as": "opraveno na"}

                # Konverze userových wildcards (* → %, ? → _) pro PBDB
                _pb_wc_name = (taxon.replace("*", "%").replace("?", "_")
                               if _has_wildcard else taxon)

                # Výchozí fallback URL — přesný odkaz na výsledky dle vzoru z dokumentace
                # URL vzory dle PBDB dokumentace:
                # Druhy rodu:  base_name=Canis&show=class&rowcount=true&datainfo=true
                # Prefix:      match_name=Can%25&show=class&rowcount=true&datainfo=true
                if _has_wildcard:
                    _pb_disp_url = (f"https://paleobiodb.org/data1.2/taxa/list.json"
                                    f"?match_name={_enc(_pb_wc_name)}"
                                    f"&show=class&rowcount=true&datainfo=true")
                else:
                    # Pro rod/druh: base_name vrátí rod + všechny podřazené taxony
                    _pb_genus_disp = taxon.split()[0]
                    _pb_disp_url = (f"https://paleobiodb.org/data1.2/taxa/list.json"
                                    f"?base_name={_enc(_pb_genus_disp)}"
                                    f"&show=class&rowcount=true&datainfo=true")
                db_result["url"] = _pb_disp_url
                _paleo_found = False

                # 1) taxa/single — přesná shoda jednoho taxonu
                try:
                    _ps_url = (f"https://paleobiodb.org/data1.2/taxa/single.json"
                               f"?name={_enc(taxon)}&show=full,classext&datainfo=false&vocab=pbdb")
                    r = _safe_get(_ps_url, timeout=TO)
                    if r.status_code == 200 and r.content:
                        _ps_recs = r.json().get("records", [])
                        if _ps_recs:
                            rec   = _ps_recs[0]
                            taxno = str(rec.get("taxon_no","") or rec.get("orig_no","") or
                                        re.sub(r"\D","", rec.get("oid","") or ""))
                            pname = (rec.get("taxon_name","") or rec.get("name","") or rec.get("nam",""))
                            if _name_matches(pname, taxon) and taxno:
                                _paleo_found = True
                                rnk = rec.get("taxon_rank","") or rec.get("rank","") or rec.get("rnk","")
                                sta = rec.get("taxon_status","") or rec.get("status","") or rec.get("sta","")
                                ext = rec.get("is_extant","") or rec.get("extant","") or rec.get("ext","")
                                db_result["found"] = True
                                # URL pro zobrazení: single dle vzoru
                                # URL: všechny druhy rodu dle vzoru base_name
                                _pb_genus_single = pname.split()[0] if pname else taxon.split()[0]
                                db_result["url"] = (f"https://paleobiodb.org/data1.2/taxa/list.json"
                                                    f"?base_name={_enc(_pb_genus_single)}"
                                                    f"&show=class&rowcount=true&datainfo=true")
                                db_result["details"] = {
                                    "název (name)":       pname,
                                    "rank":               rnk if isinstance(rnk,str) else _RANK_NUM.get(int(rnk), str(rnk)),
                                    "status":             _STATUS_MAP.get(sta, sta),
                                    "vyhynulý (extinct)": ("ne (recentní)" if str(ext) in ("1","true","True")
                                                           else "ano (fosilní)" if str(ext) in ("0","false","False") else ""),
                                    "kmen (phylum)":      rec.get("phylum","") or rec.get("phm",""),
                                    "třída (class)":      rec.get("class","") or rec.get("cll",""),
                                    "řád (order)":        rec.get("order","") or rec.get("ord",""),
                                    "čeleď (family)":     rec.get("family","") or rec.get("fml",""),
                                    "rod (genus)":        rec.get("genus","") or rec.get("gnl",""),
                                    "pbdb_id":            taxno,
                                }
                except Exception:
                    pass

                # 2) taxa/list — wildcard dotaz nebo fallback pro přesný dotaz
                if not _paleo_found or _has_wildcard:
                    try:
                        if _has_wildcard:
                            # Wildcard/prefix dotaz: match_name=X%  (dle PBDB docs)
                            _pl_url = (f"https://paleobiodb.org/data1.2/taxa/list.json"
                                       f"?match_name={_enc(_pb_wc_name)}"
                                       f"&show=class&rowcount=true&datainfo=false&vocab=pbdb&limit=200")
                        else:
                            # base_name = taxon + všechny subtaxony (dle PBDB docs)
                            _pl_url = (f"https://paleobiodb.org/data1.2/taxa/list.json"
                                       f"?base_name={_enc(taxon)}"
                                       f"&show=class&rowcount=true&datainfo=false&vocab=pbdb&limit=200")
                        r2 = _safe_get(_pl_url, timeout=TO)
                        if r2.status_code == 200 and r2.content:
                            _pl_j    = r2.json()
                            _pl_recs = _pl_j.get("records", [])
                            _pl_total = _pl_j.get("recordsFound", len(_pl_recs))
                            if _pl_recs and not _paleo_found:
                                _pl_match = next(
                                    (x for x in _pl_recs if _name_matches(
                                        x.get("taxon_name","") or x.get("nam",""), taxon)),
                                    _pl_recs[0])
                                taxno2 = str(_pl_match.get("taxon_no","") or
                                             re.sub(r"\D","", _pl_match.get("oid","") or ""))
                                pname2 = _pl_match.get("taxon_name","") or _pl_match.get("nam","")
                                rnk2   = (_pl_match.get("taxon_rank","") or _pl_match.get("rank","")
                                          or _pl_match.get("rnk",""))
                                sta2   = _pl_match.get("taxon_status","") or _pl_match.get("sta","")
                                ext2   = _pl_match.get("is_extant","") or _pl_match.get("ext","")
                                if taxno2:
                                    db_result["found"] = True
                                    # URL dle vzoru: všechny druhy rodu nebo subtaxony
                                    # URL vzory dle dokumentace
                                    if _has_wildcard:
                                        # Prefix → match_name s %
                                        db_result["url"] = (
                                            f"https://paleobiodb.org/data1.2/taxa/list.json"
                                            f"?match_name={_enc(_pb_wc_name)}"
                                            f"&show=class&rowcount=true&datainfo=true"
                                        )
                                    else:
                                        # Rod → base_name = rod (vrátí druhy rodu)
                                        _pb_genus2 = pname2.split()[0] if pname2 else taxon.split()[0]
                                        db_result["url"] = (
                                            f"https://paleobiodb.org/data1.2/taxa/list.json"
                                            f"?base_name={_enc(_pb_genus2)}"
                                            f"&show=class&rowcount=true&datainfo=true"
                                        )
                                    db_result["details"] = {
                                        "název (name)":       pname2,
                                        "rank":               rnk2 if isinstance(rnk2,str) else _RANK_NUM.get(int(rnk2), str(rnk2)),
                                        "status":             _STATUS_MAP.get(sta2, sta2),
                                        "vyhynulý (extinct)": ("ne (recentní)" if str(ext2) in ("1","true","True")
                                                               else "ano (fosilní)" if str(ext2) in ("0","false","False") else ""),
                                        "kmen (phylum)":      _pl_match.get("phylum","") or _pl_match.get("phm",""),
                                        "třída (class)":      _pl_match.get("class","") or _pl_match.get("cll",""),
                                        "řád (order)":        _pl_match.get("order","") or _pl_match.get("ord",""),
                                        "čeleď (family)":     _pl_match.get("family","") or _pl_match.get("fml",""),
                                        "rod (genus)":        _pl_match.get("genus","") or _pl_match.get("gnl",""),
                                        "pbdb_id":            taxno2,
                                        "subtaxonů celkem":   str(_pl_total),
                                    }
                            # Wildcard hits — všechny záznamy
                            if _has_wildcard and _pl_recs and len(_pl_recs) > 1:
                                db_result["wildcard_hits"] = [
                                    {"name":   (x.get("taxon_name","") or x.get("nam","")),
                                     "rank":   (x.get("taxon_rank","") or x.get("rank","") or
                                                _RANK_NUM.get(int(x.get("rnk",0) or 0), str(x.get("rnk","")))),
                                     "status": _STATUS_MAP.get(
                                         x.get("taxon_status","") or x.get("sta",""),
                                         x.get("taxon_status","") or x.get("sta","")),
                                     "pbdb_id": (x.get("taxon_no","") or
                                                 re.sub(r"\D","", x.get("oid","") or "")),
                                     "taxno":   (x.get("taxon_no","") or
                                                 re.sub(r"\D","", x.get("oid","") or ""))}
                                    for x in _pl_recs
                                ]
                    except Exception:
                        pass

                if not db_result.get("found"):
                    # Fallback URL dle vzoru: match_name pro prefix, base_name pro rod
                    if _has_wildcard:
                        db_result["url"] = (f"https://paleobiodb.org/data1.2/taxa/list.json"
                                            f"?match_name={_enc(_pb_wc_name)}"
                                            f"&show=class&rowcount=true&datainfo=true")
                    else:
                        _pb_genus_fb = taxon.split()[0]
                        db_result["url"] = (f"https://paleobiodb.org/data1.2/taxa/list.json"
                                            f"?base_name={_enc(_pb_genus_fb)}"
                                            f"&show=class&rowcount=true&datainfo=true")

            # ═══════════════════════════════════════════════════════
            # Plazi TreatmentBank
            # Strategie: 1) GgServer taxonomicName (přesné jméno)
            #            2) GgServer volný dotaz (q=)
            #            3) Genus-only dotaz
            # ═══════════════════════════════════════════════════════
            elif db == "Plazi":
                # Nový URL formát: https://tb.plazi.org/GgServer/search?fullText.ftQuery=hyolithes
                search_url = (f"https://tb.plazi.org/GgServer/search"
                              f"?fullText.ftQuery={_enc(taxon_for_db)}")
                db_result["url"] = search_url

                try:
                    rp = _safe_get(search_url,
                                   headers={"Accept": "application/json,text/xml,*/*"},
                                   timeout=TO)
                    if rp.status_code == 200:
                        # Zkus JSON parsování
                        plazi_parsed = None
                        try:
                            rj = rp.json()
                            recs = (rj.get("treatments") or rj.get("results")
                                    or (rj if isinstance(rj, list) else []))
                            total = rj.get("total", len(recs) if recs else 0)
                            if recs or total > 0:
                                plazi_parsed = {"total": total, "recs": recs}
                        except Exception:
                            pass
                        # XML fallback – ověř přítomnost jména taxonu v odpovědi
                        if plazi_parsed is None:
                            n_treat = rp.text.count("<treatment")
                            taxon_in_text = taxon.lower() in rp.text.lower()
                            if (n_treat > 0 or "treatmentId" in rp.text) and taxon_in_text:
                                plazi_parsed = {"total": n_treat, "recs": []}
                        if plazi_parsed and plazi_parsed["total"] > 0:
                            db_result["found"] = True
                            details: Dict = {"treatments": plazi_parsed["total"]}
                            if plazi_parsed["recs"]:
                                first = plazi_parsed["recs"][0]
                                tid = first.get("treatmentId") or first.get("id","")
                                if tid:
                                    db_result["url"] = f"https://treatmentbank.org/treatment/{tid}"
                                details["title"]   = str(first.get("treatmentTitle",
                                                         first.get("title","")))[:100]
                                details["authors"] = str(first.get("creators",
                                                         first.get("authors","")))[:80]
                                details["year"]    = str(first.get("date",
                                                         first.get("year","")))
                            db_result["details"] = details
                    else:
                        db_result["error"] = f"HTTP {rp.status_code}"
                except Exception as e_pl:
                    db_result["error"] = str(e_pl)[:80]

            # ═══════════════════════════════════════════════════════
            # WoRMS  –  AphiaRecordsByName REST API
            # Endpoint: /rest/AphiaRecordsByName/{name}?like={t/f}&marine_only=false
            # Strategie: 1) exact match  2) like=true  3) genus-only
            # ═══════════════════════════════════════════════════════
            elif db == "WoRMS":
                worms_base = cfg["aphia_name"]
                worms_found = False
                # Sekvence pokusů: exact → fuzzy → genus only
                worms_attempts = [
                    (taxon_for_db,  "false"),
                    (taxon_for_db,  "true"),
                    (taxon_for_db.split()[0], "true"),   # samotný rod
                ]
                for w_name, w_like in worms_attempts:
                    # WoRMS: jméno musí být v URL path, encoded
                    w_url = (f"{worms_base}{_enc(w_name)}"
                             f"?like={w_like}&marine_only=false&offset=1")
                    try:
                        r = _safe_get(w_url, timeout=TO,
                                      headers={"Accept": "application/json"})
                        if r.status_code == 200:
                            try:
                                data = r.json()
                            except Exception:
                                data = []
                            if isinstance(data, list) and data:
                                if w_like == "false":
                                    # Přesná shoda pro exact dotaz
                                    exact_w = [r for r in data
                                               if _name_matches(r.get("scientificname","") or "", taxon)]
                                    filtered_w = exact_w if exact_w else []
                                else:
                                    # like=true nebo genus-only: prefix match
                                    _q_clean = w_name.rstrip("%").lower()
                                    filtered_w = [r for r in data
                                                  if r.get("scientificname","").lower().startswith(_q_clean)]
                                    if not filtered_w:
                                        filtered_w = data
                                if filtered_w:
                                    rec = filtered_w[0]
                                    aid = rec.get("AphiaID", "")
                                    db_result["found"] = True
                                    db_result["url"]   = (f"{cfg['browse_url']}{aid}"
                                                          if aid else cfg["search_page"] + _enc(taxon_for_db))
                                    db_result["details"] = {
                                        "name":    rec.get("scientificname",""),
                                        "rank":    (rec.get("rank","") or "").lower(),
                                        "status":  rec.get("status",""),
                                        "kingdom": rec.get("kingdom",""),
                                        "phylum":  rec.get("phylum",""),
                                        "class":   rec.get("class",""),
                                        "order":   rec.get("order",""),
                                        "family":  rec.get("family",""),
                                        "genus":   rec.get("genus",""),
                                        "worms_id":str(aid),
                                    }
                                    # Wildcard: ulož všechny nalezené záznamy
                                    if _has_wildcard and len(filtered_w) > 1:
                                        db_result["wildcard_hits"] = [
                                            {"name": r.get("scientificname",""), "rank": (r.get("rank","") or "").lower(),
                                             "status": r.get("status",""), "worms_id": str(r.get("AphiaID",""))}
                                            for r in filtered_w
                                        ]
                                    worms_found = True
                                    break
                            # 200 ale prázdný list → pokračuj dalším pokusem
                        elif r.status_code == 204:
                            # 204 = databáze odepsala "nenalezeno" – zkus další variantu
                            continue
                        else:
                            db_result["error"] = f"HTTP {r.status_code}"
                            break
                    except Exception as e_w:
                        db_result["error"] = str(e_w)[:80]
                        break
                if not worms_found and not db_result.get("url"):
                    db_result["url"] = cfg["search_page"] + _enc(taxon_for_db)

            # ═══════════════════════════════════════════════════════
            # ZooBank  –  /Search.json (správný endpoint pro jména)
            # ═══════════════════════════════════════════════════════
            elif db == "ZooBank":
                # ZooBank vyžaduje CAPTCHA potvrzení při prvním přístupu z dané IP.
                # API /Search.json může vrátit prázdné výsledky dokud není CAPTCHA potvrzena.
                # Strategie: pokus o API, při prázdném výsledku vrátit search URL (ne false positive).
                zb_search_url = f"https://zoobank.org/Search#{_enc(taxon)}"
                db_result["url"] = zb_search_url
                found_zb = False
                for zb_url in [cfg["api_search"] + _enc(taxon_for_db),
                               cfg["api_names"]  + _enc(taxon)]:
                    try:
                        r = _safe_get(zb_url, headers={"Accept":"application/json"}, timeout=TO)
                        if r.status_code == 200:
                            try:
                                data = r.json()
                            except Exception:
                                # HTML odpověď = CAPTCHA blok
                                db_result["error"] = "Vyžaduje CAPTCHA potvrzení — otevři ZooBank v prohlížeči"
                                db_result["url"]   = "https://zoobank.org/"
                                break
                            items = data if isinstance(data,list) else data.get("results",[])
                            # Filtr: jméno musí odpovídat (ZooBank vrací mnoho nesouvisejících výsledků)
                            taxon_lower_zb = taxon.lower().strip()
                            matched = [it for it in items
                                       if taxon_lower_zb in (it.get("namestring","") or
                                                              it.get("label","") or
                                                              it.get("cleanprotonym","")).lower()]
                            if matched:
                                it   = matched[0]
                                uuid = (it.get("tnuuuid") or it.get("referenceuuid")
                                        or it.get("uuid",""))
                                name = (it.get("namestring") or it.get("label")
                                        or it.get("cleanprotonym",""))
                                db_result["found"] = True
                                db_result["url"]   = (f"{cfg['browse_url']}{uuid}"
                                                      if uuid else zb_search_url)
                                db_result["details"] = {
                                    "name":       name,
                                    "rank":       it.get("rankgroup","").lower(),
                                    "authors":    it.get("authors",""),
                                    "year":       str(it.get("year","")),
                                    "zoobank_id": uuid,
                                }
                                found_zb = True
                                break
                    except Exception:
                        pass
                if not found_zb and not db_result.get("error"):
                    # Žádný výsledek - může být CAPTCHA blok nebo taxon neregistrován
                    db_result["url"]   = zb_search_url
                    db_result["error"] = "Neověřeno (CAPTCHA nebo nenalezeno) — zkontroluj ručně"

            # ═══════════════════════════════════════════════════════
            # Tropicos (Missouri Botanical Garden)
            # REST API: /Name/Search?name=<n>&type=wildcard&format=json
            # ═══════════════════════════════════════════════════════
            elif db == "Tropicos":
                query_t = taxon_for_db
                trop_url = (f"{cfg['api_search']}{_enc(query_t)}{cfg['api_suffix']}")
                db_result["url"] = cfg["search_page"] + _enc(query_t)
                try:
                    r = _safe_get(trop_url, timeout=TO,
                                  headers={"Accept": "application/json"})
                    if r.status_code == 200:
                        try:
                            data  = r.json()
                            if isinstance(data, list) and data:
                                taxon_lower_t = taxon.lower().strip()
                                if _has_wildcard:
                                    # Wildcard: akceptuj vše co začíná základem dotazu
                                    _trop_base = query_t.rstrip("*%?").lower()
                                    matched_t = [x for x in data
                                                 if (x.get("ScientificName","") or "").lower().startswith(_trop_base)]
                                    if matched_t:
                                        it  = matched_t[0]
                                        nid = str(it.get("NameId",""))
                                        db_result["found"] = True
                                        db_result["url"]   = (f"https://www.tropicos.org/name/{nid}"
                                                               if nid else db_result["url"])
                                        db_result["details"] = {
                                            "name":      it.get("ScientificName",""),
                                            "rank":      it.get("RankAbbreviation","").lower(),
                                            "authors":   it.get("Author",""),
                                            "family":    it.get("Family",""),
                                            "status":    it.get("NomenclatureStatusName",""),
                                            "tropicos_id": nid,
                                        }
                                        if len(matched_t) > 1:
                                            db_result["wildcard_hits"] = [
                                                {"name":        x.get("ScientificName",""),
                                                 "rank":        x.get("RankAbbreviation","").lower(),
                                                 "status":      x.get("NomenclatureStatusName",""),
                                                 "authors":     x.get("Author",""),
                                                 "family":      x.get("Family",""),
                                                 "tropicos_id": str(x.get("NameId",""))}
                                                for x in matched_t
                                            ]
                                else:
                                    # Přesná shoda ONLY
                                    exact_t = [x for x in data
                                               if x.get("ScientificName","").lower().strip() == taxon_lower_t]
                                    if not exact_t:
                                        db_result["found"] = False
                                    else:
                                        it  = exact_t[0]
                                        nid = str(it.get("NameId",""))
                                        trop_record_url = (f"https://www.tropicos.org/name/{nid}"
                                                           if nid else db_result["url"])
                                        db_result["found"] = True
                                        db_result["url"]   = trop_record_url
                                        db_result["details"] = {
                                            "name":      it.get("ScientificName",""),
                                            "rank":      it.get("RankAbbreviation","").lower(),
                                            "authors":   it.get("Author",""),
                                            "family":    it.get("Family",""),
                                            "status":    it.get("NomenclatureStatusName",""),
                                            "tropicos_id": nid,
                                        }
                            elif isinstance(data, dict) and data.get("Error"):
                                db_result["found"] = False
                        except Exception:
                            db_result["found"] = False
                    else:
                        db_result["error"] = f"HTTP {r.status_code}"
                except Exception as e_t:
                    db_result["error"] = str(e_t)[:60]

            # ═══════════════════════════════════════════════════════
            # Global Names Verifier  –  GET + POST /api/v1/verifications
            # URL syntaxe: ?capitalize=on&names=Gompholites (funguje pro rod/druh)
            # POST syntaxe: {"nameStrings": [...], "preferredSources": [...]}
            # ═══════════════════════════════════════════════════════
            elif db == "GNverifier":
                _gnv_base = "https://verifier.globalnames.org/api/v1/verifications"
                # Odkaz pro uživatele — webové rozhraní s výsledky
                # URL pro webové zobrazení výsledku — mezery jako + (dle vzoru)
                _gnv_web_name = taxon.replace(' ', '+')
                db_result["url"] = (f"https://verifier.globalnames.org/"
                                    f"?capitalize=on&format=html&names={_gnv_web_name}")
                try:
                    # Použijeme POST s capitalize=true — stejný efekt jako ?capitalize=on v URL
                    _gnv_payload = {
                        "nameStrings": [taxon],
                        "preferredSources": cfg.get("preferred_sources", [1, 11, 172, 169, 167, 3]),
                        "withCapitalization": True,  # ekvivalent capitalize=on
                    }
                    _gnv_session = _get_http_session()
                    _gnv_resp = _gnv_session.post(
                        _gnv_base,
                        json=_gnv_payload,
                        timeout=TO,
                        headers={"Content-Type": "application/json", "Accept": "application/json"}
                    )
                    if _gnv_resp.status_code == 200 and _gnv_resp.content:
                        _gnv_data  = _gnv_resp.json()
                        _gnv_names = _gnv_data if isinstance(_gnv_data, list) else _gnv_data.get("names", [])
                        if _gnv_names:
                            _gnv_item = _gnv_names[0]
                            _gnv_best = _gnv_item.get("bestResult") or {}
                            _gnv_pref = _gnv_item.get("preferredResults") or []

                            _gnv_mt = _gnv_best.get("matchType","") if _gnv_best else ""
                            # Přijmeme Exact a Fuzzy; Partial pouze pro jednoslovný dotaz (rod)
                            _gnv_accept = (_gnv_mt in ("Exact","Fuzzy","FuzzyRelaxed",
                                                        "ExactCanonicalForm","ExactPartialMatch") or
                                           (_gnv_mt in ("Partial","PartialFuzzy","PartialFuzzyRelaxed")
                                            and len(taxon.split()) == 1))

                            if _gnv_best and _gnv_mt and _gnv_mt != "NoMatch" and _gnv_accept:
                                _gnv_name    = (_gnv_best.get("currentName","") or
                                                _gnv_best.get("matchedCanonicalSimple","") or
                                                _gnv_best.get("matchedName",""))
                                _gnv_outlink = _gnv_best.get("outlink","")
                                _gnv_score   = _gnv_best.get("score","") or _gnv_best.get("sortScore","")
                                # Překlad match type do češtiny
                                _GNV_MT_CZ = {
                                    "Exact": "přesná shoda", "Fuzzy": "přibližná shoda",
                                    "FuzzyRelaxed": "přibližná shoda (relaxed)",
                                    "Partial": "částečná shoda", "PartialFuzzy": "částečná přibližná",
                                    "ExactCanonicalForm": "přesná shoda (kanonická forma)",
                                }
                                db_result["found"] = True
                                # URL vždy ve formátu ?capitalize=on&format=html&names=...
                                # _gnv_outlink je odkaz na konkrétní záznam v DB (jiný formát) — nepoužíváme pro db_result["url"]
                                db_result["url"] = (f"https://verifier.globalnames.org/"
                                                    f"?capitalize=on&format=html&names={_gnv_web_name}")
                                db_result["details"] = {
                                    "nalezené jméno":  _gnv_name,
                                    "shoda (match)":   _GNV_MT_CZ.get(_gnv_mt, _gnv_mt),
                                    "zdroj (source)":  _gnv_best.get("dataSourceTitle","")[:50],
                                    "odkaz ve zdroji": _gnv_outlink,
                                    "skóre":           f"{float(_gnv_score):.2f}" if str(_gnv_score).replace('.','').isdigit() else str(_gnv_score),
                                    "kmen (phylum)":   _gnv_best.get("classificationPhylum",""),
                                    "třída (class)":   _gnv_best.get("classificationClass",""),
                                    "řád (order)":     _gnv_best.get("classificationOrder",""),
                                    "čeleď (family)":  _gnv_best.get("classificationFamily",""),
                                    "rod (genus)":     _gnv_best.get("classificationGenus",""),
                                    "databáze":        "100+ (Global Names)",
                                }

                            # Preferované výsledky — ze zdrojů jako PaleoDB, CoL, GBIF, WoRMS
                            if _gnv_pref:
                                db_result["gnv_preferred"] = [
                                    {"name":       (p.get("currentName","") or
                                                    p.get("matchedCanonicalSimple","")),
                                     "source":     p.get("dataSourceTitle",""),
                                     "match_type": _GNV_MT_CZ.get(p.get("matchType",""),
                                                                    p.get("matchType","")),
                                     "outlink":    p.get("outlink",""),
                                     "kingdom":    p.get("classificationKingdom",""),
                                     "family":     p.get("classificationFamily",""),
                                     "score":      str(p.get("score","") or p.get("sortScore",""))}
                                    for p in _gnv_pref
                                    if p.get("matchType","") not in ("NoMatch","")
                                ]
                    else:
                        db_result["error"] = f"HTTP {_gnv_resp.status_code}"
                except Exception as e_gnv:
                    db_result["error"] = str(e_gnv)[:80]
            # ── Garantovaný fallback URL (search link pokud žádný jiný) ──
            if not db_result.get("url") and cfg:
                if db == "GNverifier":
                    # GNverifier: vždy správný formát s ?capitalize=on&format=html&names=
                    _gnv_fb_name = taxon.replace(' ', '+')
                    db_result["url"] = (f"https://verifier.globalnames.org/"
                                        f"?capitalize=on&format=html&names={_gnv_fb_name}")
                else:
                    for key in ("search_page","search_url","browse_url","fallback_url"):
                        v = cfg.get(key)
                        if v:
                            db_result["url"] = v + _enc(taxon_for_db)
                            break
            # ── Ulož db_result do výsledků (VŽDY, bez ohledu na found) ──
            result["results"][db] = db_result

        except requests.exceptions.Timeout:
            result["results"][db] = {"found": False, "error": "Timeout",
                                     "url": (cfg.get("search_page","") or
                                             cfg.get("search_url","")) + _enc(taxon)
                                     if cfg else ""}
        except Exception as e:
            result["results"][db] = {"found": False, "error": str(e)[:80],
                                     "url": (cfg.get("search_page","") or
                                             cfg.get("search_url","")) + _enc(taxon)
                                     if cfg else ""}

    total = result["summary"]["total"]
    found = result["summary"]["found"]
    ratio = found / max(1, total)
    result["summary"]["confidence"] = (
        "high"      if ratio >= 0.45 else
        "medium"    if ratio >= 0.20 else
        "low"       if ratio  > 0   else
        "not_found"
    )
    if cache is not None:
        cache[cache_key] = result
    _disk_cache_set(cache_key, result)
    return result
def extract_taxa_from_text(text: str) -> List[str]:
    candidates = re.findall(r'\b[A-Z][a-z]{2,}\s+[a-z]{3,}\b', text)
    stop = {"This","The","In","New","Based","Table","Figure","From","After","When","Using","With","Type"}
    return sorted(set(c for c in candidates if c.split()[0] not in stop))


def fetch_pbdb_occurrences(taxon: str, timeout: int = 8) -> Optional[Dict]:
    """Načte occurrence data z PaleoDB pro taxon.
    Vrátí dict s: n_occs, strat_min, strat_max, nebo None při chybě.
    """
    try:
        url = (f"https://paleobiodb.org/data1.2/occs/list.json"
               f"?taxon_name={requests.utils.quote(taxon)}&show=loc,strat&limit=500")
        r = _safe_get(url, timeout=timeout)
        if r.status_code != 200:
            return None
        data = r.json()
        records = data.get("records", [])
        if not records:
            return None
        # Stratigrafický rozsah
        ages = []
        for rec in records:
            for field in ("early_age", "late_age", "max_ma", "min_ma"):
                v = rec.get(field)
                if v is not None:
                    try:
                        ages.append(float(v))
                    except (ValueError, TypeError):
                        pass
        result = {"n_occs": len(records)}
        if ages:
            result["strat_max_ma"] = round(max(ages), 1)
            result["strat_min_ma"] = round(min(ages), 1)
        # Interval names
        intervals = set()
        for rec in records:
            for field in ("early_interval", "late_interval"):
                v = rec.get(field, "")
                if v and str(v).strip():
                    intervals.add(str(v).strip())
        if intervals:
            result["intervals"] = ", ".join(sorted(intervals)[:5])
        return result
    except Exception:
        return None


def validation_results_to_df(results: List[Dict]) -> pd.DataFrame:
    """Konvertuje výsledky validace na DataFrame.
    Výsledek se cachuje do session state — při rerunu záložky Validace
    se DataFrame znovu nesestavuje pokud se výsledky nezměnily.
    """
    # Klíč: délka + hash prvního a posledního taxonu (rychlý fingerprint bez plného hash)
    _fprint = f"{len(results)}:{results[0]['taxon'] if results else ''}:{results[-1]['taxon'] if results else ''}"
    _cache_key = "_val_df_cache"
    _cached = st.session_state.get(_cache_key)
    if _cached is not None and _cached[0] == _fprint:
        return _cached[1]

    rows = []
    for r in results:
        row = {"Taxon": r["taxon"],
               "Confidence": r["summary"]["confidence"],
               "Found": r["summary"]["found"],
               "Total": r["summary"]["total"]}
        for db, info in r["results"].items():
            row[db] = ("✅" if info.get("found")
                       else ("⚠️ " + info.get("error","")[:30] if "error" in info else "❌"))
        rows.append(row)
    df = pd.DataFrame(rows)
    st.session_state[_cache_key] = (_fprint, df)
    return df


# ══════════════════════════════════════════════════════
# STRATIGRAFICKÁ NORMALIZACE
# ══════════════════════════════════════════════════════
def normalize_stratigraphy_local(text: str) -> Tuple[str, List[str]]:
    """Normalizuje stratigrafické termíny dle ICS nomenklatury.
    Používá pre-kompilované regexpy _ICS_PATTERNS — žádná kompilace za runtime.
    """
    changes: List[str] = []
    for pattern, ics in _ICS_PATTERNS:
        new_text, n = pattern.subn(ics, text)
        if n > 0:
            changes.append(f"'{pattern.pattern}' → '{ics}' ({n}×)")
            text = new_text
    return text, changes


def _fetch_page_text(url: str, max_chars: int = 3000) -> str:
    """Stáhne stránku a extrahuje čistý text (odstraní HTML tagy)."""
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        r = requests.get(url, headers=headers, timeout=10)
        r.raise_for_status()
        if BS4_AVAILABLE:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(r.text, "html.parser")
            # Odstraň scripty, styly, navigaci
            for tag in soup(["script","style","nav","header","footer","aside"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
        else:
            text = re.sub(r'<[^>]+>', ' ', r.text)
            text = re.sub(r'\s+', ' ', text).strip()
        return text[:max_chars]
    except Exception as e:
        return f"[Nelze načíst stránku: {e}]"


def build_translate_system(src_lang: str, tgt_lang: str,
                            glossary: Dict[str, str],
                            preserve_terms: bool = True,
                            doc_context: Dict = None) -> str:
    """Sestaví systémový prompt pro překlad."""
    pres = (" Zachovej latinská taxonomická jména, geologické jednotky a místní názvy beze změny."
            if preserve_terms else "")
    gloss_note = ""
    if glossary:
        gloss_note = ("\n\nPoužij tento terminologický slovník (originál → překlad):\n"
                      + glossary_to_text(glossary))

    # Kontext dokumentu
    ctx_note = ""
    if doc_context:
        parts = []
        if doc_context.get("author"): parts.append(f"autor: {doc_context['author']}")
        if doc_context.get("year"):   parts.append(f"rok: {doc_context['year']}")
        if doc_context.get("type") and doc_context["type"] != "—":
            parts.append(f"typ: {doc_context['type']}")
        if doc_context.get("topic"):  parts.append(f"téma: {doc_context['topic']}")
        if parts:
            ctx_note = (f"\n\nKontext dokumentu: {', '.join(parts)}. "
                        "Přizpůsob terminologii a styl tomuto oboru.")

    latin_note = ""
    if tgt_lang == "latina":
        latin_note = (
            " Překládáš do klasické vědecké latiny (latina scientifica). "
            "Používej ustálenou latinskou terminologii pro vědecké a akademické texty. "
            "Taxonomická jména ponech beze změny (jsou již latinsky). "
            "Pokud překlad do latiny není jednoznačný, upřednostni srozumitelnost.")
    elif src_lang == "latina":
        latin_note = (
            " Zdrojový text je v latině (může jít o klasickou nebo vědeckou latinu). "
            "Překládej přesně; latinské taxonomické termíny přepiš beze změny do překladu.")

    return (f"Jsi STROJOVÝ PŘEKLADAČ. Tvým jediným úkolem je přeložit zadaný text z jazyka {src_lang} do jazyka {tgt_lang}. "
            f"Nesmíš text nijak komentovat, nesmíš s uživatelem konverzovat, nesmíš odpovídat na otázky v textu. "
            f"Vrať POUZE a JENOM přeložený text bez jakýchkoliv úvodních či závěrečných frází (např. bez 'Zde je překlad:')."
            f"{pres}{latin_note}{ctx_note}{gloss_note}")


def do_translate(base_url: str, model: str, text: str, sys_msg: str,
                 temp: float = 0.1, chunk_size: int = 8000,
                 progress_cb=None, parallel: bool = False,
                 max_workers: int = 4,
                 src_lang: str = "") -> str:
    """Přeloží text s auto-chunkingem.

    parallel=True  → chunky se posílají souběžně (LM Studio continuous batching).
                     Výrazně rychlejší pro texty s mnoha bloky.
    parallel=False → sekvenční překlad s překryvem 200 znaků (lepší kontext).

    Překryv je aktivní jen v sekvenčním režimu — v paralelním není pořadí zaručeno.
    progress_cb(done, total, snippet) — volá se po každém přeloženém bloku.

    Optimalizace v23.9:
    • Adaptive max_tokens s jazykovým multiplikátorem (CJK ×1.5, slovanské ×1.2)
    • Stop sekvence: překlad nepoužívá stop (volný text)
    • Backoff s jitterem (zabrání thundering herd při souběžných retry)
    • DLQ: chybné chunky se sbírají, batch pokračuje (výsledek označen [!])
    • Deduplikace identických chunků (záhlaví, opakující se bloky)
    v23:
    • Adaptive max_tokens podle délky chunku
    • Per-chunk retry 3× s exponential backoff
    """
    text = fix_spaced_taxon_names(text)

    # Propaguj session state hodnoty do modul-level proměnných před spuštěním threadů
    # Worker thready nemají ScriptRunContext — čtou z globálních fallback proměnných
    global _THREAD_OFFLINE, _THREAD_GSP
    try:
        _THREAD_OFFLINE = st.session_state.get("offline_mode", False)
        _THREAD_GSP     = st.session_state.get("global_system_prompt", "").strip()
    except Exception:
        pass

    # Jazykový multiplikátor pro max_tokens (CJK potřebuje více výstupních tokenů)
    _cjk_langs = {"zh", "ja", "ko", "chinese", "japanese", "korean"}
    _sla_langs = {"ru", "cs", "pl", "sk", "uk", "russian", "czech", "polish"}
    _src = src_lang.lower()
    if any(l in _src for l in _cjk_langs):
        _lang_mult = 1.5
    elif any(l in _src for l in _sla_langs):
        _lang_mult = 1.2
    else:
        _lang_mult = 1.0

    def _adaptive_max_tokens(block: str) -> int:
        """Odhadni max_tokens: ~2.2 tokenu / 3 znaky × jazykový multiplikátor.
        Minimum 512, maximum _MAX_TOKENS_TRANSLATE."""
        raw = int(len(block) / 3 * 2.2 * _lang_mult)
        return max(512, min(_MAX_TOKENS_TRANSLATE, raw))

    def make_prompt(block: str, context_tail: str = "") -> str:
        ctx = f"\n\n[Kontext z předchozího bloku:…{context_tail}]\n\n" if context_tail else ""
        return f"Přelož následující text. Nevytvářej žádnou konverzaci, pouze přelož:{ctx}\n\n{block}"

    # Dead Letter Queue — chunky které selhaly po všech pokusech
    _dlq: List[Dict] = []

    # Jeden blok — přeloží přímo (bez chunking overhead)
    if len(text) <= chunk_size:
        result = chat_completion(base_url, model,
                                 [{"role":"system","content":sys_msg},
                                  {"role":"user","content":make_prompt(text)}],
                                 temp=temp, max_tokens=_adaptive_max_tokens(text),
                                 stop=_STOP_TRANSL or None)
        if progress_cb:
            progress_cb(1, 1, text[:60])
        return result

    blocks = chunk_text_smart(text, chunk_size)
    total  = len(blocks)

    if parallel and total > 1:
        # ── PARALELNÍ PŘEKLAD (continuous batching v LM Studio) ────
        # Deduplikace: stejné chunky překládáme jen jednou
        translation_cache: Dict[str, str] = {}
        cache_lock  = _threading.Lock()
        count_lock  = _threading.Lock()
        done_count  = [0]

        def _safe_progress(done, tot, snippet):
            """Volá progress_cb bezpečně — ignoruje Streamlit ScriptRunContext chyby."""
            if progress_cb:
                try:
                    progress_cb(done, tot, snippet)
                except Exception:
                    pass

        def translate_chunk_cached(args):
            idx, block = args
            cache_key = block.strip()
            with cache_lock:
                if cache_key in translation_cache:
                    with count_lock:
                        done_count[0] += 1
                        _n = done_count[0]
                    _safe_progress(_n, total, block[:60])
                    return idx, translation_cache[cache_key]

            # Per-chunk retry s exponential backoff + jitter
            # Při "Context size exceeded" automaticky zmenšíme chunk na polovinu (max 2×)
            last_exc = None
            current_block = block
            for attempt in range(3):
                try:
                    resp = chat_completion(base_url, model,
                                           [{"role":"system","content":sys_msg},
                                            {"role":"user","content":make_prompt(current_block)}],
                                           temp=temp,
                                           max_tokens=_adaptive_max_tokens(current_block),
                                           stop=_STOP_TRANSL or None)
                    with cache_lock:
                        translation_cache[cache_key] = resp
                    with count_lock:
                        done_count[0] += 1
                        _n = done_count[0]
                    _safe_progress(_n, total, current_block[:60])
                    return idx, resp
                except Exception as e:
                    last_exc = e
                    err_str = str(e).lower()
                    if "context size" in err_str or "context_length" in err_str or "context window" in err_str:
                        # Context overflow — zmenši chunk na polovinu a zkus znovu
                        if len(current_block) > 1000:
                            half = len(current_block) // 2
                            current_block = current_block[:half]
                        else:
                            break   # příliš malý chunk, vzdej se
                    elif attempt < 2:
                        _time.sleep(_backoff_jitter(attempt))
            # DLQ — chunk se nepodařilo přeložit
            _dlq.append({"idx": idx, "error": str(last_exc), "preview": block[:80]})
            with count_lock:
                done_count[0] += 1
            return idx, f"[!PŘEKLAD SELHAL blok {idx}: {last_exc}]"

        results_map: Dict[int, str] = {}
        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {pool.submit(translate_chunk_cached, (i, b)): i
                       for i, b in enumerate(blocks)}
            for fut in as_completed(futures):
                try:
                    idx, resp = fut.result()
                    results_map[idx] = resp
                except Exception as e:
                    idx = futures[fut]
                    results_map[idx] = f"[!PŘEKLAD SELHAL blok {idx}: {e}]"
                    _dlq.append({"idx": idx, "error": str(e)})

        # DLQ report do session state pro zobrazení v UI
        if _dlq:
            st.session_state["translate_dlq"] = _dlq

        return "\n\n".join(results_map[i] for i in range(total))

    else:
        # ── SEKVENČNÍ PŘEKLAD s překryvem ─────────────────────────
        # Deduplikační cache pro opakující se bloky
        _seq_cache: Dict[str, str] = {}
        translated_parts = []
        prev_tail = ""
        for i, block in enumerate(blocks):
            # Zastavení překladu – uložit částečný výsledek a ukončit
            if not st.session_state.get("translate_running", True):
                partial = "\n\n".join(translated_parts)
                st.session_state["translation_result"] = partial
                st.session_state["translation_ready"] = True
                return partial

            cache_key = block.strip()
            if cache_key in _seq_cache:
                translated_parts.append(_seq_cache[cache_key])
                if progress_cb:
                    progress_cb(i + 1, total, block[:60])
                continue
            # ... retry a překlad ...

            # Retry s jitterem + halving při context overflow
            last_exc = None
            resp = None
            current_block = block
            for attempt in range(3):
                try:
                    resp = chat_completion(base_url, model,
                                           [{"role":"system","content":sys_msg},
                                            {"role":"user","content":make_prompt(current_block, prev_tail)}],
                                           temp=temp,
                                           max_tokens=_adaptive_max_tokens(current_block),
                                           stop=_STOP_TRANSL or None)
                    break
                except Exception as e:
                    last_exc = e
                    err_str = str(e).lower()
                    if "context size" in err_str or "context_length" in err_str or "context window" in err_str:
                        if len(current_block) > 1000:
                            half = len(current_block) // 2
                            current_block = current_block[:half]
                            prev_tail = ""   # reset overlap — kratší blok nepotřebuje kontext
                        else:
                            break
                    elif attempt < 2:
                        _time.sleep(_backoff_jitter(attempt))

            if resp is None:
                # DLQ
                _dlq.append({"idx": i, "error": str(last_exc), "preview": block[:80]})
                resp = f"[!PŘEKLAD SELHAL blok {i}: {last_exc}]"

            _seq_cache[cache_key] = resp
            translated_parts.append(resp)
            prev_tail = block[-200:] if len(block) > 200 else block
            if progress_cb:
                progress_cb(i + 1, total, block[:60])

        if _dlq:
            st.session_state["translate_dlq"] = _dlq

        return "\n\n".join(translated_parts)


def do_translate_batched(base_url: str, model: str, text: str, sys_msg: str,
                          pages_per_batch: int = 5, temp: float = 0.1,
                          progress_container=None,
                          parallel: bool = False, max_workers: int = 4,
                          src_lang: str = "") -> Tuple[str, List[bytes]]:
    """Přeloží velký PDF text po dávkách N stránek.

    Vrátí (celý_přeložený_text, [docx_bytes_per_batch]).
    Každá dávka se přeloží zvlášť a zároveň je dostupná ke stažení.
    src_lang: zdrojový jazyk — předává se do do_translate pro správný adaptive max_tokens.
    """
    # Rozdělíme text po stránkách (marker "--- Strana N ---")
    page_re   = _PAGE_BATCH_RE  # pre-compiled modul-level konstanta
    parts     = page_re.split(text)
    # Sesbírej stránky jako list
    pages: List[str] = []
    current = ""
    for part in parts:
        if page_re.match(part):
            if current.strip():
                pages.append(current.strip())
            current = part
        else:
            current += part
    if current.strip():
        pages.append(current.strip())

    if not pages:
        pages = [text]

    total_pages   = len(pages)
    total_batches = max(1, (total_pages + pages_per_batch - 1) // pages_per_batch)

    all_translated: List[str] = []
    batch_docx:     List[bytes] = []

    if progress_container:
        prog = progress_container.progress(0, text=f"Dávka 1/{total_batches}…")
        log  = progress_container.empty()

    # Propaguj session_state před spuštěním paralelních překladových threadů
    global _THREAD_OFFLINE, _THREAD_GSP
    try:
        _THREAD_OFFLINE = st.session_state.get("offline_mode", False)
        _THREAD_GSP     = st.session_state.get("global_system_prompt", "").strip()
    except Exception:
        pass
    # Udržuj _temp/ čistý — zachovej jen 3 nejnovější sady pro "preklad"
    _temp_prune("preklad", keep_last=3)

    for batch_idx in range(total_batches):
        batch_pages = pages[batch_idx * pages_per_batch:(batch_idx + 1) * pages_per_batch]
        batch_text  = "\n\n".join(batch_pages)
        page_start  = batch_idx * pages_per_batch + 1
        page_end    = min(page_start + pages_per_batch - 1, total_pages)

        if progress_container:
            prog.progress(batch_idx / total_batches,
                          text=f"Dávka {batch_idx+1}/{total_batches} "
                               f"(stránky {page_start}–{page_end})…")
            log.caption(f"⏳ Překládám stránky {page_start}–{page_end}…")

        # Přeloží dávku s progress per-chunk
        chunk_size = suggest_chunk_size(len(batch_text), model)
        translated_batch = do_translate(base_url, model, batch_text, sys_msg,
                                        temp=temp, chunk_size=chunk_size,
                                        parallel=parallel, max_workers=max_workers,
                                        src_lang=src_lang)
        all_translated.append(
            f"=== Stránky {page_start}–{page_end} ===\n{translated_batch}")

        # DOCX pro tuto dávku
        _docx_bytes = to_docx_bytes(translated_batch)
        batch_docx.append(_docx_bytes)

        # ── Průběžné uložení dávky do _temp/ ─────────────────────────
        _temp_note = ""
        try:
            _batch_label = f"=== Stránky {page_start}–{page_end} ===\n{translated_batch}"
            _tp = _temp_save_txt("preklad", batch_idx + 1, _batch_label, total_batches)
            _temp_note = f" · 💾 _temp/{_tp.name}"
        except Exception as _te:
            _temp_note = f" · ⚠️ temp save chyba: {_te}"

        if progress_container:
            log.caption(f"✅ Dávka {batch_idx+1}/{total_batches} hotova "
                        f"(stránky {page_start}–{page_end}){_temp_note}")

    if progress_container:
        prog.progress(1.0, text="✅ Překlad dokončen")
        log.empty()

    full_text = "\n\n".join(all_translated)
    # Úklid _temp po úspěšném dokončení celého překladu
    try:
        _temp_cleanup("preklad")
    except Exception:
        pass
    return full_text, batch_docx


# ── Prompt profily (perzistentní) ──────────────────────────────
def load_prompt_profiles() -> Dict[str, Dict]:
    if os.path.exists(PROMPT_PROFILE_FILE):
        try:
            with open(PROMPT_PROFILE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def save_prompt_profiles(data: Dict[str, Dict]):
    with open(PROMPT_PROFILE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# ── Prompty pro čištění dat (perzistentní) ─────────────────────
def load_clean_prompts() -> Dict[str, str]:
    if os.path.exists(CLEAN_PROMPT_FILE):
        try:
            with open(CLEAN_PROMPT_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {
        "Normalizace taxonomie":
            "Normalizuj taxonomické názvy: oprav překlepy, sjednoť formát 'Rod druh Autor rok'.",
        "Čištění lokalit":
            "Sjednoť názvy lokalit: normalizuj diakritiku, oprav zkratky, sjednoť formát.",
        "Stratigrafická normalizace":
            "Normalizuj stratigrafické jednotky dle ICS nomenklatury (napr. 'spodní kambrium' → 'Terreneuvian–Series 2').",
    }

def save_clean_prompts(data: Dict[str, str]):
    with open(CLEAN_PROMPT_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# Inicializace promptů čištění do session state (jednou za session)
if not st.session_state["clean_prompts"]:
    st.session_state["clean_prompts"] = load_clean_prompts()

# ── Hyolitha DB export ─────────────────────────────────────────
# Mapování LLM klíčů (lowercase) → Access DB sloupce
HYOLITHA_FIELD_MAP: Dict[str, str] = {
    # Taxonomické základy
    "druh":          "Druh",
    "species":       "Druh",
    "rod":           "Rod",
    "genus":         "Rod",
    "autor":         "Autor",
    "author":        "Autor",
    "rok":           "Rok",
    "year":          "Rok",
    # Výskyt / lokalizace
    "lokalita":      "Lokalita",
    "locality":      "Lokalita",
    "lokalit":       "Lokalita",
    "stratigrafie":  "Stratigrafie",
    "stratigraphy":  "Stratigrafie",
    "stáří":         "Stáří",
    "age":           "Stáří",
    "formace":       "Formace",
    "formation":     "Formace",
    "stát":          "Stát",
    "country":       "Stát",
    "kontinent":     "Kontinent",
    "continent":     "Kontinent",
    "souřadnice":    "Souřadnice",
    "coordinates":   "Souřadnice",
    # Popis / morfologie
    "popis":         "Popis",
    "description":   "Popis",
    "rozměry":       "Rozměry",
    "dimensions":    "Rozměry",
    "typ":           "Typ",
    "type":          "Typ",
    # Materiál / depozitář
    "material":      "Material",
    "depozitář":     "Depozitář",
    "repository":    "Depozitář",
    "holotyp":       "Holotyp",
    "holotype":      "Holotyp",
    "katalogové_č":  "KatalČíslo",
    "catalog_no":    "KatalČíslo",
    # Meta
    "zdroj":         "Zdroj",
    "source":        "Zdroj",
    "doi":           "DOI",
    "synonymum":     "Synonymum",
    "synonym":       "Synonymum",
    "poznámka":      "Poznámka",
    "notes":         "Poznámka",
    "note":          "Poznámka",
}

# Schémata 4 tabulek dle Access DB (Hyolitha Taxonomic Database)
HYOLITHA_TAXA_COLS: list = [
    "Rod", "Druh", "Autor", "Rok",
    "Synonymum", "Typ", "Zdroj", "DOI", "Poznámka",
]
HYOLITHA_VYSKYT_COLS: list = [
    "Rod", "Druh", "Autor", "Rok",
    "Lokalita", "Stát", "Kontinent", "Souřadnice",
    "Stratigrafie", "Stáří", "Formace", "Zdroj",
]
HYOLITHA_POPIS_COLS: list = [
    "Rod", "Druh", "Autor", "Rok",
    "Popis", "Rozměry", "Typ", "Zdroj",
]
HYOLITHA_MATERIAL_COLS: list = [
    "Rod", "Druh", "Autor", "Rok",
    "Material", "Holotyp", "KatalČíslo", "Depozitář", "Zdroj",
]

def normalize_record_keys(rec: dict) -> dict:
    """Normalizuje klíče záznamu dle mapovací tabulky polí.
    _NORM_KEY_CACHE cachuje výsledky lower().strip() per klíč —
    pro batch 1500 záznamů × 20 polí se každý unikátní klíč normalizuje jen jednou.
    """
    out = {}
    for k, v in rec.items():
        nk = _NORM_KEY_CACHE.get(k)
        if nk is None:
            nk = k.lower().strip()
            _NORM_KEY_CACHE[k] = nk
        mapped = HYOLITHA_FIELD_MAP.get(nk, k)
        out[mapped] = v
    return out

# Per-klíč cache pro lower().strip() normalizaci — sdílená přes všechna volání
_NORM_KEY_CACHE: Dict[str, str] = {}

def records_to_hyolitha_dfs(records: List[dict]) -> Dict[str, pd.DataFrame]:
    """Rozdělí záznamy do 4 tabulek dle schématu Hyolitha DB.
    Používá vektorizovanou DataFrame konstrukci — pd.DataFrame(rows) je rychlejší
    než iterativní append pro velké datasety.
    """
    normalized = [normalize_record_keys(r) for r in records]

    def build_df(cols: List[str]) -> pd.DataFrame:
        # Vektorizovaná cesta: předpočítáme rows jako list of dicts jedním průchodem
        rows = [{c: rec.get(c, "") for c in cols} for rec in normalized]
        df = pd.DataFrame(rows, columns=cols)
        if "Rod" in df.columns and "Druh" in df.columns:
            mask = (df["Rod"] != "") | (df["Druh"] != "")
            df = df[mask]
        return df.reset_index(drop=True)

    return {
        "Taxa":           build_df(HYOLITHA_TAXA_COLS),
        "Taxa_Vyskyt":    build_df(HYOLITHA_VYSKYT_COLS),
        "Taxa_Popis":     build_df(HYOLITHA_POPIS_COLS),
        "Taxa_Material":  build_df(HYOLITHA_MATERIAL_COLS),
    }

def hyolitha_export_xlsx(records: List[dict]) -> bytes:
    """Exportuje záznamy do XLSX se 4 listy (jedna tabulka = jeden list)."""
    dfs = records_to_hyolitha_dfs(records)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        for sheet_name, df in dfs.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            # Nastav šířku sloupců
            ws = writer.sheets[sheet_name]
            for col_cells in ws.columns:
                max_len = max((len(str(c.value or "")) for c in col_cells), default=10)
                ws.column_dimensions[col_cells[0].column_letter].width = min(max_len + 3, 40)
    return buf.getvalue()

def hyolitha_export_csv_zip(records: List[dict]) -> bytes:
    """Exportuje záznamy do ZIP se 4 CSV soubory."""
    dfs = records_to_hyolitha_dfs(records)
    files = {f"{name}.csv": df.to_csv(index=False).encode("utf-8")
             for name, df in dfs.items()}
    return make_zip(files)


# ══════════════════════════════════════════════════════
# OFFLINE FALLBACK – lokální SQLite s PaleoDB taxony
# ══════════════════════════════════════════════════════
OFFLINE_DB_PATH = Path("offline_taxa.db")

_OFFLINE_DB_CONN: Optional[sqlite3.Connection] = None

def _get_offline_conn() -> sqlite3.Connection:
    """Sdílené SQLite spojení pro offline DB — přežije reruns, WAL mode."""
    global _OFFLINE_DB_CONN
    if _OFFLINE_DB_CONN is None:
        _OFFLINE_DB_CONN = sqlite3.connect(
            str(OFFLINE_DB_PATH), check_same_thread=False, timeout=10)
        _OFFLINE_DB_CONN.execute("PRAGMA journal_mode=WAL")
        _OFFLINE_DB_CONN.execute("PRAGMA synchronous=NORMAL")
    return _OFFLINE_DB_CONN

def offline_db_init():
    """Inicializuje prázdnou offline DB strukturu."""
    conn = _get_offline_conn()
    conn.execute("""CREATE TABLE IF NOT EXISTS taxa (
        name TEXT PRIMARY KEY,
        rank TEXT, status TEXT, kingdom TEXT,
        phylum TEXT, class_ TEXT, family TEXT,
        source TEXT, added_at TEXT
    )""")
    conn.commit()

def offline_db_lookup(taxon: str) -> Optional[Dict]:
    """Vyhledá taxon v offline DB. Vrátí dict nebo None."""
    if not OFFLINE_DB_PATH.exists():
        return None
    try:
        row = _get_offline_conn().execute(
            "SELECT name,rank,status,kingdom,phylum,class_,family,source "
            "FROM taxa WHERE name=? COLLATE NOCASE", (taxon,)).fetchone()
        if row:
            return dict(zip(["name","rank","status","kingdom","phylum","class","family","source"], row))
    except Exception:
        pass
    return None

def offline_db_insert(records: List[Dict]):
    """Vloží záznamy do offline DB (bulk insert)."""
    offline_db_init()
    conn = _get_offline_conn()
    now = datetime.now().isoformat()
    conn.executemany(
        "INSERT OR REPLACE INTO taxa VALUES (?,?,?,?,?,?,?,?,?)",
        [(r.get("name",""), r.get("rank",""), r.get("status",""),
          r.get("kingdom",""), r.get("phylum",""), r.get("class",""),
          r.get("family",""), r.get("source","manual"), now)
         for r in records if r.get("name")]
    )
    conn.commit()

def offline_db_count() -> int:
    if not OFFLINE_DB_PATH.exists(): return 0
    try:
        return _get_offline_conn().execute(
            "SELECT COUNT(*) FROM taxa").fetchone()[0]
    except Exception:
        return 0

try:
    offline_db_init()
except Exception:
    pass


# ══════════════════════════════════════════════════════
# SYSTEM PROMPT PER MODEL
# ══════════════════════════════════════════════════════
PER_MODEL_PROMPTS_FILE = "per_model_prompts.json"

def load_per_model_prompts() -> Dict[str, str]:
    if os.path.exists(PER_MODEL_PROMPTS_FILE):
        try:
            with open(PER_MODEL_PROMPTS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception: pass
    return {
        "qwen2.5":   "Odpovídej vždy v JSON pokud je to požadováno. Nepiš žádné vysvětlení mimo JSON.",
        "qwen3":     "Think step by step before answering. Output only the requested format.",
        "mistral":   "Be concise. Follow instructions exactly. No preamble.",
        "deepseek":  "Reason carefully. Output structured data without commentary.",
    }

def save_per_model_prompts(d: Dict[str, str]):
    with open(PER_MODEL_PROMPTS_FILE, "w", encoding="utf-8") as f:
        json.dump(d, f, ensure_ascii=False, indent=2)

def get_model_system_prompt(model_name: str) -> str:
    """Vrátí per-model systémový prompt podle názvu modelu."""
    prompts = load_per_model_prompts()
    mn = model_name.lower()
    for key, prompt in prompts.items():
        if key.lower() in mn:
            return prompt
    return ""


# ══════════════════════════════════════════════════════
# WEBHOOK NOTIFIKACE
# ══════════════════════════════════════════════════════
def send_webhook_notification(url: str, message: str, title: str = "SciNexus"):
    """Odešle webhook notifikaci (Slack, Teams, nebo generic JSON)."""
    if not url or not url.startswith("http"):
        return False
    try:
        payload = {"text": f"*{title}*\n{message}",
                   "content": message, "title": title,
                   "@type": "MessageCard", "summary": message}
        r = requests.post(url, json=payload, timeout=5)
        return r.status_code < 300
    except Exception:
        return False


# ══════════════════════════════════════════════════════
# ŠABLONY WORKFLOW
# ══════════════════════════════════════════════════════
WF_TEMPLATES_FILE = "workflow_templates.json"
WF_BUILTIN_TEMPLATES = {
    "Rychlá extrakce CZ→EN": {
        "steps": ["🔍 Extrakce", "🌐 Překlad"],
        "params": {"wf_lang": "angličtina", "wf_ext_p": "Extrahuj taxonomické záznamy jako JSON."}
    },
    "Plný pipeline s validací": {
        "steps": ["🔍 Extrakce", "🌐 Překlad", "🧬 Validace taxonů"],
        "params": {"wf_lang": "angličtina"}
    },
    "Pouze stratigrafická analýza": {
        "steps": ["🔍 Extrakce", "🗺️ Stratigrafická normalizace"],
        "params": {"wf_ext_p": "Extrahuj stratigrafické jednotky a věky. Výstup JSON."}
    },
    "Hyolitha kompletní": {
        "steps": ["🔍 Extrakce", "🌐 Překlad", "🧬 Validace taxonů", "🦪 Export Hyolitha DB"],
        "params": {"wf_lang": "angličtina"}
    },
}

def load_wf_templates() -> Dict:
    custom = {}
    if os.path.exists(WF_TEMPLATES_FILE):
        try:
            with open(WF_TEMPLATES_FILE, "r", encoding="utf-8") as f:
                custom = json.load(f)
        except Exception: pass
    return {**WF_BUILTIN_TEMPLATES, **custom}

def save_wf_template(name: str, steps: List[str], params: Dict):
    templates = load_wf_templates()
    templates[name] = {"steps": steps, "params": params}
    # Ulož jen custom (ne builtin)
    custom = {k: v for k, v in templates.items() if k not in WF_BUILTIN_TEMPLATES}
    with open(WF_TEMPLATES_FILE, "w", encoding="utf-8") as f:
        json.dump(custom, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════
# AUTOSAVE + RECOVERY BANNER
# ══════════════════════════════════════════════════════
# _L potřebné před main_tabs – načteme předběžně
_L = st.session_state.get("lang", "cz")

# Autosave každých 5 rerunů
_autosave_if_due()

# ── v24.29: Jednorázové úvodní akce — pouze jednou za session ────────────────
if "initial_setup_done" not in st.session_state:
    smart_auto_cleanup()
    do_nightly_backup()
    save_session_version()
    _log_startup()
    st.session_state["initial_setup_done"] = True
else:
    # Při každém rerunu: jen nightly backup (hlídá vlastní session flag)
    do_nightly_backup()

# ── Saved usernames: rychlé přihlášení tlačítky ──────────────────────────────
# Zobrazíme tlačítka PŘED presence.login_dialog() — klik = okamžité přihlášení.
if not st.session_state.get("presence_name"):
    _known = load_known_users()
    if _known:
        with st.sidebar:
            _L_login = st.session_state.get("lang", "cz")
            st.markdown("👤 **Přihlásit se jako…**" if _L_login == "cz"
                        else "👤 **Log in as…**")
            # Max 8 tlačítek — klik rovnou přihlásí
            for _ku_name in _known[:8]:
                if st.button(_ku_name, key=f"_ku_btn_{_ku_name}",
                             use_container_width=True):
                    st.session_state["presence_name"] = _ku_name
                    st.session_state["_preseed_username"] = _ku_name
                    # Přidej na začátek known_users (refresh pořadí)
                    _ku_list = load_known_users()
                    if _ku_name in _ku_list:
                        _ku_list.remove(_ku_name)
                    _ku_list.insert(0, _ku_name)
                    save_known_users(_ku_list)
                    st.rerun()
            st.markdown("---")
            if st.button("➕ " + ("Zadat nové jméno" if _L_login == "cz"
                                  else "Enter new name"),
                         key="_ku_new_name_btn",
                         use_container_width=True):
                st.session_state["_show_login_dialog"] = True

# ── Presence: přihlášení + heartbeat ─────────────────────────────────────────
# DŮLEŽITÉ: login_dialog() musí proběhnout PŘED zobrazením bannerů/session info.
# Pokud presence modul zavolá st.stop() pro nepřihlášeného uživatele,
# žádné bannery ani info o záloze se na přihlašovací stránce neobjeví.
# ── Registrace callback pro uložení jména z presence dialogu ─────────────────
# Musí být PŘED login_dialog(), aby callback byl dostupný při přihlášení.
def _save_known_user_cb(username: str):
    """Callback volaný z presence.login_dialog() při přihlášení nového uživatele."""
    _ku = load_known_users()
    if username and username not in _ku:
        _ku.insert(0, username)
        save_known_users(_ku)
st.session_state["_presence_save_user_cb"] = _save_known_user_cb

presence.login_dialog()
presence.heartbeat(st.session_state.get("presence_tab_idx", 0))

# ── Aplikace odložených nastavení (z presetů/scénářů) ─────────────────────────
# MUSÍ proběhnout PŘED inicializací widgetů, jinak Streamlit hlásí:
#   "st.session_state.<key> cannot be modified after the widget … is instantiated"
if "_pending_settings" in st.session_state:
    _pending = st.session_state.pop("_pending_settings")
    _WIDGET_ALIASES = {
        "wf_large_mode":      ["sb_large_mode_toggle"],
        "lms_max_concurrent": ["sb_lms_parallel"],
    }
    for _pk, _pv in _pending.items():
        st.session_state[_pk] = _pv
        for _wk in _WIDGET_ALIASES.get(_pk, []):
            st.session_state[_wk] = _pv

# ── Per-user profil: aktivace po přihlášení ───────────────────────────────────
# Zjistíme přihlášené jméno z presence modulu (zkusíme různé klíče)
_logged_user = (st.session_state.get("presence_name") or
                st.session_state.get("presence_user") or
                st.session_state.get("_preseed_username") or
                st.session_state.get("_current_user") or "")
if _logged_user and _logged_user != st.session_state.get("_current_user"):
    st.session_state["_current_user"] = _logged_user
    # Aktivace per-user souborových cest
    activate_user_profile(_logged_user)
    # Uložení jména do known_users.json
    _ku = load_known_users()
    if _logged_user not in _ku:
        _ku.insert(0, _logged_user)
        save_known_users(_ku)
    # Reload per-user dat (šablony, slovníky, presets)
    st.session_state["prompt_templates"] = load_templates()
    st.session_state["prompt_templates"] = st.session_state["prompt_templates"] or {}
    st.session_state["quick_presets"]    = load_quick_presets()
elif _logged_user and st.session_state.get("_current_user") == _logged_user:
    # Již přihlášen — jen zajistíme per-user cesty (po rerunu)
    activate_user_profile(_logged_user)

# ── Bannery a session info — zobrazují se AŽ PO přihlášení ───────────────────

# Recovery banner — zobrazí se jen pokud existuje autosave a session je prázdná
# Příznak "dismissed" zabraňuje opětovnému zobrazení po rerun bez dat
_has_active_work = (
    bool(st.session_state.get("last_extracted_text")) or
    bool(st.session_state.get("translation_result")) or
    bool(st.session_state.get("last_validation_results"))
)
_banner_dismissed = st.session_state.get("_autosave_banner_dismissed", False)
_gsb_running      = st.session_state.get("gsb_active", False)
# Zobraz banner jen pokud neprobíhá žádná operace (gsb_active=False)
# Dismiss/Obnovit při probíhající operaci by způsobil rerun a přerušení
_TEMP_DIR        = Path("_temp")   # průběžné ukládání překladu / extrakce / validace
# ── Banner: obnovení z přerušené operace (_temp soubory) ──────────────
_temp_orphans: dict = {}
if _TEMP_DIR.exists() and not _gsb_running:
    for _pfx in ("preklad", "preklad_final", "extrakce", "validace"):
        _files = sorted(_TEMP_DIR.glob(f"{_pfx}_b*"))
        if _files:
            _temp_orphans[_pfx] = _files
if _temp_orphans and not st.session_state.get("_temp_banner_dismissed"):
    with st.container():
        _tp_lines = []
        for _pfx, _flist in _temp_orphans.items():
            _tp_lines.append(f"**{_pfx}**: {len(_flist)} souborů "
                             f"(poslední: `{_flist[-1].name}`)")
        st.warning(
            tt("⚠️ **Nalezeny soubory z přerušené operace** v `_temp/`:\n" +
               "\n".join(f"  • {l}" for l in _tp_lines) +
               "\n\nSoubory jsou k dispozici pro ruční obnovu.",
               "⚠️ **Found files from interrupted operation** in `_temp/`:\n" +
               "\n".join(f"  • {l}" for l in _tp_lines) +
               "\n\nFiles are available for manual recovery.", _L),
            icon="⚠️"
        )
        _tc1, _tc2 = st.columns([1, 4])
        with _tc1:
            if st.button(tt("🗑️ Smazat _temp", "🗑️ Clear _temp", _L),
                         key="btn_clear_temp_orphans"):
                for _flist in _temp_orphans.values():
                    for _tf in _flist:
                        try: _tf.unlink()
                        except Exception: pass
                st.session_state["_temp_banner_dismissed"] = True
                st.rerun()
        with _tc2:
            if st.button(tt("✕ Zavřít", "✕ Dismiss", _L),
                         key="btn_dismiss_temp_banner"):
                st.session_state["_temp_banner_dismissed"] = True
                st.rerun()

# ── v24.20: Auto-Resume banner (checkpoint-based přerušené operace) ───────────
render_auto_resume_banner()

# ── v24.24: Partial resume sekce ─────────────────────────────────────────────
render_partial_resume_section()

# ── v24.37: CSS — banner nepřekrývá záložky (opraveno pro sloupce) ─────────────
st.html("""<style>
/* Bannery i jejich sloupce (Obnovit/Zavřít) zaberou plnou šířku */
[data-testid="stVerticalBlock"] > [data-testid="element-container"] > [data-testid="stAlert"],
[data-testid="stVerticalBlock"] > [data-testid="stNotification"],
[data-testid="stVerticalBlock"] > [data-testid="element-container"] > [data-testid="stHorizontalBlock"] {
    width: 100% !important;
    box-sizing: border-box !important;
    margin-bottom: 6px !important;
}
/* Záložky se vykreslí vždy POD bannery */
[data-testid="stTabs"] {
    margin-top: 4px !important;
    clear: both !important;
}
/* Svislé zarovnání tlačítek v banneru */
[data-testid="stHorizontalBlock"] > [data-testid="column"] {
    align-self: center !important;
}
</style>""")

if not _has_active_work and not _banner_dismissed and not _gsb_running and os.path.exists(SESSION_AUTOSAVE_FILE):
    _autosave_info = session_persist_info(SESSION_AUTOSAVE_FILE)
    if _autosave_info:
        try:
            from datetime import datetime as _dt
            _saved_dt = _dt.fromisoformat(_autosave_info)
            _age_min  = int((datetime.now() - _saved_dt).total_seconds() / 60)
            _age_str  = (f"před {_age_min} min" if _age_min < 60
                         else f"před {_age_min // 60} hod {_age_min % 60} min")
        except Exception:
            _age_str = _autosave_info[:16]

        # Zjisti co záloha obsahuje (před sloupci)
        _bak_has = []
        try:
            import json as _json_bk
            with open(SESSION_AUTOSAVE_FILE, "r", encoding="utf-8") as _fbk:
                _bkd = _json_bk.load(_fbk).get("data", {})
            if _bkd.get("translation_result"): _bak_has.append("překlad" if _L=="cz" else "translation")
            if _bkd.get("last_extracted_text") or _bkd.get("last_extracted_df_json"): _bak_has.append("extrakci" if _L=="cz" else "extraction")
            if _bkd.get("last_validation_results"): _bak_has.append("validaci" if _L=="cz" else "validation")
            if _bkd.get("tr_src_text"): _bak_has.append("text k překladu" if _L=="cz" else "input text")
        except Exception:
            pass
        _bak_desc = (", ".join(_bak_has) if _bak_has else
                     ("extrakci, překlad nebo validaci" if _L=="cz" else "extraction, translation or validation"))
        # ── Banner: plná šířka (info), pod ním tlačítka vedle sebe ──────────
        st.info(
            f"**Nalezena automatická záloha** z {_age_str} — obsahuje: {_bak_desc}.",
            icon="💾")
        _rc2, _rc3, _ = st.columns([2, 1, 4])
        with _rc2:
            if st.button("🔄 Obnovit předchozí session", key="restore_autosave",
                         type="primary", width="stretch"):
                loaded = load_session_from_disk(SESSION_AUTOSAVE_FILE,
                                                SESSION_AUTOSAVE_KEYS)
                n = len([k for k in loaded if not k.startswith("_")])
                # ── v24.31: Synchronizace aliasů po obnově ───────────────────
                # Zajistíme, že translation_ready=True pokud máme výsledek
                if st.session_state.get("translation_result"):
                    st.session_state["translation_ready"] = True
                # Sync extraction aliases
                _ext_txt = (st.session_state.get("last_extracted_text") or
                            st.session_state.get("last_extraction_text") or "")
                if _ext_txt:
                    st.session_state["last_extracted_text"] = _ext_txt
                    st.session_state["last_extraction_text"] = _ext_txt
                # Přejdi na záložku kde jsou obnovená data
                if st.session_state.get("translation_ready") and st.session_state.get("translation_result"):
                    st.session_state["_gsb_jump_to_tab"] = 1
                elif st.session_state.get("last_validation_results"):
                    st.session_state["_gsb_jump_to_tab"] = 2
                elif st.session_state.get("last_extracted_text") or st.session_state.get("last_extracted_df_json"):
                    st.session_state["_gsb_jump_to_tab"] = 0
                # ── v24.31: Detailní hláška o úspěšné obnově session ─────────
                _restored_items = []
                if st.session_state.get("translation_result"):    _restored_items.append("překlad" if _L=="cz" else "translation")
                if st.session_state.get("last_extracted_text"):   _restored_items.append("extrakce" if _L=="cz" else "extraction")
                if st.session_state.get("last_validation_results"): _restored_items.append("validace" if _L=="cz" else "validation")
                if st.session_state.get("tr_src_text"):           _restored_items.append("text k překladu" if _L=="cz" else "input text")
                _items_str = ", ".join(_restored_items) if _restored_items else str(n) + (" klíčů" if _L=="cz" else " keys")
                st.session_state["_restore_success_msg"] = (
                    f"✅ Session obnovena z {_age_str} — obnoveno: {_items_str}" if _L == "cz"
                    else f"✅ Session restored from {_age_str} — restored: {_items_str}"
                )
                st.session_state["_show_restore_toast"] = True
                st.session_state["_autosave_banner_dismissed"] = True
                st.rerun()
        with _rc3:
            if st.button("✕ Zavřít", key="dismiss_autosave",
                         help="Zavřít toto upozornění", width="stretch"):
                st.session_state["_autosave_banner_dismissed"] = True
                try:
                    os.remove(SESSION_AUTOSAVE_FILE)
                except Exception:
                    pass
                st.rerun()

# ══════════════════════════════════════════════════════
# GLOBÁLNÍ STATUS BAR
# ══════════════════════════════════════════════════════
# Inicializace status state
for _sk, _sv in {
    "gsb_active":    False,   # probíhá nějaká operace?
    "gsb_label":     "",      # popis operace
    "gsb_detail":    "",      # detail (soubor, blok…)
    "gsb_progress":  0.0,     # 0.0–1.0
    "gsb_start":     0.0,     # _time.time() při startu
    "gsb_tab":       "",      # záložka kde operace běží
    "gsb_warn":      "",      # varování (prázdné = žádné)
}.items():
    if _sk not in st.session_state:
        st.session_state[_sk] = _sv

def gsb_start(label: str, tab: str = ""):
    """Aktivuj globální status bar. Volej před spuštěním dlouhé operace."""
    st.session_state.update({
        "gsb_active":   True,
        "gsb_label":    label,
        "gsb_detail":   "",
        "gsb_progress": 0.0,
        "gsb_start":    _time.time(),
        "gsb_tab":      tab,
        "gsb_warn":     "",
    })

def gsb_update(detail: str = "", progress: float = None, warn: str = ""):
    """Aktualizuj detail a progress (0.0–1.0). Volej v průběhu operace."""
    if detail:
        st.session_state["gsb_detail"] = detail
    if progress is not None:
        st.session_state["gsb_progress"] = max(0.0, min(1.0, progress))
    if warn:
        st.session_state["gsb_warn"] = warn

def gsb_done(label: str = ""):
    """Ukonči operaci. Nastaví _gsb_jump_to_tab aby po rerunu zůstala správná záložka.
    NEPŘEPÍNÁ na záložku 0 — zůstane tam kde úkol běžel."""
    _TAB_INDEX_DONE = {
        "🔍 Extrakce": 0, "🌐 Překlad": 1, "🧬 Validace": 2,
        "🤖 DeepSeek": 3, "💬 Chat": 3,
        "🧹 Čištění dat": 4, "✍️ Stylistika": 5, "⚙️ Workflow": 6,
        "📜 Historie": 7, "🧪 A/B Test": 8, "⚙️ Nastavení": 9,
        "💬 LMS Chat": 10,
        "🔍 Extraction": 0, "🌐 Translation": 1, "🧬 Validation": 2,
        "🧹 Data Cleaning": 4, "✍️ Style Polish": 5, "📜 History": 7,
        "⚙️ Settings": 9,
    }
    _tab = st.session_state.get("gsb_tab", "")
    _tab_idx = _TAB_INDEX_DONE.get(_tab, -1)
    # Fallback: pokud název záložky neodpovídá, použij current_tab_index
    if _tab_idx < 0:
        _tab_idx = st.session_state.get("current_tab_index", 0)
    st.session_state.update({
        "gsb_active":   False,
        "gsb_label":    label or "",
        "gsb_detail":   "",
        "gsb_progress": 1.0,
        "gsb_warn":     "",
    })
    # Nastav jump target — zabrání přepnutí na záložku 0 po rerunu
    st.session_state["_gsb_jump_to_tab"] = _tab_idx
    keep_current_tab(_tab_idx)
    # Žádný agresivní st.rerun() zde — volající kód ho provede sám pokud je třeba


def gsb_done_stay_on_tab(message: str, tab_index: int = None):
    """Bezpečná verze gsb_done — vždy zůstane na aktuální záložce."""
    if tab_index is None:
        tab_index = st.session_state.get("current_tab_index", 0)
    _TAB_NAMES = {
        0: "🔍 Extrakce", 1: "🌐 Překlad", 2: "🧬 Validace", 3: "🤖 DeepSeek",
        4: "🧹 Čištění dat", 5: "✍️ Stylistika", 6: "⚙️ Workflow",
        7: "📜 Historie", 8: "🧪 A/B Test", 9: "⚙️ Nastavení",
        10: "💬 LMS Chat", 11: "❓ Nápověda",
    }
    _tab_name = st.session_state.get("gsb_tab", "") or _TAB_NAMES.get(tab_index, "")
    st.session_state.update({
        "gsb_active":        False,
        "gsb_label":         message or "",
        "gsb_detail":        "",
        "gsb_progress":      1.0,
        "gsb_warn":          "",
        "gsb_tab":           _tab_name,
        "_gsb_jump_to_tab":  tab_index,
    })
    keep_current_tab(tab_index)
    save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)

def gsb_render():
    """Vyrenderuj globální status bar. Volej těsně před main_tabs.
    Po dokončení operace se zobrazí klikatelný banner → přeskočí na příslušnou záložku.
    """
    _L = st.session_state.get("lang", "cz")
    active   = st.session_state.get("gsb_active", False)
    label    = st.session_state.get("gsb_label", "")
    detail   = st.session_state.get("gsb_detail", "")
    progress = st.session_state.get("gsb_progress", 0.0)
    tab      = st.session_state.get("gsb_tab", "")
    warn     = st.session_state.get("gsb_warn", "")

    if not active and not label:
        return  # nic nezobrazovat

    # Mapování názvu záložky → index (musí odpovídat pořadí v main_tabs)
    _TAB_INDEX = {
        "🔍 Extrakce": 0, "🌐 Překlad": 1, "🧬 Validace": 2,
        "🤖 DeepSeek": 3, "💬 Chat": 3,
        "🧹 Čištění dat": 4, "✍️ Stylistika": 5, "⚙️ Workflow": 6,
        "📜 Historie": 7, "🧪 A/B Test": 8, "⚙️ Nastavení": 9,
        "💬 LMS Chat": 10, "❓ Nápověda": 11,
        "🔍 Extraction": 0, "🌐 Translation": 1, "🧬 Validation": 2,
        "🧹 Data Cleaning": 4, "✍️ Style Polish": 5, "📜 History": 7,
        "⚙️ Settings": 9,
    }
    tab_idx = _TAB_INDEX.get(tab, -1)

    if active:
        _color = "#f59e0b" if warn else "#2563b0"
        _bg    = "rgba(245,158,11,.1)" if warn else "rgba(37,99,176,.08)"
        # Spinner CSS animace
        _spin_css = ('<style>@keyframes _gsb_spin{to{transform:rotate(360deg)}}'
                     '._gsb_spinner{display:inline-block;animation:_gsb_spin 1s linear infinite;'
                     'margin-right:6px}</style>')
        _spinner  = '<span class="_gsb_spinner">⚙️</span>'
        _label_lc = label.lower()
        # Přeložená popisná zpráva (bez časomíry)
        if _L == "cz":
            if "validac" in _label_lc:
                _status_note = "Probíhá validace…"
            elif "překl" in _label_lc or "translat" in _label_lc:
                _status_note = "Probíhá překlad…"
            elif "extrakc" in _label_lc or "extract" in _label_lc:
                _status_note = "Probíhá extrakce…"
            elif "workflow" in _label_lc:
                _status_note = "Probíhá workflow…"
            else:
                _status_note = "Probíhá operace…"
        else:
            if "validac" in _label_lc or "validat" in _label_lc:
                _status_note = "Validation in progress…"
            elif "překl" in _label_lc or "translat" in _label_lc:
                _status_note = "Translation in progress…"
            elif "extrakc" in _label_lc or "extract" in _label_lc:
                _status_note = "Extraction in progress…"
            elif "workflow" in _label_lc:
                _status_note = "Workflow in progress…"
            else:
                _status_note = "Operation in progress…"
        st.markdown(
            _spin_css +
            f'<div style="background:{_bg};border:1px solid {_color};border-radius:8px;'
            f'padding:8px 14px;margin-bottom:8px;display:flex;align-items:center;gap:12px">'
            f'{_spinner}'
            f'<div style="flex:1">'
            f'<div style="font-weight:600;font-size:.9rem;color:{_color}">{label}</div>'
            f'<div style="font-size:.8rem;color:#64748b">{_status_note}'
            + (f' · {detail}' if detail else '') +
            '</div>'
            f'{"<div style=font-size:.8rem;color:#f59e0b>⚠️ " + warn + "</div>" if warn else ""}'
            f'</div>'
            f'</div>',
            unsafe_allow_html=True
        )
        if progress > 0:
            st.progress(progress)
    elif label:
        # Dokončená operace — klikatelný banner; zobrazí se dokud uživatel nezavře (✕)
        if tab_idx >= 0:
            _goto_lbl = (f"Přejít na {tab}" if _L == "cz" else f"Go to {tab}")
            _col_msg, _col_btn, _col_cls = st.columns([6, 2, 1])
            with _col_msg:
                st.success(label)
            with _col_btn:
                if st.button(f"👉 {_goto_lbl}", key="gsb_goto_tab", type="primary"):
                    st.session_state["_gsb_jump_to_tab"] = tab_idx
                    st.query_params["tab"] = str(tab_idx)
                    keep_current_tab(tab_idx)
                    st.rerun()
            with _col_cls:
                if st.button("✕", key="gsb_dismiss", help=("Zavřít" if _L == "cz" else "Dismiss")):
                    st.session_state["gsb_label"] = ""
                    st.session_state["gsb_tab"]   = ""
                    _cur_tab_dismiss = st.session_state.get("presence_tab_idx", tab_idx)
                    st.session_state["_gsb_jump_to_tab"] = _cur_tab_dismiss
                    st.query_params["tab"] = str(_cur_tab_dismiss)
                    st.rerun()
        else:
            # Bez known tab → zobraz zprávu s ✕
            _col_msg2, _col_cls2 = st.columns([8, 1])
            with _col_msg2:
                st.success(label)
            with _col_cls2:
                if st.button("✕", key="gsb_dismiss_notab", help=("Zavřít" if _L == "cz" else "Dismiss")):
                    st.session_state["gsb_label"] = ""
                    # Zkus zachovat aktuální tab z presence_tab_idx
                    _cur_tab = st.session_state.get("presence_tab_idx", 0)
                    if _cur_tab > 0:
                        st.session_state["_gsb_jump_to_tab"] = _cur_tab
                    st.rerun()

# ══════════════════════════════════════════════════════
# DESKTOP NOTIFIKACE
# ══════════════════════════════════════════════════════
def desktop_notify(title: str, body: str = ""):
    """Spustí browser desktop notifikaci přes window._lmuNotify (definováno v JS)."""
    safe_title = title.replace("'", "\\'").replace("\n", " ")
    safe_body  = body.replace("'", "\\'").replace("\n", " ")
    st.html(f"<script>if(window._lmuNotify)window._lmuNotify('{safe_title}','{safe_body}');</script>")


# ══════════════════════════════════════════════════════
# CHECKPOINT SYSTÉM — pozastavení a obnovení dávkové extrakce
# ══════════════════════════════════════════════════════
_CHECKPOINT_DIR = Path("checkpoints")

# ── _temp helper funkce ───────────────────────────────────────────────
def _temp_save_txt(prefix: str, batch_idx: int, text: str, total: int = 0) -> "Path":
    """Atomicky uloží dávku jako TXT do _temp/. Vrátí cestu."""
    _TEMP_DIR.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    suf = f"_of{total}" if total else ""
    p   = _TEMP_DIR / f"{prefix}_b{batch_idx:04d}{suf}_{ts}.txt"
    tmp = str(p) + ".tmp"
    with open(tmp, "w", encoding="utf-8") as fh:
        fh.write(text)
    os.replace(tmp, p)
    # v24.35: zachovat jen posledních 10 souborů tohoto prefixu
    _temp_prune(prefix, keep_last=10)
    return p

def _temp_save_docx(prefix: str, batch_idx: int, text: str, total: int = 0) -> "Path":
    """Atomicky uloží dávku jako DOCX do _temp/. Vrátí cestu.
    Při chybě python-docx padne zpět na .txt (data se neztratí)."""
    _TEMP_DIR.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    suf = f"_of{total}" if total else ""
    try:
        p   = _TEMP_DIR / f"{prefix}_b{batch_idx:04d}{suf}_{ts}.docx"
        tmp = str(p) + ".tmp"
        with open(tmp, "wb") as fh:
            fh.write(to_docx_bytes(text))
        os.replace(tmp, p)
        # v24.35: zachovat jen posledních 10 souborů tohoto prefixu
        _temp_prune(prefix, keep_last=10)
        return p
    except Exception:
        # Fallback: ulož jako .txt aby data nezanikla
        p2  = _TEMP_DIR / f"{prefix}_b{batch_idx:04d}{suf}_{ts}_fallback.txt"
        tmp2 = str(p2) + ".tmp"
        with open(tmp2, "w", encoding="utf-8") as fh:
            fh.write(text)
        os.replace(tmp2, p2)
        # v24.35: zachovat jen posledních 10 souborů tohoto prefixu
        _temp_prune(prefix, keep_last=10)
        return p2

def _temp_save_json(prefix: str, batch_idx: int, data: Any, total: int = 0) -> "Path":
    """Atomicky uloží dávku jako JSON do _temp/. Vrátí cestu."""
    _TEMP_DIR.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    suf = f"_of{total}" if total else ""
    p   = _TEMP_DIR / f"{prefix}_b{batch_idx:04d}{suf}_{ts}.json"
    tmp = str(p) + ".tmp"
    with open(tmp, "w", encoding="utf-8") as fh:
        json.dump(data, fh, ensure_ascii=False, indent=2)
    os.replace(tmp, p)
    # v24.35: zachovat jen posledních 10 souborů tohoto prefixu
    _temp_prune(prefix, keep_last=10)
    return p

def _temp_cleanup(prefix: str):
    """Smaže temp soubory daného prefixu (volá se po úspěšném dokončení)."""
    if not _TEMP_DIR.exists():
        return
    for f in list(_TEMP_DIR.glob(f"{prefix}_b*")):
        try:
            f.unlink()
        except Exception:
            pass

def _temp_prune(prefix: str, keep_last: int = 3):
    """Zachová jen posledních N souborů daného prefixu — zabrání plnění disku
    při vícenásobném překladu/extrakci bez úspěšného dokončení.
    Volá se na začátku každé operace."""
    if not _TEMP_DIR.exists():
        return
    files = sorted(_TEMP_DIR.glob(f"{prefix}_b*"), key=lambda p: p.stat().st_mtime)
    for old in files[:-keep_last] if len(files) > keep_last else []:
        try:
            old.unlink()
        except Exception:
            pass


def _temp_file_browser(lang: str = "cz", key_prefix: str = "tb") -> None:
    """Prohlížeč souborů _temp/ přímo ve Streamlit UI.
    key_prefix zajišťuje unikátní klíče při volání z více záložek."""
    temp_dir = _TEMP_DIR
    label = "📁 Prohlížeč dočasných souborů (_temp/)" if lang == "cz" else "📁 Temp file browser (_temp/)"
    with st.expander(label, expanded=False):
        if not temp_dir.exists() or not any(temp_dir.iterdir()):
            st.info("Adresář _temp/ je prázdný." if lang == "cz" else "The _temp/ directory is empty.")
            return
        all_files = sorted([p for p in temp_dir.rglob("*") if p.is_file()],
                           key=lambda p: p.stat().st_mtime, reverse=True)
        if not all_files:
            st.info("Adresář _temp/ je prázdný." if lang == "cz" else "The _temp/ directory is empty.")
            return
        extensions = sorted({p.suffix.lower() for p in all_files})
        ext_filter = st.multiselect(
            "Filtr přípony" if lang == "cz" else "Extension filter",
            options=extensions, default=extensions, key=f"{key_prefix}_ext_filter")
        filtered = [p for p in all_files if p.suffix.lower() in ext_filter]
        st.caption(f"Nalezeno {len(filtered)} / {len(all_files)} souborů v _temp/")
        for fp in filtered:
            try:
                size_kb = fp.stat().st_size / 1024
                mtime = datetime.fromtimestamp(fp.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                rel_path = fp.relative_to(temp_dir.parent)
            except Exception:
                continue
            col_info, col_dl, col_del = st.columns([5, 1, 1])
            with col_info:
                st.markdown(f"`{rel_path}` &nbsp; <span style='color:gray;font-size:0.85em'>"
                            f"{size_kb:.1f} kB · {mtime}</span>", unsafe_allow_html=True)
            with col_dl:
                try:
                    st.download_button("⬇️", data=fp.read_bytes(), file_name=fp.name,
                        mime="application/octet-stream", key=f"{key_prefix}_dl_{fp}",
                        help="Stáhnout" if lang == "cz" else "Download")
                except Exception as e:
                    st.error(str(e))
            with col_del:
                if st.button("🗑️", key=f"{key_prefix}_del_{fp}",
                             help="Smazat" if lang == "cz" else "Delete"):
                    try:
                        fp.unlink()
                        st.toast(f"Smazáno: {fp.name}" if lang == "cz" else f"Deleted: {fp.name}")
                        st.rerun()
                    except Exception as e:
                        st.error(str(e))
        st.markdown("---")
        if st.button("🗑️ Smazat vše zobrazené" if lang == "cz" else "🗑️ Delete all listed",
                     key=f"{key_prefix}_del_all", type="secondary"):
            deleted = errors = 0
            for fp in filtered:
                try:
                    fp.unlink(); deleted += 1
                except Exception:
                    errors += 1
            st.toast(f"Smazáno {deleted}" + (f", {errors} chyb." if errors else "."))
            st.rerun()

def checkpoint_save(name: str, data: dict):
    """Uloží checkpoint dávkové operace na disk (atomický zápis)."""
    _CHECKPOINT_DIR.mkdir(exist_ok=True)
    path     = _CHECKPOINT_DIR / f"{name}.json"
    tmp_path = str(path) + ".tmp"
    payload  = {
        "name":     name,
        "saved_at": datetime.now().isoformat(),
        "data":     data,
    }
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, path)
    return str(path)

def checkpoint_load(name: str) -> Optional[dict]:
    """Načte checkpoint ze souboru. Vrátí None pokud neexistuje."""
    path = _CHECKPOINT_DIR / f"{name}.json"
    if not path.exists():
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None

def checkpoint_delete(name: str):
    """Smaže checkpoint soubor."""
    path = _CHECKPOINT_DIR / f"{name}.json"
    if path.exists():
        path.unlink()

def checkpoint_list() -> List[dict]:
    """Vrátí seznam všech dostupných checkpointů."""
    if not _CHECKPOINT_DIR.exists():
        return []
    result = []
    for f in sorted(_CHECKPOINT_DIR.glob("*.json"), reverse=True):
        try:
            with open(f, "r", encoding="utf-8") as fh:
                meta = json.load(fh)
            files_done = len(meta.get("data", {}).get("completed", {}))
            files_total = meta.get("data", {}).get("total_files", "?")
            result.append({
                "name":       meta.get("name", f.stem),
                "saved_at":   meta.get("saved_at", "")[:16].replace("T", " "),
                "files_done": files_done,
                "files_total": files_total,
                "path":       str(f),
            })
        except Exception:
            pass
    return result


# ══════════════════════════════════════════════════════
# FTS — full-text search v historii výsledků extrakce
# ══════════════════════════════════════════════════════
_FTS_DB_PATH = Path("extraction_fts.db")
_FTS_CONN: Optional[sqlite3.Connection] = None

def _fts_conn() -> sqlite3.Connection:
    global _FTS_CONN
    if _FTS_CONN is None:
        _FTS_CONN = sqlite3.connect(str(_FTS_DB_PATH), check_same_thread=False, timeout=10)
        _FTS_CONN.execute("PRAGMA journal_mode=WAL")
        _FTS_CONN.execute("PRAGMA synchronous=NORMAL")
        _FTS_CONN.execute("PRAGMA cache_size=2000")
        _FTS_CONN.execute("""CREATE TABLE IF NOT EXISTS ext_records (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            source_file TEXT,
            timestamp   TEXT,
            content     TEXT,
            result_json TEXT
        )""")
        # FTS5 virtual table
        _FTS_CONN.execute("""CREATE VIRTUAL TABLE IF NOT EXISTS ext_fts
            USING fts5(content, source_file, timestamp, content='ext_records', content_rowid='id')
        """)
        _FTS_CONN.commit()
    return _FTS_CONN

def fts_index_results(results: list, source_file: str = ""):
    """Zaindexuje výsledky extrakce do FTS databáze.
    Používá batch insert (executemany) s jediným commitem — rychlejší pro N souborů.
    """
    try:
        conn = _fts_conn()
        ts = datetime.now().isoformat()
        records_rows = []
        fts_rows = []
        for rec in results:
            if not rec.get("ok"):
                continue
            fname = rec.get("file", source_file)
            content_text = rec.get("result", "")
            try:
                recs_json = json.loads(_JSON_FENCE_RE.sub("", content_text).strip())
                if isinstance(recs_json, list):
                    content_text = " ".join(
                        " ".join(str(v) for v in r.values() if v and str(v).strip())
                        for r in recs_json if isinstance(r, dict)
                    )
            except Exception:
                pass
            records_rows.append((fname, ts, content_text, rec.get("result", "")))
            fts_rows.append((fname, ts, content_text))

        if not records_rows:
            return

        # Batch insert do ext_records, pak okamžitě FTS pro každý záznam
        cursor = conn.cursor()
        for row in records_rows:
            row_id = cursor.execute(
                "INSERT INTO ext_records (source_file, timestamp, content, result_json) VALUES (?,?,?,?)",
                row
            ).lastrowid
            # row: (fname, ts, content_text, result_json)
            cursor.execute(
                "INSERT INTO ext_fts(rowid, content, source_file, timestamp) VALUES (?,?,?,?)",
                (row_id, row[2], row[0], row[1])
            )
        conn.commit()
    except Exception:
        pass

def fts_search(query: str, limit: int = 20) -> List[dict]:
    """Prohledá FTS index. Vrátí seznam nalezených záznamů."""
    if not query.strip():
        return []
    try:
        conn = _fts_conn()
        rows = conn.execute("""
            SELECT r.id, r.source_file, r.timestamp, r.result_json,
                   highlight(ext_fts, 0, '<mark>', '</mark>') as snippet
            FROM ext_fts f
            JOIN ext_records r ON r.id = f.rowid
            WHERE ext_fts MATCH ?
            ORDER BY rank
            LIMIT ?
        """, (query, limit)).fetchall()
        return [
            {"id": r[0], "source_file": r[1], "timestamp": r[2][:16].replace("T"," "),
             "result_json": r[3], "snippet": r[4]}
            for r in rows
        ]
    except Exception:
        return []

def fts_record_count() -> int:
    try:
        return _fts_conn().execute("SELECT COUNT(*) FROM ext_records").fetchone()[0]
    except Exception:
        return 0


# ══════════════════════════════════════════════════════
# HLAVNÍ ZÁLOŽKY


# ══════════════════════════════════════════════════════
# SIDEBAR – cachované helpery (šetří SQLite a subprocess volání)
# ══════════════════════════════════════════════════════
@st.cache_data(ttl=5.0, show_spinner=False)
def _cached_sidebar_stats() -> tuple:
    """Vrací (disk_cache_count, fts_record_count, checkpoint_count).
    Cachováno na 5 s, aby se při každém rerun nevolaly SQLite dotazy."""
    try:
        dc  = _disk_cache_count()
    except Exception:
        dc = 0
    try:
        fts = fts_record_count()
    except Exception:
        fts = 0
    try:
        ck  = len(checkpoint_list())
    except Exception:
        ck = 0
    return dc, fts, ck


@st.cache_data(ttl=30.0, show_spinner=False)
def _cached_lms_available() -> bool:
    """Cache pro _lms_available — subprocess volání je drahé na každý rerun.
    TTL 30 s pokryje běžnou uživatelskou aktivitu."""
    try:
        return _lms_available()
    except Exception:
        return False


# ══════════════════════════════════════════════════════
# HW MONITORING — CPU / RAM / GPU
# ══════════════════════════════════════════════════════
try:
    import psutil as _psutil
    _PSUTIL_OK = True
except ImportError:
    _PSUTIL_OK = False
    _psutil = None

_PYNVML_OK   = False
_PYNVML_INIT = False

def _init_pynvml() -> bool:
    """Inicializuje pynvml jednou — lazy, cachuje výsledek.
    Potlačí FutureWarning z deprecated pynvml balíčku
    (správný balíček je nvidia-ml-py, ale API je stejné).
    """
    global _PYNVML_OK, _PYNVML_INIT
    if _PYNVML_INIT:
        return _PYNVML_OK
    _PYNVML_INIT = True
    try:
        import warnings
        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", category=FutureWarning, module="pynvml")
            warnings.filterwarnings("ignore", message=".*pynvml.*deprecated.*")
            import pynvml as _pynvml_mod
            _pynvml_mod.nvmlInit()
        _PYNVML_OK = True
    except Exception:
        _PYNVML_OK = False
    return _PYNVML_OK


def _get_hw_stats() -> dict:
    """Vrátí aktuální vytížení CPU, RAM a GPU.
    Bezpečné volání — všechny výjimky jsou zachyceny.
    Výsledek:
      cpu_pct    : float 0–100
      ram_pct    : float 0–100
      ram_used_mb: int
      ram_total_mb: int
      gpu_pct    : float 0–100 | None (pokud GPU není k dispozici)
      vram_pct   : float 0–100 | None
      vram_used_mb: int | None
      vram_total_mb: int | None
      gpu_temp_c : int | None
      gpu_name   : str | None
    """
    stats = {
        "cpu_pct": 0.0, "ram_pct": 0.0,
        "ram_used_mb": 0, "ram_total_mb": 0,
        "gpu_pct": None, "vram_pct": None,
        "vram_used_mb": None, "vram_total_mb": None,
        "gpu_temp_c": None, "gpu_name": None,
    }

    # ── CPU + RAM ────────────────────────────────────
    if _PSUTIL_OK:
        try:
            stats["cpu_pct"]     = _psutil.cpu_percent(interval=None)
            vm = _psutil.virtual_memory()
            stats["ram_pct"]     = vm.percent
            stats["ram_used_mb"] = vm.used  // 1_048_576
            stats["ram_total_mb"]= vm.total // 1_048_576
        except Exception:
            pass

    # ── GPU přes pynvml ──────────────────────────────
    if _init_pynvml():
        try:
            import pynvml as _pynvml_mod
            h    = _pynvml_mod.nvmlDeviceGetHandleByIndex(0)
            name = _pynvml_mod.nvmlDeviceGetName(h)
            # nvmlDeviceGetName vrací bytes nebo str dle verze pynvml
            if isinstance(name, bytes):
                name = name.decode("utf-8", errors="replace")
            util = _pynvml_mod.nvmlDeviceGetUtilizationRates(h)
            mem  = _pynvml_mod.nvmlDeviceGetMemoryInfo(h)
            try:
                temp = _pynvml_mod.nvmlDeviceGetTemperature(
                    h, _pynvml_mod.NVML_TEMPERATURE_GPU)
            except Exception:
                temp = None
            stats["gpu_name"]     = name
            stats["gpu_pct"]      = float(util.gpu)
            stats["vram_used_mb"] = mem.used  // 1_048_576
            stats["vram_total_mb"]= mem.total // 1_048_576
            stats["vram_pct"]     = (mem.used / mem.total * 100) if mem.total else 0.0
            stats["gpu_temp_c"]   = temp
        except Exception:
            pass

    # ── GPU fallback: nvidia-smi subprocess ──────────
    if stats["gpu_pct"] is None:
        try:
            import subprocess as _sp
            r = _sp.run(
                ["nvidia-smi",
                 "--query-gpu=name,utilization.gpu,memory.used,memory.total,temperature.gpu",
                 "--format=csv,noheader,nounits"],
                capture_output=True, text=True, timeout=2)
            if r.returncode == 0 and r.stdout.strip():
                parts = [p.strip() for p in r.stdout.strip().split(",")]
                if len(parts) >= 5:
                    stats["gpu_name"]     = parts[0]
                    stats["gpu_pct"]      = float(parts[1])
                    stats["vram_used_mb"] = int(parts[2])
                    stats["vram_total_mb"]= int(parts[3])
                    stats["vram_pct"]     = (int(parts[2]) / int(parts[3]) * 100
                                             if int(parts[3]) else 0.0)
                    stats["gpu_temp_c"]   = int(parts[4]) if parts[4].isdigit() else None
        except Exception:
            pass

    return stats


@st.cache_data(ttl=2.0, show_spinner=False)
def _cached_hw_stats() -> dict:
    """Cache HW stats na 2 s — zabrání zbytečným syscallům při každém rerun."""
    return _get_hw_stats()


def _hw_bar_html(pct: float, color: str = "#2563b0", height: int = 6) -> str:
    """Vrátí HTML mini progress bar pro dané procento (0–100)."""
    pct_c = max(0.0, min(100.0, pct))
    bg    = "rgba(0,0,0,.07)"
    return (
        f'<div style="background:{bg};border-radius:{height}px;'
        f'height:{height}px;overflow:hidden;margin:2px 0 4px 0">'
        f'<div style="width:{pct_c:.1f}%;background:{color};'
        f'height:100%;border-radius:{height}px;'
        f'transition:width .4s ease"></div></div>'
    )


def _bar_color(pct: float) -> str:
    """Barva podle vytížení: zelená → žlutá → červená."""
    if pct < 60:
        return "#22c55e"   # zelená
    if pct < 85:
        return "#f59e0b"   # žlutá / oranžová
    return "#ef4444"       # červená

# ══════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════
with st.sidebar:
    _L = st.session_state["lang"]

    # ── Logo + lang toggle ────────────────────────────
    col_logo, col_lang = st.columns([3, 2])
    with col_logo:
        st.markdown(f"**🧠 v{VERSION}**")
    with col_lang:
        lang_lbl  = "🇬🇧 **English**" if _L == "cz" else "🇨🇿 **Čeština**"
        lang_help = t("lang_toggle", _L)
        if st.button(lang_lbl, help=lang_help, width='stretch', key="sb_lang"):
            st.session_state["lang"] = "en" if _L == "cz" else "cz"
            st.rerun()

    st.divider()

    # ── Quick Presets ─────────────────────────────────
    _preset_label = "🚀 Rychlé nastavení" if _L == "cz" else "🚀 Quick Preset"

    # v24.27: sloučení vestavěných + uživatelských presetů
    _builtin_presets = [
        "— vyber scénář —" if _L == "cz" else "— select scenario —",
        "Hyolitha velký PDF (300+ stran)",
        "Historický text překlad (CZ/RU/DE → EN)",
        "Rychlá extrakce + validace",
        "Plný Workflow – velký dokument",
    ]
    _user_presets = list(st.session_state.get("quick_presets", {}).keys())
    _all_preset_options = _builtin_presets + (["─── Moje nastavení ───"] + _user_presets if _user_presets else [])

    # Aktivní preset — tučný černý nápis
    _active_pname = st.session_state.get("_active_preset_name", "")
    if _active_pname:
        st.markdown(
            f"<div style='font-size:15px;font-weight:700;color:#000;margin-bottom:4px'>"
            f"✅ Aktivní: {_active_pname}</div>" if _L == "cz" else
            f"<div style='font-size:15px;font-weight:700;color:#000;margin-bottom:4px'>"
            f"✅ Active: {_active_pname}</div>",
            unsafe_allow_html=True)

    _preset = st.selectbox(
        _preset_label,
        options=_all_preset_options,
        key="quick_preset",
        label_visibility="visible",
    )

    # Mapování vestavěných presetů na nastavení
    _BUILTIN_PRESET_SETTINGS = {
        "Hyolitha velký PDF (300+ stran)": {
            "wf_large_mode": True, "wf_chunk_target": 4200, "lms_max_concurrent": 4,
        },
        "Historický text překlad (CZ/RU/DE → EN)": {
            "wf_large_mode": False, "wf_chunk_target": 2200, "temp_t": 0.1,
        },
        "Rychlá extrakce + validace": {
            "wf_large_mode": False, "lms_max_concurrent": 4,
        },
        "Plný Workflow – velký dokument": {
            "wf_large_mode": True, "wf_chunk_target": 3500, "lms_max_concurrent": 4,
        },
    }
    _BUILTIN_PRESET_LABELS = {
        "Hyolitha velký PDF (300+ stran)": "✅ Preset: Large Document Mode + Hyolitha optimalizace",
        "Historický text překlad (CZ/RU/DE → EN)": "ℹ️ Preset: překlad historického textu (sekvenční, nižší teplota)",
        "Rychlá extrakce + validace": "ℹ️ Preset: rychlá extrakce + validace — doporučené výchozí nastavení",
        "Plný Workflow – velký dokument": "ℹ️ Preset: plný pipeline pro velký dokument (Large Mode)",
    }

    if _preset in _BUILTIN_PRESET_SETTINGS:
        # Zobrazení popisu zvoleného presetu
        _lbl = _BUILTIN_PRESET_LABELS[_preset]
        if "✅" in _lbl:
            st.success(_lbl)
        else:
            st.info(_lbl)
        # Tlačítko Použít — aplikace přes defer (bezpečné i pro temp_t a jiné widgety)
        if st.button(tt("🚀 Použít tento scénář", "🚀 Apply this scenario", _L),
                     key="sb_apply_builtin_preset", width="stretch"):
            defer_apply_settings(_BUILTIN_PRESET_SETTINGS[_preset])
            st.toast(tt(f"✅ Preset '{_preset}' bude aplikován",
                        f"✅ Preset '{_preset}' will be applied", _L))
            st.rerun()
    elif _preset in _user_presets:
        # Uživatelský preset — zobraz popis a použij
        _up_data = st.session_state["quick_presets"].get(_preset, {})
        if _up_data.get("description"):
            st.caption(_up_data["description"])
        if st.button(tt("🚀 Použít nastavení", "🚀 Apply preset", _L),
                     key="sb_apply_user_preset", width="stretch"):
            _up_settings = {}
            if "temperature" in _up_data:
                _up_settings.update({"ext_temp": _up_data["temperature"],
                                     "tr_temp":  _up_data["temperature"],
                                     "temp_t":   _up_data["temperature"]})
            if "chunk_size"  in _up_data: _up_settings.update({"tr_chunk_target": _up_data["chunk_size"], "ext_chunk_size": _up_data["chunk_size"]})
            if "max_tokens"  in _up_data: _up_settings["ext_max_tokens"] = _up_data["max_tokens"]
            if "large_mode"  in _up_data: _up_settings["wf_large_mode"]  = _up_data["large_mode"]
            if "parallel"    in _up_data: _up_settings["tr_parallel"]    = _up_data["parallel"]
            if _up_data.get("prompt"):   _up_settings["ext_prompt"]      = _up_data["prompt"]
            _up_settings["quick_preset"] = "— vyber scénář —" if _L=="cz" else "— select scenario —"
            defer_apply_settings(_up_settings)
            st.session_state["_active_preset_name"] = _preset
            st.toast(tt(f"✅ Nastavení '{_preset}' aplikováno",
                        f"✅ Preset '{_preset}' applied", _L))
            st.rerun()

    # ── Large Document Mode toggle ────────────────────
    _large_mode = st.toggle(
        "📄 Obsáhlý dokument" if _L == "cz" else "📄 Large Document Mode",
        value=st.session_state.get("wf_large_mode", False),
        key="sb_large_mode_toggle",
        help=("Optimalizuje dělení na bloky a paralelismus pro dokumenty 100+ stran."
              if _L == "cz" else
              "Optimizes chunking and parallelism for 100+ page documents."),
    )
    if _large_mode != st.session_state.get("wf_large_mode", False):
        safe_set("wf_large_mode", _large_mode)

    st.divider()
    base_url = st.text_input(t("api_url", _L), value=DEFAULT_BASE_URL, key="sb_api_url")

    # ── LMS CLI dostupnost (cachováno 30s) ─────────────
    _lms_ok = _cached_lms_available()

    # ── Tlačítka: Připojit + Načíst seznam modelů ────
    _rc1, _rc2 = st.columns(2)
    with _rc1:
        if st.button(t("sb_connect", _L), width='stretch',
                     key="sb_load_models", help=t("sb_connect_help", _L)):
            with st.spinner(t("sb_connecting", _L)):
                try:
                    models = list_models(base_url)
                    st.session_state["available_models"] = models
                    st.session_state["model_status"] = "ok"
                except Exception as e:
                    st.session_state["model_status"] = "error"
                    st.error(f"❌ {e}")
    with _rc2:
        if st.button(t("sb_list_models", _L), width='stretch', key="sb_lms_ls",
                     disabled=not _lms_ok,
                     help=t("sb_list_models_on", _L) if _lms_ok
                          else t("sb_list_models_off", _L)):
            with st.spinner(t("sb_loading_list", _L)):
                _lms_result, _lms_raw = lms_ls()
                st.session_state["lms_model_list"] = _lms_result
                st.session_state["lms_ls_raw"] = _lms_raw
                _active = lms_ps()
                if _active:
                    st.session_state["lms_active_model"] = _active
                if _lms_result:
                    st.toast(t("sb_loaded_n_models", _L, n=len(_lms_result)))

    models = st.session_state["available_models"]

    # ── Zobrazení aktivního modelu ────────────────────
    if models:
        selected_model = models[0]
        st.session_state["sb_model_sel"] = selected_model
        # Automaticky aplikuj per-model settings (teplota, max_tokens) při výběru modelu
        _pms_all = st.session_state.get("per_model_settings", {})
        _pms_match = next(
            (v for k, v in _pms_all.items() if k.lower() in selected_model.lower()),
            None)
        if _pms_match:
            defer_apply_settings({
                "temp_ext": float(_pms_match.get("temperature", 0.1)),
                "temp_t":   float(_pms_match.get("temperature", 0.1)),
            })
            if "max_tokens" in _pms_match:
                st.session_state["ext_max_tokens"] = int(_pms_match["max_tokens"])
        status_icon = {"ok": "🟢", "error": "🔴", "unknown": "⚪"}.get(
            st.session_state.get("model_status", "unknown"), "⚪")
        model_short = selected_model.split("/")[-1]
        _active_lbl = t("sb_active_model", _L)
        st.markdown(
            f'<div style="background:#f0fdf4;border:1px solid #86efac;border-radius:6px;'
            f'padding:6px 12px;margin:4px 0">'
            f'<div style="font-size:.75rem;color:#6b7280;margin-bottom:2px">{_active_lbl}</div>'
            f'<div style="font-weight:700;font-size:.9rem;color:#166534">{status_icon} {model_short}</div>'
            f'</div>',
            unsafe_allow_html=True)
    else:
        selected_model = None
        st.warning(t("no_model", _L))

    # ── Přepínač modelů přes lms CLI ─────────────────
    st.markdown(f"**{t('sb_switch_model', _L)}**")

    _lms_list = st.session_state.get("lms_model_list", [])
    _lms_raw  = st.session_state.get("lms_ls_raw", "")

    if not _lms_ok:
        st.caption(t("sb_lms_missing", _L))
    elif not _lms_list:
        st.caption(t("sb_click_list", _L))
        if _lms_raw:
            with st.expander(t("sb_debug_lms", _L), expanded=True):
                st.code(_lms_raw[:1500], language=None)
                st.caption(t("sb_debug_hint", _L))
    else:
        # Vyber model z rozbalovací nabídky
        _model_ids = [m["id"] for m in _lms_list if m.get("id")]
        _active_id = st.session_state.get("lms_active_model", selected_model or "")

        _selected_for_load = st.selectbox(
            t("sb_select_load", _L),
            _model_ids,
            index=_model_ids.index(_active_id) if _active_id in _model_ids else 0,
            key="sb_lms_select",
            format_func=lambda x: x.split("/")[-1] if "/" in x else x
        )

        # Detaily vybraného modelu
        _sel_meta = next((m for m in _lms_list if m.get("id") == _selected_for_load), {})
        if _sel_meta:
            _meta_parts = []
            if _sel_meta.get("size_gb"):      _meta_parts.append(f"💾 {_sel_meta['size_gb']}")
            if _sel_meta.get("architecture"): _meta_parts.append(f"🏗️ {_sel_meta['architecture']}")
            if _sel_meta.get("context"):      _meta_parts.append(f"📏 ctx {_sel_meta['context']}")
            if _meta_parts:
                st.caption(" · ".join(_meta_parts))

        # Je to aktuálně načtený model? — vždy bool
        _is_active = bool(
            _selected_for_load and _active_id and (
                _selected_for_load == _active_id or
                _selected_for_load.split("/")[-1] in (_active_id or "")
            )
        )

        if _is_active:
            st.markdown(f'<span style="color:#15803d;font-size:.82rem">'
                        f'{t("sb_model_active", _L)}</span>',
                        unsafe_allow_html=True)

        _lcol1, _lcol2 = st.columns([3, 2])
        with _lcol1:
            _load_disabled = bool(st.session_state.get("lms_loading", False))
            _btn_load_lbl  = t("sb_loading", _L) if _load_disabled else t("sb_load_model", _L)
            if st.button(_btn_load_lbl, type="primary", width='stretch',
                         key="sb_lms_load",
                         disabled=bool(_load_disabled or _is_active),
                         help=t("sb_load_model_help", _L)):
                st.session_state["lms_loading"] = True
                st.rerun()

        with _lcol2:
            if st.button(t("sb_unload", _L), width='stretch', key="sb_lms_unload",
                         help=t("sb_unload_help", _L)):
                with st.spinner(t("sb_unloading", _L)):
                    _ok, _msg = lms_unload()
                    if _ok:
                        st.session_state["available_models"] = []
                        st.session_state["lms_active_model"] = None
                        st.session_state["sb_model_sel"] = None
                        st.toast(t("sb_unloaded", _L))
                    else:
                        st.error(f"❌ {_msg}")
                st.rerun()

    # ── Skutečné načtení modelu (mimo button kvůli spinnerům) ──
    if st.session_state.get("lms_loading"):
        _model_to_load = st.session_state.get("sb_lms_select")
        if _model_to_load:
            _short = _model_to_load.split('/')[-1]
            with st.status(t("sb_loading_status", _L, m=_short),
                           expanded=True) as _load_status:
                st.write(t("sb_unloading_current", _L))
                lms_unload()
                st.write(t("sb_loading_new", _L, m=_model_to_load))
                st.write(t("sb_loading_takes", _L))
                _ok, _msg = lms_load(_model_to_load)
                if _ok:
                    _load_status.update(label=t("sb_loaded_ok", _L), state="complete")
                    st.session_state["lms_active_model"] = _model_to_load
                    try:
                        _new_models = list_models(base_url)
                        st.session_state["available_models"] = _new_models
                        st.session_state["model_status"] = "ok"
                    except Exception:
                        pass
                    st.toast(t("sb_model_active_toast", _L, m=_short))
                else:
                    _load_status.update(label=t("sb_load_error", _L), state="error")
                    st.error(f"❌ {_msg}")
        st.session_state["lms_loading"] = False
        st.rerun()

    # ── Paralelní requesty — dobře viditelné pod modelem ──
    st.markdown(f"**{t('sb_parallel', _L)}**")
    _cur_par = st.session_state.get("lms_max_concurrent", 1)
    lms_par = st.select_slider(
        t("sb_max_concurrent", _L),
        options=[1, 2, 3, 4, 6, 8],
        value=_cur_par,
        key="sb_lms_parallel",
        help=t("sb_par_help", _L))
    if lms_par != _cur_par:
        st.session_state["lms_max_concurrent"] = lms_par
        _llm_queue.set_max(lms_par)
    elif _llm_queue.max_concurrent != _cur_par:
        _llm_queue.set_max(_cur_par)

    _par_serial_label = t("sb_par_serial", _L)
    _par_rec_label    = t("sb_par_recommended", _L)
    _par_labels = {1: (_par_serial_label,  "#9ca3af"),
                   2: ("🟡 2×",             "#d97706"),
                   3: ("🟡 3×",             "#d97706"),
                   4: (_par_rec_label,     "#15803d"),
                   6: ("🟢 6×",             "#15803d"),
                   8: ("🟢 8×",             "#15803d")}
    _par_label, _par_color = _par_labels.get(lms_par, ("⚡", "#374151"))
    st.markdown(f'<span style="color:{_par_color};font-size:.85rem;font-weight:600">'
                f'{_par_label}</span>', unsafe_allow_html=True)

    st.divider()

    # ── Settings ──────────────────────────────────────
    st.session_state["offline_mode"] = st.checkbox(
        t("offline_mode", _L),
        value=st.session_state["offline_mode"],
        help=t("offline_help", _L), key="sb_offline")
    new_timeout = st.slider(t("val_timeout", _L), 5, 60,
                             st.session_state["val_timeout"], 1, key="val_to_slider")
    if new_timeout != st.session_state["val_timeout"]:
        st.session_state["val_timeout"] = new_timeout

    st.divider()

    # ── Global system prompt ──────────────────────────
    with st.expander(t("global_sys_prompt", _L)):
        new_gsp = st.text_area(t("gsp_help", _L),
                               value=st.session_state["global_system_prompt"],
                               height=90, key="gsp_input")
        if new_gsp != st.session_state["global_system_prompt"]:
            st.session_state["global_system_prompt"] = new_gsp

    # ── Prompt profiles ───────────────────────────────
    with st.expander(t("prompt_profiles", _L)):
        st.caption(t("profiles_help", _L))
        pp_name = st.text_input(t("profile_name", _L), key="pp_name_sb",
                                placeholder=t("profile_placeholder", _L))
        if st.button(t("save_profile", _L), key="pp_save_sb") and pp_name.strip():
            profiles = load_prompt_profiles()
            profiles[pp_name] = {
                "global_system_prompt": st.session_state["global_system_prompt"],
                "prompt_templates":     st.session_state["prompt_templates"],
                "glossaries":           st.session_state["glossaries"],
                "saved_at":             datetime.now().isoformat(),
            }
            save_prompt_profiles(profiles)
            st.session_state["prompt_profiles"] = profiles
            st.success(t("profile_saved", _L, n=pp_name))
        existing = list(load_prompt_profiles().keys())
        if existing:
            pp_load = st.selectbox(t("load_profile", _L), ["—"] + existing, key="pp_load_sb")
            if st.button(t("load_btn", _L), key="pp_load_btn") and pp_load != "—":
                profiles = load_prompt_profiles()
                p = profiles[pp_load]
                st.session_state["global_system_prompt"] = p.get("global_system_prompt",
                    st.session_state["global_system_prompt"])
                st.session_state["prompt_templates"].update(p.get("prompt_templates", {}))
                st.session_state["glossaries"].update(p.get("glossaries", {}))
                save_templates(st.session_state["prompt_templates"])
                save_glossaries(st.session_state["glossaries"])
                st.success(t("profile_loaded", _L, n=pp_load)); st.rerun()

    st.divider()

    # ── Session persistence ────────────────────────────
    st.markdown(f"**{t('sb_session_persist', _L)}**")
    _sp_info = session_persist_info()
    if _sp_info:
        _sp_dt = _sp_info[:16].replace("T"," ")
        st.caption(t("sb_last_backup", _L, d=_sp_dt))
    sp_col1, sp_col2 = st.columns(2)
    with sp_col1:
        if st.button(t("sb_save", _L),
                     key="sb_save_session", width='stretch',
                     help=t("sb_save_help", _L)):
            ok1 = save_session_to_disk()
            ok2 = save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
            if ok1:
                st.toast(t("sb_session_saved", _L))
            else:
                st.toast(t("sb_save_error", _L))
    with sp_col2:
        if st.button(t("sb_load", _L),
                     key="sb_load_session", width='stretch',
                     help=t("sb_load_help", _L)):
            loaded = load_session_from_disk()
            n = len([k for k in loaded if not k.startswith("_")])
            saved_at = loaded.get("_saved_at","")[:16]
            if n > 0:
                # ── v24.31: Synchronizace aliasů po obnově ───────────────────
                if st.session_state.get("translation_result"):
                    st.session_state["translation_ready"] = True
                _ext_txt = (st.session_state.get("last_extracted_text") or
                            st.session_state.get("last_extraction_text") or "")
                if _ext_txt:
                    st.session_state["last_extracted_text"] = _ext_txt
                    st.session_state["last_extraction_text"] = _ext_txt
                # Zvolit záložku podle obnovených dat
                if st.session_state.get("translation_ready") and st.session_state.get("translation_result"):
                    st.session_state["_gsb_jump_to_tab"] = 1
                elif st.session_state.get("last_validation_results"):
                    st.session_state["_gsb_jump_to_tab"] = 2
                elif st.session_state.get("last_extracted_text") or st.session_state.get("last_extracted_df_json"):
                    st.session_state["_gsb_jump_to_tab"] = 0
                st.session_state["_restore_success_msg"] = (
                    f"✅ Session obnovena z {saved_at} ({n} klíčů)" if st.session_state.get("lang","cz")=="cz"
                    else f"✅ Session restored from {saved_at} ({n} keys)"
                )
                st.session_state["_show_restore_toast"] = True
                st.rerun()
            else:
                st.toast(t("sb_no_session", _L))

    # ── Autosave status ───────────────────────────────
    _as_info = session_persist_info(SESSION_AUTOSAVE_FILE)
    if _as_info:
        try:
            from datetime import datetime as _dta
            _as_dt  = _dta.fromisoformat(_as_info)
            _as_age = int((datetime.now() - _as_dt).total_seconds() / 60)
            if _as_age < 120:
                st.caption(t("sb_autosave_ago", _L, n=_as_age))
            else:
                st.caption(t("sb_autosave_time", _L, t=_as_info[11:16]))
        except Exception:
            st.caption(t("sb_autosave_time", _L, t=_as_info[:16]))
    else:
        st.caption(t("sb_autosave_none", _L))

    st.divider()

    # ── Kompaktní stats grid (cachováno 5s — šetří SQLite volání) ──
    _dc, _fts, _ck = _cached_sidebar_stats()
    _mem = len(st.session_state["validation_cache"])
    _gl  = len(st.session_state["glossaries"])
    _tp  = len(st.session_state["prompt_templates"])
    _lbl_vc   = t("sb_stat_valcache",    _L)
    _lbl_gl   = t("sb_stat_glossaries",  _L)
    _lbl_tp   = t("sb_stat_templates",   _L)
    _lbl_fts  = t("sb_stat_fts",         _L)
    _lbl_ck   = t("sb_stat_checkpoints", _L)
    st.markdown(
        f'<div style="display:grid;grid-template-columns:1fr 1fr;gap:3px 10px;'
        f'font-size:.73rem;color:#64748b;padding:2px 0">'
        f'<span>{_lbl_vc}</span><span style="text-align:right;font-weight:600">{_mem} / {_dc}</span>'
        f'<span>{_lbl_gl}</span><span style="text-align:right;font-weight:600">{_gl}</span>'
        f'<span>{_lbl_tp}</span><span style="text-align:right;font-weight:600">{_tp}</span>'
        f'<span>{_lbl_fts}</span><span style="text-align:right;font-weight:600">{_fts:,}</span>'
        f'{"<span>" + _lbl_ck + "</span><span style=text-align:right;font-weight:600>" + str(_ck) + "</span>" if _ck else ""}'
        f'</div>',
        unsafe_allow_html=True
    )
    st.divider()

    # ── HW monitoring: CPU / RAM / GPU ───────────────
    _hw = _cached_hw_stats()
    _cpu = _hw["cpu_pct"]
    _ram = _hw["ram_pct"]
    _gpu = _hw["gpu_pct"]
    _vram= _hw["vram_pct"]

    _hw_rows = []

    # CPU
    _cpu_bar = _hw_bar_html(_cpu, _bar_color(_cpu))
    _hw_rows.append(
        f'<div style="font-size:.72rem;color:#64748b;display:flex;justify-content:space-between">'
        f'<span>CPU</span>'
        f'<span style="font-weight:600;color:{_bar_color(_cpu)}">{_cpu:.0f} %</span></div>'
        + _cpu_bar
    )

    # RAM
    _ram_lbl = (f'{_hw["ram_used_mb"]:,} / {_hw["ram_total_mb"]:,} MB'
                if _hw["ram_total_mb"] else f'{_ram:.0f} %')
    _ram_bar = _hw_bar_html(_ram, _bar_color(_ram))
    _hw_rows.append(
        f'<div style="font-size:.72rem;color:#64748b;display:flex;justify-content:space-between">'
        f'<span>RAM</span>'
        f'<span style="font-weight:600;color:{_bar_color(_ram)}">{_ram_lbl}</span></div>'
        + _ram_bar
    )

    # GPU (jen pokud dostupné)
    if _gpu is not None:
        _gpu_name_short = (_hw.get("gpu_name") or "GPU")
        # Zkrátit název: "NVIDIA GeForce RTX 4090" → "RTX 4090"
        import re as _re_hw
        _gpu_name_short = _re_hw.sub(r"NVIDIA (GeForce )?", "", _gpu_name_short).strip()
        _temp_str = (f" · {_hw['gpu_temp_c']}°C" if _hw.get("gpu_temp_c") is not None else "")
        _gpu_bar  = _hw_bar_html(_gpu, _bar_color(_gpu))
        _hw_rows.append(
            f'<div style="font-size:.72rem;color:#64748b;display:flex;justify-content:space-between">'
            f'<span>GPU <span style="opacity:.7;font-size:.68rem">{_gpu_name_short}{_temp_str}</span></span>'
            f'<span style="font-weight:600;color:{_bar_color(_gpu)}">{_gpu:.0f} %</span></div>'
            + _gpu_bar
        )

    # VRAM
    if _vram is not None:
        _vram_lbl = (f'{_hw["vram_used_mb"]:,} / {_hw["vram_total_mb"]:,} MB'
                     if _hw.get("vram_total_mb") else f'{_vram:.0f} %')
        _vram_bar = _hw_bar_html(_vram, _bar_color(_vram), height=5)
        _hw_rows.append(
            f'<div style="font-size:.72rem;color:#64748b;display:flex;justify-content:space-between">'
            f'<span>VRAM</span>'
            f'<span style="font-weight:600;color:{_bar_color(_vram)}">{_vram_lbl}</span></div>'
            + _vram_bar
        )

    st.markdown(
        '<div style="padding:4px 0">' + "".join(_hw_rows) + '</div>',
        unsafe_allow_html=True
    )

    st.divider()
    st.markdown(
        f'<p class="sb-hint">{t("sb_hint", _L)}</p>',
        unsafe_allow_html=True
    )

    # ── Presence widget ────────────────────────────────
    presence.sidebar_widget()

    # ── Aktuální operace v sidebaru ───────────────────
    if st.session_state.get("gsb_active"):
        _gsb_lbl  = st.session_state.get("gsb_label", "")
        _gsb_det  = st.session_state.get("gsb_detail", "")
        _gsb_prog = st.session_state.get("gsb_progress", 0.0)
        st.divider()
        st.markdown(
            f'<div style="font-size:.78rem;color:#f59e0b;font-weight:600">⚙️ {_gsb_lbl}</div>'
            + (f'<div style="font-size:.72rem;color:#94a3b8">{_gsb_det}</div>' if _gsb_det else ""),
            unsafe_allow_html=True
        )
        st.progress(_gsb_prog)

    # ── OFF tlačítko – ukončí aplikaci (v24.18: zavře i záložku) ──
    st.divider()
    # ── v24.24: Aktivní úlohy (task queue) ───────────────────────
    render_task_queue()
    if st.button("⛔ OFF – zavřít záložku + ukončit aplikaci", key="kill_app",
                 help="Zavře záložku v prohlížeči a ukončí Streamlit server",
                 type="secondary"):
        st.warning(tt("Ukončuji aplikaci…", "Shutting down…",
                      st.session_state.get("lang", "cz")))
        shutdown_app()


# ══════════════════════════════════════════════════════
# KEYBOARD SHORTCUTS + PWA MANIFEST + KEYBOARD NAVIGATION
# ══════════════════════════════════════════════════════
st.html("""
<script>
(function() {
  // ── Ctrl+Enter → odeslání ──────────────────────────
  function addCtrlEnterShortcut() {
    document.querySelectorAll('textarea').forEach(function(ta) {
      if (ta.dataset.ctrlEnterBound) return;
      ta.dataset.ctrlEnterBound = "1";
      ta.addEventListener('keydown', function(e) {
        if ((e.ctrlKey || e.metaKey) && e.key === 'Enter') {
          e.preventDefault();
          var form = ta.closest('[data-testid="stForm"]') || ta.closest('.stVerticalBlock');
          if (form) {
            var btns = form.querySelectorAll('button[kind="primaryFormSubmit"], button.stButton');
            if (btns.length > 0) { btns[0].click(); return; }
          }
          var primary = document.querySelector('button[data-testid="baseButton-primary"]');
          if (primary) primary.click();
        }
      });
    });
  }

  // ── Keyboard navigation: accesskeys ───────────────
  // Alt+E = Extrakce, Alt+T = Překlad, Alt+V = Validace,
  // Alt+C = Chat, Alt+W = Workflow
  function addAccessKeys() {
    var tabs = document.querySelectorAll('[data-testid="stTabs"] button[role="tab"]');
    var keys = ['e','t','v','c','d','s','r','w','h'];
    tabs.forEach(function(tab, i) {
      if (i < keys.length && !tab.dataset.accesskeySet) {
        tab.dataset.accesskeySet = "1";
        tab.setAttribute('accesskey', keys[i]);
        tab.title = (tab.title || tab.textContent) + ' (Alt+' + keys[i].toUpperCase() + ')';
      }
    });
  }

  // ── Tabindex pro validační checkboxy ──────────────
  function addValidationTabindex() {
    var checkboxes = document.querySelectorAll('[data-testid="stCheckbox"] input[type="checkbox"]');
    checkboxes.forEach(function(cb, i) {
      if (!cb.dataset.tabindexSet) {
        cb.dataset.tabindexSet = "1";
        cb.setAttribute('tabindex', String(i + 1));
      }
    });
  }

  addCtrlEnterShortcut();
  addAccessKeys();
  addValidationTabindex();
  var obs = new MutationObserver(function() {
    addCtrlEnterShortcut();
    addAccessKeys();
    addValidationTabindex();
  });
  obs.observe(document.body, { childList: true, subtree: true });

  // ── Desktop notifikace API ─────────────────────────
  window._lmuNotify = function(title, body) {
    if (!("Notification" in window)) return;
    if (Notification.permission === "granted") {
      new Notification(title, {body: body, icon: "https://www.svgrepo.com/show/530459/brain.svg"});
    } else if (Notification.permission !== "denied") {
      Notification.requestPermission().then(function(p) {
        if (p === "granted") {
          new Notification(title, {body: body, icon: "https://www.svgrepo.com/show/530459/brain.svg"});
        }
      });
    }
  };
  // Požádej o povolení při prvním načtení (tiše)
  if ("Notification" in window && Notification.permission === "default") {
    Notification.requestPermission();
  }

  // ── PWA manifest (Progressive Web App) ────────────
  if (!document.querySelector('link[rel="manifest"]')) {
    var manifest = {
      "name": "SciNexus",
      "short_name": "SciNexus",
      "description": "AI systém s lokálními LLM",
      "start_url": "/",
      "display": "standalone",
      "background_color": "#0e1117",
      "theme_color": "#4f46e5",
      "lang": "cs",
      "icons": [
        {"src": "https://www.svgrepo.com/show/530459/brain.svg",
         "sizes": "192x192", "type": "image/svg+xml"},
        {"src": "https://www.svgrepo.com/show/530459/brain.svg",
         "sizes": "512x512", "type": "image/svg+xml"}
      ]
    };
    var blob = new Blob([JSON.stringify(manifest)], {type: 'application/json'});
    var url  = URL.createObjectURL(blob);
    var link = document.createElement('link');
    link.rel  = 'manifest';
    link.href = url;
    document.head.appendChild(link);

    // Theme color meta tag
    if (!document.querySelector('meta[name="theme-color"]')) {
      var meta = document.createElement('meta');
      meta.name    = 'theme-color';
      meta.content = '#4f46e5';
      document.head.appendChild(meta);
    }

    // Apple touch icon meta
    if (!document.querySelector('meta[name="apple-mobile-web-app-capable"]')) {
      var apMeta = document.createElement('meta');
      apMeta.name    = 'apple-mobile-web-app-capable';
      apMeta.content = 'yes';
      document.head.appendChild(apMeta);
    }
  }
})();
</script>
""")

# ══════════════════════════════════════════════════════
# HLAVNÍ ZÁLOŽKY
# ══════════════════════════════════════════════════════
_L = st.session_state.get("lang", "cz")

# ── Globální status bar — zobrazí se nad záložkami ───
gsb_render()

if st.session_state.pop("_show_restore_toast", False):
    _restore_msg = st.session_state.pop("_restore_success_msg", None)
    if _restore_msg:
        st.success(_restore_msg, icon="💾")
    else:
        st.success(tt("✅ Relace byla úspěšně obnovena.", "✅ Session successfully restored.", _L), icon="💾")
    st.toast("💾 " + tt("Relace obnovena", "Session restored", _L))

main_tabs = st.tabs([
    t("tab_extract",  _L),      # 0
    t("tab_translate", _L),     # 1
    t("tab_validate", _L),      # 2
    t("tab_chat",     _L),      # 3
    t("tab_clean",    _L),      # 4
    t("tab_style",    _L),      # 5
    t("tab_workflow", _L),      # 6
    t("tab_history",  _L),      # 7
    t("tab_abtest",   _L),      # 8
    t("tab_settings", _L),      # 9
    t("tab_deepseek", _L),      # 10
    t("tab_help",     _L),      # 11
])

# ── Presence: tracking aktivní záložky ───────────────────────────────────────
_tab_param = st.query_params.get("tab", "0")
try:
    _tab_idx = int(_tab_param)
except ValueError:
    _tab_idx = 0
st.session_state["presence_tab_idx"] = _tab_idx
# Synchronizuj current_tab_index s URL — zajistí správný fallback i při form-submit rerunu
if not st.session_state.get("gsb_active", False):
    st.session_state["current_tab_index"] = _tab_idx
elif "current_tab_index" not in st.session_state:
    st.session_state["current_tab_index"] = _tab_idx

# ── JS: zachování záložky + autoclick při cílené navigaci ────────────────────
_jump_target = st.session_state.pop("_gsb_jump_to_tab", None)
# Fallback bez JS sandboxu: current_tab_index je synchronizován s URL výše,
# JS vždy dostane platné číslo — nikdy nečte window.parent (sandbox selže).
if _jump_target is None:
    _jump_target = st.session_state.get("current_tab_index", _tab_idx)
_js_tab_target = _jump_target

st.html(f"""
<script>
(function(){{
  var TARGET_TAB = {_js_tab_target};

  function getDoc(){{
    try {{ return window.parent.document; }} catch(e) {{}}
    try {{ return window.top.document; }} catch(e) {{}}
    return document;
  }}
  function getWin(){{
    try {{ return window.parent; }} catch(e) {{}}
    try {{ return window.top; }} catch(e) {{}}
    return window;
  }}

  function getTabs(){{
    return getDoc().querySelectorAll('[data-baseweb="tab"]');
  }}

  function watchTabs(){{
    var tabs = getTabs();
    tabs.forEach(function(tab, idx){{
      if (tab.dataset.lmuWatched) return;
      tab.dataset.lmuWatched = "1";
      tab.addEventListener('click', function(){{
        var w = getWin();
        try {{
          var url = new URL(w.location.href);
          url.searchParams.set('tab', idx);
          w.history.replaceState({{}}, '', url);
        }} catch(e) {{}}
        try {{ w.sessionStorage.setItem('lmu_active_tab', idx); }} catch(e) {{}}
      }});
    }});
  }}

  function clickTab(attempt){{
    attempt = attempt || 0;
    if (TARGET_TAB < 0) return;
    if (attempt > 25) return;
    var tabs = getTabs();
    if (tabs.length > TARGET_TAB) {{
      var alreadyActive = tabs[TARGET_TAB].getAttribute('aria-selected') === 'true';
      if (!alreadyActive) {{
        tabs[TARGET_TAB].click();
      }}
      // ověření za 50ms
      setTimeout(function(){{
        var tabs2 = getTabs();
        if (tabs2.length > TARGET_TAB && tabs2[TARGET_TAB].getAttribute('aria-selected') !== 'true') {{
          clickTab(attempt + 1);
        }}
      }}, 50);
    }} else {{
      setTimeout(function(){{ clickTab(attempt + 1); }}, 60);
    }}
  }}

  function addKeyboardShortcuts(){{
    var w = getWin();
    if (w._lmuKbBound) return;
    w._lmuKbBound = true;
    getDoc().addEventListener('keydown', function(e){{
      if (!e.altKey && !(e.ctrlKey && e.shiftKey)) return;
      var n = parseInt(e.key);
      if (isNaN(n) || n < 1 || n > 9) return;
      var tabs = getTabs();
      var idx = n - 1;
      if (tabs.length > idx) {{
        e.preventDefault();
        tabs[idx].click();
      }}
    }});
  }}

  function init(){{
    watchTabs();
    addKeyboardShortcuts();
    if (TARGET_TAB >= 0) {{
      // Minimální zpoždění — 0ms (asynchronní, ale okamžité)
      setTimeout(function(){{ clickTab(0); }}, 0);
    }}
  }}

  if (document.readyState === 'loading') {{
    document.addEventListener('DOMContentLoaded', init);
  }} else {{
    setTimeout(init, 0);
  }}

  try {{
    var obs = new MutationObserver(function(){{ watchTabs(); addKeyboardShortcuts(); }});
    obs.observe(getDoc().body, {{ childList: true, subtree: true }});
  }} catch(e) {{
    var obs2 = new MutationObserver(function(){{ watchTabs(); addKeyboardShortcuts(); }});
    obs2.observe(document.body, {{ childList: true, subtree: true }});
  }}
}})();
</script>
""")


# ══════════════════════════════════════════════════════
# KONFIDENCE PER-POLE – vizuální označení nejistých polí
# ══════════════════════════════════════════════════════
def inject_confidence_prompt(base_prompt: str) -> str:
    """Přidá instrukci pro _confidence pole do extrakčního promptu."""
    conf_instruction = (
        '\n\nKaždý objekt v JSON poli musí obsahovat pole "_confidence" jako slovník '
        's hodnotami 0.0–1.0 pro každé extrahované pole. Příklad: '
        '"_confidence": {"druh": 0.95, "stratigrafie": 0.4, "lokalita": 0.7}. '
        'Hodnota 1.0 = absolutní jistota (přímo v textu), 0.5 = odvozeno, 0.0 = odhadnuto.'
    )
    return base_prompt + conf_instruction


def render_confidence_table(records: List[Dict]) -> str:
    """Vytvoří HTML tabulku s barevně označenými poli podle konfidence."""
    if not records:
        return ""
    # Zjisti sloupce (bez _confidence)
    all_keys = []
    for rec in records:
        for k in rec.keys():
            if k != "_confidence" and k not in all_keys:
                all_keys.append(k)

    def conf_color(v: float) -> str:
        if v >= 0.8: return "background:rgba(0,180,0,0.15)"
        if v >= 0.5: return "background:rgba(255,200,0,0.18)"
        return "background:rgba(255,60,60,0.18)"

    header = "".join(f'<th style="padding:4px 8px;border:1px solid #333;font-size:.8rem">{k}</th>'
                     for k in all_keys)
    rows = ""
    for rec in records[:50]:
        conf = rec.get("_confidence", {})
        cells = ""
        for k in all_keys:
            v = rec.get(k, "")
            c = float(conf.get(k, 1.0)) if isinstance(conf, dict) else 1.0
            style = conf_color(c)
            title = f'konfidence: {c:.2f}' if isinstance(conf, dict) and k in conf else ''
            cells += (f'<td style="padding:3px 6px;border:1px solid #333;font-size:.78rem;'
                      f'{style}" title="{title}">{str(v)[:60]}</td>')
        rows += f"<tr>{cells}</tr>"

    return (f'<div style="overflow-x:auto"><table style="border-collapse:collapse;'
            f'font-size:.82rem;width:100%"><thead><tr>{header}</tr></thead>'
            f'<tbody>{rows}</tbody></table>'
            f'<p style="font-size:.75rem;color:#888;margin-top:4px">'
            f'🟢 vysoká konfidence (&gt;0.8) &nbsp; 🟡 střední (0.5–0.8) &nbsp; 🔴 nízká (&lt;0.5)</p></div>')


# ══════════════════════════════════════════════════════
# POROVNÁNÍ EXTRAKCÍ – diff nových vs. starých záznamů
# ══════════════════════════════════════════════════════
def diff_extraction_records(old_json: str, new_json: str) -> Dict:
    """Porovná dva JSON výsledky extrakce. Vrátí přidané/odebrané/změněné záznamy."""
    def parse(j: str) -> List[Dict]:
        try:
            clean = _JSON_FENCE_RE.sub("", j).strip()
            parsed = json.loads(clean)
            return parsed if isinstance(parsed, list) else [parsed]
        except Exception:
            return []

    old_recs = parse(old_json)
    new_recs = parse(new_json)

    def rec_key(r: Dict) -> str:
        return "|".join(str(r.get(k,"")) for k in ("druh","species","taxon","rod","genus"))

    old_keys = {rec_key(r): r for r in old_recs}
    new_keys = {rec_key(r): r for r in new_recs}

    added   = [new_keys[k] for k in new_keys if k not in old_keys]
    removed = [old_keys[k] for k in old_keys if k not in new_keys]
    common  = [k for k in new_keys if k in old_keys]
    changed = []
    for k in common:
        o, n = old_keys[k], new_keys[k]
        diffs = {fld: (o.get(fld,""), n.get(fld,""))
                 for fld in set(list(o.keys())+list(n.keys()))
                 if fld != "_confidence" and str(o.get(fld,"")) != str(n.get(fld,""))}
        if diffs:
            changed.append({"key": k, "diffs": diffs})

    return {"added": added, "removed": removed, "changed": changed,
            "total_old": len(old_recs), "total_new": len(new_recs)}


def _merge_json_chunks(chunks: List[str]) -> str:
    """Sloučí více JSON odpovědí (z chunked extrakce) do jednoho JSON pole.
    Pokud chunky nejsou validní JSON, spojí je prostým oddělovačem."""
    all_records: list = []
    for chunk in chunks:
        try:
            clean  = _JSON_FENCE_RE.sub("", chunk).strip()
            parsed = json.loads(clean)
            if isinstance(parsed, list):
                all_records.extend(parsed)
            elif isinstance(parsed, dict):
                all_records.append(parsed)
            else:
                return "\n\n---\n\n".join(chunks)
        except Exception:
            return "\n\n---\n\n".join(chunks)
    return json.dumps(all_records, ensure_ascii=False, indent=2)


def _parse_iter_json(iter_results: List[str]) -> List[List[dict]]:
    """Parsuje seznam JSON stringů iterací na seznam listů záznamů."""
    parsed = []
    for r in iter_results:
        try:
            clean = _JSON_FENCE_RE.sub("", r).strip()
            data  = json.loads(clean)
            if isinstance(data, list):
                parsed.append(data)
            elif isinstance(data, dict):
                parsed.append([data])
            else:
                parsed.append([])
        except Exception:
            parsed.append([])
    return parsed


def _merge_union_dedup(iter_results: List[str]) -> str:
    """Union merge: sloučí záznamy ze všech iterací, deduplikuje podle klíčových polí.
    Pro kolizní textová pole zachová nejdelší/nejúplnější hodnotu."""
    parsed = _parse_iter_json(iter_results)
    all_recs = [rec for group in parsed for rec in group]
    if not all_recs:
        return _merge_json_chunks(iter_results)

    KEY_FIELDS = ("druh", "rod", "autor", "rok")

    def _rec_key(rec: dict) -> tuple:
        return tuple(str(rec.get(f, "") or "").strip().lower() for f in KEY_FIELDS)

    def _merge_two(base: dict, other: dict) -> dict:
        """Sloučí dva záznamy — vybere delší/neprázdnou hodnotu pro každé pole."""
        merged = dict(base)
        for k, v in other.items():
            bv = merged.get(k)
            if not bv and v:
                merged[k] = v
            elif bv and v and isinstance(bv, str) and isinstance(v, str):
                merged[k] = bv if len(bv) >= len(v) else v
        return merged

    seen: Dict[tuple, dict] = {}
    for rec in all_recs:
        key = _rec_key(rec)
        if key in seen:
            seen[key] = _merge_two(seen[key], rec)
        else:
            seen[key] = rec

    return json.dumps(list(seen.values()), ensure_ascii=False, indent=2)


def _merge_consensus(iter_results: List[str]) -> str:
    """Consensus merge: pro každé pole vybere hodnotu která se vyskytuje nejčastěji.
    Pro textová pole (popis, lokalita) preferuje nejdelší variantu z majority."""
    parsed = _parse_iter_json(iter_results)
    if not parsed or all(len(g) == 0 for g in parsed):
        return _merge_json_chunks(iter_results)

    # Seskup záznamy napříč iteracemi podle indexu (předpoklad: stejné pořadí)
    max_recs = max(len(g) for g in parsed)
    result = []

    for idx in range(max_recs):
        candidates = [g[idx] for g in parsed if idx < len(g)]
        if not candidates:
            continue
        if len(candidates) == 1:
            result.append(candidates[0])
            continue

        # Všechna pole ze všech kandidátů
        all_keys = set(k for c in candidates for k in c.keys())
        merged = {}
        for field in all_keys:
            values = [str(c.get(field, "") or "").strip() for c in candidates]
            non_empty = [v for v in values if v and v.lower() not in ("null","none","n/a","")]
            if not non_empty:
                merged[field] = None
                continue
            # Voting: vyber nejčastější hodnotu
            counts = Counter(non_empty)
            top_val, top_cnt = counts.most_common(1)[0]
            # Pokud je shoda jasná (>50 % iterací), použij ji
            if top_cnt > len(candidates) / 2:
                merged[field] = top_val
            else:
                # Jinak vyber nejdelší hodnotu (pro textová pole)
                merged[field] = max(non_empty, key=len)
        result.append(merged)

    return json.dumps(result, ensure_ascii=False, indent=2)


def _merge_llm_judge(iter_results: List[str], base_url: str, model: str,
                      src_text_preview: str = "", max_tokens: int = 8000) -> str:
    """LLM Judge: pošle všechny iterace modelu, který vybere/sloučí nejlepší výsledek."""
    variants = "\n\n".join(
        f"=== ITERACE {i+1} ===\n{r}" for i, r in enumerate(iter_results))
    system_judge = (
        "Jsi expert na extrakci taxonomických dat z paleontologické literatury. "
        "Dostaneš několik pokusů o extrakci dat ze stejného textu jako JSON. "
        "Tvůj úkol:\n"
        "1. Zachovej záznamy které se opakují ve více iteracích (vyšší spolehlivost).\n"
        "2. Doplň záznamy nebo pole která jsou v některých iteracích, ale chybí v jiných.\n"
        "3. Pro kolizní hodnoty polí vyber tu nejpravděpodobnější (nejpřesnější, nejkonzistentnější).\n"
        "4. Odstraň zjevné LLM artefakty (null, None, ?, neznámý).\n"
        "Výstup je POUZE čistý JSON seznam záznamů, bez komentářů, bez markdown."
    )
    user_msg = ""
    if src_text_preview:
        user_msg += f"PŮVODNÍ TEXT (výňatek pro referenci):\n{src_text_preview[:2000]}\n\n"
    user_msg += f"ITERACE K SLOUČENÍ:\n{variants}"

    try:
        result = chat_completion_queued(
            base_url, model,
            [{"role": "system", "content": system_judge},
             {"role": "user",   "content": user_msg}],
            temp=0.05, max_tokens=max_tokens)
        # Ověř že výsledek je JSON
        clean = _JSON_FENCE_RE.sub("", result).strip()
        json.loads(clean)  # raises if invalid
        return clean
    except Exception:
        # Fallback na union merge
        return _merge_union_dedup(iter_results)


def _merge_llm_judge_translation(iter_results: List[str], base_url: str, model: str,
                                   src_text: str, tgt_lang: str) -> str:
    """LLM Judge pro překlad: diff-based sloučení s LLM rozhodnutím u divergentních míst."""
    n = len(iter_results)
    variants_text = "\n\n".join(
        f"=== Varianta {i+1} ===\n{v}" for i, v in enumerate(iter_results))

    merge_sys = (
        f"Jsi expert na překlady vědeckých textů do jazyka: {tgt_lang}. "
        f"Dostaneš {n} variant překladu stejného textu. "
        "Postup:\n"
        "1. Porovnej varianty odstavec po odstavci.\n"
        "2. Kde se varianty shodují, použij společnou formulaci.\n"
        "3. Kde se liší (divergují), vyber nejpřesnější, nejplynulejší a vědecky nejsprávnější verzi.\n"
        "4. Zachovej latinské taxonomické názvy beze změny.\n"
        "Výstup je POUZE výsledný překlad, bez komentářů, bez označení variant."
    )
    user_msg = f"ORIGINÁL (výňatek):\n{src_text[:2000]}\n\n{variants_text}"

    try:
        return chat_completion_queued(
            base_url, model,
            [{"role": "system", "content": merge_sys},
             {"role": "user",   "content": user_msg}],
            temp=0.1, max_tokens=_MAX_TOKENS_TRANSLATE)
    except Exception:
        return iter_results[0]


def _merge_consensus_translation(iter_results: List[str]) -> str:
    """Consensus merge pro překlad: porovná odstavce napříč iteracemi,
    pro každý odstavec vybere verzi která se nejčastěji opakuje (nebo nejdelší při shodě).
    Funguje bez LLM — čistě lokálně."""
    from collections import Counter

    # Rozděl každou iteraci na odstavce
    split_iters = [r.strip().split("\n\n") for r in iter_results]
    max_paras   = max(len(s) for s in split_iters)
    result_paras = []

    for idx in range(max_paras):
        candidates = [s[idx].strip() for s in split_iters if idx < len(s) and s[idx].strip()]
        if not candidates:
            continue
        if len(candidates) == 1:
            result_paras.append(candidates[0])
            continue
        # Voting: vyber nejčastější odstavec (přesná shoda)
        counts = Counter(candidates)
        top_para, top_cnt = counts.most_common(1)[0]
        if top_cnt > len(candidates) / 2:
            result_paras.append(top_para)
        else:
            # Při nerozhodném hlasování vyber nejdelší (nejúplnější) variantu
            result_paras.append(max(candidates, key=len))

    return "\n\n".join(result_paras) if result_paras else iter_results[0]


def _make_iterations_zip(iter_results: List[str], prefix: str = "iterace") -> bytes:
    """Zabalí všechny iterace do ZIP souboru (každá jako .txt nebo .json)."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, r in enumerate(iter_results):
            # Detekuj jestli je to JSON
            is_json = False
            try:
                clean = _JSON_FENCE_RE.sub("", r).strip()
                json.loads(clean)
                is_json = True
            except Exception:
                pass
            ext  = ".json" if is_json else ".txt"
            name = f"{prefix}_{i+1:02d}{ext}"
            zf.writestr(name, r.encode("utf-8") if isinstance(r, str) else r)
    return buf.getvalue()

# ══════════════════════════════════════════════════════
# 0 – EXTRAKCE  (auto-chunking + prompt šablony)
# ══════════════════════════════════════════════════════
with main_tabs[0]:
    keep_current_tab(0)   # ← v24.17: zachování záložky Extrakce
    _L = st.session_state.get("lang","cz")
    st.markdown(t("ext_title", _L))
    with st.expander(tt("ℹ️ Jak funguje extrakce?",
                         "ℹ️ How does extraction work?", _L), expanded=False):
        st.markdown(tt(
            """
**Extrakce** načte PDF/DOCX/TXT soubor, rozdělí ho na bloky a pomocí LLM vytáhne strukturovaná data (taxonomické záznamy, stratigrafii, lokality atd.) jako JSON.

**Postup:**
1. Nahraj jeden nebo více souborů
2. Vyber **Režim extrakce** (Standard = doporučeno pro hyolity)
3. Klikni **▶️ Extrahovat** — výsledky se zobrazují průběžně soubor po souboru
4. Stáhni výsledky jako CSV, JSON nebo XLSX

**Tipy:**
- **Two-pass extrakce** — druhý průchod doplní pole která první přeskočil; doporučeno pro složité tabulky
- **Konfidence** — přidá skóre jistoty ke každému poli (zelená ≥ 0.8, žlutá ≥ 0.5, červená < 0.5)
- **Chunk size** — pro Qwen2.5-72B nastav 20 000+; pro Mistral 8 000
- **Prompt Lab** — otestuj prompt na 500 znacích před plnou extrakcí
- **Paralelní chunky** — nastav ⚡ v sidebaru na 4 pro 4× rychlejší zpracování
            """,
            """
**Extraction** reads a PDF/DOCX/TXT file, splits it into chunks, and uses the LLM to pull structured data (taxonomic records, stratigraphy, localities, etc.) as JSON.

**How to:**
1. Upload one or more files
2. Pick an **Extraction mode** (Standard = recommended for hyoliths)
3. Click **▶️ Extract** — results appear progressively, file by file
4. Download results as CSV, JSON, or XLSX

**Tips:**
- **Two-pass extraction** — a second pass fills in fields the first skipped; recommended for dense tables
- **Confidence** — adds a certainty score to each field (green ≥ 0.8, yellow ≥ 0.5, red < 0.5)
- **Chunk size** — set 20 000+ for Qwen2.5-72B; 8 000 for Mistral
- **Prompt Lab** — test the prompt on 500 chars before a full run
- **Parallel chunks** — set ⚡ in the sidebar to 4 for 4× faster processing
            """,
            _L
        ))

    col_left, col_right = st.columns([3, 2])
    with col_left:
        files = st.file_uploader(t("ext_upload",_L),
                                 type=["pdf","docx","txt"], accept_multiple_files=True)
        # ── PDF preview po nahrání ──────────────────────
        if files:
            for pf in files[:3]:  # max 3 náhledy
                if pf.name.lower().endswith(".pdf"):
                    try:
                        from PyPDF2 import PdfReader as _PR
                        _rdr = _PR(io.BytesIO(pf.read()))
                        pf.seek(0)  # resetuj pozici pro další čtení
                        _preview_txt = ""
                        for _pg in _rdr.pages[:2]:
                            _preview_txt += _pg.extract_text() or ""
                            if len(_preview_txt) > 400:
                                break
                        if _preview_txt.strip():
                            with st.expander(tt("👁️ Náhled: ", "👁️ Preview: ", _L) + pf.name):
                                st.code(_preview_txt[:400] + ("…" if len(_preview_txt)>400 else ""),
                                        language=None)
                        else:
                            st.caption(tt(f"⚠️ {pf.name}: text nelze extrahovat (naskenované PDF?)",
                                          f"⚠️ {pf.name}: text not extractable (scanned PDF?)", _L))
                    except Exception:
                        pass
    with col_right:
        mode = st.selectbox(t("ext_mode",_L), list(EXTRACTION_MODES.keys()),
                            help="\n".join(f"**{k}**: {v['desc']}" for k,v in EXTRACTION_MODES.items()))
        st.caption(f"_{EXTRACTION_MODES[mode]['desc']}_")
        temp_ext = st.slider(t("ext_temp",_L), 0.0, 1.0, 0.1, 0.05, key="temp_ext",
                             help=(
                                 "**Teplota** řídí kreativitu modelu.\n\n"
                                 "- **0.0–0.1** = deterministický ✅ doporučeno pro extrakci\n"
                                 "- **0.2–0.4** = mírně variabilní\n"
                                 "- **0.5+** = kreativní — může vymýšlet data!"
                                 if _L=="cz" else
                                 "**Temperature** — model creativity.\n\n"
                                 "0.0–0.1 = deterministic ✅ recommended. 0.5+ may hallucinate."
                             ))
        max_tok_ext = st.number_input(t("ext_maxtok",_L), 500, 16000, _MAX_TOKENS_EXTRACT, 500, key="mtok_ext",
                                      help=(
                                          "Max tokenů výstupu na blok.\n\n"
                                          "2 000–4 000: rychlé / málo záznamů\n"
                                          "6 000: standard ✅\n"
                                          "8 000+: husté tabulky\n\n"
                                          "Příliš nízko = useknutý JSON."
                                          if _L=="cz" else
                                          "Max output tokens per block. Low = truncated JSON."
                                      ))
        use_two_pass = st.checkbox(
            "🔄 Dva průchody extrakce" if _L=="cz" else "🔄 Two-pass extraction",
            value=False, key="ext_two_pass",
            help=(
                "Spustí LLM **dvakrát** na každý dokument:\n\n"
                "1. První průchod — standardní extrakce\n"
                "2. Druhý průchod — dostane výsledek + originál, doplní chybějící pole\n\n"
                "✅ Doporučeno pro husté tabulky a historické texty\n"
                "⚠️ ~2× pomalejší · 📈 ~15–25 % lepší data"
                if _L=="cz" else
                "Runs LLM twice. 2nd pass fills missing fields. ~2× slower, ~15-25% better recall."
            ))
        use_confidence = st.checkbox(
            "📊 Skóre jistory pole" if _L=="cz" else "📊 Field confidence scores",
            value=False, key="ext_confidence",
            help=(
                "Přidá ke každému záznamu pole `_confidence` (0.0–1.0):\n\n"
                "🟢 ≥ 0.8 = přímo v textu · 🟡 ≥ 0.5 = odvozeno · 🔴 < 0.5 = odhadnuto\n\n"
                "Výsledky se zobrazí jako barevná tabulka. ⚠️ ~10 % pomalejší."
                if _L=="cz" else
                "Adds _confidence per field. 🟢≥0.8 certain, 🟡≥0.5 inferred, 🔴<0.5 guessed."
            ))
        use_json_schema = st.checkbox(
            "🗂️ Strukturované JSON schéma" if _L=="cz" else "🗂️ Enforce JSON schema",
            value=False, key="ext_json_schema",
            help=(
                "Přidá **explicitní JSON schema** do system promptu — výstup bude konzistentnější.\n\n"
                "✅ Doporučeno pro: Qwen2.5, DeepSeek\n"
                "⚠️ U starších modelů může způsobit odmítnutí extrakce"
                if _L=="cz" else
                "Adds JSON schema to system prompt. Best with Qwen2.5/DeepSeek."
            ))

        # ── Paralelní extrakce ─────────────────────────
        _ext_par_max = st.session_state.get("lms_max_concurrent", 1)
        _ext_par_default = _ext_par_max > 1
        use_ext_parallel = st.checkbox(
            "⚡ Paralelní extrakce bloků" if _L=="cz" else "⚡ Parallel chunk extraction",
            value=_ext_par_default,
            key="ext_parallel",
            help=(
                f"Odesílá všechny bloky **současně** jako souběžné LLM požadavky.\n\n"
                f"⚡ ~{_ext_par_max}× rychlejší pro dokumenty s ≥3 bloky\n"
                f"⚠️ Bez sdíleného kontextu mezi bloky\n"
                f"💡 Aktuální limit: **{_ext_par_max}** souběžných req. (nastav v sidebaru)\n\n"
                f"LM Studio → Developer → Max Concurrent Predictions = {_ext_par_max}"
                if _L=="cz" else
                f"Sends all chunks simultaneously. ~{_ext_par_max}× faster. "
                f"Set LM Studio Max Concurrent Predictions = {_ext_par_max}."
            ))
        if use_ext_parallel and _ext_par_max < 2:
            st.caption("⚠️ Pro paralelní extrakci nastav **Max souběžných požadavků ≥ 2** v sidebaru." if _L=="cz" else "⚠️ Set Max concurrent requests ≥ 2 in sidebar.")
        elif use_ext_parallel:
            st.caption(f"⚡ Aktivní: {_ext_par_max} souběžných req." if _L=="cz" else f"⚡ Active: {_ext_par_max} concurrent requests")


    # ── Prompt šablony ──────────────────────────────
    st.markdown("#### 📝 Prompt" if _L=="cz" else "#### 📝 Prompt")
    templates_all = st.session_state["prompt_templates"]
    tpl_names = [t("ext_own_prompt",_L)] + list(templates_all.keys())
    tpl_sel = st.selectbox(t("ext_template",_L), tpl_names, key="tpl_sel")

    if tpl_sel != t("ext_own_prompt",_L) and tpl_sel in templates_all:
        tpl_data = templates_all[tpl_sel]
        st.caption(f"_{tpl_data.get('desc','')}_ | {'doporučený jazyk' if _L=='cz' else 'recommended language'}: **{tpl_data.get('lang_hint','')}**")
        default_prompt = tpl_data.get("prompt","")
    else:
        default_prompt = ("Extrahuj taxonomické údaje: rod, druh, autor, rok, lokalita, stratigrafie. "
                          "Výstup jako JSON seznam objektů."
                          if _L=="cz" else
                          "Extract taxonomic data: genus, species, author, year, locality, stratigraphy. "
                          "Output as JSON array of objects.")

    prompt = st.text_area(t("ext_prompt_label",_L), value=default_prompt, height=120, key="ext_prompt")

    # ── Inline nápověda ──────────────────────────────
    with st.expander("ℹ️ Jak na to?" if _L=="cz" else "ℹ️ Quick guide"):
        st.markdown(
            "1. **Nahraj PDF/DOCX/TXT** soubory vlevo\n"
            "2. Vyber **režim extrakce** a nastav **teplotu** (0.0–0.1 doporučeno)\n"
            "3. Vlevo dole nastav **prompt** — buď vlastní nebo ze šablony\n"
            "4. Zapni **Dva průchody** pro lepší data u složitých tabulek\n"
            "5. Stiskni **Spustit extrakci** — výsledky se zobrazí průběžně"
            if _L=="cz" else
            "1. **Upload PDF/DOCX/TXT** files on the left\n"
            "2. Select **extraction mode** and set **temperature** (0.0–0.1 recommended)\n"
            "3. Configure **prompt** — use own or select from template\n"
            "4. Enable **Two-pass** for better recall on complex tables\n"
            "5. Click **Run extraction** — results appear progressively"
        )

    # ── PDF nastavení ────────────────────────────────
    with st.expander(t("ext_pdf_settings",_L)):
        page_spec = st.text_input(t("ext_pages",_L), "")
        use_ocr   = st.checkbox(t("ext_ocr",_L), value=False, disabled=not OCR_AVAILABLE)
        if use_ocr and OCR_AVAILABLE:
            ocr_dpi = st.slider(
                "OCR DPI" if _L=="cz" else "OCR DPI (resolution)", 150, 600, 300, 50, key="ext_ocr_dpi",
                help=tt("300 DPI = dobrý kompromis rychlost/kvalita. "
                        "Nižší = rychlejší ale horší OCR. Vyšší = pomalejší ale přesnější.",
                        "300 DPI = good speed/quality trade-off. "
                        "Lower = faster but worse OCR. Higher = slower but more accurate.", _L))
        else:
            ocr_dpi = 300

        ext_taxa_batch_size = st.slider(
            tt("Taxony na batch při extrakci", "Taxa per batch in extraction", _L),
            1, 5, 2, key="ext_taxa_batch_size"
        )

        if not OCR_AVAILABLE:
            st.caption("OCR: `pip install pytesseract pdf2image Pillow`")

        # Auto chunk size
        _auto_chunk = st.checkbox(
            "🤖 Automatická velikost bloku" if _L=="cz" else "🤖 Auto block size",
            value=True, key="ext_auto_chunk",
            help=(
                "Automaticky vypočítá optimální velikost bloku podle:\n\n"
                "- délky dokumentu\n"
                "- modelu (Qwen2.5-72B → ~24 000 znaků, Mistral-24B → ~10 000)\n\n"
                "Pro ruční nastavení odškrtni."
                if _L=="cz" else
                "Auto-calculates optimal chunk size based on document length and model."
            ))
        if _auto_chunk:
            chunk_size = None
            st.caption("Velikost bloku bude určena automaticky při zpracování." if _L=="cz" else "Chunk size will be calculated automatically.")
        else:
            chunk_size = st.slider(
                t("ext_chunk_size",_L), 4000, 30000, 12000, 1000,
                help=(
                    "Velikost jednoho bloku textu odeslaného LLM.\n\n"
                    "**Doporučené hodnoty:**\n"
                    "- Qwen2.5-72B: **20 000–30 000**\n"
                    "- Mistral-24B: **8 000–12 000**\n"
                    "- Llama-3-8B: **4 000–6 000**\n\n"
                    "Větší blok = méně volání, ale vyšší riziko přetečení kontextu."
                    if _L=="cz" else
                    "Qwen2.5-72B: 20k-30k · Mistral-24B: 8k-12k · Llama-3-8B: 4k-6k"
                ))

        # OCR preprocessing
        use_preprocess = st.checkbox(
            "🧹 OCR zpracování" if _L=="cz" else "🧹 Text preprocessing",
            value=True, key="ext_preprocess",
            help=(
                "Lokálně vyčistí text před odesláním do LLM:\n\n"
                "- Odstraní záhlaví stran (`--- Strana N ---`)\n"
                "- Odstraní izolovaná čísla stránek\n"
                "- Odstraní opakující se záhlaví tabulek (≥3× stejný řádek)\n\n"
                "✅ Snižuje zbytečné tokeny o **15–30 %** → rychlejší + levnější\n"
                "⚠️ Vypni pokud text vypadá špatně ořezaný."
                if _L=="cz" else
                "Removes page headers, page numbers, repeated table headers. "
                "Reduces tokens by 15-30%."
            ))
        if use_preprocess:
            pp_remove_bibref = st.checkbox(
                "Odstranit bibliografické reference" if _L=="cz" else "Remove bibliography references",
                value=False, key="ext_pp_bibref",
                help=(
                    "Odstraní citace ve formátu DOI:..., rok. Autor. Časopis.\n"
                    "Užitečné pro dokumenty s dlouhými bibliografiemi na konci.\n"
                    "⚠️ Může omylem odstranit validní data — používej opatrně."
                    if _L=="cz" else
                    "Removes DOI citations and bibliography entries. Use carefully."
                ))
        else:
            pp_remove_bibref = False

        st.caption("Auto-chunking: dlouhé dokumenty se automaticky rozdělí na bloky." if _L=="cz" else "Auto-chunking: long documents are split into blocks automatically.")

        # Overlap mezi bloky
        _ext_overlap = st.slider(
            "↔️ Překryv mezi bloky (znaky)" if _L=="cz" else "↔️ Block overlap (chars)",
            0, 800, 200, 100,
            key="ext_overlap",
            help=(
                "Přidá N znaků z konce předchozího bloku na začátek dalšího bloku.\n\n"
                "**Proč je to důležité:**\n"
                "Záznamy na hranici dvou bloků se mohou ztratit, protože věta začala\n"
                "v jednom bloku a skončila v dalším.\n\n"
                "**Doporučené hodnoty:**\n"
                "- **200–400**: tabulkový text, husté záznamy ✅\n"
                "- **0**: plynulá próza (překlad)\n"
                "- **400–600**: velmi husté tabulky na celou stránku"
                if _L=="cz" else
                "Adds N chars from previous block start. Prevents losing records at chunk boundaries.\n"
                "200-400: tabular text ✅ · 0: prose · 400-600: very dense tables"
            )
        )

        # ── v24.18: Taxony na chunk ───────────────────────────────
        st.slider(
            "🦪 Taxony na blok (1–5)" if _L=="cz" else "🦪 Taxa per chunk (1–5)",
            1, 5, st.session_state.get("taxa_per_chunk", 3), 1,
            key="taxa_per_chunk",
            help=(
                "Kolik taxonomických záznamů se zpracuje najednou v jednom LLM volání.\n"
                "**1–2**: pomalejší, vyšší přesnost pro složité záznamy\n"
                "**3**: doporučeno (výchozí) ✅\n"
                "**4–5**: rychlejší, vhodné pro jednoduché záznamy"
                if _L=="cz" else
                "How many taxonomic records to process per LLM call.\n"
                "1–2: slower, higher accuracy · 3: recommended ✅ · 4–5: faster, simple records"
            )
        )

    # ── Prompt Lab – testování promptu na vzorku ─────
    with st.expander(tt("🧪 Prompt Lab — testování bez spuštění celé extrakce",
                         "🧪 Prompt Lab — test without running full extraction", _L)):
        st.caption(tt("Otestuj prompt na malém vzorku textu bez plné extrakce.",
                      "Test prompt on a small sample text without running full extraction.", _L))
        plab_col1, plab_col2 = st.columns([3, 2])
        with plab_col1:
            plab_text = st.text_area(
                tt("Vzorový text (nebo nahrané PDF se načte automaticky)",
                   "Sample text (or uploaded PDF loads automatically)", _L),
                height=120, key="plab_text",
                value=st.session_state.get("last_extracted_text","")[:500]
                      if st.session_state.get("last_extracted_text") else "",
                placeholder=tt("Vlož 200–500 znaků z dokumentu pro test…",
                               "Paste 200–500 chars from document for testing…", _L))
        with plab_col2:
            plab_prompt = st.text_area(
                tt("Testovací prompt", "Test prompt", _L), height=120, key="plab_prompt",
                value=st.session_state.get("ext_prompt",
                      "Extrahuj taxonomické údaje jako JSON seznam."))
            plab_temp = st.slider(tt("Teplota", "Temperature", _L), 0.0, 1.0, 0.1, 0.05, key="plab_temp")

        if st.button(tt("🧪 Spustit test promptu", "🧪 Run prompt test", _L), key="plab_run_btn",
                     disabled=not (selected_model and plab_text.strip())):
            _plab_sys = (st.session_state.get("ext_sys_prompt", "").strip()
                         or EXTRACTION_MODES.get("Standard", {}).get("system", ""))
            with st.spinner(tt("Testuji prompt…", "Testing prompt…", _L)):
                plab_result = chat_completion_queued(
                    base_url, selected_model,
                    [{"role": "system", "content": _plab_sys},
                     {"role": "user",   "content": f"{plab_prompt}\n\n{plab_text[:500]}"}],
                    temp=plab_temp, max_tokens=2000)
            st.text_area(tt("📤 Výsledek testu", "📤 Test result", _L),
                         plab_result, height=200, key="plab_result")
            # Pokus o parsování JSON
            try:
                clean_r = _JSON_FENCE_RE.sub("", plab_result).strip()
                parsed_r = json.loads(clean_r)
                if isinstance(parsed_r, list):
                    st.success(tt(f"✅ Validní JSON: {len(parsed_r)} záznamů",
                                  f"✅ Valid JSON: {len(parsed_r)} records", _L))
                    if parsed_r:
                        st.dataframe(pd.DataFrame(parsed_r[:5]), width='stretch')
                else:
                    st.info(tt("ℹ️ JSON objekt (ne pole)", "ℹ️ JSON object (not array)", _L))
            except Exception:
                st.warning(tt("⚠️ Výsledek není validní JSON — zkontroluj prompt",
                              "⚠️ Output is not valid JSON — check your prompt", _L))

    # ── Správa šablon (editace) ─────────────────────
    with st.expander(tt("🗂️ Správa prompt šablon", "🗂️ Manage prompt templates", _L)):
        st.caption(tt("Uložené šablony se zachovají mezi relacemi.",
                      "Saved templates persist between sessions.", _L))

        _tpl_tab_new, _tpl_tab_edit, _tpl_tab_del = st.tabs([
            tt("➕ Nová šablona", "➕ New template", _L),
            tt("✏️ Upravit šablonu", "✏️ Edit template", _L),
            tt("🗑️ Smazat šablonu", "🗑️ Delete template", _L),
        ])

        # ── Záložka: Nová šablona ──────────────────────
        with _tpl_tab_new:
            new_tpl_name = st.text_input(tt("Název nové šablony", "New template name", _L), key="new_tpl_name")
            new_tpl_desc = st.text_input(tt("Popis", "Description", _L), key="new_tpl_desc")
            new_tpl_lang = st.selectbox(
                "Doporučený zdrojový jazyk" if _L=="cz" else "Recommended source language",
                LANGUAGE_OPTIONS, key="new_tpl_lang")
            new_tpl_prompt = st.text_area(
                "Text promptu" if _L=="cz" else "Prompt text", height=100, key="new_tpl_prompt")
            if st.button("💾 Uložit šablonu" if _L=="cz" else "💾 Save template",
                         key="btn_save_new_tpl"):
                if new_tpl_name.strip():
                    st.session_state["prompt_templates"][new_tpl_name] = {
                        "desc": new_tpl_desc, "lang_hint": new_tpl_lang, "prompt": new_tpl_prompt}
                    save_templates(st.session_state["prompt_templates"])
                    st.success(f"Šablona '{new_tpl_name}' uložena" if _L=="cz" else f"Template '{new_tpl_name}' saved")
                    st.rerun()
                else:
                    st.warning(tt("Zadej název šablony.", "Enter a template name.", _L))

        # ── Záložka: Upravit existující šablonu ───────
        with _tpl_tab_edit:
            if not templates_all:
                st.caption(tt("Zatím žádné šablony.", "No templates yet.", _L))
            else:
                _edit_tpl_sel = st.selectbox(
                    tt("Vyber šablonu k úpravě", "Select template to edit", _L),
                    options=list(templates_all.keys()),
                    key="edit_tpl_sel"
                )
                if _edit_tpl_sel:
                    _edit_tpl_data = templates_all[_edit_tpl_sel]
                    _edit_new_name = st.text_input(
                        tt("Název", "Name", _L),
                        value=_edit_tpl_sel,
                        key="edit_tpl_name"
                    )
                    _edit_new_desc = st.text_input(
                        tt("Popis", "Description", _L),
                        value=_edit_tpl_data.get("desc", ""),
                        key="edit_tpl_desc"
                    )
                    # Vybereme správný index jazyka
                    _edit_lang_val = _edit_tpl_data.get("lang_hint", "")
                    _edit_lang_idx = LANGUAGE_OPTIONS.index(_edit_lang_val) if _edit_lang_val in LANGUAGE_OPTIONS else 0
                    _edit_new_lang = st.selectbox(
                        "Doporučený zdrojový jazyk" if _L=="cz" else "Recommended source language",
                        LANGUAGE_OPTIONS, index=_edit_lang_idx,
                        key="edit_tpl_lang"
                    )
                    _edit_new_prompt = st.text_area(
                        "Text promptu" if _L=="cz" else "Prompt text",
                        value=_edit_tpl_data.get("prompt", ""),
                        height=120, key="edit_tpl_prompt"
                    )
                    _ec1, _ec2 = st.columns(2)
                    with _ec1:
                        if st.button(tt("💾 Uložit změny", "💾 Save changes", _L),
                                     key="btn_edit_tpl_save", use_container_width=True):
                            if _edit_new_name.strip():
                                _templates = st.session_state["prompt_templates"]
                                if _edit_new_name != _edit_tpl_sel:
                                    _templates.pop(_edit_tpl_sel, None)
                                _templates[_edit_new_name.strip()] = {
                                    "desc": _edit_new_desc,
                                    "lang_hint": _edit_new_lang,
                                    "prompt": _edit_new_prompt,
                                }
                                save_templates(_templates)
                                st.success(tt("✅ Šablona uložena", "✅ Template saved", _L))
                                st.rerun()
                            else:
                                st.warning(tt("Název nesmí být prázdný.", "Name cannot be empty.", _L))
                    with _ec2:
                        if st.button(tt("📋 Načíst do editoru", "📋 Load to editor", _L),
                                     key="btn_edit_tpl_load", use_container_width=True):
                            # Předvyplní hlavní prompt textarea hodnotou šablony
                            st.session_state["ext_prompt"] = _edit_tpl_data.get("prompt", "")
                            st.toast(tt(f"✅ Šablona '{_edit_tpl_sel}' načtena do promptu",
                                        f"✅ Template '{_edit_tpl_sel}' loaded into prompt", _L))
                            st.rerun()

        # ── Záložka: Smazat šablonu ────────────────────
        with _tpl_tab_del:
            if not templates_all:
                st.caption(tt("Zatím žádné šablony.", "No templates yet.", _L))
            else:
                del_name = st.selectbox(
                    "Smazat šablonu" if _L=="cz" else "Delete template",
                    ["—"] + list(templates_all.keys()), key="del_tpl")
                if del_name != "—":
                    if st.checkbox(tt(f"Opravdu smazat šablonu '{del_name}'?",
                                      f"Really delete template '{del_name}'?", _L),
                                   key="chk_del_tpl"):
                        if st.button("🗑️ Ano, smazat" if _L=="cz" else "🗑️ Yes, delete",
                                     key="btn_del_tpl", type="secondary"):
                            st.session_state["prompt_templates"].pop(del_name, None)
                            save_templates(st.session_state["prompt_templates"])
                            st.success(f"Šablona '{del_name}' smazána" if _L=="cz" else f"Template '{del_name}' deleted")
                            st.rerun()

    with st.expander(t("ext_sys_prompt",_L)):
        ext_sys_prompt = st.text_area(
            tt("Systémový prompt extrakce", "Extraction system prompt", _L),
            value=EXTRACTION_MODES[mode]["system"],
            height=80, key="ext_sys_prompt",
            help=tt("Přepíše výchozí systémový prompt pro tento režim extrakce.",
                    "Overrides the default system prompt for this extraction mode.", _L))

    # ── Checkpoint banner ─────────────────────────────
    _ckpts = checkpoint_list()
    if _ckpts:
        with st.expander(tt(f"💾 Dostupné checkpointy ({len(_ckpts)})",
                             f"💾 Available checkpoints ({len(_ckpts)}) — click to restore", _L)):
            for _ck in _ckpts:
                _cck1, _cck2, _cck3 = st.columns([4, 2, 1])
                _cck1.markdown(
                    f"**{_ck['name']}** — {_ck['saved_at']}  \n"
                    + tt(f"Dokončeno: {_ck['files_done']}/{_ck['files_total']} souborů",
                         f"Completed: {_ck['files_done']}/{_ck['files_total']} files", _L))
                if _cck2.button(tt("▶ Obnovit", "▶ Restore", _L),
                                key=f"ck_restore_{_ck['name']}",
                                width='stretch'):
                    _ck_data = checkpoint_load(_ck["name"])
                    if _ck_data:
                        st.session_state["ext_resume_state"] = _ck_data["data"].get("completed", {})
                        st.toast(tt(f"✅ Checkpoint '{_ck['name']}' načten — spusť extrakci znovu",
                                    f"✅ Checkpoint '{_ck['name']}' loaded — run extraction again", _L))
                if _cck3.button("🗑️", key=f"ck_del_{_ck['name']}",
                                help=tt("Smazat checkpoint", "Delete checkpoint", _L)):
                    checkpoint_delete(_ck["name"])
                    st.rerun()

    _ext_btn_col, _pause_col = st.columns([3, 1])
    with _ext_btn_col:
        run_btn = st.button(t("ext_run",_L), type="primary",
                            disabled=not (selected_model and files),
                            width='stretch')
    with _pause_col:
        if st.button(
                tt("⏸ Pozastavit", "⏸ Pause", _L),
                key="ext_pause_btn",
                help=tt("Po dokončení aktuálního souboru uloží checkpoint na disk a zastaví.\n"
                        "Checkpoint lze obnovit i po restartu aplikace.",
                        "Saves checkpoint after current file. Resume anytime, even after restart.", _L),
                width='stretch'):
            st.session_state["ext_pause_requested"] = True
            st.toast(tt("⏸ Pozastavení naplánováno — dokončí aktuální soubor",
                        "⏸ Pause scheduled — will finish current file", _L))
    ext_iterations = st.slider(
        "🔁 Počet iterací extrakce" if _L=="cz" else "🔁 Extraction iterations",
        1, 5, 1, 1, key="ext_iter",
        help=(
            "Spustí extrakci N-krát a výsledky sloučí zvolenou metodou.\n\n"
            "**Kdy použít:**\n"
            "- **1**: rychlé, standardní dokumenty\n"
            "- **2–3**: lepší recall, složité tabulky\n"
            "- **4–5**: maximální recall, husté dokumenty\n\n"
            "⚠️ N iterací = N× delší čas. Při >1 iteraci vyber metodu slučování."
            if _L=="cz" else
            "Runs extraction N times and merges results.\n"
            "2-3 = better recall. N iterations = N× longer processing."
        ))
    if st.session_state.get("ext_iter", 1) > 1:
        _merge_options = {
            "🤖 LLM Judge": "llm_judge",
            "📊 Consensus (voting)": "consensus",
            "🔗 Union + dedup": "union",
        }
        _merge_sel = st.selectbox(
            "Metoda slučování iterací" if _L=="cz" else "Iteration merge method",
            list(_merge_options.keys()),
            key="ext_merge_method",
            help=(
                "**🤖 LLM Judge** — pošle všechny iterace modelu, který vybere nejlepší výsledek.\n"
                "Nejinteligentnější, ale potřebuje extra LLM volání.\n\n"
                "**📊 Consensus** — pro každé pole vybere hodnotu z majority iterací (lokální, bez LLM).\n"
                "Spolehlivé pro strukturovaná data.\n\n"
                "**🔗 Union + dedup** — sloučí vše, odstraní duplicity. Maximální recall."
                if _L=="cz" else
                "**🤖 LLM Judge** — model picks best merge (extra LLM call).\n"
                "**📊 Consensus** — majority vote per field (local, no LLM).\n"
                "**🔗 Union + dedup** — merge all, remove duplicates. Max recall."
            )
        )

    # ── Stop / Pause tlačítka extrakce — viditelná VŽDY pokud extrakce běží/pozastavena (v24.31) ──
    _ext_is_running = st.session_state.get("extraction_running", False)
    _ext_is_paused  = st.session_state.get("ext_pause_requested", False)
    if _ext_is_running or _ext_is_paused:
        with st.container():
            _exc1, _exc2, _exc3 = st.columns([1, 1, 4])
            with _exc1:
                if st.button(tt("⏹️ Zastavit extrakci", "⏹️ Stop extraction", _L),
                             key="btn_stop_extraction_outer", type="secondary"):
                    st.session_state["extraction_running"] = False
                    st.session_state["ext_pause_requested"] = False
                    keep_current_tab(0)
                    st.session_state["_gsb_jump_to_tab"] = 0
                    st.rerun()
            with _exc2:
                if st.button(
                    tt("▶️ Pokračovat" if _ext_is_paused else "⏸️ Pozastavit",
                       "▶️ Resume"     if _ext_is_paused else "⏸️ Pause", _L),
                    key="btn_pause_extraction_outer", type="secondary"):
                    if _ext_is_paused:
                        st.session_state["ext_pause_requested"] = False
                        st.session_state["extraction_running"] = True
                    else:
                        st.session_state["ext_pause_requested"] = True
                    keep_current_tab(0)
                    st.session_state["_gsb_jump_to_tab"] = 0
                    st.rerun()
            with _exc3:
                if _ext_is_paused:
                    st.warning(tt("⏸️ Extrakce pozastavena — klikni Pokračovat",
                                  "⏸️ Extraction paused — click Resume", _L))
                elif _ext_is_running:
                    st.info(tt("⚙️ Extrakce probíhá…", "⚙️ Extraction running…", _L))

    if run_btn:
        system_msg = ext_sys_prompt if ext_sys_prompt.strip() else EXTRACTION_MODES[mode]["system"]
        # Přidej per-model prompt prefix
        _model_pfx = get_model_system_prompt(selected_model or "")
        if _model_pfx:
            system_msg = _model_pfx + "\n\n" + system_msg
        # Konfidence per-pole — injektuj do promptu
        _eff_prompt = prompt
        if st.session_state.get("ext_confidence", False):
            _eff_prompt = inject_confidence_prompt(_eff_prompt)
        # JSON schema v system promptu — přidá se do system_msg, prompt se nemění
        if st.session_state.get("ext_json_schema", False):
            system_msg += (
                "\n\nVýstup MUSÍ být validní JSON pole objektů. Každý objekt musí mít klíče: "
                "druh, rod, autor, rok, lokalita, stratigrafie, popis. "
                "Pokud hodnota není v textu, použij null. Žádný text mimo JSON.")

        results    = []
        # Udržuj _temp/ čistý před novou extrakcí (zachovej 5 nejnovějších sad)
        _temp_prune("extrakce", keep_last=5)
        prog_ext   = st.progress(0, text="Zpracovávám soubory…" if _L=="cz" else "Processing files…")
        n_iter     = st.session_state.get("ext_iter", 1)
        # Načti nastavení jednou mimo smyčku (výkon)
        _ext_parallel_on = st.session_state.get("ext_parallel", False)
        _ext_workers     = st.session_state.get("lms_max_concurrent", 1) if _ext_parallel_on else 1
        _ext_auto_chunk  = st.session_state.get("ext_auto_chunk", True)
        _ext_overlap     = st.session_state.get("ext_overlap", 200)
        _ext_preprocess  = st.session_state.get("ext_preprocess", True)
        _ext_pp_bibref   = st.session_state.get("ext_pp_bibref", False)
        _merge_sel_label = st.session_state.get("ext_merge_method", "🤖 LLM Judge")
        _merge_method_map = {"🤖 LLM Judge": "llm_judge",
                             "📊 Consensus (voting)": "consensus",
                             "🔗 Union + dedup": "union"}
        _merge_method = _merge_method_map.get(_merge_sel_label, "llm_judge")

        # Resume state: pokud existuje nedokončená extrakce, nabídneme pokračování
        _resume_key = "ext_resume_state"
        _resume = st.session_state.pop(_resume_key, None)

        # Časomíra
        _ext_start_time = _time.time()
        _time_col = st.empty()

        # Aktivuj globální status bar
        gsb_start(
            "🔍 Extrakce dat" if _L=="cz" else "🔍 Extracting data",
            tab=t("tab_extract", _L)
        )

        # ── příznak aktivní extrakce + živé okno ──────────────────
        st.session_state["extraction_running"] = True
        _ext_live_box = st.empty()

        # ── Stop / Pause tlačítka jsou nyní renderována VŽDY nad run_btn (v24.31) ──
        # Duplikáty odstraněny — tlačítka jsou definována v bloku výše (outside run_btn)

        # Inkrementální výstup: výsledky se zobrazují průběžně
        inc_container = st.container()

        for fi, f in enumerate(files):
            prog_ext.progress(fi / max(1, len(files)), text=f"📄 {f.name}")
            # Aktualizuj globální status bar i časomíru
            gsb_update(
                detail=f"{'Soubor' if _L=='cz' else 'File'} {fi+1}/{len(files)}: {f.name}",
                progress=fi / max(1, len(files))
            )
            _elapsed = _time.time() - _ext_start_time
            _remaining_est = (_elapsed / max(1, fi)) * (len(files) - fi) if fi > 0 else 0
            _time_col.caption(
                f"⏱ {_elapsed:.0f}s"
                + (f" · odhadovaný zbytek: ~{_remaining_est:.0f}s" if fi > 0 else "")
            )
            try:
                full_text = read_uploaded_file(f, page_spec or None, use_ocr, ocr_dpi)

                # OCR preprocessing
                preprocess_stats = {}
                if _ext_preprocess:
                    full_text, preprocess_stats = preprocess_text_for_llm(
                        full_text,
                        remove_page_markers=True,
                        remove_page_numbers=True,
                        remove_bibref=_ext_pp_bibref)

                # Auto chunk size
                _eff_chunk = (suggest_chunk_size(len(full_text), selected_model)
                              if _ext_auto_chunk
                              else chunk_size or 12000)
                blocks = chunk_text(full_text, _eff_chunk, overlap=_ext_overlap)
                blocks = deduplicate_chunks(blocks)   # odstraní opakující se záhlaví/patičky (OCR PDF)
                iter_results = []

                # Resume: přeskoč soubory které jsou již hotové
                if _resume and f.name in _resume:
                    results.append(_resume[f.name])
                    with inc_container:
                        st.success(f"⏭️ {f.name} — obnoveno z předchozí session ({_resume[f.name].get('chunks',0)} bloků)" if _L=="cz" else f"⏭️ {f.name} — restored from previous session ({_resume[f.name].get('chunks',0)} blocks)")
                    continue

                for iteration in range(n_iter):
                    all_chunks_out = []

                    if _ext_workers > 1 and len(blocks) > 1:
                        # ── Paralelní extrakce chunků ─────────────────
                        prog_ext.progress(
                            fi / max(1, len(files)),
                            text=f"📄 {f.name} – iter {iteration+1}/{n_iter} ⚡ {len(blocks)} bloků paralelně…")
                        chunk_results: Dict[int, str] = {}
                        _ext_lock = _threading.Lock()
                        _ext_done = [0]

                        # Propaguj session state před spuštěním threadů
                        # (modul-level proměnné — global deklarace není potřeba)
                        try:
                            _THREAD_OFFLINE = st.session_state.get("offline_mode", False)
                            _THREAD_GSP     = st.session_state.get("global_system_prompt", "").strip()
                        except Exception:
                            pass
                        _ext_count_lock = _threading.Lock()

                        def _extract_chunk(args, _sys=system_msg, _prompt=_eff_prompt,
                                           _temp=temp_ext, _mtok=max_tok_ext,
                                           _mode=mode, _base=base_url, _mdl=selected_model):
                            idx, blk = args
                            last_err = None
                            cur_blk  = blk
                            for _attempt in range(3):  # max 3 pokusy per chunk
                                try:
                                    resp = chat_completion(
                                        _base, _mdl,
                                        [{"role": "system", "content": _sys},
                                         {"role": "user",   "content": f"{_prompt}\n\n{cur_blk}"}],
                                        temp=_temp + _attempt * 0.05,
                                        max_tokens=_mtok,
                                        stop=_STOP_JSON if _mode == "JSON" else None)
                                    with _ext_count_lock:
                                        _ext_done[0] += 1
                                    return idx, resp
                                except Exception as err:
                                    last_err = err
                                    err_s = str(err).lower()
                                    if ("context size" in err_s or "context_length" in err_s
                                            or "context window" in err_s):
                                        # Context overflow — zmenši blok na polovinu
                                        if len(cur_blk) > 800:
                                            cur_blk = cur_blk[:len(cur_blk) // 2]
                                        else:
                                            break
                                    elif _attempt < 2:
                                        _time.sleep(_backoff_jitter(_attempt))
                            with _ext_count_lock:
                                _ext_done[0] += 1
                            return idx, f"[CHYBA po 3 pokusech: {last_err}]"

                        with ThreadPoolExecutor(max_workers=_ext_workers) as _ext_pool:
                            _ext_futures = {
                                _ext_pool.submit(_extract_chunk, (bi, blk)): bi
                                for bi, blk in enumerate(blocks)
                            }
                            for fut in as_completed(_ext_futures):
                                try:
                                    idx, resp = fut.result()
                                    chunk_results[idx] = resp
                                except Exception as e_ch:
                                    chunk_results[_ext_futures[fut]] = f"[CHYBA: {e_ch}]"
                                with _ext_count_lock:
                                    _done_now = _ext_done[0]
                                try:
                                    prog_ext.progress(
                                        (fi + _done_now / len(blocks)) / max(1, len(files)),
                                        text=f"📄 {f.name} – blok {_done_now}/{len(blocks)}")
                                except Exception:
                                    pass

                        all_chunks_out = [chunk_results[i] for i in range(len(blocks))]
                    else:
                        # ── Sekvenční extrakce ────────────────────────
                        _ext_seq_cache: Dict[str, str] = {}   # deduplikace identických bloků
                        for bi, block in enumerate(blocks):
                            prog_ext.progress(
                                (fi + (bi+1)/len(blocks)) / max(1, len(files)),
                                text=f"📄 {f.name} – iter {iteration+1}/{n_iter} blok {bi+1}/{len(blocks)}")
                            # Deduplikace: přeskoč identické bloky (záhlaví, opakující se části)
                            _blk_key = block.strip()
                            if _blk_key in _ext_seq_cache:
                                all_chunks_out.append(_ext_seq_cache[_blk_key])
                            else:
                                # Retry s jitter backoff (3 pokusy)
                                _seq_resp = None
                                _seq_err  = None
                                for _seq_attempt in range(3):
                                    try:
                                        _seq_resp = chat_completion_queued(
                                            base_url, selected_model,
                                            [{"role": "system", "content": system_msg},
                                             {"role": "user",   "content": f"{_eff_prompt}\n\n{block}"}],
                                            temp=temp_ext + _seq_attempt * 0.05,
                                            max_tokens=max_tok_ext,
                                            stop=_STOP_JSON if mode == "JSON" else None)
                                        break
                                    except Exception as _seq_e:
                                        _seq_err = _seq_e
                                        if _seq_attempt < 2:
                                            _time.sleep(_backoff_jitter(_seq_attempt))
                                if _seq_resp is None:
                                    _seq_resp = f"[!EXTRAKCE SELHALA blok {bi}: {_seq_err}]"
                                _ext_seq_cache[_blk_key] = _seq_resp
                                all_chunks_out.append(_seq_resp)
                            # ── v24.19: živé okno extrakce ───────────────
                            try:
                                _ext_preview = "\n---\n".join(
                                    [str(r)[:300] for r in all_chunks_out[-3:] if r]
                                )
                                _ext_preview_html = (
                                    _ext_preview.replace("&","&amp;")
                                               .replace("<","&lt;")
                                               .replace(">","&gt;")
                                               .replace("\n","<br>")
                                )
                                _ext_live_box.markdown(
                                    f'''<div id="lmu-ext-live" style="background:#f0fdf4;border:1px solid
#86efac;border-radius:6px;padding:.5rem .8rem;font-size:.78em;max-height:220px;
overflow-y:auto;color:#166534;font-family:monospace">
<b>⚙️ Blok {bi+1}/{len(blocks)} — {f.name}</b><br><br>{_ext_preview_html}
</div>
<script>(function(){{var b=document.getElementById("lmu-ext-live");if(b)b.scrollTop=b.scrollHeight;}})();</script>''',
                                    unsafe_allow_html=True,
                                )
                            except Exception:
                                pass
                            # ── Průběžný checkpoint každých N chunků ─────
                            if (bi + 1) % _EXT_CHUNK_CHECKPOINT_INTERVAL == 0 and bi + 1 < len(blocks):
                                _pc_completed = {r["file"]: r for r in results if r.get("ok")}
                                _pc_name = f"ext_progress_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                                checkpoint_save(_pc_name, {
                                    "completed":                   _pc_completed,
                                    "total_files":                 len(files),
                                    "in_progress_file":            f.name,
                                    "in_progress_chunks_done":     bi + 1,
                                    "in_progress_chunks_total":    len(blocks),
                                    "in_progress_partial_results": all_chunks_out[:],
                                    "iteration":                   iteration,
                                    "mode":                        mode,
                                    "prompt":                      prompt,
                                    "partial":                     True,
                                })

                    merged_iter = (all_chunks_out[0] if len(all_chunks_out) == 1
                                   else _merge_json_chunks(all_chunks_out))
                    iter_results.append(merged_iter)

                # ── Two-pass: druhý průchod doplní/opraví chybějící pole ──
                if st.session_state.get("ext_two_pass", False):
                    prog_ext.progress(
                        (fi + 0.5) / len(files),
                        text=f"📄 {f.name} – 2. průchod (oprava a doplnění)…")
                    pass2_sys = (
                        "Jsi expert na taxonomická data. Dostaneš výsledek první extrakce "
                        "a původní text. Tvým úkolem je: "
                        "1) Doplnit chybějící pole (NULL/prázdná) pokud jsou dostupná v textu. "
                        "2) Opravit zjevně chybné hodnoty. "
                        "3) Zachovat vše co je správné. "
                        "Výstup je POUZE opravený/doplněný JSON seznam, žádné komentáře."
                    )
                    first_pass_result = iter_results[-1] if iter_results else ""
                    src_preview = full_text[:4000]
                    pass2_resp = chat_completion_queued(
                        base_url, selected_model,
                        [{"role": "system", "content": pass2_sys},
                         {"role": "user", "content":
                          f"PŮVODNÍ TEXT (výňatek):\n{src_preview}\n\n"
                          f"VÝSLEDEK 1. PRŮCHODU:\n{first_pass_result}"}],
                        temp=0.05, max_tokens=max_tok_ext)
                    iter_results.append(pass2_resp)

                # Sloučení iterací
                if n_iter == 1 and not st.session_state.get("ext_two_pass", False):
                    merged = iter_results[0]
                else:
                    with st.spinner(tt(f"Slučuji {len(iter_results)} iterací ({_merge_sel_label})…",
                                       f"Merging {len(iter_results)} iterations ({_merge_sel_label})…", _L)):
                        if _merge_method == "llm_judge":
                            merged = _merge_llm_judge(
                                iter_results, base_url, selected_model,
                                src_text_preview=full_text[:2000],
                                max_tokens=max_tok_ext)
                        elif _merge_method == "consensus":
                            merged = _merge_consensus(iter_results)
                        else:
                            merged = _merge_union_dedup(iter_results)

                rec = {"file": f.name, "text": full_text, "result": merged,
                       "iter_results": iter_results,   # uložíme všechny iterace pro stažení
                       "chunks": len(blocks), "ok": True,
                       "preprocess_stats": preprocess_stats}
                results.append(rec)

                # ── Průběžný checkpoint po každém dokončeném souboru ─
                _pf_completed = {r["file"]: r for r in results if r.get("ok")}
                _pf_ck_name = f"ext_progress_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                checkpoint_save(_pf_ck_name, {
                    "completed":   _pf_completed,
                    "total_files": len(files),
                    "files_done":  fi + 1,
                    "mode":        mode,
                    "prompt":      prompt,
                    "partial":     (fi + 1) < len(files),
                })
                # ── _temp uložení výsledku tohoto souboru ────────────
                try:
                    _safe_fn = re.sub(r"[^\w\-.]", "_", f.name)
                    _temp_save_json(
                        f"extrakce_{_safe_fn}", fi + 1,
                        {"file": f.name, "result": merged,
                         "chunks": len(blocks), "saved_at": datetime.now().isoformat()},
                        len(files)
                    )
                except Exception:
                    pass

                # ── Inkrementální zobrazení výsledku ─────────────
                with inc_container:
                    pp_note = ""
                    if preprocess_stats.get("token_reduction_pct"):
                        pp_note = (f" · 🧹 preprocessing: "
                                   f"−{preprocess_stats['token_reduction_pct']}% tokenů")
                    chunk_note = f" ({len(blocks)} bloků)" if len(blocks) > 1 else ""
                    iter_note  = f" · {len(iter_results)} iterací" if len(iter_results) > 1 else ""
                    with st.expander(f"✅ {f.name}{chunk_note}{iter_note}{pp_note}", expanded=False):
                        st.text(preview_text(merged))
                        if preprocess_stats:
                            detail = ", ".join(f"{k}: {v}" for k,v in preprocess_stats.items()
                                               if k != "token_reduction_pct")
                            if detail:
                                st.caption(f"Preprocessing: {detail}")
                        # Stažení jednotlivých iterací inline
                        if len(iter_results) > 1:
                            _safe_fname = re.sub(r"[^\w\-.]", "_", f.name)
                            st.markdown(tt("**⬇️ Stáhnout iterace:**",
                                           "**⬇️ Download iterations:**", _L))
                            _icols = st.columns(len(iter_results) + 1)
                            for _ii, _iv in enumerate(iter_results):
                                _is_json = False
                                try:
                                    json.loads(_JSON_FENCE_RE.sub("", _iv).strip())
                                    _is_json = True
                                except Exception:
                                    pass
                                _iext = ".json" if _is_json else ".txt"
                                with _icols[_ii]:
                                    st.download_button(
                                        f"Iter {_ii+1}{_iext}",
                                        _iv.encode("utf-8"),
                                        f"{_safe_fname}_iter{_ii+1:02d}{_iext}",
                                        key=f"iter_single_{fi}_{_ii}"
                                    )
                            with _icols[-1]:
                                st.download_button(
                                    tt("📦 ZIP (vše)", "📦 ZIP (all)", _L),
                                    _make_iterations_zip(iter_results, prefix=_safe_fname),
                                    f"{_safe_fname}_iterace.zip",
                                    mime="application/zip",
                                    key=f"iter_zip_inline_{fi}"
                                )

                # ── Checkpoint / Pause ────────────────────────────
                if st.session_state.get("ext_pause_requested"):
                    completed_so_far = {r["file"]: r for r in results if r.get("ok")}
                    _ck_name = f"ext_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                    checkpoint_save(_ck_name, {
                        "completed":   completed_so_far,
                        "total_files": len(files),
                        "mode":        mode,
                        "prompt":      prompt,
                    })
                    st.session_state["ext_pause_requested"] = False
                    prog_ext.progress((fi+1)/len(files), text=tt("⏸ Pozastaveno", "⏸ Paused", _L))
                    st.warning(tt(
                        f"⏸ Extrakce pozastavena po souboru **{f.name}** "
                        f"({fi+1}/{len(files)}). "
                        f"Checkpoint uložen jako `{_ck_name}`. "
                        f"Klikni **▶ Obnovit** nad tlačítkem Spustit extrakci.",
                        f"⏸ Extraction paused after file **{f.name}** "
                        f"({fi+1}/{len(files)}). "
                        f"Checkpoint saved as `{_ck_name}`. "
                        f"Click **▶ Restore** above the Run extraction button.", _L)
                    )
                    desktop_notify("⏸ Extrakce pozastavena",
                                   f"Checkpoint: {fi+1}/{len(files)} souborů")
                    break  # zastav smyčku

            except Exception as e:
                rec_err = {"file": f.name, "result": str(e), "ok": False, "chunks": 0}
                results.append(rec_err)
                with inc_container:
                    st.error(f"❌ {f.name}: {str(e)[:100]}")
                # ── Auto-checkpoint při chybě ─────────────────────
                completed = {r["file"]: r for r in results if r.get("ok")}
                if completed:
                    _err_ck_name = f"ext_err_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                    checkpoint_save(_err_ck_name, {
                        "completed":   completed,
                        "total_files": len(files),
                        "mode":        mode,
                        "prompt":      prompt,
                        "error_on":    f.name,
                    })
                    st.session_state["ext_resume_state"] = completed
                    st.warning(
                        f"⚠️ Extrakce selhala na souboru **{f.name}**. "
                        f"Dosavadní výsledky ({len(completed)} souborů) uloženy jako "
                        f"checkpoint `{_err_ck_name}`. "
                        f"Klikni **▶ Obnovit** nebo spusť extrakci znovu."
                    )
                    desktop_notify("❌ Extrakce selhala",
                                   f"Checkpoint uložen: {len(completed)} souborů OK")

        prog_ext.progress(1.0, text="✅ Hotovo")

        # ── Úklid průběžných checkpointů a _temp po dokončení ──
        if _CHECKPOINT_DIR.exists():
            for _old_ck in _CHECKPOINT_DIR.glob("ext_progress_*.json"):
                try:
                    _old_ck.unlink()
                except Exception:
                    pass
        # Úklid _temp extrakce jen pokud vše OK — při chybách zachováme temp soubory
        _ext_ok_count  = sum(1 for r in results if r.get("ok"))
        _ext_err_count = sum(1 for r in results if not r.get("ok"))
        if _ext_err_count == 0 and _ext_ok_count > 0:
            try:
                _temp_cleanup("extrakce")
            except Exception:
                pass
        # Pokud jsou chyby, temp soubory zůstanou v _temp/ pro ruční obnovu

        if results:
            combined = "\n\n".join(f"=== {r['file']} ===\n{r['result']}" for r in results)
            st.session_state["last_extraction_text"]  = combined   # pro tab Překlad
            st.session_state["last_extracted_text"]   = combined   # pro UI + autosave + snapshoty
            st.session_state["last_extraction_taxa"]  = extract_taxa_from_text(combined)
            # ── Okamžitý autosave dokončené extrakce na disk ─────────
            save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)

            # ── Automatický snapshot dokončené extrakce ─────
            try:
                _auto_snap_label = (
                    f"auto_{datetime.now().strftime('%Y%m%d_%H%M')}"
                    f"_{len([r for r in results if r.get('ok')])}files"
                )
                save_extraction_version(combined, _auto_snap_label)
            except Exception:
                pass

            # ── FTS indexování ───────────────────────────────
            fts_index_results(results)

            ok_n   = sum(1 for r in results if r.get("ok"))
            taxa_n = len(st.session_state["last_extraction_taxa"])

            # ── Globální status bar — hotovo ─────────────────
            gsb_done_stay_on_tab(f"✅ Extrakce: {ok_n} souborů, {taxa_n} taxonů"
                     if _L=="cz" else
                     f"✅ Extraction: {ok_n} files, {taxa_n} taxa",
                     tab_index=0)

            # ── Desktop notifikace ───────────────────────────
            desktop_notify(
                "✅ Extrakce dokončena",
                f"{ok_n} souborů, {taxa_n} taxonů"
            )

            # ── Pydantic validace extrahovaných záznamů ──────
            all_json_for_pydantic = []
            for r in results:
                if r.get("ok"):
                    try:
                        clean_p = _JSON_FENCE_RE.sub("", r["result"]).strip()
                        parsed_p = json.loads(clean_p)
                        if isinstance(parsed_p, list):
                            all_json_for_pydantic.extend(parsed_p)
                    except Exception:
                        pass
            if all_json_for_pydantic:
                validated_recs, pydantic_errors = validate_extraction_records(all_json_for_pydantic)
                if pydantic_errors:
                    with st.expander(tt(f"⚠️ JSON validace: {len(pydantic_errors)} problémů "
                                         f"(z {len(all_json_for_pydantic)} záznamů)",
                                         f"⚠️ JSON validation: {len(pydantic_errors)} issues "
                                         f"(of {len(all_json_for_pydantic)} records)", _L)):
                        for err in pydantic_errors[:10]:
                            st.warning(err)
                        if not PYDANTIC_AVAILABLE:
                            st.caption(tt("💡 Nainstaluj `pydantic` pro detailnější validaci: `pip install pydantic`",
                                          "💡 Install `pydantic` for detailed validation: `pip install pydantic`", _L))
                else:
                    st.success(tt(f"✅ JSON validace: všech {len(all_json_for_pydantic)} záznamů OK",
                                  f"✅ JSON validation: all {len(all_json_for_pydantic)} records OK", _L)
                               + (" (Pydantic)" if PYDANTIC_AVAILABLE else ""))

        ok_count = sum(1 for r in results if r.get("ok"))
        _n_blocks = sum(r.get('chunks',1) for r in results)
        _n_taxa   = len(st.session_state['last_extraction_taxa'])
        st.success(tt(f"Zpracováno {ok_count}/{len(results)} souborů / "
                      f"{_n_blocks} bloků. Taxonů: {_n_taxa}",
                      f"Processed {ok_count}/{len(results)} files / "
                      f"{_n_blocks} blocks. Taxa: {_n_taxa}", _L))

        if results:
            combined_text = "\n\n".join(f"=== {r['file']} ===\n{r['result']}" for r in results)
            st.markdown(t("ext_export",_L))
            render_export_buttons(combined_text, "extrakce", {"mode": mode, "files": len(results)})

            # ── Stažení jednotlivých iterací ─────────────────────
            _results_with_iters = [r for r in results if r.get("ok") and len(r.get("iter_results", [])) > 1]
            if _results_with_iters:
                with st.expander(tt(f"⬇️ Stáhnout jednotlivé iterace ({len(_results_with_iters)} souborů)",
                                     f"⬇️ Download individual iterations ({len(_results_with_iters)} files)", _L)):
                    st.caption(tt("Každý soubor obsahuje výsledky každé iterace zvlášť.",
                                  "Each file contains results of each iteration separately — "
                                  "for manual comparison or post-processing.", _L))
                    _iter_cols = st.columns(min(3, len(_results_with_iters)))
                    for _ri, _r in enumerate(_results_with_iters):
                        _iters = _r["iter_results"]
                        _safe  = re.sub(r"[^\w\-.]", "_", _r["file"])
                        with _iter_cols[_ri % len(_iter_cols)]:
                            # ZIP se všemi iteracemi daného souboru
                            st.download_button(
                                f"📦 {_r['file']}\n({len(_iters)} iter.)",
                                _make_iterations_zip(_iters, prefix=_safe),
                                f"{_safe}_iterace.zip",
                                mime="application/zip",
                                key=f"iter_zip_export_{_ri}",
                                width='stretch'
                            )
                    # Pokud je více souborů, nabídni mega-ZIP se vším
                    if len(_results_with_iters) > 1:
                        st.divider()
                        _all_files: Dict[str, bytes] = {}
                        for _r in _results_with_iters:
                            _s = re.sub(r"[^\w\-.]", "_", _r["file"])
                            for _ii, _iv in enumerate(_r["iter_results"]):
                                _ext = ".json" if _iv.strip().startswith("[") else ".txt"
                                _all_files[f"{_s}/iterace_{_ii+1:02d}{_ext}"] = _iv.encode("utf-8")
                        _mega_buf = io.BytesIO()
                        with zipfile.ZipFile(_mega_buf, "w", zipfile.ZIP_DEFLATED) as _mz:
                            for _fn, _fc in _all_files.items():
                                _mz.writestr(_fn, _fc)
                        st.download_button(
                            tt(f"📦 Všechny iterace — mega ZIP ({len(_all_files)} souborů)",
                               f"📦 All iterations — mega ZIP ({len(_all_files)} files)", _L),
                            _mega_buf.getvalue(),
                            "vsechny_iterace.zip",
                            mime="application/zip",
                            key="iter_mega_zip"
                        )

            # ── Konfidence per-pole tabulka ──────────────────
            if st.session_state.get("ext_confidence", False):
                all_json_recs = []
                for r in results:
                    if r.get("ok"):
                        try:
                            clean = _JSON_FENCE_RE.sub("", r["result"]).strip()
                            parsed = json.loads(clean)
                            if isinstance(parsed, list): all_json_recs.extend(parsed)
                        except Exception: pass
                if all_json_recs:
                    with st.expander(tt(f"📊 Konfidence per-pole ({len(all_json_recs)} záznamů)",
                                         f"📊 Field confidence ({len(all_json_recs)} records)", _L),
                                     expanded=True):
                        st.markdown(render_confidence_table(all_json_recs),
                                    unsafe_allow_html=True)

            # ── Porovnání s předchozí extrakcí ────────────────
            prev_ext = st.session_state.get("prev_extraction_text","")
            if prev_ext:
                if st.button(tt("🔀 Porovnat s předchozí extrakcí",
                                "🔀 Compare with previous extraction", _L), key="ext_diff_btn"):
                    diff = diff_extraction_records(prev_ext, combined_text)
                    with st.expander(tt("🔀 Diff: nová vs. předchozí extrakce",
                                         "🔀 Diff: new vs. previous extraction", _L), expanded=True):
                        d1, d2, d3 = st.columns(3)
                        d1.metric(tt("➕ Přidáno",  "➕ Added",   _L), len(diff["added"]))
                        d2.metric(tt("➖ Odebráno", "➖ Removed", _L), len(diff["removed"]))
                        d3.metric(tt("✏️ Změněno",  "✏️ Changed", _L), len(diff["changed"]))
                        if diff["added"]:
                            st.markdown(tt("**Nové záznamy:**", "**New records:**", _L))
                            for rec in diff["added"][:10]:
                                st.code(json.dumps(rec, ensure_ascii=False)[:200])
                        if diff["removed"]:
                            st.markdown("**Chybějící záznamy:**" if _L=="cz" else "**Missing records:**")
                            for rec in diff["removed"][:10]:
                                st.code(json.dumps(rec, ensure_ascii=False)[:200])
                        if diff["changed"]:
                            st.markdown("**Změněné záznamy:**" if _L=="cz" else "**Changed records:**")
                            for ch in diff["changed"][:10]:
                                st.markdown(f"- `{ch['key']}`: " +
                                           ", ".join(f"{k}: `{v[0]}` → `{v[1]}`"
                                                     for k,v in list(ch["diffs"].items())[:4]))
            # Uložíme aktuální jako "předchozí" pro příští porovnání
            st.session_state["prev_extraction_text"] = combined_text

            # CSV pokud JSON
            try:
                all_records = []
                for r in results:
                    if r["ok"]:
                        clean  = _JSON_FENCE_RE.sub("", r["result"]).strip()
                        parsed = json.loads(clean)
                        items  = parsed if isinstance(parsed, list) else [parsed]
                        for rec in items:
                            rec["_zdroj"] = r["file"]; all_records.append(rec)
                if all_records:
                    batch_size = st.session_state.get("ext_taxa_batch_size", 2)
                    for b in range(0, len(all_records), batch_size):
                        batch = all_records[b:b+batch_size]
                        try:
                            _temp_save_json("extraction_partial", b // batch_size + 1, batch)
                            _temp_save_txt(
                                "extraction_partial",
                                b // batch_size + 1,
                                json.dumps(batch, ensure_ascii=False, indent=2)
                            )
                            _temp_save_docx(
                                "extraction_partial",
                                b // batch_size + 1,
                                json.dumps(batch, ensure_ascii=False, indent=2)
                            )
                        except Exception:
                            pass

                    df_ext = pd.DataFrame(all_records)
                    st.markdown(t("ext_table",_L))
                    st.dataframe(df_ext, width='stretch')
                    col_ec1, col_ec2 = st.columns(2)
                    with col_ec1:
                        st.download_button("⬇️ CSV", df_ext.to_csv(index=False).encode("utf-8"),
                                           "extrakce.csv", "text/csv", width='stretch')
                    with col_ec2:
                        if XLSX_AVAILABLE:
                            st.download_button("⬇️ Excel (.xlsx)", to_xlsx_bytes(df_ext),
                                               "extrakce.xlsx", width='stretch')
            except Exception:
                st.caption(t("ext_no_json",_L))
            # ── Export pro databázi Hyolitha ─────────────────────
            st.markdown(t("ext_hyolitha",_L))
            st.caption(t("ext_hyolitha_cap",_L))
            if all_records:
                col_hx1, col_hx2 = st.columns(2)
                with col_hx1:
                    if XLSX_AVAILABLE:
                        xlsx_h = hyolitha_export_xlsx(all_records)
                        st.download_button(
                            t("ext_hyolitha_xlsx",_L),
                            xlsx_h, "hyolitha_export.xlsx",
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            width='stretch')
                with col_hx2:
                    zip_h = hyolitha_export_csv_zip(all_records)
                    st.download_button(
                        t("ext_hyolitha_csv",_L),
                        zip_h, "hyolitha_export_csv.zip", "application/zip",
                        width='stretch')
                # Uložit do session state + _temp pro jistotu
                st.session_state["hyolitha_export_records"] = all_records
                try:
                    _temp_save_json("hyolitha_export", 1, all_records)
                except Exception:
                    pass
                # Náhled tabulky Taxa
                with st.expander(t("ext_preview_taxa",_L)):
                    dfs_prev = records_to_hyolitha_dfs(all_records)
                    st.dataframe(dfs_prev["Taxa"].head(10), width='stretch')
            else:
                st.caption(t("ext_hyolitha_none",_L))
            save_to_history("extrakce", {"mode": mode, "files": [r["file"] for r in results]})

            # ── Verzování výsledků extrakce ───────────────────
            _ext_data = st.session_state.get("last_extracted_text","")
            if _ext_data:
                ev_col1, ev_col2 = st.columns([1,3])
                with ev_col1:
                    ev_label = st.text_input(
                        "Popis snapshotu" if _L=="cz" else "Snapshot label",
                        key="ext_version_label",
                        placeholder="např. Malinky2004_v1" if _L=="cz" else "e.g. Malinky2004_v1")
                with ev_col2:
                    st.write("")
                    st.write("")
                    if st.button("📸 Uložit snapshot" if _L=="cz"
                                 else "📸 Save extraction snapshot",
                                 key="ext_save_version"):
                        vpath = save_extraction_version(_ext_data, ev_label)
                        if vpath:
                            st.toast(tt(f"✅ Snapshot uložen: {Path(vpath).name}",
                                        f"✅ Snapshot saved: {Path(vpath).name}", _L))
                        else:
                            st.toast(tt("❌ Uložení snapshotu selhalo",
                                        "❌ Saving snapshot failed", _L))

        # ── Přehled uložených verzí extrakce ─────────────
        ext_versions = list_extraction_versions()
        if ext_versions:
            with st.expander(f"📚 Uložené snapshoty extrakce ({len(ext_versions)})" if _L=="cz"
                             else f"📚 Saved extraction snapshots ({len(ext_versions)})"):
                for ev in ext_versions:
                    ev_size_kb = round(ev["size"]/1024, 1)
                    evc1, evc2 = st.columns([3,1])
                    with evc1:
                        label_str = f" — *{ev['label']}*" if ev["label"] else ""
                        st.markdown(f"🕐 **{ev['timestamp']}**{label_str} ({ev_size_kb} KB)")
                    with evc2:
                        if st.button("📂 Načíst" if _L=="cz" else "📂 Load",
                                     key=f"ev_load_{ev['file']}"):
                            try:
                                with open(ev["file"], "r", encoding="utf-8") as fh:
                                    ev_data = json.load(fh)
                                loaded_snap = ev_data.get("data","")
                                st.session_state["last_extracted_text"] = loaded_snap
                                st.session_state["last_extraction_text"] = loaded_snap
                                save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
                                st.toast(tt(f"✅ Snapshot načten: {ev['timestamp']}",
                                            f"✅ Snapshot loaded: {ev['timestamp']}", _L))
                                st.rerun()
                            except Exception as ex:
                                st.toast(f"❌ {ex}")


    _temp_file_browser(lang=st.session_state.get("lang", "cz"), key_prefix="tb_ext")


# ══════════════════════════════════════════════════════
# 1 – PŘEKLAD
# ══════════════════════════════════════════════════════
with main_tabs[1]:
    keep_current_tab(1)   # ← v24.17: zachování záložky Překlad
    _L = st.session_state.get("lang","cz")
    st.markdown(t("tr_title",_L))
    with st.expander("ℹ️ Jak funguje překlad?" if _L=="cz" else "ℹ️ How does translation work?", expanded=False):
        st.markdown("""
**Translation** splits any text into blocks and translates them sequentially or in parallel.

**Modes:**
- **Single text** — translate one file or manually entered text
- **Batch** — translate multiple files, download as ZIP
- **Compare** — translate same text with two different settings

**Performance:**
- Set ⚡ **Max concurrent requests** to 4 in sidebar
- Enable **⚡ Parallel** for 3–4× faster processing on long texts
""" if _L == "en" else """
**Překlad** zpracovává libovolně dlouhé texty — automaticky je rozdělí na bloky a přeloží je sekvenčně nebo paralelně.

**Režimy:**
- **Jeden text** — přeloží jeden soubor nebo ručně zadaný text
- **Dávkový** — přeloží více souborů najednou a zabalí je do ZIP
- **Porovnat** — přeloží stejný text dvěma různými modely/nastaveními

**Velké PDF (100+ stran):**
1. Nastav **Dávky po N stránkách** = 5–10
2. Každá dávka se přeloží zvlášť a je ke stažení jako DOCX
3. Na konci stáhni ZIP se všemi dávkami nebo celý spojený překlad

**Výkon:**
- V sidebaru nastav ⚡ **Max souběžných požadavků** na 4
- Zaškrtni **⚡ Paralelně** v záložce překladu
- U delších textů (10+ bloků) je paralelní překlad 3–4× rychlejší

**Doladění překladu** (po překladu):
- **Styl** — přizpůsobí text stylu zvoleného časopisu
- **Terminologie** — zkontroluje odborné pojmy v dané doméně
- **Zpětný překlad** — přeloží výsledek zpět a porovná s originálem
        """)

    trans_mode = st.radio(
        "Režim překladu" if _L=="cz" else "Translation mode",
        ["📄 Jeden text", "📦 Dávkový (více souborů)", "⚖️ Porovnat dva překlady"]
        if _L=="cz" else
        ["📄 Single text", "📦 Batch (multiple files)", "⚖️ Compare two translations"],
        horizontal=True, key="trans_mode_sel")

    # ── Sdílená nastavení jazyků s vlajkami ─────────
    if "tr_src_lang_val" not in st.session_state:
        st.session_state["tr_src_lang_val"] = lang_label("čeština")
    if "tr_tgt_lang_val" not in st.session_state:
        st.session_state["tr_tgt_lang_val"] = lang_label("angličtina")

    # Zaruč platné hodnoty po upgrade
    all_labeled_with_auto = [LANGUAGE_AUTO_LABELED] + LANGUAGE_OPTIONS_LABELED
    if st.session_state["tr_src_lang_val"] not in all_labeled_with_auto:
        st.session_state["tr_src_lang_val"] = LANGUAGE_AUTO_LABELED
    if st.session_state["tr_tgt_lang_val"] not in LANGUAGE_OPTIONS_LABELED:
        st.session_state["tr_tgt_lang_val"] = lang_label("angličtina")

    # ── Swap příznak: spracujeme PŘED vytvořením widgetů ──────────
    if st.session_state.pop("_do_lang_swap", False):
        _os = st.session_state.get("tr_src_lang_val", lang_label("čeština"))
        _ot = st.session_state.get("tr_tgt_lang_val", lang_label("angličtina"))
        st.session_state["tr_src_lang_val"] = _ot if _ot in LANGUAGE_OPTIONS_LABELED else lang_label("čeština")
        st.session_state["tr_tgt_lang_val"] = _os if _os in LANGUAGE_OPTIONS_LABELED else lang_label("angličtina")

    col_l1, col_swap, col_l2, col_l3 = st.columns([4, 1, 4, 3])
    with col_l1:
        # Bez key= — řízeno čistě přes index, bez konfliktu s widget session state
        _src_opts = [LANGUAGE_AUTO_LABELED] + LANGUAGE_OPTIONS_LABELED
        _src_idx  = (_src_opts.index(st.session_state["tr_src_lang_val"])
                     if st.session_state["tr_src_lang_val"] in _src_opts else 0)
        src_lang_labeled = st.selectbox(
            t("tr_src_lang",_L) if _L=="cz" else "Source language",
            _src_opts, index=_src_idx)
        st.session_state["tr_src_lang_val"] = src_lang_labeled
        src_lang_t = (AUTO_LANGUAGE_LABEL if src_lang_labeled == LANGUAGE_AUTO_LABELED
                      else lang_from_label(src_lang_labeled))
    with col_swap:
        st.write(""); st.write("")
        if st.button("⇄", key="btn_swap_langs",
                     help=tt("Prohodit zdrojový a cílový jazyk (+ přehodí text výsledku do vstupu)",
                             "Swap source and target language (+ moves result text to input)", _L)):
            # Nastavíme příznak — swap se provede na začátku příštího runu (před widgety)
            st.session_state["_do_lang_swap"] = True
            # Přehodit text výsledku do vstupního pole
            _result = st.session_state.get("translation_result", "")
            if _result:
                st.session_state["tr_src_text"] = _result
                st.session_state["translation_result"] = ""
                st.session_state["translation_ready"]  = False
            # Zůstat na záložce Překlad (tab index 1)
            st.session_state["_gsb_jump_to_tab"] = 1
            st.rerun()
    with col_l2:
        _tgt_idx = (LANGUAGE_OPTIONS_LABELED.index(st.session_state["tr_tgt_lang_val"])
                    if st.session_state["tr_tgt_lang_val"] in LANGUAGE_OPTIONS_LABELED else 1)
        tgt_lang_labeled = st.selectbox(
            t("tr_tgt_lang",_L),
            LANGUAGE_OPTIONS_LABELED, index=_tgt_idx)
        st.session_state["tr_tgt_lang_val"] = tgt_lang_labeled
        tgt_lang_t = lang_from_label(tgt_lang_labeled)
    with col_l3:
        temp_t = st.slider(
            "Teplota" if _L=="cz" else "Temperature",
            0.0, 0.8, 0.15, 0.05, key="temp_t",
            help=(
                "**Teplota překladu**\n\n"
                "- **0.1–0.2**: přesný, vědecký text ✅\n"
                "- **0.3–0.5**: vyváženější, přirozenější\n"
                "- **0.6+**: kreativní styl\n\n"
                "Pro paleontologickou literaturu doporučujeme **0.15**."
                if _L=="cz" else
                "0.1-0.2: precise scientific ✅ · 0.3-0.5: balanced · 0.6+: creative"
            ))

    col_opt1, col_opt2, col_opt3, col_opt4 = st.columns([3, 1, 1, 2])
    with col_opt1:
        preserve_terms_t = st.checkbox(
            "Zachovat latinské termíny beze změny" if _L=="cz" else "Preserve Latin terms unchanged",
            value=True,
            help=(
                "Instrukce pro LLM, aby latinská taxonomická jména (Hyolithes, Cambrian atd.) "
                "nepřekládal a zanechal beze změny.\n\n"
                "✅ Vždy nechej zapnuto pro paleontologické texty."
                if _L=="cz" else
                "Tells LLM to keep Latin taxonomic names unchanged. Always enable for paleontology."
            ))
    with col_opt2:
        tr_iterations = st.slider(
            "🔁 Iterace" if _L=="cz" else "🔁 Iterations",
            1, 3, 1, 1, key="tr_iter",
            help=(
                "Přeloží text N-krát a sloučí výsledky.\n\n"
                "- **1**: rychlé, vědecké záznamy\n"
                "- **2–3**: lepší kvalita pro 19. stol. literaturu\n"
                "⚠️ N iterací = N× delší čas"
                if _L=="cz" else
                "2-3 iterations = better quality for historical texts. N× longer."
            ))
    with col_opt3:
        _par_enabled = st.session_state.get("lms_max_concurrent", 1) > 1
        tr_parallel = st.checkbox(
            "⚡ Paralelně" if _L=="cz" else "⚡ Parallel",
            value=_par_enabled,
            key="tr_parallel",
            help=(
                "Odesílá bloky textu současně jako souběžné LLM požadavky.\n\n"
                "⚡ 3–4× rychlejší pro texty s ≥5 bloky\n"
                "⚠️ Bez sdíleného kontextu — pro historické texty použij sekvenční\n"
                "✅ Ideální pro tabulkové záznamy (hyolity)\n\n"
                "Nastav v sidebaru: Max souběžných požadavků"
                if _L=="cz" else
                "3-4× faster for ≥5 blocks. No shared context — use sequential for prose."
            ))
    with col_opt4:
        pass

    if st.session_state.get("tr_iter", 1) > 1:
        _tr_merge_options = {
            "🤖 LLM Judge (diff-based)": "llm_judge",
            "📊 Consensus (voting po odstavcích)": "consensus",
            "🔗 Union (prostá konkatenace)": "union",
        }
        st.selectbox(
            "Metoda slučování překladů" if _L=="cz" else "Merge method",
            list(_tr_merge_options.keys()),
            key="tr_merge_method",
            help=(
                "**LLM Judge** — model porovná varianty odstavec po odstavci a vybere nejlepší ✅\n\n"
                "**Consensus** — pro každý odstavec vybere variantu která se nejčastěji opakuje (rychlé, bez LLM)\n\n"
                "**Union** — varianty se sloučí prostě za sebou"
                if _L=="cz" else
                "**LLM Judge** — compares variants paragraph by paragraph ✅\n\n"
                "**Consensus** — picks most common variant per paragraph (fast, no LLM)\n\n"
                "**Union** — concatenates all variants"
            )
        )

    st.subheader(tt("📦 Nastavení bloků překladu", "📦 Translation chunking settings", _L))
    tr_chunk_col1, tr_chunk_col2, tr_chunk_col3 = st.columns([3, 2, 2])
    with tr_chunk_col1:
        tr_chunk_target = st.slider(
            tt("Cílová velikost bloku (znaky)", "Target chunk size (chars)", _L),
            1200, 6000, 2200, 100, key="tr_chunk_target"
        )
    with tr_chunk_col2:
        tr_use_smart_chunk = st.checkbox(
            tt("Chytré dělení na bloky (konec věty/odstavce)", "Smart chunking (sentence/paragraph)", _L),
            value=True, key="tr_use_smart_chunk"
        )
    with tr_chunk_col3:
        tr_chunk_overlap = st.slider(
            tt("Překryv (znaky)", "Overlap (chars)", _L),
            200, 600, 350, 50, key="tr_chunk_overlap"
        )

    tr_live_preview_box = st.empty()

    # ── v24.18: JS auto-scroll živého okna překladu dolů ─────────
    st.markdown("""
<script>
(function() {
  function scrollLiveBoxToBottom() {
    var boxes = document.querySelectorAll('[data-testid="stTextArea"] textarea');
    boxes.forEach(function(box) {
      box.scrollTop = box.scrollHeight;
    });
  }
  // Opakuj každých 800ms po dobu překladu
  var _liveScrollInterval = setInterval(scrollLiveBoxToBottom, 800);
  // Zastav po 30 minutách (bezpečnostní pojistka)
  setTimeout(function() { clearInterval(_liveScrollInterval); }, 1800000);
})();
</script>""", unsafe_allow_html=True)
    with st.expander(tt("📄 Kontext dokumentu (volitelné, zlepšuje terminologii)",
                         "📄 Document context (optional, improves terminology)", _L)):
        tr_doc_author = st.text_input(tt("Autor", "Author", _L), key="tr_doc_author",
                                      placeholder=tt("např. Malinky, J.M.", "e.g. Malinky, J.M.", _L))
        tr_doc_year   = st.text_input(tt("Rok", "Year", _L), key="tr_doc_year",
                                      placeholder=tt("např. 2004", "e.g. 2004", _L))
        tr_doc_type   = st.selectbox(tt("Typ dokumentu", "Document type", _L), key="tr_doc_type",
                                     options=(["—","vědecký článek","monografie","diplomová práce",
                                               "zpráva o výzkumu","kniha","sborník"] if _L=="cz"
                                              else ["—","journal article","monograph","thesis",
                                                    "research report","book","proceedings"]))
        tr_doc_topic  = st.text_input(tt("Téma/obor", "Topic/field", _L), key="tr_doc_topic",
                                      placeholder="např. Hyolitha, Cambrian stratigraphy")

    # ── Inline nápověda ──────────────────────────────
    with st.expander(tt("ℹ️ Jak na to?", "ℹ️ Quick guide", _L)):
        st.markdown(
            "1. Vyber **zdrojový a cílový jazyk** (nebo nech Auto-detect)\n"
            "2. Vyplň volitelně **kontext dokumentu** — autor, rok, téma zlepší terminologii\n"
            "3. Aktivuj **slovník** pro konzistentní překlad odborných termínů\n"
            "4. Zadej text ručně, ze souboru nebo z extrakce\n"
            "5. Překlad se zobrazuje **živě** (streaming) token po tokenu\n"
            "6. Po překladu se automaticky **ověří latinská jména** — chybějící jsou vyznačena\n"
            "7. Export DOCX track changes umožňuje ruční korekci v editoru"
            if _L=="cz" else
            "1. Select **source and target language** (or leave Auto-detect)\n"
            "2. Optionally fill **document context** — author, year, topic improve terminology\n"
            "3. Activate **glossary** for consistent translation of technical terms\n"
            "4. Enter text manually, from file or from extraction\n"
            "5. Translation appears **live** (streaming) token by token\n"
            "6. After translation, **Latin names** are automatically checked\n"
            "7. DOCX track changes export enables manual correction in Word"
        )

    # ── Výběr slovníku ──────────────────────────────
    gloss_names = list(st.session_state["glossaries"].keys())
    _no_gloss = tt("— bez slovníku —", "— no glossary —", _L)
    active_gloss_name = st.selectbox(
        tt("Aktivní slovník (volitelně)", "Active glossary (optional)", _L),
        [_no_gloss] + gloss_names, key="active_gloss")
    active_glossary = (st.session_state["glossaries"].get(active_gloss_name, {})
                       if active_gloss_name != _no_gloss else {})
    if active_glossary:
        terms_preview = ", ".join(list(active_glossary.keys())[:8])
        _more = "…" if len(active_glossary) > 8 else ""
        st.info(tt(f"📚 Aktivní slovník **{active_gloss_name}**: "
                   f"{len(active_glossary)} termínů – {terms_preview}{_more}",
                   f"📚 Active glossary **{active_gloss_name}**: "
                   f"{len(active_glossary)} terms – {terms_preview}{_more}", _L))

    # ════════════════════════════════════
    # REŽIM A – jeden text
    # ════════════════════════════════════
    if trans_mode == "📄 Jeden text":
        input_mode_t = st.radio("Zdroj textu",
                                ["✏️ Ručně","📄 Ze souboru","🔗 Z extrakce"],
                                horizontal=True, key="t_input_single")


# --- načtení zdrojového textu ---
        src_text_t = ""
        do_translate_btn = False

        if input_mode_t == "✏️ Ručně":
            # Formulář: text se uloží do session_state při submit, překlad pak čte ze session_state
            with st.form("tr_manual_form"):
                src_text_t_widget = st.text_area(
                    "Text k překladu",
                    height=220,
                    value=st.session_state.get("tr_src_text", ""),
                    placeholder="Vložte nebo napište text k překladu…" if _L=="cz" else "Paste or type text to translate…")
                do_translate_btn = st.form_submit_button("🌐 Přeložit", type="primary", disabled=not selected_model)

            # Po submit okamžitě ulož text do session_state – musí proběhnout PŘED rerunem
            if do_translate_btn and src_text_t_widget.strip():
                st.session_state["tr_src_text"] = src_text_t_widget

            # Vždy čteme src_text_t ze session_state (přežije rerun)
            src_text_t = st.session_state.get("tr_src_text", "")

            # Tlačítko pro smazání vstupního textu
            if src_text_t.strip():
                if st.button(tt("🗑️ Smazat vstup", "🗑️ Clear input", _L),
                             key="btn_clear_tr_input",
                             help=tt("Smaže vstupní text k překladu",
                                     "Clears the input text", _L)):
                    st.session_state["tr_src_text"] = ""
                    st.session_state["_gsb_jump_to_tab"] = st.session_state.get("presence_tab_idx", 1)
                    st.rerun()


        elif input_mode_t == "📄 Ze souboru":
            t_file = st.file_uploader("Soubor", type=["pdf","docx","txt"], key="t_file_single")

            # Nastavení rozsahu stran a dávkování (jen pro PDF)
            tr_page_spec   = ""
            tr_batch_pages = 0
            if t_file and t_file.name.lower().endswith(".pdf"):
                pspec_col1, pspec_col2, pspec_col3 = st.columns(3)
                with pspec_col1:
                    tr_page_spec = st.text_input(
                        tt("Rozsah stran (volitelné)", "Page range (optional)", _L),
                        key="tr_page_spec",
                        placeholder="1-10, 15, 20-30",
                        help=tt("Přeloží jen vybrané stránky. Prázdné = celý dokument.",
                                "Translates only selected pages. Empty = entire document.", _L))
                with pspec_col2:
                    tr_batch_pages = st.number_input(
                        tt("Dávky po N stránkách", "Batches of N pages", _L),
                        min_value=0, max_value=50, value=0, step=5,
                        key="tr_batch_pages",
                        help=tt("0 = bez dávkování\n5–10 = doporučeno pro velká PDF",
                                "0 = no batching\n5–10 = recommended for large PDFs", _L))
                with pspec_col3:
                    if tr_batch_pages > 0:
                        st.info(tt(f"📦 Režim dávek: po {tr_batch_pages} stranách",
                                   f"📦 Batch mode: {tr_batch_pages} pages per batch", _L))

            # Tlačítko vždy viditelné hned po nahrání souboru
            if t_file:
                st.caption(tt(f"📄 **{t_file.name}** — připraveno ke čtení",
                              f"📄 **{t_file.name}** — ready to read", _L))
            do_translate_btn = st.button(
                tt("🌐 Přeložit", "🌐 Translate", _L),
                type="primary",
                disabled=not (selected_model and t_file),
                key="btn_translate_main_file",
                help=tt("Nejdříve načte soubor, pak spustí překlad",
                        "First reads the file, then runs translation", _L) if t_file
                     else tt("Nejdříve nahraj soubor",
                             "Upload a file first", _L))

            # Čtení souboru proběhne AŽ PO kliknutí — uživatel vidí spinner
            src_text_t = ""
            if do_translate_btn and t_file:
                with st.spinner(tt(f"📖 Načítám {t_file.name}…",
                                   f"📖 Reading {t_file.name}…", _L)):
                    try:
                        src_text_t = read_uploaded_file(
                            t_file,
                            page_spec=tr_page_spec or None,
                            ocr_dpi=st.session_state.get("ext_ocr_dpi", 300))
                    except Exception as e_tf:
                        st.error(tt(f"Chyba čtení souboru: {e_tf}",
                                    f"File read error: {e_tf}", _L))
                        do_translate_btn = False
                if src_text_t:
                    n_pages_est = src_text_t.count("--- Strana ") or 1
                    st.caption(tt(f"✅ Načteno {len(src_text_t):,} znaků "
                                  f"(~{estimate_tokens(src_text_t):,} tokenů, "
                                  f"~{n_pages_est} stran)",
                                  f"✅ Loaded {len(src_text_t):,} chars "
                                  f"(~{estimate_tokens(src_text_t):,} tokens, "
                                  f"~{n_pages_est} pages)", _L))
            elif not t_file:
                do_translate_btn = False


        else:
            src_text_t = st.session_state.get("last_extraction_text","")
            if src_text_t:
                st.text_area(tt("Náhled extrakce", "Extraction preview", _L),
                             preview_text(src_text_t,800), height=100, disabled=True)
            else:
                st.warning(tt("Nejprve proveď extrakci.", "Run extraction first.", _L))
                src_text_t = ""

            do_translate_btn = st.button(tt("🌐 Přeložit", "🌐 Translate", _L), type="primary",
                                         disabled=not (selected_model and src_text_t.strip()),
                                         key="btn_translate_main_ext")

        # ══ v24.35: Stop / Pause / Resume — chunk-per-rerun architektura ═══════
        # Každý chunk = jeden Streamlit rerun → UI reaguje na tlačítka mezi chunky
        _tr_is_running = st.session_state.get("translate_running", False)
        _tr_is_paused  = st.session_state.get("translate_paused",  False)
        _tr_active     = _tr_is_running or _tr_is_paused or do_translate_btn

        if _tr_active:
            with st.container():
                _col_stop, _col_pause, _col_status = st.columns([1, 1, 4])
                with _col_stop:
                    if st.button(tt("⏹️ Zastavit", "⏹️ Stop", _L),
                                 key="btn_stop_translate", type="secondary"):
                        st.session_state["translate_running"] = False
                        st.session_state["translate_paused"]  = False
                        st.session_state["tr_live_partial"]   = ""
                        st.session_state["tr_chunk_queue"]    = None
                        # Zůstaneme na záložce Překlad
                        keep_current_tab(1)
                        st.query_params["tab"] = "1"
                        st.session_state["_gsb_jump_to_tab"] = 1
                        st.rerun()
                with _col_pause:
                    _btn_label = (
                        tt("▶️ Pokračovat", "▶️ Resume", _L) if _tr_is_paused
                        else tt("⏸️ Pozastavit", "⏸️ Pause", _L)
                    )
                    if st.button(_btn_label, key="btn_pause_translate", type="secondary"):
                        if _tr_is_paused:
                            st.session_state["translate_running"] = True
                            st.session_state["translate_paused"]  = False
                        else:
                            st.session_state["translate_paused"]  = True
                            st.session_state["translate_running"] = False
                        # Zůstaneme na záložce Překlad
                        keep_current_tab(1)
                        st.query_params["tab"] = "1"
                        st.session_state["_gsb_jump_to_tab"] = 1
                        st.rerun()
                with _col_status:
                    _q   = st.session_state.get("tr_chunk_queue") or []
                    _tot = st.session_state.get("tr_chunk_total", 0)
                    _done = _tot - len(_q)
                    if _tr_is_paused:
                        _prog_txt = (f" ({_done}/{_tot} bloků)" if _tot > 1 else "")
                        st.warning(tt(f"⏸️ Pozastaveno{_prog_txt} — klikni Pokračovat",
                                      f"⏸️ Paused{_prog_txt} — click Resume", _L))
                    elif _tr_is_running or do_translate_btn:
                        _prog_txt = (f" {_done}/{_tot}" if _tot > 1 else "")
                        st.info(tt(f"🌐 Překládám…{_prog_txt}", f"🌐 Translating…{_prog_txt}", _L))

        # ── Živé okno — vždy viditelné pokud máme průběžný text ─────────────
        tr_live_preview_box = st.empty()
        _tr_partial = st.session_state.get("tr_live_partial", "")
        if _tr_partial:
            _ph = (_tr_partial[-6000:]
                   .replace("&","&amp;").replace("<","&lt;")
                   .replace(">","&gt;").replace("\n","<br>"))
            _live_lbl = (
                tt("⏸️ Přeloženo dosud (pozastaveno):", "⏸️ Translated so far (paused):", _L)
                if _tr_is_paused else
                tt("🌐 Živý překlad:", "🌐 Live translation:", _L)
            )
            st.caption(_live_lbl)
            tr_live_preview_box.markdown(
                f'<div id="lmu-live-box" style="background:#f8fafc;border:1px solid #cbd5e1;'
                f'border-radius:6px;padding:.6rem 1rem;font-size:.82em;line-height:1.55;'
                f'max-height:320px;overflow-y:auto;color:#1e293b;white-space:pre-wrap;'
                f'font-family:inherit">{_ph}</div>',
                unsafe_allow_html=True,
            )

        # ══ Spuštění / pokračování překladu ══════════════════════════════════
        # Případ A: nový překlad (do_translate_btn) — inicializuj frontu chunků
        if do_translate_btn and src_text_t.strip():
            keep_current_tab(1)                         # ← Udrží záložku 1
            st.query_params["tab"] = "1"
            st.session_state["_gsb_jump_to_tab"] = 1
            actual_src = src_lang_t
            if src_lang_t == AUTO_LANGUAGE_LABEL:
                actual_src = detect_language(src_text_t) or "neznámý"
            _doc_ctx = {
                "author": st.session_state.get("tr_doc_author",""),
                "year":   st.session_state.get("tr_doc_year",""),
                "type":   st.session_state.get("tr_doc_type","—"),
                "topic":  st.session_state.get("tr_doc_topic",""),
            }
            sys_msg = build_translate_system(actual_src, tgt_lang_t, active_glossary,
                                             preserve_terms_t, doc_context=_doc_ctx)
            auto_chunk = suggest_chunk_size(len(src_text_t), selected_model)
            # Chunkuj text do fronty
            if len(src_text_t) <= auto_chunk:
                _chunks = [src_text_t]   # 1 chunk → streaming
            elif st.session_state.get("tr_use_smart_chunk", True):
                _chunks = smart_chunk_text(src_text_t,
                                           target_chars=st.session_state.get("tr_chunk_target", auto_chunk),
                                           overlap=st.session_state.get("tr_chunk_overlap", 350))
            else:
                _chunks = chunk_text_smart(src_text_t,
                                           max_chars=st.session_state.get("tr_chunk_target", auto_chunk))
            st.session_state.update({
                "translate_running":  True,
                "translate_paused":   False,
                "tr_live_partial":    "",
                "tr_chunk_queue":     _chunks,
                "tr_chunk_total":     len(_chunks),
                "tr_sys_msg":         sys_msg,
                "tr_actual_src":      actual_src,
                "tr_tgt_lang_run":    tgt_lang_t,
                "tr_temp_run":        temp_t,
                "tr_src_text_run":    src_text_t,
                # Reset starého výsledku — nový překlad začíná
                "translation_ready":      False,
                "translation_result":     "",
                "translation_original":   "",
                "translation_src_lang":   "",
                "translation_tgt_lang":   "",
                "translation_iterations": [],
            })
            gsb_start(
                f"🌐 Překlad {actual_src} → {tgt_lang_t}" if _L=="cz"
                else f"🌐 Translating {actual_src} → {tgt_lang_t}",
                tab=t("tab_translate", _L)
            )
            st.rerun()

        # Případ B: fronta chunků existuje a překlad běží → zpracuj 1 chunk
        _chunk_queue = st.session_state.get("tr_chunk_queue")
        if (_chunk_queue is not None and len(_chunk_queue) > 0
                and st.session_state.get("translate_running", False)):
            _chunk    = _chunk_queue[0]
            _sys_msg  = st.session_state.get("tr_sys_msg", "")
            _temp_run = st.session_state.get("tr_temp_run", temp_t)
            _src_run  = st.session_state.get("tr_src_text_run", src_text_t)
            _tgt_run  = st.session_state.get("tr_tgt_lang_run", tgt_lang_t)
            _src_lang_run = st.session_state.get("tr_actual_src", "")
            _tot      = st.session_state.get("tr_chunk_total", len(_chunk_queue))
            _done_n   = _tot - len(_chunk_queue) + 1

            gsb_update(
                detail=tt(f"Blok {_done_n}/{_tot}", f"Block {_done_n}/{_tot}", _L),
                progress=(_done_n - 1) / max(1, _tot)
            )

            # Přelož jeden chunk
            if len(_chunk_queue) == 1 and _tot == 1:
                # 1 chunk → zkus streaming
                _stream_box = tr_live_preview_box
                try:
                    _translated = chat_completion_stream(
                        base_url, selected_model,
                        [{"role":"system","content":_sys_msg},
                         {"role":"user","content":f"Přelož následující text:\n\n{_chunk}"}],
                        temp=_temp_run,
                        max_tokens=min(_MAX_TOKENS_TRANSLATE, max(1024, int(len(_chunk)/3*2.2))),
                        placeholder=_stream_box,
                        stop=_STOP_TRANSL or None)
                except Exception:
                    _translated = do_translate(base_url, selected_model, _chunk, _sys_msg,
                                               temp=_temp_run, src_lang=_src_lang_run)
            else:
                _translated = do_translate(base_url, selected_model, _chunk, _sys_msg,
                                           temp=_temp_run, src_lang=_src_lang_run)

            _translated = (_translated or "").strip()
            # Připoj k dosavadnímu překladu
            _prior = st.session_state.get("tr_live_partial", "")
            _current = (_prior + "\n\n" + _translated).strip() if _prior else _translated
            st.session_state["tr_live_partial"] = _current
            # Odeber zpracovaný chunk z fronty
            st.session_state["tr_chunk_queue"] = _chunk_queue[1:]
            # Průběžně uložit do _temp
            try:
                _temp_save_txt("preklad_live", _done_n, _current, _tot)
            except Exception:
                pass

            if len(_chunk_queue) == 1:
                # Poslední chunk — překlad dokončen
                _src_orig = st.session_state.get("tr_src_text_run", src_text_t)
                st.session_state.update({
                    "translation_original":  _src_orig,
                    "translation_result":    _current,
                    "translation_src_lang":  _src_lang_run,
                    "translation_tgt_lang":  _tgt_run,
                    "translation_iterations":[_current],
                    "translation_ready":     True,
                    "tr_batch_docx_files":   [],
                    "translate_running":     False,
                    "translate_paused":      False,
                    "tr_chunk_queue":        None,
                })
                save_to_history("překlad", {"src":_src_lang_run,"tgt":_tgt_run,
                                             "chars":len(_src_orig),"batch_pages":0,"iters":1})
                save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
                gsb_done_stay_on_tab(
                    tt(f"✅ Přeloženo {len(_src_orig):,} znaků → {_tgt_run}",
                       f"✅ Translated {len(_src_orig):,} chars → {_tgt_run}", _L),
                    tab_index=1)
                keep_current_tab(1)   # zajistí záložku Překlad po rerunu
                desktop_notify(tt("✅ Překlad dokončen","✅ Translation complete",_L))
                st.query_params["tab"] = "1"
                st.session_state["_gsb_jump_to_tab"] = 1
                st.rerun()
            else:
                # Další chunk čeká — rerun pro zpracování dalšího
                keep_current_tab(1)                         # ← Udrží záložku 1
                st.query_params["tab"] = "1"
                st.session_state["_gsb_jump_to_tab"] = 1
                st.rerun()

        # původní kód překladu nahrazen výše — přeskočit na zobrazení výsledku
        if False:
            actual_src = src_lang_t

            if src_lang_t == AUTO_LANGUAGE_LABEL:
                actual_src = detect_language(src_text_t) or "neznámý"
            _doc_ctx = {
                "author": st.session_state.get("tr_doc_author",""),
                "year": st.session_state.get("tr_doc_year",""),
                "type": st.session_state.get("tr_doc_type","—"),
                "topic": st.session_state.get("tr_doc_topic",""),
            }
            sys_msg = build_translate_system(actual_src, tgt_lang_t, active_glossary,
                                             preserve_terms_t, doc_context=_doc_ctx)
            n_iter = st.session_state.get("tr_iter", 1)
            batch_pages = st.session_state.get("tr_batch_pages", 0)
            batch_docx_files: List[bytes] = []
            iteration_outputs: List[str] = []
            translation_out = ""

            # Aktivuj globální status bar
            gsb_start(
                f"🌐 Překlad {actual_src} → {tgt_lang_t}" if _L=="cz"
                else f"🌐 Translating {actual_src} → {tgt_lang_t}",
                tab=t("tab_translate", _L)
            )

            # ── DÁVKOVÝ REŽIM (velké PDF po N stránkách) ────────────
            if batch_pages > 0 and input_mode_t == "📄 Ze souboru":
                _use_par = st.session_state.get("tr_parallel", False)
                _max_w   = st.session_state.get("lms_max_concurrent", 4)
                st.info(tt(f"📦 Dávkový překlad: po {batch_pages} stránkách",
                           f"📦 Batched translation: {batch_pages} pages per batch", _L)
                        + (tt(f" · ⚡ paralelně ({_max_w} req.)",
                              f" · ⚡ parallel ({_max_w} req.)", _L) if _use_par else ""))
                prog_container = st.container()
                translation_out, batch_docx_files = do_translate_batched(
                    base_url, selected_model, src_text_t, sys_msg,
                    pages_per_batch=batch_pages, temp=temp_t,
                    progress_container=prog_container,
                    parallel=_use_par, max_workers=_max_w,
                    src_lang=actual_src)

            else:
                # ── STANDARDNÍ PŘEKLAD s progress per-chunk ─────────
                iteration_outputs = []
                for it in range(n_iter):
                    t_it = min(temp_t + it * 0.05, 0.9) if n_iter > 1 else temp_t
                    iter_label = (tt(f"Překládám… (iterace {it+1}/{n_iter})",
                                     f"Translating… (iteration {it+1}/{n_iter})", _L)
                                  if n_iter > 1
                                  else tt("Překládám…", "Translating…", _L))
                    st.caption(f"🌐 {iter_label}")

                    if it == 0:
                        _temp_prune("preklad_final", keep_last=2)
                    # Progress bar + aktuální chunk
                    prog_bar  = st.progress(0, text=tt("Připravuji…", "Preparing…", _L))
                    chunk_box = st.empty()

                    def _progress(done, total, snippet, _pb=prog_bar, _cb=chunk_box, _L=_L):
                        pct = done / max(1, total)
                        _pb.progress(pct, text=tt(f"Blok {done}/{total} — {snippet[:50]}…",
                                                   f"Block {done}/{total} — {snippet[:50]}…", _L))
                        _cb.caption(tt(f"⚙️ Přeloženo {done}/{total} bloků",
                                       f"⚙️ Translated {done}/{total} blocks", _L))

                    # Odhadni chunk size a pracovníky
                    auto_chunk   = suggest_chunk_size(len(src_text_t), selected_model)
                    use_parallel = st.session_state.get("tr_parallel", False)
                    max_w        = st.session_state.get("lms_max_concurrent", 4)

                    # Pokus o streaming (jen pro 1 blok)
                    if len(src_text_t) <= auto_chunk:
                        stream_box = st.empty()
                        stream_box.markdown(
                            '<div style="background:#f8f9fa;border:1px solid #e0e0e0;'
                            'border-radius:6px;padding:.5rem 1rem;font-size:.85em;'
                            'max-height:180px;overflow-y:auto;color:#333">'
                            '⏳ Čekám na model…</div>', unsafe_allow_html=True)
                        try:
                            out_it = chat_completion_stream(
                                base_url, selected_model,
                                [{"role":"system","content":sys_msg},
                                 {"role":"user","content":
                                  f"Přelož následující text:\n\n{src_text_t}"}],
                                temp=t_it,
                                max_tokens=min(_MAX_TOKENS_TRANSLATE,
                                               max(1024, int(len(src_text_t)/3*2.2))),
                                placeholder=stream_box,
                                stop=_STOP_TRANSL or None)
                        except Exception:
                            out_it = translate_text_live(
                                base_url=base_url,
                                model=selected_model,
                                text=src_text_t,
                                sys_msg=sys_msg,
                                temp=t_it,
                                target_chars=st.session_state.get("tr_chunk_target", 2200),
                                overlap=st.session_state.get("tr_chunk_overlap", 350),
                                use_smart=st.session_state.get("tr_use_smart_chunk", True),
                                live_placeholder=tr_live_preview_box,
                                progress_cb=_progress,
                                src_lang=actual_src,
                                temp_prefix=f"preklad_iter{it+1}_live"
                            )
                        stream_box.empty()
                        prog_bar.progress(1.0, text="✅ Hotovo")
                    else:
                        # Více bloků — paralelní nebo sekvenční
                        if use_parallel:
                            chunk_box.caption(
                                f"⚡ Paralelní překlad ({max_w} souběžných požadavků)…")
                        out_it = translate_text_live(
                            base_url=base_url,
                            model=selected_model,
                            text=src_text_t,
                            sys_msg=sys_msg,
                            temp=t_it,
                            target_chars=st.session_state.get("tr_chunk_target", auto_chunk),
                            overlap=st.session_state.get("tr_chunk_overlap", 350),
                            use_smart=st.session_state.get("tr_use_smart_chunk", True),
                            live_placeholder=tr_live_preview_box,
                            progress_cb=_progress,
                            src_lang=actual_src,
                            temp_prefix=f"preklad_iter{it+1}_live"
                        )
                        prog_bar.progress(1.0, text=tt("✅ Hotovo", "✅ Done", _L))

                    chunk_box.empty()
                    iteration_outputs.append(out_it)
                    # ── Průběžné uložení iterace do _temp/ ──────────
                    try:
                        _ti_label = f"iter{it+1}_of{n_iter}"
                        _temp_save_txt(f"preklad_{_ti_label}", it + 1, out_it, n_iter)
                        _temp_save_docx(f"preklad_{_ti_label}", it + 1, out_it, n_iter)
                    except Exception:
                        pass

                if n_iter == 1:
                    translation_out = iteration_outputs[0]
                else:
                    _tr_merge_map = {
                        "🤖 LLM Judge (diff-based)": "llm_judge",
                        "📊 Consensus (voting po odstavcích)": "consensus",
                        "🔗 Union (prostá konkatenace)": "union",
                    }
                    _tr_merge_label = st.session_state.get("tr_merge_method", "🤖 LLM Judge (diff-based)")
                    _tr_merge_method = _tr_merge_map.get(_tr_merge_label, "llm_judge")

                    with st.spinner(tt(f"Slučuji {n_iter} iterací překladu ({_tr_merge_label})…",
                                       f"Merging {n_iter} translation iterations ({_tr_merge_label})…", _L)):
                        if _tr_merge_method == "llm_judge":
                            translation_out = _merge_llm_judge_translation(
                                iteration_outputs, base_url, selected_model,
                                src_text=src_text_t, tgt_lang=tgt_lang_t)
                        elif _tr_merge_method == "consensus":
                            translation_out = _merge_consensus_translation(iteration_outputs)
                        else:
                            translation_out = "\n\n".join(iteration_outputs)

            # ── Uložit finální překlad do _temp/ ─────────────────
            try:
                _temp_save_txt("preklad_final", 1, translation_out)
                _temp_save_docx("preklad_final", 1, translation_out)
                # Uklidíme iterační temp soubory jen pokud finál existuje
                if translation_out.strip():
                    for _ii in range(n_iter):
                        _temp_cleanup(f"preklad_iter{_ii+1}_of{n_iter}")
            except Exception:
                pass

            st.session_state.update({
                "translation_original":   src_text_t,
                "translation_result":     translation_out,
                "translation_src_lang":   actual_src,
                "translation_tgt_lang":   tgt_lang_t,
                "translation_iterations": iteration_outputs if batch_pages == 0 else [],
                "translation_ready":      True,
                "tr_batch_docx_files":    batch_docx_files,
            })
            save_to_history("překlad", {"src":actual_src,"tgt":tgt_lang_t,
                                         "chars":len(src_text_t),
                                         "batch_pages": batch_pages,
                                         "iters": n_iter if batch_pages == 0 else 1})
            # ── Okamžitý autosave překladu na disk ───────────────────
            st.session_state["translate_running"] = False
            st.session_state["translate_paused"]  = False
            save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
            gsb_done_stay_on_tab(tt(f"✅ Přeloženo {len(src_text_t):,} znaků → {tgt_lang_t}",
                        f"✅ Translated {len(src_text_t):,} chars → {tgt_lang_t}", _L), tab_index=1)
            keep_current_tab(1)   # zajistí záložku Překlad po rerunu
            desktop_notify(tt("✅ Překlad dokončen", "✅ Translation complete", _L),
                           tt(f"{len(src_text_t):,} znaků → {tgt_lang_t}",
                              f"{len(src_text_t):,} chars → {tgt_lang_t}", _L))
            st.session_state["_gsb_jump_to_tab"] = 1
            st.query_params["tab"] = "1"
            st.rerun()

        # --- zobrazení výsledku ---
        # Pokud máme výsledek ale ready=False (např. po obnově session), opravíme to
        if st.session_state.get("translation_result") and not st.session_state.get("translation_ready"):
            st.session_state["translation_ready"] = True
        if st.session_state.get("translation_ready") and st.session_state.get("translation_result"):
            translation_out = st.session_state["translation_result"]
            actual_src      = st.session_state.get("translation_src_lang","")
            tgt_lang_disp   = st.session_state.get("translation_tgt_lang","")

            st.divider()
            col_res_hdr, col_res_clr = st.columns([5, 1])
            with col_res_hdr:
                _n_chars = len(st.session_state.get('translation_original',''))
                st.success(tt(f"{_n_chars:,} znaků ({actual_src} → {tgt_lang_disp})",
                              f"{_n_chars:,} chars ({actual_src} → {tgt_lang_disp})", _L))
            with col_res_clr:
                if st.button(tt("🗑️ Smazat", "🗑️ Clear", _L), key="btn_clear_translation",
                             help=tt("Smaže zobrazený výsledek překladu",
                                     "Clears the displayed translation result", _L)):
                    st.session_state["translation_ready"] = False
                    st.session_state["translation_result"] = ""
                    st.session_state["_gsb_jump_to_tab"] = st.session_state.get("presence_tab_idx", 1)
                    st.rerun()

            # ── DLQ varování — chunky které selhaly po všech pokusech ──
            _tr_dlq = st.session_state.pop("translate_dlq", None)
            if _tr_dlq:
                with st.expander(
                    tt(f"⚠️ {len(_tr_dlq)} bloků selhalo při překladu (DLQ)",
                       f"⚠️ {len(_tr_dlq)} blocks failed during translation (DLQ)", _L),
                    expanded=True
                ):
                    for _dlq_item in _tr_dlq:
                        st.error(
                            tt(f"Blok {_dlq_item.get('idx','?')}: {_dlq_item.get('error','?')} "
                               f"| náhled: {_dlq_item.get('preview','')}",
                               f"Block {_dlq_item.get('idx','?')}: {_dlq_item.get('error','?')} "
                               f"| preview: {_dlq_item.get('preview','')}", _L)
                        )
                    st.caption(tt(
                        "Tyto bloky jsou v překladu označeny [!PŘEKLAD SELHAL]. "
                        "Zkus snížit velikost bloku nebo zvýšit timeout.",
                        "These blocks are marked [!PŘEKLAD SELHAL] in the output. "
                        "Try reducing chunk size or increasing timeout.", _L))

            preview_tr = translation_out[:4000] + ("…" if len(translation_out) > 4000 else "")
            st.text_area(tt("📄 Výsledek překladu (náhled 4000 znaků)",
                            "📄 Translation result (preview 4000 chars)", _L),
                         value=preview_tr, height=300, key="t_result_area")

            # ── v24.17: Porovnat originál × překlad (side-by-side) ───────────
            if st.button(tt("↔️ Porovnat originál × překlad (side-by-side)",
                            "↔️ Compare original × translation (side-by-side)", _L),
                         key="btn_compare_sidebyside", type="secondary"):
                st.session_state["tr_show_compare"] = not st.session_state.get("tr_show_compare", False)
            if st.session_state.get("tr_show_compare"):
                _orig_cmp = st.session_state.get("translation_original", "")
                _col_cmp_o, _col_cmp_t = st.columns(2)
                with _col_cmp_o:
                    st.text_area(tt("📄 Originál", "📄 Original", _L),
                                 value=_orig_cmp, height=600, key="cmp_orig_area")
                with _col_cmp_t:
                    st.text_area(tt("🌐 Překlad", "🌐 Translation", _L),
                                 value=translation_out, height=600, key="cmp_trans_area")

            # Segment-level diff (odstavce 1:1)
            with st.expander(tt("↔️ Segment-level diff (odstavce 1:1 s barevným zvýrazněním)",
                                 "↔️ Segment-level diff (paragraphs 1:1 with colour highlighting)", _L)):
                st.markdown(segment_diff_html(
                    st.session_state.get("translation_original",""), translation_out),
                    unsafe_allow_html=True)

            # Iterace — zobrazení a stažení
            iters = st.session_state.get("translation_iterations",[])
            if len(iters) > 1:
                with st.expander(tt(f"🔁 Jednotlivé iterace ({len(iters)}) — zobrazení a stažení",
                                     f"🔁 Individual iterations ({len(iters)}) — view and download", _L)):
                    for ii, iv in enumerate(iters):
                        c1, c2 = st.columns([4, 1])
                        with c1:
                            st.markdown(tt(f"**Varianta {ii+1}:**",
                                           f"**Variant {ii+1}:**", _L))
                            st.text(iv[:800] + ("…" if len(iv)>800 else ""))
                        with c2:
                            st.download_button(
                                f"⬇️ Iter {ii+1}",
                                iv.encode("utf-8"),
                                f"preklad_iter{ii+1:02d}.txt",
                                key=f"tr_iter_dl_{ii}",
                                width='stretch'
                            )
                    st.divider()
                    st.download_button(
                        tt(f"📦 ZIP všech iterací ({len(iters)})",
                           f"📦 ZIP of all iterations ({len(iters)})", _L),
                        _make_iterations_zip(iters, prefix="preklad"),
                        "preklad_iterace.zip",
                        mime="application/zip",
                        key="tr_iter_zip_all"
                    )

            render_export_buttons(translation_out, "preklad",
                                  {"src":actual_src,"tgt":tgt_lang_disp})
            # Track changes DOCX export
            orig_for_tc = st.session_state.get("translation_original","")
            if orig_for_tc:
                st.download_button(
                    tt("📝 Track Changes DOCX", "📝 Track Changes DOCX", _L),
                    to_docx_track_changes(orig_for_tc, translation_out),
                    "preklad_track_changes.docx",
                    help=tt("DOCX s přeškrtnutým originálem (červeně) a podtrženým překladem (zeleně)",
                            "DOCX with struck-through original (red) and underlined translation (green)", _L))

            # ── Stažení dávek (pro dávkový překlad velkého PDF) ──
            _batch_files = st.session_state.get("tr_batch_docx_files", [])
            if _batch_files:
                with st.expander(tt(f"📦 Stáhnout jednotlivé dávky ({len(_batch_files)} souborů)",
                                     f"📦 Download individual batches ({len(_batch_files)} files)", _L)):
                    batch_pages_n = st.session_state.get("tr_batch_pages", 5)
                    dcols = st.columns(min(4, len(_batch_files)))
                    for bi, bdocx in enumerate(_batch_files):
                        page_start = bi * batch_pages_n + 1
                        page_end   = page_start + batch_pages_n - 1
                        with dcols[bi % len(dcols)]:
                            st.download_button(
                                tt(f"⬇️ Str. {page_start}–{page_end}",
                                   f"⬇️ Pages {page_start}–{page_end}", _L),
                                bdocx,
                                f"preklad_str{page_start:04d}-{page_end:04d}.docx",
                                key=f"batch_dl_{bi}")
                    # ZIP všech dávek
                    if len(_batch_files) > 1:
                        all_zip = make_zip({
                            f"preklad_davka_{i+1:02d}.docx": d
                            for i, d in enumerate(_batch_files)
                        })
                        st.download_button(
                            tt(f"📦 Stáhnout vše jako ZIP ({len(_batch_files)} dávek)",
                               f"📦 Download all as ZIP ({len(_batch_files)} batches)", _L),
                            all_zip, "preklad_davky.zip", "application/zip",
                            key="batch_dl_zip")

            # ── DOLADĚNÍ ─────────────────────────────────
            st.divider()
            st.markdown(tt("#### ✨ Doladění překladu",
                           "#### ✨ Translation Refinement", _L))
            r1, r2, r3 = st.tabs([tt("✍️ Styl / čistota",  "✍️ Style / polish",     _L),
                                   tt("🔬 Terminologie",    "🔬 Terminology",         _L),
                                   tt("🔄 Zpětný překlad", "🔄 Back-translation",   _L)])

            with r1:
                st.markdown(tt("**Stylizace překladu**", "**Translation style**", _L))
                # Interní klíče česky (používají se jako lookup do style_map),
                # zobrazované labely přes format_func
                style_map = {
                    "🔬 Vědecká (konzervativní)":
                        ("Oprav pouze zjevné chyby, překlepy, nadbytečné mezery a interpunkci. "
                         "NEZMĚŇ žádný věcný obsah, terminologii ani strukturu vět. "
                         "Zachovej latinská jména beze změny."),
                    "📖 Obecná (čtivost)":
                        ("Lehce dolaď překlad – zlepši plynulost a čtivost. "
                         "Zachovej všechna fakta a terminologii, ale formulace mohou být přirozenější."),
                    "✂️ Silná (redakce)":
                        ("Výrazně přepracuj styl: zkracej redundantní pasáže, zjednoduš složitá souvětí, "
                         "sjednoť terminologii, zlepši strukturu odstavců. Věcný obsah zachovej beze změny."),
                }
                _style_labels = {
                    "🔬 Vědecká (konzervativní)": tt("🔬 Vědecká (konzervativní)",
                                                      "🔬 Scientific (conservative)", _L),
                    "📖 Obecná (čtivost)":        tt("📖 Obecná (čtivost)",
                                                      "📖 General (readability)",    _L),
                    "✂️ Silná (redakce)":         tt("✂️ Silná (redakce)",
                                                      "✂️ Strong (editorial)",       _L),
                }
                style_sel = st.radio(tt("Typ úpravy", "Edit type", _L),
                                     list(style_map.keys()),
                                     format_func=lambda k: _style_labels[k],
                                     key="r_style_type", horizontal=True)
                extra_note = st.text_input(tt("Doplňující instrukce (volitelné)",
                                               "Additional instructions (optional)", _L),
                                            key="r_extra_style")
                fix_gram   = st.checkbox(tt("Opravit gramatiku a interpunkci",
                                             "Fix grammar and punctuation", _L),
                                         value=True, key="r_gram")
                style_iter = st.slider(tt("Iterace stylu", "Style iterations", _L),
                                       1, 3, 1, 1, key="r_style_iter")

                if st.button(tt("✍️ Aplikovat styl", "✍️ Apply style", _L), key="btn_style"):
                    extras = (" Oprav gramatiku a interpunkci." if fix_gram else "")
                    extras += (f" Navíc: {extra_note}" if extra_note.strip() else "")
                    sys_s = (f"Jsi redaktor vědeckých překladů do jazyka {tgt_lang_disp}. "
                             f"{style_map[style_sel]}{extras}")
                    _style_chunk = suggest_chunk_size(len(translation_out), selected_model)
                    _style_par   = st.session_state.get("lms_max_concurrent", 1) > 1
                    _style_mw    = st.session_state.get("lms_max_concurrent", 1)
                    with st.spinner(tt("Upravuji styl…", "Applying style…", _L)):
                        style_out = translation_out
                        for _si in range(style_iter):
                            style_out = do_translate(
                                base_url, selected_model, style_out, sys_s,
                                temp=0.1, chunk_size=_style_chunk,
                                parallel=_style_par, max_workers=_style_mw,
                                src_lang=tgt_lang_disp)
                    # Uložit style_out do _temp ihned — nezávisí na kliknutí "Přijmout"
                    try:
                        _temp_save_txt("preklad_styl", 1, style_out)
                    except Exception:
                        pass
                    st.text_area(tt("Výsledek stylu", "Styled result", _L),
                                 style_out[:4000], height=260, key="r_style_out")
                    col_b1, col_b2 = st.columns(2)
                    with col_b1:
                        if st.button(tt("✅ Přijmout jako výsledek",
                                         "✅ Accept as result", _L), key="acc_style"):
                            st.session_state["translation_result"] = style_out
                            st.session_state["translation_ready"] = True
                            st.session_state["_gsb_jump_to_tab"] = 1
                            save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
                            st.rerun()
                    with col_b2:
                        if st.button("↔️ Diff", key="diff_style"):
                            st.markdown(simple_diff_html(translation_out, style_out),
                                        unsafe_allow_html=True)
                    render_export_buttons(style_out, "preklad_styl", {"typ": style_sel})

            with r2:
                # Interní hodnoty česky (používají se v sys promptu), zobrazované přes format_func
                domain_opts = ["paleontologie / taxonomie","geologie / stratigrafie",
                               "biologie / ekologie","obecná věda","vlastní"]
                _domain_labels = {
                    "paleontologie / taxonomie": tt("paleontologie / taxonomie",
                                                     "paleontology / taxonomy",    _L),
                    "geologie / stratigrafie":   tt("geologie / stratigrafie",
                                                     "geology / stratigraphy",     _L),
                    "biologie / ekologie":       tt("biologie / ekologie",
                                                     "biology / ecology",          _L),
                    "obecná věda":               tt("obecná věda",
                                                     "general science",            _L),
                    "vlastní":                   tt("vlastní",
                                                     "custom",                     _L),
                }
                domain_sel  = st.selectbox(tt("Doména", "Domain", _L), domain_opts,
                                           format_func=lambda k: _domain_labels[k], key="r_domain")
                custom_gloss_text = ""
                if domain_sel == "vlastní":
                    custom_gloss_text = st.text_area(
                        tt("Slovník (originál → překlad)",
                           "Glossary (original → translation)", _L),
                        height=100, key="r_cgloss")
                check_cons = st.checkbox(tt("Zkontroluj konzistenci",
                                             "Check consistency", _L), value=True, key="r_cons")
                if st.button(tt("🔬 Doladit terminologii",
                                 "🔬 Refine terminology", _L), key="btn_term"):
                    g_note = (f"\n\nSlovník:\n{custom_gloss_text}" if custom_gloss_text.strip() else "")
                    cons   = " Zkontroluj konzistenci pojmů napříč textem." if check_cons else ""
                    # Přidáme kontext originálu do system promptu (limitujeme na 1000 znaků)
                    orig_ctx = st.session_state.get('translation_original','')[:1000]
                    orig_note = (f"\n\nKontext originálu ({actual_src}):\n{orig_ctx}"
                                 if orig_ctx else "")
                    sys_t = (f"Jsi terminologický editor oboru {domain_sel}. "
                             f"Uprav přeložený text – pouze terminologie, obsah nezměněn."
                             f"{cons}{g_note}{orig_note}")
                    _term_chunk = suggest_chunk_size(len(translation_out), selected_model)
                    _term_par   = st.session_state.get("lms_max_concurrent", 1) > 1
                    _term_mw    = st.session_state.get("lms_max_concurrent", 1)
                    with st.spinner(tt("Doladuji terminologii…",
                                       "Refining terminology…", _L)):
                        ref = do_translate(base_url, selected_model, translation_out,
                                           sys_t, temp=0.1,
                                           chunk_size=_term_chunk,
                                           parallel=_term_par, max_workers=_term_mw,
                                           src_lang=tgt_lang_disp)
                    # Uložit ref do _temp ihned — nezávisí na kliknutí "Přijmout"
                    try:
                        _temp_save_txt("preklad_terminologie", 1, ref)
                    except Exception:
                        pass
                    st.text_area(tt("Výsledek terminologie",
                                    "Terminology result", _L),
                                 ref[:4000], height=260, key="r_term_out")
                    col_a1, col_a2 = st.columns(2)
                    with col_a1:
                        if st.button(tt("✅ Přijmout", "✅ Accept", _L), key="acc_term"):
                            st.session_state["translation_result"] = ref
                            st.session_state["translation_ready"] = True
                            st.session_state["_gsb_jump_to_tab"] = 1
                            save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
                            st.rerun()
                    with col_a2:
                        if st.button("↔️ Diff", key="diff_term"):
                            st.markdown(simple_diff_html(translation_out, ref), unsafe_allow_html=True)
                    render_export_buttons(ref, "preklad_terminologie", {"typ":"terminologie"})

            with r3:
                back_lang = st.selectbox(tt("Jazyk zpětného překladu",
                                             "Back-translation language", _L),
                                          LANGUAGE_OPTIONS,
                    index=LANGUAGE_OPTIONS.index(actual_src) if actual_src in LANGUAGE_OPTIONS else 0,
                    key="r_back_lang")
                show_analysis = st.checkbox(tt("LLM analýza rozdílů",
                                                "LLM diff analysis", _L),
                                             value=True, key="r_analysis")
                if st.button(tt("🔄 Zpětně přeložit", "🔄 Back-translate", _L), key="btn_back"):
                    _back_chunk = suggest_chunk_size(len(translation_out), selected_model)
                    _back_par   = st.session_state.get("lms_max_concurrent", 1) > 1
                    _back_mw    = st.session_state.get("lms_max_concurrent", 1)
                    sys_b = (f"Přelož přesně z {tgt_lang_disp} do {back_lang}. "
                             "Zachovej latinská taxonomická jména.")
                    with st.spinner(tt("Zpětný překlad…", "Back-translating…", _L)):
                        back = do_translate(base_url, selected_model, translation_out,
                                            sys_b, temp=0.1,
                                            chunk_size=_back_chunk,
                                            parallel=_back_par, max_workers=_back_mw,
                                            src_lang=tgt_lang_disp)
                    # Uložit zpětný překlad do _temp ihned
                    try:
                        _temp_save_txt("preklad_zpetny", 1, back)
                    except Exception:
                        pass
                    st.markdown(tt("##### ↔️ Originál vs. zpětný překlad",
                                   "##### ↔️ Original vs. back-translation", _L))
                    st.markdown(simple_diff_html(
                        st.session_state.get("translation_original",""), back),
                        unsafe_allow_html=True)
                    st.text_area(tt("Zpětný překlad", "Back-translation", _L),
                                 back[:4000], height=180, key="r_back_out")
                    if show_analysis:
                        with st.spinner(tt("Analyzuji…", "Analysing…", _L)):
                            # Bezpečné oříznutí — analýza nepotřebuje celý text
                            orig_snip = st.session_state.get('translation_original','')[:2000]
                            back_snip = back[:2000]
                            ana = chat_completion(base_url, selected_model,
                                [{"role":"system","content":
                                  "Porovnej originál a zpětný překlad. Identifikuj: "
                                  "1) ztracené informace, 2) terminologické nepřesnosti, "
                                  "3) hodnocení kvality 1–5. Odpovídej česky stručně."},
                                 {"role":"user","content":
                                  f"ORIGINÁL (výňatek):\n{orig_snip}"
                                  f"\n\nZPĚTNÝ PŘEKLAD (výňatek):\n{back_snip}"}],
                                temp=0.2, max_tokens=800)
                        st.markdown(tt("##### 🔎 Analýza kvality",
                                       "##### 🔎 Quality analysis", _L))
                        st.markdown(ana)
                    render_export_buttons(back, "zpetny_preklad", {"lang":back_lang})

    # ════════════════════════════════════
    # REŽIM B – dávkový překlad → ZIP
    # ════════════════════════════════════
    elif trans_mode == "📦 Dávkový (více souborů)":
        st.markdown(tt("Přeloží každý soubor zvlášť a zabalí výsledky do ZIP archivu.",
                       "Translates each file separately and packages results into a ZIP archive.", _L))
        batch_files = st.file_uploader(
            tt("Nahrát soubory k překladu (PDF, DOCX, TXT)",
               "Upload files for translation (PDF, DOCX, TXT)", _L),
            type=["pdf","docx","txt"], accept_multiple_files=True, key="batch_files")
        out_fmt = st.selectbox(tt("Formát výstupních souborů", "Output format", _L),
                               ["TXT", "DOCX"], key="batch_fmt")
        bc1, bc2, bc3 = st.columns(3)
        with bc1:
            add_orig = st.checkbox(tt("Přidat originál vedle překladu",
                                       "Include original alongside translation", _L),
                                   value=False, key="batch_add_orig")
        with bc2:
            use_dedup = st.checkbox(
                tt("🔁 Deduplikovat odstavce", "🔁 Deduplicate paragraphs", _L),
                value=True, key="batch_dedup",
                help=tt("Identické odstavce přes soubory se přeloží jen jednou (MD5 hash). "
                        "Výrazně zrychlí dávky s opakujícím se obsahem.",
                        "Identical paragraphs across files are translated only once (MD5 hash). "
                        "Significantly speeds up batches with repetitive content.", _L))
        with bc3:
            use_shared_gloss = st.checkbox(
                tt("📚 Sdílet slovník přes všechny soubory",
                   "📚 Share glossary across all files", _L),
                value=True, key="batch_shared_gloss",
                help=tt("Používá jeden slovník pro celou dávku — konzistentní terminologie.",
                        "Uses one glossary for the whole batch — consistent terminology.", _L))

        if st.button(tt("📦 Spustit dávkový překlad", "📦 Run batch translation", _L),
                     type="primary",
                     disabled=not (selected_model and batch_files)):
            actual_src_b = src_lang_t
            zip_contents: Dict[str, bytes] = {}
            prog_b = st.progress(0, text=tt("Dávkový překlad…", "Batch translation…", _L))
            log_b  = []

            # Sdílený slovník (stejný pro všechny soubory pokud zapnutý)
            shared_gloss = active_glossary if use_shared_gloss else {}

            # Deduplikace: sesbírej všechny odstavce, přelož unikátní, mapuj zpět
            dedup_cache: Dict[str, str] = {}  # hash → přeložený text

            for bi, bf in enumerate(batch_files):
                prog_b.progress(bi / max(1, len(batch_files)), text=f"📄 {bf.name}")
                try:
                    raw_text = read_uploaded_file(bf)
                    if src_lang_t == AUTO_LANGUAGE_LABEL:
                        actual_src_b = detect_language(raw_text) or "neznámý"
                    _doc_ctx_b = {
                        "author": st.session_state.get("tr_doc_author",""),
                        "year":   st.session_state.get("tr_doc_year",""),
                        "type":   st.session_state.get("tr_doc_type","—"),
                        "topic":  st.session_state.get("tr_doc_topic",""),
                    }
                    sys_msg_b = build_translate_system(
                        actual_src_b, tgt_lang_t, shared_gloss, preserve_terms_t,
                        doc_context=_doc_ctx_b)

                    if use_dedup:
                        # Rozděl na odstavce, přelož jen unikátní
                        paragraphs = [p for p in raw_text.split("\n") if p.strip()]
                        _bd_par  = st.session_state.get("lms_max_concurrent", 1)
                        # Identifikuj nové odstavce (mimo cache)
                        new_paras = [(i, p) for i, p in enumerate(paragraphs)
                                     if str(hash(p.strip()))
                                     not in dedup_cache]
                        reused = len(paragraphs) - len(new_paras)
                        if reused > 0:
                            log_b.append(f"  ♻️ Deduplikace: {reused} odstavců z cache")

                        if new_paras and _bd_par > 1:
                            # Paralelní překlad nových odstavců
                            _bd_lock = _threading.Lock()
                            def _translate_para(args, _sys=sys_msg_b):
                                idx, para = args
                                h = str(hash(para.strip()))
                                resp = chat_completion(
                                    base_url, selected_model,
                                    [{"role":"system","content":_sys},
                                     {"role":"user","content":
                                      f"Přelož následující odstavec:\n\n{para}"}],
                                    temp=temp_t, max_tokens=4000)
                                with _bd_lock:
                                    dedup_cache[h] = resp
                                return idx, h, resp

                            with ThreadPoolExecutor(max_workers=_bd_par) as _bd_pool:
                                _bd_futures = [_bd_pool.submit(_translate_para, ip)
                                               for ip in new_paras]
                                for fut in as_completed(_bd_futures):
                                    try:
                                        fut.result()
                                    except Exception as e_bd:
                                        log_b.append(f"  ⚠️ Chyba odstavce: {e_bd}")
                        elif new_paras:
                            for _, para in new_paras:
                                h = str(hash(para.strip()))
                                tr_para = chat_completion_queued(
                                    base_url, selected_model,
                                    [{"role":"system","content":sys_msg_b},
                                     {"role":"user","content":
                                      f"Přelož následující odstavec:\n\n{para}"}],
                                    temp=temp_t, max_tokens=4000)
                                dedup_cache[h] = tr_para

                        # Sestavit výsledek v původním pořadí
                        translated_parts = []
                        for para in paragraphs:
                            h = str(hash(para.strip()))
                            translated_parts.append(dedup_cache.get(h, para))
                        translated_b = "\n\n".join(translated_parts)
                    else:
                        _bdo_chunk = suggest_chunk_size(len(raw_text), selected_model)
                        _bdo_par   = st.session_state.get("lms_max_concurrent", 1) > 1
                        _bdo_mw    = st.session_state.get("lms_max_concurrent", 1)
                        translated_b = do_translate(base_url, selected_model, raw_text,
                                                    sys_msg_b, temp_t,
                                                    chunk_size=_bdo_chunk,
                                                    parallel=_bdo_par,
                                                    max_workers=_bdo_mw,
                                                    src_lang=actual_src_b)

                    stem = Path(bf.name).stem
                    content = (f"=== ORIGINÁL ({actual_src_b}) ===\n{raw_text}"
                               f"\n\n=== PŘEKLAD ({tgt_lang_t}) ===\n{translated_b}"
                               if add_orig else translated_b)
                    if out_fmt == "TXT":
                        zip_contents[f"{stem}_preklad.txt"] = to_txt_bytes(content)
                    else:
                        zip_contents[f"{stem}_preklad.docx"] = to_docx_bytes(content)
                    log_b.append(f"✅ {bf.name} → {len(translated_b):,} znaků")
                except Exception as e:
                    log_b.append(f"❌ {bf.name}: {e}")

            prog_b.progress(1.0, text=tt("✅ Hotovo", "✅ Done", _L))
            log_text = (f"Dávkový překlad {datetime.now():%Y-%m-%d %H:%M}\n"
                        f"{actual_src_b} → {tgt_lang_t}\n\n" + "\n".join(log_b))
            zip_contents["_log.txt"] = to_txt_bytes(log_text)
            zip_bytes = make_zip(zip_contents)
            ok_count  = sum(1 for l in log_b if l.startswith("✅"))
            st.success(tt(f"Přeloženo {ok_count}/{len(batch_files)} souborů "
                          f"(z cache: {len(dedup_cache)} unikátních odstavců)",
                          f"Translated {ok_count}/{len(batch_files)} files "
                          f"(from cache: {len(dedup_cache)} unique paragraphs)", _L))
            for line in log_b:
                (st.success if line.startswith("✅") else
                 st.info if line.startswith("  ♻️") else st.error)(line)
            st.download_button(
                tt(f"⬇️ Stáhnout ZIP ({len(zip_contents)-1} souborů)",
                   f"⬇️ Download ZIP ({len(zip_contents)-1} files)", _L),
                zip_bytes, f"preklady_{tgt_lang_t}.zip", "application/zip",
                width='stretch')
            save_to_history("dávkový překlad",
                            {"src":actual_src_b,"tgt":tgt_lang_t,"files":len(batch_files)})


    # ════════════════════════════════════
    # REŽIM C – porovnat dva překlady
    # ════════════════════════════════════
    else:
        st.markdown(tt("Zadej originál a dvě verze překladu. LLM rozhodne, který je lepší.",
                       "Enter original and two translation versions. LLM decides which is better.", _L))
        orig_cmp = st.text_area(tt("Originální text", "Original text", _L),
                                height=150, key="cmp_orig")
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            st.markdown(tt("**Překlad A**", "**Translation A**", _L))
            trans_a = st.text_area(tt("Překlad A", "Translation A", _L),
                                    height=200, key="cmp_a",
                                   placeholder=tt("Vložte první variantu překladu…",
                                                   "Paste first translation variant…", _L))
            a_label = st.text_input(tt("Popis A", "Label A", _L),
                                    value=tt("Varianta A", "Variant A", _L), key="cmp_a_label")
        with col_c2:
            st.markdown(tt("**Překlad B**", "**Translation B**", _L))
            trans_b = st.text_area(tt("Překlad B", "Translation B", _L),
                                    height=200, key="cmp_b",
                                   placeholder=tt("Vložte druhou variantu překladu…",
                                                   "Paste second translation variant…", _L))
            b_label = st.text_input(tt("Popis B", "Label B", _L),
                                    value=tt("Varianta B", "Variant B", _L), key="cmp_b_label")

        # Vnitřní hodnoty česky (používají se v sys promptu), zobrazované přes format_func
        _eval_opts = ["Věrnost originálu","Terminologická přesnost","Plynulost a čtivost",
                      "Zachování odborného stylu","Konzistence pojmů"]
        _eval_labels = {
            "Věrnost originálu":          tt("Věrnost originálu",
                                              "Fidelity to original",          _L),
            "Terminologická přesnost":    tt("Terminologická přesnost",
                                              "Terminological accuracy",       _L),
            "Plynulost a čtivost":        tt("Plynulost a čtivost",
                                              "Fluency and readability",       _L),
            "Zachování odborného stylu":  tt("Zachování odborného stylu",
                                              "Preservation of scholarly style", _L),
            "Konzistence pojmů":          tt("Konzistence pojmů",
                                              "Term consistency",              _L),
        }
        eval_criteria = st.multiselect(
            tt("Hodnotící kritéria", "Evaluation criteria", _L),
            _eval_opts,
            default=["Věrnost originálu","Terminologická přesnost","Plynulost a čtivost"],
            format_func=lambda k: _eval_labels[k],
            key="cmp_criteria")

        if st.button(tt("⚖️ Porovnat překlady", "⚖️ Compare translations", _L),
                     type="primary",
                     disabled=not (selected_model and orig_cmp and trans_a and trans_b)):
            criteria_str = ", ".join(eval_criteria)
            sys_cmp = (
                f"Jsi expert na překlady odborných paleontologických textů do jazyka {tgt_lang_t}. "
                f"Porovnej dvě verze překladu podle: {criteria_str}. Výstup:\n"
                f"1) Hodnocení {a_label} (1-5 bodů)\n2) Hodnocení {b_label} (1-5 bodů)\n"
                "3) Konkrétní ukázky\n4) Celkový vítěz\n5) Sloučená nejlepší verze\n"
                "Odpovídej česky.")
            user_cmp = (f"ORIGINÁL:\n{orig_cmp}\n\n"
                        f"=== {a_label} ===\n{trans_a}\n\n=== {b_label} ===\n{trans_b}")
            with st.spinner(tt("Porovnávám překlady…", "Comparing translations…", _L)):
                evaluation = chat_completion(
                    base_url, selected_model,
                    [{"role":"system","content":sys_cmp},
                     {"role":"user","content":user_cmp}],
                    temp=0.2, max_tokens=3000)
            st.markdown(tt("### ⚖️ Výsledek porovnání", "### ⚖️ Comparison result", _L))
            st.markdown(evaluation)
            render_export_buttons(evaluation, "porovnani_prekladu",
                                  {"A":a_label,"B":b_label,"criteria":eval_criteria})
            save_to_history("porovnání překladů",
                            {"A":a_label,"B":b_label,"criteria":eval_criteria})

    # ── Správa slovníků ───────────────────────────────
    st.divider()
    st.markdown(tt("#### 📚 Správa slovníků", "#### 📚 Glossary management", _L))
    gl_tabs = st.tabs([tt("Zobrazit / editovat", "View / edit", _L),
                        tt("Nový slovník",       "New glossary", _L)])
    with gl_tabs[0]:
        if not st.session_state["glossaries"]:
            st.info(tt("Žádné slovníky. Vytvoř nový v záložce 'Nový slovník'.",
                       "No glossaries. Create one in the 'New glossary' tab.", _L))
        else:
            gl_sel = st.selectbox(tt("Vybrat slovník", "Select glossary", _L),
                                  list(st.session_state["glossaries"].keys()), key="gl_view_sel")
            gl_text = st.text_area(
                tt("Obsah (originál → překlad, jeden pár na řádek)",
                   "Content (original → translation, one pair per line)", _L),
                glossary_to_text(st.session_state["glossaries"].get(gl_sel, {})),
                height=200, key="gl_edit_text")
            col_gl1, col_gl2 = st.columns(2)
            with col_gl1:
                if st.button(tt("💾 Uložit změny", "💾 Save changes", _L), key="gl_save"):
                    st.session_state["glossaries"][gl_sel] = text_to_glossary(gl_text)
                    save_glossaries(st.session_state["glossaries"])
                    st.success(tt("Uloženo", "Saved", _L))
                    st.session_state["_gsb_jump_to_tab"] = 1
                    st.rerun()
            with col_gl2:
                if st.checkbox(tt(f"Opravdu smazat slovník '{gl_sel}'?",
                                  f"Really delete glossary '{gl_sel}'?", _L),
                               key="chk_del_gl"):
                    if st.button(tt("🗑️ Ano, smazat", "🗑️ Yes, delete glossary", _L),
                                 key="gl_del", type="secondary"):
                        st.session_state["glossaries"].pop(gl_sel, None)
                        save_glossaries(st.session_state["glossaries"])
                        st.success(tt(f"'{gl_sel}' smazán", f"'{gl_sel}' deleted", _L))
                        st.session_state["_gsb_jump_to_tab"] = 1
                        st.rerun()
            st.download_button(tt("⬇️ Exportovat JSON", "⬇️ Export JSON", _L),
                to_json_bytes(st.session_state["glossaries"].get(gl_sel, {})),
                f"glossary_{gl_sel}.json", key="gl_export")

    with gl_tabs[1]:
        new_gl_name = st.text_input(tt("Název slovníku", "Glossary name", _L), key="new_gl_name",
                                    placeholder=tt("např. Hyolitha_CZ-EN", "e.g. Hyolitha_CZ-EN", _L))
        new_gl_src = st.selectbox(tt("Zdrojový jazyk slovníku", "Glossary source language", _L),
                                  LANGUAGE_OPTIONS, key="new_gl_src")
        new_gl_tgt = st.selectbox(tt("Cílový jazyk slovníku", "Glossary target language", _L),
                                  LANGUAGE_OPTIONS, index=1, key="new_gl_tgt")
        new_gl_text = st.text_area(
            tt("Termíny (originál → překlad, jeden pár na řádek)",
               "Terms (original → translation, one pair per line)", _L),
            height=200, key="new_gl_text",
            placeholder=("Hyolithida → Hyolithida\nsclerite → sklerit" if _L=="cz" else "Hyolithida → Hyolithida\nsclerite → sclerite")
        )
        gl_import = st.file_uploader(tt("Nebo importovat z JSON",
                                         "Or import from JSON", _L),
                                      type=["json"], key="gl_import")
        if gl_import:
            try:
                imported    = json.loads(gl_import.read().decode("utf-8"))
                new_gl_text = glossary_to_text(imported)
                st.success(tt(f"Importováno {len(imported)} termínů",
                              f"Imported {len(imported)} terms", _L))
            except Exception as e:
                st.error(tt(f"Chyba importu: {e}", f"Import error: {e}", _L))
        if st.button(tt("🤖 Generovat slovník (LLM)",
                         "🤖 Generate glossary (LLM)", _L),
                     key="gl_gen", disabled=not selected_model):
            context = st.session_state.get("last_extraction_text","")[:4000]
            if not context:
                st.warning(tt("Nejprve proveď extrakci.", "Run extraction first.", _L))
            else:
                with st.spinner(tt("Generuji slovník…", "Generating glossary…", _L)):
                    gen_sys = (f"Z textu extrahuj terminologický slovník pro překlad "
                               f"z {new_gl_src} do {new_gl_tgt}. "
                               f"Výstup jako JSON {{\"originál\": \"překlad\"}}. "
                               "Pouze odborné termíny, taxonomická jména, stratigrafické pojmy.")
                    gen_resp = chat_completion(base_url, selected_model,
                        [{"role":"system","content":gen_sys},
                         {"role":"user","content":context}],
                        temp=0.1, max_tokens=2000)
                try:
                    gen_dict  = json.loads(_JSON_FENCE_RE.sub("", gen_resp).strip())
                    new_gl_text = glossary_to_text(gen_dict)
                    st.success(tt(f"Vygenerováno {len(gen_dict)} termínů",
                                  f"Generated {len(gen_dict)} terms", _L))
                    st.text_area(tt("Vygenerovaný slovník", "Generated glossary", _L),
                                 new_gl_text, height=150)
                except Exception:
                    st.text_area(tt("Surový výstup LLM", "Raw LLM output", _L),
                                 gen_resp, height=150)
        if (st.button(tt("💾 Uložit nový slovník", "💾 Save new glossary", _L),
                      type="primary", key="save_new_gl")
                and new_gl_name.strip()):
            gl_data = text_to_glossary(new_gl_text)
            st.session_state["glossaries"][new_gl_name] = gl_data
            save_glossaries(st.session_state["glossaries"])
            st.success(tt(f"Slovník '{new_gl_name}' uložen ({len(gl_data)} termínů)",
                          f"Glossary '{new_gl_name}' saved ({len(gl_data)} terms)", _L))
            st.session_state["_gsb_jump_to_tab"] = 1
            st.rerun()

    _temp_file_browser(lang=st.session_state.get("lang", "cz"), key_prefix="tb_tr")


# ══════════════════════════════════════════════════════
# 2 – TAXONOMICKÁ VALIDACE  (+ duplicate checker)
# ══════════════════════════════════════════════════════
with main_tabs[2]:
    keep_current_tab(2)   # ← v24.17
    _L = st.session_state.get("lang","cz")
    # Zpracuj příznak smazání pole taxonů PŘED renderem widgetů
    # Správný vzor: nastavit val_taxa_manual přímo v session_state PŘED renderem widgetu.
    # Streamlit widget s key= čte svou hodnotu z session_state[key] při každém renderu.
    if st.session_state.pop("_val_clear_pending", False):
        st.session_state["val_taxa_input"] = ""
        st.session_state["val_taxa_manual"] = ""
    st.markdown(t("val_title",_L))
    with st.expander(tt("ℹ️ Jak funguje validace taxonů?",
                         "ℹ️ How does taxon validation work?", _L), expanded=False):
        st.markdown(tt(
            """
**Validace** ověří každé taxonomické jméno ve vybraných online databázích a zobrazí výsledky s přímými odkazy.

**Databáze pro paleontologii** (doporučeno):
- **PaleoDB** — Paleobiology Database, nejdůležitější pro fosilní taxony
- **GBIF** — Global Biodiversity Information Facility, rozsáhlá globální DB
- **IRMNG** — Interim Register of Marine and Nonmarine Genera
- **Fossilworks** — veřejné rozhraní PaleoDB

**Výsledky:**
- 🟢 **high** — nalezen ve ≥75 % vybraných databázích
- 🟡 **medium** — nalezen v 50–74 %
- 🟠 **low** — nalezen v <50 %
- 🔴 **not found** — nenalezen nikde

**Tipy:**
- Použij **preset** "Hyolitha (doporučeno)" pro předdefinovaný výběr databází
- **⚡ Paralelně** výrazně urychlí validaci — HTTP dotazy na DB jsou nezávislé
- **Fuzzy oprava** automaticky opraví OCR záměny (0/O, l/I, rn/m)
- **Offline DB** — lokální SQLite záloha pro práci bez internetu
            """,
            """
**Validation** checks each taxonomic name against selected online databases and shows results with direct links.

**Paleontology databases** (recommended):
- **PaleoDB** — Paleobiology Database, the key resource for fossil taxa
- **GBIF** — Global Biodiversity Information Facility, broad global DB
- **IRMNG** — Interim Register of Marine and Nonmarine Genera
- **Fossilworks** — public interface to PaleoDB

**Results:**
- 🟢 **high** — found in ≥75 % of selected DBs
- 🟡 **medium** — found in 50–74 %
- 🟠 **low** — found in <50 %
- 🔴 **not found** — not found anywhere

**Tips:**
- Use the **Hyolitha (recommended)** preset for a pre-defined DB selection
- **⚡ Parallel** dramatically speeds up validation — HTTP queries are independent
- **Fuzzy fix** auto-corrects OCR confusions (0/O, l/I, rn/m)
- **Offline DB** — local SQLite fallback for work without internet
            """,
            _L
        ))

    _v_src_map = {
        "📝 Ruční":      tt("📝 Ruční",      "📝 Manual",         _L),
        "📄 Ze souboru": tt("📄 Ze souboru", "📄 From file",      _L),
        "🔗 Z extrakce": tt("🔗 Z extrakce", "🔗 From extraction", _L),
    }
    input_mode_v = st.radio(tt("Zdroj taxonů", "Taxon source", _L),
                            list(_v_src_map.keys()),
                            format_func=lambda k: _v_src_map[k],
                            horizontal=True)
    taxa_list: List[str] = []

    if input_mode_v == "📝 Ruční":
        st.caption(tt(
            "💡 Wildcards: `?` = jeden libovolný znak, `*` = libovolný počet znaků  ·  "
            "**Podporují wildcards + genus expansion:** WoRMS, IRMNG, PaleoDB, Fossilworks, ITIS, CoL, GBIF, IPNI, Tropicos, IFPNI  ·  "
            "**Bez podpory** (použije čistý název): ZooBank, Plazi, BioLib, Mikrotax",
            "💡 Wildcards: `?` = any single char, `*` = any number of chars  ·  "
            "**Wildcard + genus expansion:** WoRMS, IRMNG, PaleoDB, Fossilworks, ITIS, CoL, GBIF, IPNI, Tropicos, IFPNI  ·  "
            "**No support** (uses clean name): ZooBank, Plazi, BioLib, Mikrotax",
            _L))
        # Tlačítko Smazat
        tc1, tc2 = st.columns([6, 1])
        with tc1:
            taxa_input = st.text_area(tt("Taxony (jeden na řádek)",
                                         "Taxa (one per line)", _L), height=150,
                                      key="val_taxa_manual",
                                      placeholder=("Hyolithes acutus\nHyolith*\nAlfaites rome?\n*itheca"))
        with tc2:
            st.write("")
            st.write("")
            if st.button(tt("🗑️ Smazat", "🗑️ Clear", _L), key="val_taxa_clear",
                         help=tt("Vymaže celé pole taxonů",
                                 "Clears the entire taxa field", _L)):
                # Nastavíme příznak — widget key se vymaže v příštím runu PŘED renderem widgetu
                st.session_state["_val_clear_pending"] = True
                st.session_state["_gsb_jump_to_tab"] = 2
                st.query_params["tab"] = "2"
                st.rerun()
        # Synchronizuj session state (widget → náš stav)
        if taxa_input != st.session_state.get("val_taxa_input",""):
            st.session_state["val_taxa_input"] = taxa_input
        taxa_list = [t.strip() for t in taxa_input.split("\n") if t.strip()]
        # Vizuálně označit wildcard dotazy
        wc_taxa = [tx for tx in taxa_list if "?" in tx or "*" in tx]
        if wc_taxa:
            _wc_list = ', '.join(wc_taxa[:5]) + (" …" if len(wc_taxa) > 5 else "")
            st.info(tt(f"🔍 Wildcard dotazy ({len(wc_taxa)}): {_wc_list}",
                       f"🔍 Wildcard queries ({len(wc_taxa)}): {_wc_list}", _L))

    elif input_mode_v == "📄 Ze souboru":
        v_file = st.file_uploader(tt("TXT nebo CSV", "TXT or CSV", _L),
                                  type=["txt","csv"], key="v_file")
        if v_file:
            if v_file.name.endswith(".csv"):
                df_v = pd.read_csv(v_file)
                col_sel = st.selectbox(tt("Sloupec s taxony", "Taxon column", _L), df_v.columns)
                taxa_list = df_v[col_sel].dropna().astype(str).str.strip().tolist()
            else:
                taxa_list = [l.strip() for l in
                             v_file.read().decode("utf-8").split("\n") if l.strip()]
            st.info(tt(f"Načteno {len(taxa_list)} taxonů",
                       f"Loaded {len(taxa_list)} taxa", _L))
    else:
        if st.session_state["last_extraction_taxa"]:
            _n = len(st.session_state['last_extraction_taxa'])
            edited_taxa = st.text_area(
                tt(f"Taxony z extrakce ({_n} nalezeno) – uprav",
                   f"Taxa from extraction ({_n} found) – edit", _L),
                "\n".join(st.session_state["last_extraction_taxa"]), height=200)
            taxa_list = [t.strip() for t in edited_taxa.split("\n") if t.strip()]
        else:
            st.warning(tt("Nejprve proveď extrakci.", "Run extraction first.", _L))
            taxa_list = []

    # ── Duplicate checker ──────────────────────────
    if taxa_list:
        dup_groups = find_duplicate_taxa(taxa_list)
        if dup_groups:
            with st.expander(tt(f"⚠️ Nalezeny možné duplikáty / překlepy ({len(dup_groups)} skupin)",
                                f"⚠️ Possible duplicates / typos found ({len(dup_groups)} groups)", _L)):
                for grp in dup_groups:
                    variants = " | ".join(f"**{t}**" if t == grp["canonical"] else t
                                          for t in grp["duplicates"])
                    st.markdown(
                        f'<div class="dup-warn">🔁 {variants} '
                        f'<small>(canonical: <b>{grp["canonical"]}</b>)</small></div>',
                        unsafe_allow_html=True)
                if st.checkbox(tt("Automaticky deduplikovat (ponechat canonical)",
                                  "Auto-deduplicate (keep canonical)", _L), key="auto_dedup"):
                    seen_norm = set()
                    deduped   = []
                    for t in taxa_list:
                        n = _normalize_for_dup(t)
                        if n not in seen_norm:
                            seen_norm.add(n); deduped.append(t)
                    taxa_list = deduped
                    st.success(tt(f"Deduplikováno: {len(taxa_list)} unikátních taxonů",
                                  f"Deduplicated: {len(taxa_list)} unique taxa", _L))
        else:
            st.caption(tt("✅ Žádné podezřelé duplikáty nalezeny.",
                          "✅ No suspicious duplicates found.", _L))

    # ── Databáze – výběr ve 4 sloupcích, abecedně ──────
    st.markdown(tt("#### 🗄️ Databáze", "#### 🗄️ Databases", _L))

    # Tlačítka: Označ vše / Vymaž vše
    btn_col1, btn_col2, btn_col3 = st.columns([1, 1, 6])
    with btn_col1:
        if st.button(tt("☑️ Vše", "☑️ All", _L), key="db_sel_all",
                     help=tt("Označ všechny databáze", "Select all databases", _L)):
            for db_key in TAXONOMIC_DATABASES:
                st.session_state[f"db_{db_key}"] = True
            st.session_state["_gsb_jump_to_tab"] = 2
            st.query_params["tab"] = "2"
            st.rerun()
    with btn_col2:
        if st.button(tt("☐ Nic", "☐ None", _L), key="db_sel_none",
                     help=tt("Zrušit výběr všech", "Clear selection", _L)):
            for db_key in TAXONOMIC_DATABASES:
                st.session_state[f"db_{db_key}"] = False
            st.session_state["_gsb_jump_to_tab"] = 2
            st.query_params["tab"] = "2"
            st.rerun()

    # Checkboxy seřazené abecedně ve 4 sloupcích
    # Inicializace session state pro DB checkboxy (jednou, pokud klíč neexistuje)
    db_keys_sorted = sorted(TAXONOMIC_DATABASES.keys())
    for db_key in db_keys_sorted:
        ss_key = f"db_{db_key}"
        if ss_key not in st.session_state:
            st.session_state[ss_key] = TAXONOMIC_DATABASES[db_key].get("default", False)

    cols_db = st.columns(4)
    selected_dbs: List[str] = []
    for i, db_key in enumerate(db_keys_sorted):
        db_info = TAXONOMIC_DATABASES[db_key]
        with cols_db[i % 4]:
            # BEZ value= parametru — Streamlit čte stav přímo ze session_state přes key=
            checked = st.checkbox(
                f"{db_info['icon']} {db_info['name']}",
                key=f"db_{db_key}",
                help=db_info.get("desc",""))
            if checked:
                selected_dbs.append(db_key)

    max_taxa      = st.slider(tt("Max. taxonů", "Max. taxa", _L), 5, 500, 100, 5)
    fcol1, fcol2, fcol3 = st.columns([1,2,2])
    with fcol1:
        parallel_val  = st.checkbox(tt("⚡ Paralelní", "⚡ Parallel", _L), value=True,
                                    help=tt("Paralelní validace (rychlejší)",
                                            "Parallel validation (faster)", _L))
    with fcol2:
        use_fuzzy_fix = st.checkbox(
            tt("🔤 Fuzzy OCR korekce", "🔤 Fuzzy OCR fix", _L),
            value=False,
            help=tt("Automaticky opraví časté OCR záměny (0/O, l/I, rn/m) v taxonomických jménech",
                    "Auto-corrects common OCR confusions (0/O, l/I, rn/m) in taxon names", _L))
    with fcol3:
        _odc = offline_db_count()
        use_offline_fallback = st.checkbox(
            tt(f"🦕 Offline fallback ({_odc:,} taxonů)",
               f"🦕 Offline fallback ({_odc:,} taxa)", _L),
            value=_odc > 0, disabled=_odc == 0, key="val_offline_fallback",
            help=tt("Nejprve zkusí lokální SQLite DB, teprve pak jde na internet. "
                    "Rychlé pro časté hyolitové taxony.",
                    "Tries local SQLite DB first, then falls back to the internet. "
                    "Fast for common hyolith taxa.", _L))

    # ── Responsivní checkboxy (JS detekce šířky okna) ──
    st.html("""
<script>
(function() {
  function adjustDbCols() {
    var w = window.innerWidth;
    var cols = w < 800 ? 2 : w < 1200 ? 3 : 4;
    // Streamlit re-renders automatically; zde jen evidujeme šířku
    if (window._lastWidth !== w) {
      window._lastWidth = w;
    }
  }
  adjustDbCols();
  window.addEventListener('resize', adjustDbCols);
})();
</script>
""")

    # ── Fuzzy preview ──────────────────────────────────────
    if use_fuzzy_fix and taxa_list:
        all_fixes = []
        fixed_taxa = []
        for tx in taxa_list[:max_taxa]:
            fixed, fixes = fuzzy_fix_taxon_name(tx)
            fixed_taxa.append(fixed)
            if fixes:
                all_fixes.append((tx, fixed, fixes))
        if all_fixes:
            with st.expander(tt(f"🔤 Fuzzy opravy – {len(all_fixes)} jmen upraveno (náhled)",
                                f"🔤 Fuzzy fixes – {len(all_fixes)} names corrected (preview)", _L)):
                for orig, fixed, fixes_list in all_fixes[:20]:
                    st.markdown(f"- `{orig}` → **`{fixed}`** — {'; '.join(fixes_list)}")
                if len(all_fixes) > 20:
                    st.caption(tt(f"… a dalších {len(all_fixes)-20} oprav",
                                  f"… and {len(all_fixes)-20} more fixes", _L))
        else:
            st.success(tt("✅ Fuzzy kontrola: žádné OCR chyby nenalezeny",
                          "✅ Fuzzy check: no OCR errors found", _L))
    else:
        fixed_taxa = taxa_list[:max_taxa]

    # ── Stop / Pause tlačítka validace — viditelná VŽDY pokud validace běží/pozastavena (v24.31) ──
    _val_is_running = st.session_state.get("validation_running", False)
    _val_is_paused  = st.session_state.get("validation_paused", False)
    if _val_is_running or _val_is_paused:
        with st.container():
            _vc1, _vc2, _vc3 = st.columns([1, 1, 4])
            with _vc1:
                if st.button(tt("⏹️ Zastavit validaci", "⏹️ Stop validation", _L),
                             key="btn_stop_validation_outer", type="secondary"):
                    st.session_state["validation_running"] = False
                    st.session_state["validation_paused"]  = False
                    keep_current_tab(2)
                    st.session_state["_gsb_jump_to_tab"] = 2
                    st.rerun()
            with _vc2:
                if st.button(
                    tt("▶️ Pokračovat" if _val_is_paused else "⏸️ Pozastavit",
                       "▶️ Resume"     if _val_is_paused else "⏸️ Pause", _L),
                    key="btn_pause_validation_outer", type="secondary"):
                    if _val_is_paused:
                        st.session_state["validation_paused"] = False
                        st.session_state["validation_running"] = True
                    else:
                        st.session_state["validation_paused"] = True
                    keep_current_tab(2)
                    st.session_state["_gsb_jump_to_tab"] = 2
                    st.rerun()
            with _vc3:
                if _val_is_paused:
                    st.warning(tt("⏸️ Validace pozastavena — klikni Pokračovat",
                                  "⏸️ Validation paused — click Resume", _L))
                elif _val_is_running:
                    st.info(tt("🧬 Validace probíhá…", "🧬 Validation running…", _L))

    if st.button(tt("🚀 Spustit validaci", "🚀 Run validation", _L), type="primary",
                 disabled=not (taxa_list and selected_dbs)):
        to_validate = (fixed_taxa if use_fuzzy_fix else taxa_list[:max_taxa])
        results_v   = []

        # ── Automaticky vymazat CELOU cache před každou validací ──
        st.session_state["validation_cache"] = {}
        try:
            _disk_cache_clear()
        except Exception:
            pass

        # Aktivuj globální status bar
        gsb_start(
            tt(f"🧬 Validace {len(to_validate)} taxonů",
               f"🧬 Validating {len(to_validate)} taxa", _L),
            tab=t("tab_validate", _L)
        )

        # ── Stop/Pause tlačítka jsou nyní renderována VŽDY nad run button (v24.31) ──
        st.session_state["validation_running"] = True
        st.session_state["validation_paused"]  = False

        with st.status(tt(f"🔍 Validuji {len(to_validate)} taxonů v {len(selected_dbs)} databázích…",
                          f"🔍 Validating {len(to_validate)} taxa in {len(selected_dbs)} databases…", _L),
                       expanded=True) as val_status:
            prog_v = st.progress(0)

            _use_off = st.session_state.get("val_offline_fallback", False)
            # Udržuj _temp/ čistý před novou validací
            _temp_prune("validace", keep_last=3)
            # Snapshot validation_cache před spuštěním threadů (dict je thread-safe pro čtení)
            _val_cache_snap = dict(st.session_state.get("validation_cache", {}))
            if parallel_val and len(to_validate) > 1:
                from concurrent.futures import ThreadPoolExecutor, as_completed
                # HTTP requesty na DB: paralelizujeme na taxon úrovni
                # Vnitřně validate_taxon_name již paralelizuje per-DB (8 workerů)
                # Proto omezíme na max 4 souběžné taxony — celkem max 32 HTTP spojení
                _val_workers = min(4, len(to_validate))
                _val_result_lock = _threading.Lock()
                _val_done_count  = [0]
                futures = {}
                with ThreadPoolExecutor(max_workers=_val_workers) as pool:
                    for tx in to_validate:
                        fut = pool.submit(validate_taxon_name, tx, selected_dbs,
                                          _val_cache_snap, _use_off)
                        futures[fut] = tx
                    for fut in as_completed(futures):
                        try:
                            _vr = fut.result()
                        except Exception as _ve:
                            _vr = {"taxon": futures[fut],
                                   "summary": {"found": 0, "total": len(selected_dbs)},
                                   "results": {}, "error": str(_ve)}
                        with _val_result_lock:
                            results_v.append(_vr)
                            _val_done_count[0] += 1
                            done = _val_done_count[0]
                        pct = done / len(to_validate)
                        try:
                            prog_v.progress(pct,
                                text=f"✅ {futures[fut]}  ({done}/{len(to_validate)})")
                            val_status.write(f"{'✅' if _vr['summary']['found']>0 else '❌'} "
                                             f"**{futures[fut]}** — "
                                             f"{_vr['summary']['found']}/{_vr['summary']['total']} DB")
                        except Exception:
                            pass
                        # ── _temp uložení po každých 10 taxonech nebo na konci ──
                        if done % 10 == 0 or done == len(to_validate):
                            try:
                                _temp_save_json(
                                    "validace", done,
                                    [{"taxon": r["taxon"],
                                      "found": r["summary"]["found"],
                                      "total": r["summary"]["total"],
                                      "results": r.get("results", {})}
                                     for r in results_v],
                                    len(to_validate)
                                )
                            except Exception:
                                pass
                order = {t: i for i, t in enumerate(to_validate)}
                results_v.sort(key=lambda r: order.get(r["taxon"], 999))
                # Batch flush nových výsledků do disk cache (1 transakce místo N)
                _batch_cache_items = []
                for _rv in results_v:
                    _ck = _rv["taxon"] + '::' + ','.join(sorted(selected_dbs))
                    _batch_cache_items.append((_ck, _rv))
                _disk_cache_set_batch(_batch_cache_items)
            else:
                for i, tx in enumerate(to_validate):
                    prog_v.progress((i+1)/len(to_validate), text=f"🔍 {tx}  ({i+1}/{len(to_validate)})")
                    val_status.write(f"🔍 Vyhledávám **{tx}**…")
                    res = validate_taxon_name(tx, selected_dbs,
                                              cache=st.session_state["validation_cache"],
                                              use_offline=_use_off)
                    results_v.append(res)
                    val_status.write(f"{'✅' if res['summary']['found']>0 else '❌'} "
                                     f"**{tx}** — {res['summary']['found']}/{res['summary']['total']} DB")
                    # ── _temp uložení po každých 10 taxonech nebo na konci ──
                    if (i + 1) % 10 == 0 or (i + 1) == len(to_validate):
                        try:
                            _temp_save_json(
                                "validace", i + 1,
                                [{"taxon": r["taxon"],
                                  "found": r["summary"]["found"],
                                  "total": r["summary"]["total"],
                                  "results": r.get("results", {})}
                                 for r in results_v],
                                len(to_validate)
                            )
                        except Exception:
                            pass

            val_status.update(
                label=f"✅ Hotovo — nalezeno {sum(1 for r in results_v if r['summary']['found']>0)}"
                      f"/{len(results_v)} taxonů",
                state="complete", expanded=False)

        # Ulož výsledky do session state — zobrazí se níže mimo button blok
        _found_n_done = sum(1 for r in results_v if r["summary"]["found"] > 0)
        gsb_done_stay_on_tab(f"✅ Validace: {_found_n_done}/{len(results_v)} nalezeno"
                 if _L == "cz" else
                 f"✅ Validation: {_found_n_done}/{len(results_v)} found",
                 tab_index=2)
        st.session_state["last_validation_results"] = results_v
        st.session_state["val_selected_dbs"]        = selected_dbs
        st.session_state["_gsb_jump_to_tab"]        = 2
        # ── Okamžitý autosave validace na disk ───────────────────────
        save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
        # Úklid _temp validace — vše hotovo, výsledky uloženy na disk
        try:
            _temp_cleanup("validace")
        except Exception:
            pass
        st.query_params["tab"] = "2"
        st.rerun()

    # ── Zobrazení výsledků (MIMO button blok – přetrvává po rerunu) ──
    results_v   = st.session_state.get("last_validation_results", [])

    # Detekce starých výsledků bez db_result dat (bug fix v16.3+)
    if results_v:
        broken = any(
            len(r.get("results", {})) == 0
            for r in results_v
        )
        if broken:
            st.warning(
                tt("⚠️ Výsledky jsou z cache před opravou bugu — databázové detaily nejsou k dispozici. "
                   "Klikni **🗑️ Vymazat cache** v sidebaru a spusť validaci znovu.",
                   "⚠️ Results are from pre-bugfix cache — DB details unavailable. "
                   "Click **🗑️ Clear cache** in the sidebar and re-run validation.", _L),
                icon="⚠️")

    selected_dbs_disp = st.session_state.get("val_selected_dbs", selected_dbs)

    if results_v:
        found_count    = sum(1 for r in results_v if r["summary"]["found"] > 0)
        notfound_count = len(results_v) - found_count

        # ── Metriky ─────────────────────────────────────
        m1, m2, m3 = st.columns(3)
        m1.metric(tt("✅ Nalezeno (alespoň 1 DB)", "✅ Found (≥ 1 DB)", _L), found_count)
        m2.metric(tt("❌ Nenalezeno", "❌ Not found", _L), notfound_count)
        m3.metric(tt("🗄️ Databáze", "🗄️ Databases", _L), len(selected_dbs_disp))

        # ── Batch re-validate jen nenalezené ────────────
        not_found_taxa = [r["taxon"] for r in results_v if r["summary"]["found"] == 0]
        if not_found_taxa:
            if st.button(tt(f"🔄 Znovu validovat nenalezené ({len(not_found_taxa)})",
                            f"🔄 Re-validate not-found ({len(not_found_taxa)})", _L),
                         key="revalidate_notfound"):
                new_results = []
                with st.status(tt(f"🔍 Revaliduji {len(not_found_taxa)} nenalezených…",
                                  f"🔍 Re-validating {len(not_found_taxa)} not-found…", _L),
                               expanded=True) as rv_status:
                    prog_rv = st.progress(0)
                    for i, tx in enumerate(not_found_taxa):
                        cache_key_rv = tx + '::' + ','.join(sorted(selected_dbs_disp))
                        st.session_state["validation_cache"].pop(cache_key_rv, None)
                        res = validate_taxon_name(tx, selected_dbs_disp,
                                                   cache=st.session_state["validation_cache"])
                        new_results.append(res)
                        prog_rv.progress((i+1)/len(not_found_taxa))
                        rv_status.write(f"{'✅' if res['summary']['found']>0 else '❌'} **{tx}**")
                    rv_status.update(label=tt("✅ Hotovo", "✅ Done", _L), state="complete")
                new_map = {r["taxon"]: r for r in new_results}
                updated = [new_map.get(r["taxon"], r) for r in results_v]
                st.session_state["last_validation_results"] = updated
                st.session_state["_gsb_jump_to_tab"] = 2
                st.query_params["tab"] = "2"
                save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
                st.rerun()

        # ── Přehledová tabulka ───────────────────────────
        df_val = validation_results_to_df(results_v)
        st.dataframe(df_val, width='stretch')

        st.markdown("---")
        st.markdown(tt("#### 📋 Podrobné výsledky", "#### 📋 Detailed results", _L))

        # Konstanty
        CONF_COLOR = {"high":"#059669","medium":"#d97706","low":"#4f46e5","not_found":"#dc2626"}
        CONF_BG    = {"high":"#f0fdf4","medium":"#fffbeb","low":"#eef2ff","not_found":"#fef2f2"}
        CONF_ICON  = {"high":"🟢","medium":"🟡","low":"🟠","not_found":"🔴"}
        HIER_KEYS  = ["kingdom","phylum","class","order","family","genus","species","rank","status"]
        DB_PREF    = ("PaleoDB","GBIF","IRMNG","WoRMS","CoL","ITIS",
                      "ZooBank","Fossilworks","Plazi","BioLib","IFPNI","Tropicos")

        for idx, r in enumerate(results_v):
            conf    = r["summary"].get("confidence", "not_found")
            found_n = r["summary"]["found"]
            total_n = r["summary"]["total"]
            taxon   = r["taxon"]

            # Collect details from found DBs
            all_details = {}
            for db, info in r["results"].items():
                det = info.get("details") or {}
                if info.get("found") and det:
                    all_details[db] = det

            # Best taxonomic position
            best_det = {}
            for pref in DB_PREF:
                if pref in all_details and all_details[pref]:
                    best_det = all_details[pref]
                    break
            if not best_det and all_details:
                best_det = next(iter(all_details.values()))

            # Confidence styling
            conf_colors = {
                "high":      ("#059669", "#f0fdf4", "🟢"),
                "medium":    ("#d97706", "#fffbeb", "🟡"),
                "low":       ("#4f46e5", "#eef2ff", "🟠"),
                "not_found": ("#dc2626", "#fef2f2", "🔴"),
            }
            border_c, bg_c, conf_ico = conf_colors.get(conf, conf_colors["not_found"])

            # Taxonomy line
            tax_parts = []
            for k in HIER_KEYS:
                v = str(best_det.get(k, "")).strip()
                if v and v.lower() not in ("", "-", "?", "none", "unknown", "null"):
                    tax_parts.append(f"{k}: {v}")
            tax_line = "  ·  ".join(tax_parts)

            # ── Záhlaví taxonu (st.markdown — funguje spolehlivě) ──
            st.markdown(
                f'<div style="border-left:5px solid {border_c};background:{bg_c};'
                f'border-radius:8px;padding:.6rem 1rem .4rem 1rem;margin:.5rem 0">'
                f'<span style="font-size:1.1em;font-weight:700">{conf_ico} {taxon}</span>'
                f'&nbsp;&nbsp;<span style="color:#6b7280;font-size:.85em">'
                f'{found_n}/{total_n} databází</span>'
                + (f'<br><span style="font-size:.8rem;color:#555;font-style:italic">'
                   f'{tax_line}</span>' if tax_line else '')
                + '</div>',
                unsafe_allow_html=True
            )

            # ── Tabulka DB výsledků — skutečné HTML hypertextové odkazy ──
            for db, info in r["results"].items():
                db_cfg  = TAXONOMIC_DATABASES.get(db, {})
                icon    = db_cfg.get("icon", "")
                url     = info.get("url", "")
                err     = info.get("error", "")
                found   = info.get("found", False)

                a_found  = 'color:#1d4ed8;font-weight:600;text-decoration:underline;text-underline-offset:2px'
                a_search = 'color:#6b7280;text-decoration:underline;text-underline-offset:2px'

                if found and url:
                    status_html = '<span style="color:#15803d;font-weight:600">&#x2705; nalezen</span>'
                    link_html   = f'<a href="{url}" target="_blank" style="{a_found}">&#x1F517; Otev&#345;&#237;t v datab&#225;zi</a>'
                    bg = "#f0fdf4"; border = "#059669"
                elif found:
                    status_html = '<span style="color:#15803d;font-weight:600">&#x2705; nalezen</span>'
                    link_html   = '<span style="color:#6b7280;font-style:italic">bez p&#345;&#237;m&#233;ho odkazu</span>'
                    bg = "#f0fdf4"; border = "#059669"
                elif err:
                    safe_err = err[:80].replace("<","&lt;").replace('"',"'")
                    status_html = '<span style="color:#b45309">&#x26A0;&#xFE0F; chyba</span>'
                    link_part = (f'&nbsp;<a href="{url}" target="_blank" style="{a_search}">&#x1F50D; Hledat</a>'
                                 if url else "")
                    link_html = f'<span style="color:#78716c;font-size:.82rem">{safe_err}</span>{link_part}'
                    bg = "#fefce8"; border = "#d97706"
                else:
                    status_html = '<span style="color:#9ca3af">&#x274C; nenalezen</span>'
                    link_html   = (f'<a href="{url}" target="_blank" style="{a_search}">&#x1F50D; Hledat v datab&#225;zi</a>'
                                   if url else "&#8212;")
                    bg = "#fafafa"; border = "#e5e7eb"

                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:0;background:{bg};'
                    f'border-left:3px solid {border};border-radius:4px;'
                    f'padding:5px 12px;margin:2px 0;font-size:.88rem">'
                    f'<span style="font-weight:600;min-width:160px;white-space:nowrap">{icon} {db}</span>'
                    f'<span style="min-width:130px">{status_html}</span>'
                    f'<span>{link_html}</span>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            # ── Podrobné strukturované detaily (expander) ──
            if all_details or r.get("pbdb_occs") is not None:
                with st.expander(tt(f"📋 Podrobné detaily — {taxon}",
                                    f"📋 Detailed info — {taxon}", _L)):
                    # PBDB occurrences
                    pbdb = r.get("pbdb_occs")
                    if pbdb is None and r["results"].get("PaleoDB", {}).get("found"):
                        pbdb = fetch_pbdb_occurrences(taxon)
                        r["pbdb_occs"] = pbdb
                    if pbdb:
                        p = [tt(f"**{pbdb['n_occs']}** výskytů v PaleoDB",
                                f"**{pbdb['n_occs']}** occurrences in PaleoDB", _L)]
                        if "strat_max_ma" in pbdb:
                            p.append(f"{pbdb['strat_min_ma']}–{pbdb['strat_max_ma']} Ma")
                        if "intervals" in pbdb:
                            p.append(pbdb["intervals"])
                        st.info("🦕 " + " · ".join(p))

                    for db_name, det in all_details.items():
                        db_url  = r["results"][db_name].get("url", "")
                        db_ico  = TAXONOMIC_DATABASES.get(db_name, {}).get("icon", "")
                        wc_hits = r["results"][db_name].get("wildcard_hits", [])
                        if db_url:
                            st.markdown(f'**{db_ico} {db_name}** — '
                                        f'[{tt("🔗 Otevřít", "🔗 Open", _L)}]({db_url})')
                        else:
                            st.markdown(f'**{db_ico} {db_name}**')
                        clean = {k: v for k, v in det.items()
                                 if v not in (None, "", {}, "unknown", "-", "none", "null")}
                        if clean:
                            cols = st.columns(3)
                            for ci, (k, v) in enumerate(list(clean.items())[:9]):
                                cols[ci % 3].caption(f"**{k}**: {v}")
                        # Wildcard: zobraz všechny nalezené záznamy
                        if wc_hits:
                            st.caption(tt(f"🔍 Wildcard: nalezeno {len(wc_hits)} záznamů (zobrazeno prvních 50)",
                                          f"🔍 Wildcard: found {len(wc_hits)} records (showing first 50)", _L))
                            _wc_rows = []
                            for h in wc_hits[:50]:
                                _wc_name   = h.get("name", "")
                                _wc_rank   = h.get("rank", "")
                                _wc_status = h.get("status", "")
                                _wc_id_key = next((k for k in ("pbdb_id","irmng_id","worms_id") if k in h), None)
                                _wc_id     = h.get(_wc_id_key, "") if _wc_id_key else ""
                                # Odkaz přímo na záznam pokud máme ID
                                _wc_link   = ""
                                if db_name == "PaleoDB" and _wc_id:
                                    _wc_id_clean = re.sub(r'\D', '', str(_wc_id))
                                    _wc_link = f"https://paleobiodb.org/classic/displayTaxonInfo?taxon_no={_wc_id_clean}"
                                elif db_name == "IRMNG" and _wc_id:
                                    _wc_link = f"https://www.irmng.org/aphia.php?p=taxdetails&id={_wc_id}"
                                elif db_name == "WoRMS" and _wc_id:
                                    _wc_link = f"https://www.marinespecies.org/aphia.php?p=taxdetails&id={_wc_id}"
                                _wc_rows.append({
                                    tt("Jméno","Name",_L): _wc_name,
                                    tt("Rank","Rank",_L): _wc_rank,
                                    tt("Status","Status",_L): _wc_status,
                                    "ID": _wc_id,
                                    tt("Odkaz","Link",_L): _wc_link,
                                })
                            if _wc_rows:
                                _wc_df = pd.DataFrame(_wc_rows)
                                # Zobraz jako HTML tabulku s klikacími odkazy
                                _wc_html = "<table style='font-size:.82rem;border-collapse:collapse;width:100%'>"
                                _wc_html += "<tr>" + "".join(f"<th style='padding:3px 8px;border-bottom:1px solid #e5e7eb;text-align:left'>{c}</th>" for c in _wc_df.columns) + "</tr>"
                                for _, row in _wc_df.iterrows():
                                    _wc_html += "<tr>"
                                    for col, val in row.items():
                                        if col == tt("Odkaz","Link",_L) and val:
                                            _wc_html += f"<td style='padding:2px 8px'><a href='{val}' target='_blank'>🔗</a></td>"
                                        else:
                                            _wc_html += f"<td style='padding:2px 8px;border-bottom:1px solid #f3f4f6'>{val}</td>"
                                    _wc_html += "</tr>"
                                _wc_html += "</table>"
                                st.markdown(_wc_html, unsafe_allow_html=True)
                        # GNverifier: zobraz preferované výsledky (ze zdrojů jako PaleoDB, CoL, GBIF)
                        gnv_pref = r["results"][db_name].get("gnv_preferred", []) if db_name == "GNverifier" else []
                        if gnv_pref:
                            st.caption(tt(f"🌐 GNverifier — výsledky z preferovaných zdrojů ({len(gnv_pref)}):",
                                          f"🌐 GNverifier — results from preferred sources ({len(gnv_pref)}):", _L))
                            _gnv_html = "<table style='font-size:.82rem;border-collapse:collapse;width:100%'>"
                            _gnv_html += "<tr>" + "".join(
                                f"<th style='padding:3px 8px;border-bottom:1px solid #e5e7eb;text-align:left'>{c}</th>"
                                for c in [tt("Jméno","Name",_L), tt("Zdroj","Source",_L),
                                          "Match", tt("Rodina","Family",_L), tt("Odkaz","Link",_L)]
                            ) + "</tr>"
                            for _gp in gnv_pref[:20]:
                                _gp_link = _gp.get("outlink","")
                                _gp_link_td = f"<a href='{_gp_link}' target='_blank'>🔗</a>" if _gp_link else "—"
                                _gnv_html += (
                                    f"<tr>"
                                    f"<td style='padding:2px 8px'>{_gp.get('name','')}</td>"
                                    f"<td style='padding:2px 8px'>{_gp.get('source','')[:30]}</td>"
                                    f"<td style='padding:2px 8px'>{_gp.get('match_type','')}</td>"
                                    f"<td style='padding:2px 8px'>{_gp.get('family','')}</td>"
                                    f"<td style='padding:2px 8px'>{_gp_link_td}</td>"
                                    f"</tr>"
                                )
                            _gnv_html += "</table>"
                            st.markdown(_gnv_html, unsafe_allow_html=True)
                        st.write("")

            # ── Genus expansion: hledat druhy v rodu ────────────────
            # Zobrazíme jen pro rod-level taxony (nebo pokud jméno je jednoslovné)
            _genus_name = taxon.split()[0] if taxon else ""
            _taxon_has_wc = ("?" in taxon or "*" in taxon)
            _is_genus_query = (len(taxon.split()) == 1 and not _taxon_has_wc)
            if _is_genus_query and _genus_name:
                st.divider()
                _ge_key = f"genus_expand_{taxon}_{idx}"
                _ge_res_key = f"genus_expand_results_{taxon}"
                _ge_btn_lbl = (tt(f"🔍 Hledat druhy rodu {_genus_name} (wildcard {_genus_name}*)",
                                  f"🔍 Search species in genus {_genus_name} (wildcard {_genus_name}*)", _L))
                if st.button(_ge_btn_lbl, key=_ge_key):
                    _wc_taxon = f"{_genus_name}*"
                    _ge_dbs   = [db for db in selected_dbs_disp
                                 if db in ("PaleoDB","IRMNG","WoRMS","GBIF","CoL","ITIS","Fossilworks")]
                    if _ge_dbs:
                        with st.spinner(tt(f"Hledám druhy v rodu {_genus_name}…",
                                           f"Searching species in genus {_genus_name}…", _L)):
                            _ge_result = validate_taxon_name(
                                _wc_taxon, _ge_dbs,
                                cache=st.session_state["validation_cache"])
                        # Sbírka všech wildcard hitů ze všech DB
                        _all_hits = {}
                        for _ge_db, _ge_info in _ge_result.get("results", {}).items():
                            _ge_hits = _ge_info.get("wildcard_hits", [])
                            for _h in _ge_hits:
                                _h_name = _h.get("name","")
                                if _h_name and _h_name not in _all_hits:
                                    _all_hits[_h_name] = {**_h, "db": _ge_db}
                        st.session_state[_ge_res_key] = sorted(_all_hits.values(), key=lambda x: x.get("name",""))
                        st.session_state["_gsb_jump_to_tab"] = 2
                        st.query_params["tab"] = "2"
                        st.rerun()

                # Zobraz uložené výsledky genus expansion
                _ge_saved = st.session_state.get(_ge_res_key)
                if _ge_saved:
                    st.caption(tt(f"Druhy rodu **{_genus_name}** — nalezeno {len(_ge_saved)} záznamů:",
                                  f"Species of genus **{_genus_name}** — found {len(_ge_saved)} records:", _L))
                    _ge_html = "<table style='font-size:.82rem;border-collapse:collapse;width:100%'>"
                    _ge_html += "<tr>" + "".join(
                        f"<th style='padding:3px 8px;border-bottom:1px solid #e5e7eb;text-align:left'>{c}</th>"
                        for c in [tt("Jméno","Name",_L), tt("Rank","Rank",_L),
                                  tt("Status","Status",_L), "DB", tt("Odkaz","Link",_L)]) + "</tr>"
                    for _h in _ge_saved[:100]:
                        _h_name = _h.get("name","")
                        _h_rank = _h.get("rank","")
                        _h_stat = _h.get("status","")
                        _h_db   = _h.get("db","")
                        # Odkaz z ID
                        _h_link = ""
                        if _h_db == "PaleoDB":
                            _h_taxno = re.sub(r"\D","",str(_h.get("pbdb_id","")))
                            if _h_taxno: _h_link = f"https://paleobiodb.org/classic/displayTaxonInfo?taxon_no={_h_taxno}"
                        elif _h_db == "IRMNG":
                            _h_id = _h.get("irmng_id","")
                            if _h_id: _h_link = f"https://www.irmng.org/aphia.php?p=taxdetails&id={_h_id}"
                        elif _h_db == "WoRMS":
                            _h_id = _h.get("worms_id","")
                            if _h_id: _h_link = f"https://www.marinespecies.org/aphia.php?p=taxdetails&id={_h_id}"
                        _link_td = f"<a href='{_h_link}' target='_blank'>🔗</a>" if _h_link else "—"
                        _ge_html += (f"<tr><td style='padding:2px 8px'>{_h_name}</td>"
                                     f"<td style='padding:2px 8px'>{_h_rank}</td>"
                                     f"<td style='padding:2px 8px'>{_h_stat}</td>"
                                     f"<td style='padding:2px 8px'>{_h_db}</td>"
                                     f"<td style='padding:2px 8px'>{_link_td}</td></tr>")
                    _ge_html += "</table>"
                    st.markdown(_ge_html, unsafe_allow_html=True)
                    if len(_ge_saved) > 100:
                        st.caption(tt(f"… zobrazeno 100 z {len(_ge_saved)} záznamů",
                                      f"… showing 100 of {len(_ge_saved)} records", _L))
                    if st.button(tt("🗑️ Smazat výsledky rodu", "🗑️ Clear genus results", _L),
                                 key=f"ge_clear_{taxon}_{idx}"):
                        st.session_state.pop(_ge_res_key, None)
                        st.session_state["_gsb_jump_to_tab"] = 2
                        st.query_params["tab"] = "2"
                        st.rerun()


        # ── Správa offline DB ─────────────────────────
        st.divider()
        with st.expander(tt(f"🦕 Offline DB – správa lokální databáze taxonů ({offline_db_count():,} záznamů)",
                            f"🦕 Offline DB – manage local taxon database ({offline_db_count():,} records)", _L)):
            st.caption(tt("Lokální SQLite databáze pro rychlou validaci bez internetu. "
                          "Naplň ji ručně nebo importem z validačních výsledků.",
                          "Local SQLite database for fast validation without internet. "
                          "Populate it manually or by importing validation results.", _L))
            odb_tab1, odb_tab2 = st.tabs([tt("📥 Import z výsledků", "📥 Import from results", _L),
                                           tt("✏️ Ruční přidání", "✏️ Manual add", _L)])
            with odb_tab1:
                if results_v:
                    found_results = [r for r in results_v if r["summary"]["found"] > 0]
                    st.caption(tt(f"Nalezených taxonů připravených k importu: {len(found_results)}",
                                  f"Found taxa ready to import: {len(found_results)}", _L))
                    if st.button(tt(f"📥 Importovat {len(found_results)} nalezených do offline DB",
                                    f"📥 Import {len(found_results)} found into offline DB", _L),
                                 key="odb_import_btn", disabled=not found_results):
                        to_insert = []
                        for r in found_results:
                            det = r["summary"].get("details", {})
                            to_insert.append({
                                "name":    r["taxon"],
                                "rank":    det.get("rank",""),
                                "status":  det.get("status",""),
                                "kingdom": det.get("kingdom",""),
                                "phylum":  det.get("phylum",""),
                                "class":   det.get("class",""),
                                "family":  det.get("family",""),
                                "source":  "validation_import",
                            })
                        offline_db_insert(to_insert)
                        st.success(tt(f"✅ Importováno {len(to_insert)} taxonů do offline DB",
                                      f"✅ Imported {len(to_insert)} taxa into offline DB", _L))
                        st.session_state["_gsb_jump_to_tab"] = 2
                        st.query_params["tab"] = "2"
                        st.rerun()
                else:
                    st.info(tt("Nejprve spusť validaci aby bylo co importovat.",
                               "Run validation first so there's something to import.", _L))
            with odb_tab2:
                odb_manual = st.text_area(
                    tt("Taxony k přidání (jeden per řádek, nebo JSON)",
                       "Taxa to add (one per line, or JSON)", _L),
                    height=100, key="odb_manual_input",
                    placeholder="Hyolithes obtusus\nHyolithes carinatus\n...")
                if st.button(tt("➕ Přidat do offline DB", "➕ Add to offline DB", _L),
                             key="odb_manual_btn",
                             disabled=not odb_manual.strip()):
                    lines_odb = [l.strip() for l in odb_manual.split("\n") if l.strip()]
                    offline_db_insert([{"name": l, "source": "manual"} for l in lines_odb])
                    st.success(tt(f"✅ Přidáno {len(lines_odb)} taxonů",
                                  f"✅ Added {len(lines_odb)} taxa", _L))
                    st.session_state["_gsb_jump_to_tab"] = 2
                    st.query_params["tab"] = "2"
                    st.rerun()

        # ── ZooBank name registration check ──────────
        st.divider()
        with st.expander(tt("🔖 ZooBank – ověření registrace nových druhů",
                            "🔖 ZooBank – verify registration of new species", _L)):
            st.caption(tt("Ověří, zda jsou nově popsané taxony registrovány v ZooBank (ICZN platnost).",
                          "Checks whether newly described taxa are registered in ZooBank (ICZN validity).", _L))
            zb_taxa_input = st.text_area(
                tt("Taxony k ověření (jeden per řádek)",
                   "Taxa to verify (one per line)", _L), height=100, key="zb_taxa_input",
                value="\n".join(
                    r["taxon"] for r in results_v
                    if r["summary"]["found"] > 0
                )[:500] if results_v else "")
            if st.button(tt("🔖 Ověřit v ZooBank", "🔖 Verify in ZooBank", _L), key="zb_check_btn",
                         disabled=not zb_taxa_input.strip()):
                zb_taxa = [t.strip() for t in zb_taxa_input.split("\n") if t.strip()]
                zb_results = {}
                with st.status(tt(f"Ověřuji {len(zb_taxa)} taxonů v ZooBank…",
                                   f"Verifying {len(zb_taxa)} taxa in ZooBank…", _L),
                                expanded=True) as zbs:
                    for zbt in zb_taxa:
                        res_zb = check_zoobank_registration(zbt)
                        zb_results[zbt] = res_zb
                        if res_zb and res_zb.get("found"):
                            zbs.write(tt(f"✅ {zbt} — registrován (LSID: {res_zb.get('lsid','')})",
                                         f"✅ {zbt} — registered (LSID: {res_zb.get('lsid','')})", _L))
                        elif res_zb and res_zb.get("found") is False:
                            zbs.write(tt(f"❌ {zbt} — NENÍ registrován v ZooBank",
                                         f"❌ {zbt} — NOT registered in ZooBank", _L))
                        else:
                            zbs.write(tt(f"⚠️ {zbt} — chyba: {res_zb.get('error','') if res_zb else 'N/A'}",
                                         f"⚠️ {zbt} — error: {res_zb.get('error','') if res_zb else 'N/A'}", _L))
                    zbs.update(label=tt("✅ ZooBank check dokončen",
                                        "✅ ZooBank check complete", _L), state="complete")
                st.session_state["zoobank_results"] = zb_results

            if st.session_state.get("zoobank_results"):
                zb_rows = []
                for taxon, res in st.session_state["zoobank_results"].items():
                    zb_rows.append({
                        tt("Taxon",        "Taxon",        _L): taxon,
                        tt("Registrován",  "Registered",   _L): "✅" if res and res.get("found") else "❌",
                        "LSID": res.get("lsid","") if res else "",
                        "URL":  res.get("url","")  if res else "",
                        tt("Datum pub.", "Pub. date", _L): res.get("pub_date","") if res else "",
                    })
                st.dataframe(pd.DataFrame(zb_rows), width='stretch')

        # ── DOI resolver ──────────────────────────────
        with st.expander(tt("🔗 DOI resolver – ověření a doplnění metadat",
                            "🔗 DOI resolver – verify and fill in metadata", _L)):
            st.caption(tt("Ověří DOI přes doi.org a doplní chybějící bibliografická metadata.",
                          "Verifies DOIs via doi.org and fills in missing bibliographic metadata.", _L))
            doi_input = st.text_area(tt("DOI ke kontrole (jeden per řádek)",
                                         "DOIs to check (one per line)", _L), height=80,
                                     key="doi_resolver_input",
                                     placeholder=("10.1016/j.palaeo.2004.01.003\nhttps://doi.org/10.1080/…"))
            if st.button(tt("🔗 Ověřit DOI", "🔗 Verify DOIs", _L), key="doi_resolve_btn",
                         disabled=not doi_input.strip()):
                dois = [d.strip() for d in doi_input.split("\n") if d.strip()]
                doi_rows = []
                with st.spinner(tt(f"Ověřuji {len(dois)} DOI…",
                                   f"Verifying {len(dois)} DOIs…", _L)):
                    for doi in dois:
                        meta = resolve_doi(doi)
                        if meta:
                            doi_rows.append({
                                "DOI":                         meta.get("doi", doi),
                                tt("Platný",  "Valid",   _L):  "✅" if meta.get("valid") else "❌",
                                tt("Titul",   "Title",   _L):  str(meta.get("title",""))[:60],
                                tt("Autor",   "Author",  _L):  str(meta.get("author",""))[:40],
                                tt("Rok",     "Year",    _L):  meta.get("year",""),
                                tt("Časopis", "Journal", _L):  str(meta.get("journal",""))[:40],
                            })
                        else:
                            doi_rows.append({
                                "DOI":                         doi,
                                tt("Platný", "Valid", _L):     "⚠️",
                                tt("Titul",  "Title", _L):     tt("Chyba spojení",
                                                                   "Connection error", _L),
                            })
                if doi_rows:
                    st.dataframe(pd.DataFrame(doi_rows), width='stretch')
                    st.download_button("⬇️ CSV", pd.DataFrame(doi_rows).to_csv(index=False).encode(),
                                       "doi_check.csv", key="doi_dl")

        # ── Hromadná editace validovaných taxonů ─────
        with st.expander(tt("✏️ Hromadná editace a re-validace",
                            "✏️ Bulk edit and re-validate", _L)):
            st.caption(tt("Oprav jméno taxonu a znovu ho zvaliduj bez restartu celé validace.",
                          "Fix a taxon name and re-validate without restarting the whole run.", _L))
            edit_rows = []
            _lbl_taxon_orig = tt("Taxon (původní)", "Taxon (original)", _L)
            _lbl_conf       = tt("Confidence",       "Confidence",       _L)
            _lbl_found      = tt("Nalezeno",         "Found",            _L)
            for ridx, r in enumerate(results_v):
                conf = r["summary"].get("confidence","not_found")
                edit_rows.append({
                    "idx": ridx,
                    _lbl_taxon_orig: r["taxon"],
                    _lbl_conf:       conf,
                    _lbl_found:      f"{r['summary']['found']}/{r['summary']['total']}",
                })
            df_edit = pd.DataFrame(edit_rows)
            edited_df = st.data_editor(
                df_edit[[_lbl_taxon_orig, _lbl_conf, _lbl_found]],
                column_config={_lbl_taxon_orig: st.column_config.TextColumn(width="large")},
                key="val_bulk_edit", width='stretch', num_rows="fixed")

            if st.button(tt("🔄 Re-validovat upravené řádky",
                            "🔄 Re-validate edited rows", _L), key="val_bulk_revalidate"):
                changed = []
                for i, (orig_row, new_row) in enumerate(
                        zip(df_edit[_lbl_taxon_orig], edited_df[_lbl_taxon_orig])):
                    if orig_row != new_row and new_row.strip():
                        changed.append((i, new_row.strip()))
                if changed:
                    with st.status(tt(f"Re-validuji {len(changed)} upravených taxonů…",
                                       f"Re-validating {len(changed)} edited taxa…", _L),
                                   expanded=True) as brs:
                        for idx, new_name in changed:
                            res = validate_taxon_name(new_name, selected_dbs_disp,
                                                      cache=st.session_state["validation_cache"])
                            results_v[idx] = res
                            brs.write(f"{'✅' if res['summary']['found']>0 else '❌'} {new_name}")
                        brs.update(label=tt("✅ Hotovo", "✅ Done", _L), state="complete")
                    st.session_state["last_validation_results"] = results_v
                    st.session_state["_gsb_jump_to_tab"] = 2
                    st.query_params["tab"] = "2"
                    save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
                    st.rerun()
                else:
                    st.info(tt("Žádné změny nebyly provedeny.", "No changes were made.", _L))

        # ── Export ────────────────────────────────────
        st.divider()
        st.markdown(tt("#### ⬇️ Export", "#### ⬇️ Export", _L))

        # Filtr exportu  (vnitřní hodnoty zachovány kvůli srovnávání níže)
        _exp_map = {
            "Vše":                tt("Vše",                "All",             _L),
            "Jen nalezené ✅":    tt("Jen nalezené ✅",    "Found only ✅",   _L),
            "Jen nenalezené ❌":  tt("Jen nenalezené ❌",  "Not found only ❌", _L),
        }
        exp_filter = st.radio(
            tt("Exportovat:", "Export:", _L),
            list(_exp_map.keys()),
            format_func=lambda k: _exp_map[k],
            horizontal=True, key="val_export_filter")
        if exp_filter == "Jen nalezené ✅":
            results_exp = [r for r in results_v if r["summary"]["found"] > 0]
        elif exp_filter == "Jen nenalezené ❌":
            results_exp = [r for r in results_v if r["summary"]["found"] == 0]
        else:
            results_exp = results_v

        st.caption(tt(f"Exportuje se {len(results_exp)} z {len(results_v)} taxonů.",
                      f"Exporting {len(results_exp)} of {len(results_v)} taxa.", _L))

        EXPORT_TAX_COLS = ["kingdom","phylum","class","order","family","genus",
                           "species","rank","status","extinct"]
        rows_exp = []
        for r in results_exp:
            conf = r["summary"].get("confidence","not_found")
            row  = {"Taxon": r["taxon"], "Confidence": conf,
                    "Nalezeno": r["summary"]["found"],
                    "Celkem_DB": r["summary"]["total"]}
            for db, info in r["results"].items():
                row[f"{db}_found"] = ("✅" if info.get("found")
                                      else ("⚠️" + info.get("error","")[:20]
                                            if "error" in info else "❌"))
                row[f"{db}_url"] = info.get("url","")
            det = r["summary"].get("details",{})
            for k in EXPORT_TAX_COLS:
                row[k] = det.get(k,"")
            # PBDB occurrences do exportu
            pbdb = r.get("pbdb_occs") or {}
            row["pbdb_n_occs"]    = pbdb.get("n_occs","")
            row["pbdb_strat_min"] = pbdb.get("strat_min_ma","")
            row["pbdb_strat_max"] = pbdb.get("strat_max_ma","")
            rows_exp.append(row)
        df_exp = pd.DataFrame(rows_exp)

        def make_val_txt(rv: List[Dict]) -> str:
            out = ["VALIDACE TAXONOMICKÝCH JMEN",
                   f"Datum: {datetime.now():%Y-%m-%d %H:%M}",
                   f"Databáze: {', '.join(selected_dbs_disp)}",
                   f"Taxonů: {len(rv)}", "="*60, ""]
            for r in rv:
                conf = r["summary"].get("confidence","not_found")
                out += [f"{'='*40}",
                        f"TAXON: {r['taxon']}",
                        f"Confidence: {conf.upper()} ({r['summary']['found']}/{r['summary']['total']} databází)"]
                det = r["summary"].get("details",{})
                if det:
                    out.append("Taxonomická pozice:")
                    for k in EXPORT_TAX_COLS:
                        if det.get(k):
                            out.append(f"  {k:12s}: {det[k]}")
                pbdb = r.get("pbdb_occs") or {}
                if pbdb:
                    out.append(f"  PBDB: {pbdb.get('n_occs','')} nálezů "
                               f"({pbdb.get('strat_min_ma','')}–{pbdb.get('strat_max_ma','')} Ma)")
                out.append("")
                out.append("Výsledky databází:")
                out.append(f"  {'Databáze':<22} {'Status':<12} URL")
                out.append(f"  {'-'*22} {'-'*12} {'-'*40}")
                for db, info in r["results"].items():
                    if info.get("found"):
                        status = "NALEZEN"
                        url = info.get("url","—")
                    elif info.get("error"):
                        status = "CHYBA"
                        url = info.get("error","")[:60]
                    else:
                        status = "NENALEZEN"
                        url = "—"
                    out.append(f"  {db:<22} {status:<12} {url}")
                out.append("")
            return "\n".join(out)

        ec1, ec2, ec3, ec4, ec5 = st.columns(5)
        with ec1:
            st.download_button("⬇️ CSV", df_exp.to_csv(index=False).encode("utf-8"),
                               "validace.csv","text/csv", width='stretch')
        with ec2:
            if XLSX_AVAILABLE:
                st.download_button("⬇️ Excel", to_xlsx_bytes(df_exp),
                                   "validace.xlsx", width='stretch')
        with ec3:
            st.download_button("⬇️ JSON", to_json_bytes(results_exp),
                               "validace.json", width='stretch')
        with ec4:
            st.download_button("⬇️ TXT", make_val_txt(results_exp).encode("utf-8"),
                               "validace.txt","text/plain", width='stretch')
        with ec5:
            _dwc_bytes = export_to_darwin_core(results_exp)
            if _dwc_bytes:
                st.download_button(
                    tt("⬇️ Darwin Core", "⬇️ Darwin Core", _L),
                    _dwc_bytes, "validace_dwc.csv", "text/csv",
                    help=tt("Export ve formátu Darwin Core Archive (GBIF standard)",
                            "Export in Darwin Core Archive format (GBIF standard)", _L),
                    width='stretch'
                )
        save_to_history("validace", {"count": len(results_v),
                                     "databases": selected_dbs_disp,
                                     "found": found_count})
        # gsb_done je voláno před st.rerun() výše — zde pouze fallback pro případ
        # kdy výsledky existují z minulé session a gsb_active je stale True
        if st.session_state.get("gsb_active") and results_v:
            _fv = sum(1 for r in results_v if r.get("summary",{}).get("found",0) > 0)
            gsb_done_stay_on_tab(f"✅ Validace: {_fv}/{len(results_v)} nalezeno"
                     if _L=="cz" else
                     f"✅ Validation: {_fv}/{len(results_v)} found",
                     tab_index=2)


# ══════════════════════════════════════════════════════
# 3 – DEEPSEEK CHAT
# ══════════════════════════════════════════════════════
with main_tabs[3]:
    keep_current_tab(3)
    _L = st.session_state.get("lang", "cz")
    st.markdown(tt("## 🤖 DeepSeek Chat",
                   "## 🤖 DeepSeek Chat", _L))
    st.caption(tt(
        "Přímé propojení na DeepSeek API (api.deepseek.com). OpenAI-kompatibilní. "
        "Klíč z [platform.deepseek.com](https://platform.deepseek.com). "
        "Klíč se ukládá **pouze do session** — nikdy na disk.",
        "Direct connection to DeepSeek API (api.deepseek.com). OpenAI-compatible. "
        "Key from [platform.deepseek.com](https://platform.deepseek.com). "
        "Key stored **in session only** — never to disk.",
        _L))

    # ── Konfigurace API ──────────────────────────────────────────────────────
    with st.expander(
            tt("🔑 API klíč, model a nastavení",
               "🔑 API key, model & settings", _L),
            expanded=not st.session_state.get("ds_api_key")):

        _ds_key_input = st.text_input(
            tt("DeepSeek API klíč", "DeepSeek API key", _L),
            value=st.session_state.get("ds_api_key", ""),
            type="password", placeholder="sk-...",
            key="ds_api_key_input",
            help=tt("Uloženo jen v session — není persistentní.",
                    "Stored in session only — not persistent.", _L))
        if _ds_key_input != st.session_state.get("ds_api_key", ""):
            st.session_state["ds_api_key"] = _ds_key_input

        # Aktuální modely k 2025/2026 — deepseek-v4-flash a deepseek-v4-pro
        # POZOR: deepseek-chat a deepseek-reasoner budou deprecated 2026-07-24
        # (nyní jsou aliasy pro deepseek-v4-flash non-thinking / thinking mode)
        _DS_MODELS = [
            "deepseek-v4-flash",        # rychlý, levný, non-thinking (ex deepseek-chat)
            "deepseek-v4-flash:thinking", # thinking mode V4-Flash (ex deepseek-reasoner)
            "deepseek-v4-pro",           # nejvýkonnější, 1M kontext
            "deepseek-v4-pro:thinking",  # thinking mode V4-Pro
            "deepseek-chat",             # legacy alias → v4-flash non-thinking (do 2026-07-24)
            "deepseek-reasoner",         # legacy alias → v4-flash thinking (do 2026-07-24)
        ]
        _dc1, _dc2, _dc3 = st.columns(3)
        with _dc1:
            _ds_model_default = st.session_state.get("ds_model", "deepseek-v4-flash")
            if _ds_model_default not in _DS_MODELS:
                _ds_model_default = "deepseek-v4-flash"
            _ds_model = st.selectbox(
                tt("Model", "Model", _L), options=_DS_MODELS,
                index=_DS_MODELS.index(_ds_model_default),
                key="ds_model_sel",
                help=tt(
                    "v4-flash: rychlý $0.14/M. v4-pro: nejsilnější $1.74/M. "
                    ":thinking = aktivuje Chain-of-Thought reasoning. "
                    "Legacy aliasy (chat/reasoner) zaniknou 2026-07-24.",
                    "v4-flash: fast $0.14/M. v4-pro: strongest $1.74/M. "
                    ":thinking = enables Chain-of-Thought reasoning. "
                    "Legacy aliases (chat/reasoner) retire 2026-07-24.",
                    _L))
            st.session_state["ds_model"] = _ds_model
        with _dc2:
            _ds_temp = st.slider(tt("Teplota", "Temperature", _L),
                                 0.0, 1.5,
                                 float(st.session_state.get("ds_temp", 0.7)), 0.05,
                                 key="ds_temp_sl")
            st.session_state["ds_temp"] = _ds_temp
        with _dc3:
            _ds_maxtok = st.number_input("Max tokens", 256, 131072,
                                         int(st.session_state.get("ds_max_tokens", 8192)),
                                         256, key="ds_maxtok_inp")
            st.session_state["ds_max_tokens"] = _ds_maxtok

        # Ceník (vstup cache miss / výstup) — aktuální k 2025/2026
        _DS_PRICING = {
            "deepseek-v4-flash":         ("$0.14",  "$0.28"),
            "deepseek-v4-flash:thinking": ("$0.14",  "$0.28"),
            "deepseek-v4-pro":           ("$1.74",  "$3.48"),
            "deepseek-v4-pro:thinking":  ("$1.74",  "$3.48"),
            "deepseek-chat":             ("$0.14",  "$0.28"),   # alias v4-flash
            "deepseek-reasoner":         ("$0.14",  "$0.28"),   # alias v4-flash thinking
        }
        _in_p, _out_p = _DS_PRICING.get(_ds_model, ("?", "?"))
        _is_legacy = _ds_model in ("deepseek-chat", "deepseek-reasoner")
        st.caption(tt(
            f"💰 {_in_p}/1M vstupních · {_out_p}/1M výstupních tokenů"
            + (" · ⚠️ legacy alias, zanikne 2026-07-24" if _is_legacy else ""),
            f"💰 {_in_p}/1M input · {_out_p}/1M output tokens"
            + (" · ⚠️ legacy alias, retires 2026-07-24" if _is_legacy else ""),
            _L))

    # ── Systémový prompt ─────────────────────────────────────────────────────
    _ds_sys_default = tt(
        "Jsi vědecký asistent specializovaný na paleontologii a taxonomii.",
        "You are a scientific assistant specialising in palaeontology and taxonomy.",
        _L)
    _ds_sysprompt = st.text_area(
        tt("Systémový prompt (volitelný)", "System prompt (optional)", _L),
        value=st.session_state.get("ds_sysprompt", _ds_sys_default),
        height=70, key="ds_sysprompt_area")
    st.session_state["ds_sysprompt"] = _ds_sysprompt

    # ── Injekce kontextu extrakce ────────────────────────────────────────────
    _dinj1, _dinj2 = st.columns([1, 3])
    with _dinj1:
        if st.button(tt("📎 Vložit extrakci", "📎 Inject extraction", _L),
                     key="ds_inject_ctx",
                     disabled=not st.session_state.get("last_extracted_text", "")):
            _ext_ctx = st.session_state.get("last_extracted_text", "")[:2000]
            st.session_state["ds_injected_context"] = (
                tt("Kontext z poslední extrakce:", "Context from last extraction:", _L)
                + f"\n\n{_ext_ctx}\n\n"
                + tt("Odpovídej na otázky o těchto datech.",
                     "Answer questions about this data.", _L))
            st.toast(tt("✅ Kontext vložen", "✅ Context injected", _L))
    with _dinj2:
        _ds_inj_ctx = st.session_state.get("ds_injected_context", "")
        if _ds_inj_ctx:
            st.caption(tt(f"📌 Kontext: {_ds_inj_ctx[:80]}…",
                          f"📌 Context: {_ds_inj_ctx[:80]}…", _L))
            if st.button(tt("✖️ Odebrat", "✖️ Remove", _L), key="ds_remove_ctx"):
                st.session_state["ds_injected_context"] = ""
                st.rerun()

    # ── Chat historie ────────────────────────────────────────────────────────
    if "ds_chat_history" not in st.session_state:
        st.session_state["ds_chat_history"] = []
    _ds_hist = st.session_state["ds_chat_history"]

    _ds_chat_container = st.container(height=420, border=True)
    with _ds_chat_container:
        if not _ds_hist:
            st.caption(tt("Konverzace se zobrazí zde…",
                          "Conversation will appear here…", _L))
        for _msg in _ds_hist:
            with st.chat_message(_msg["role"],
                                 avatar="🧬" if _msg["role"] == "user" else "🤖"):
                st.markdown(_msg["content"])

    # Statistiky + export
    if _ds_hist:
        _ds_turns = len([m for m in _ds_hist if m["role"] == "user"])
        _ds_chars = sum(len(m["content"]) for m in _ds_hist)
        _stat_col, _exp_col = st.columns([3, 1])
        with _stat_col:
            st.caption(tt(f"💬 {_ds_turns} otázek · {_ds_chars:,} znaků",
                          f"💬 {_ds_turns} turns · {_ds_chars:,} chars", _L))
        with _exp_col:
            _ds_export_txt = "\n\n".join(
                f"{'TY' if m['role']=='user' else 'DEEPSEEK'}:\n{m['content']}"
                for m in _ds_hist)
            st.download_button(
                tt("⬇️ Export", "⬇️ Export", _L),
                to_txt_bytes(_ds_export_txt),
                "deepseek_chat.txt", key="ds_exp_btn",
                use_container_width=True)

    # ── Vstup ────────────────────────────────────────────────────────────────
    _dsi1, _dsi2 = st.columns([6, 1])
    with _dsi1:
        _ds_user_input = st.text_area(
            tt("Zpráva", "Message", _L), height=80,
            key="ds_user_input",
            placeholder=tt("Napiš dotaz pro DeepSeek…",
                           "Enter your message for DeepSeek…", _L),
            label_visibility="collapsed")
    with _dsi2:
        _ds_send = st.button(
            "🚀 " + tt("Odeslat", "Send", _L),
            key="ds_send_btn", use_container_width=True,
            type="primary",
            disabled=not st.session_state.get("ds_api_key"))
        if st.button("🗑️ " + tt("Smazat", "Clear", _L),
                     key="ds_clear_btn", use_container_width=True):
            if st.session_state.get("ds_confirm_clear"):
                st.session_state["ds_chat_history"] = []
                st.session_state.pop("ds_confirm_clear", None)
                st.session_state.pop("ds_injected_context", None)
                st.rerun()
            else:
                st.session_state["ds_confirm_clear"] = True
        if st.session_state.get("ds_confirm_clear"):
            st.caption(tt("⚠️ Klikni znovu", "⚠️ Click again", _L))

    if not st.session_state.get("ds_api_key"):
        st.info(tt(
            "ℹ️ Zadej DeepSeek API klíč výše. "
            "Zdarma na platform.deepseek.com.",
            "ℹ️ Enter your DeepSeek API key above. "
            "Free at platform.deepseek.com.", _L))

    # ── Odeslání na DeepSeek API ─────────────────────────────────────────────
    # Dvoustupňový pattern:
    #   Klik Odeslat → stupeň 1: přidej zprávu do hist + _ds_waiting=True + rerun
    #                  stupeň 2: zobraz spinner + volej API + rerun s odpovědí
    # Tím uživatel okamžitě vidí svoji zprávu a spinner místo zmrzlé stránky.

    if st.session_state.pop("_ds_waiting", False) and st.session_state.get("ds_api_key"):
        # Stupeň 2: zpráva je v historii, teď volej API
        _ds_hist = st.session_state.get("ds_chat_history", [])
        _ds_api_key     = st.session_state["ds_api_key"]
        _ds_mod_raw     = st.session_state.get("ds_model", "deepseek-v4-flash")
        _ds_temperature = float(st.session_state.get("ds_temp", 0.7))
        _ds_max_t       = int(st.session_state.get("ds_max_tokens", 8192))
        _ds_sys         = st.session_state.get("ds_sysprompt", "").strip()
        _ds_inj         = st.session_state.get("ds_injected_context", "")
        _ds_thinking    = _ds_mod_raw.endswith(":thinking") or _ds_mod_raw == "deepseek-reasoner"
        _ds_mod         = _ds_mod_raw.replace(":thinking", "")
        _ds_messages    = []
        _full_sys       = (_ds_sys + "\n\n" + _ds_inj).strip() if _ds_inj else _ds_sys
        if _full_sys:
            _ds_messages.append({"role": "system", "content": _full_sys})
        _ds_messages.extend(_ds_hist[-20:])
        _spinner_lbl = tt(
            "DeepSeek přemýšlí… 🤔" if _ds_thinking else "DeepSeek odpovídá…",
            "DeepSeek is thinking… 🤔" if _ds_thinking else "DeepSeek is responding…", _L)
        with st.spinner(_spinner_lbl):
            try:
                import openai as _openai_ds
                _ds_client = _openai_ds.OpenAI(
                    api_key=_ds_api_key,
                    base_url="https://api.deepseek.com")
                _ds_kwargs = dict(model=_ds_mod, messages=_ds_messages,
                                  temperature=_ds_temperature, max_tokens=_ds_max_t,
                                  stream=False)
                if _ds_thinking:
                    _ds_kwargs["extra_body"] = {"thinking": {"type": "enabled"}}
                _ds_resp  = _ds_client.chat.completions.create(**_ds_kwargs)
                _ds_msg   = _ds_resp.choices[0].message
                _ds_answer = _ds_msg.content or ""
                _ds_reasoning = getattr(_ds_msg, "reasoning_content", None)
                if _ds_reasoning:
                    _ds_answer = (
                        tt("**🧠 Úvaha modelu:**", "**🧠 Model reasoning:**", _L)
                        + f"\n\n{_ds_reasoning}\n\n"
                        + tt("**💬 Odpověď:**", "**💬 Answer:**", _L)
                        + f"\n\n{_ds_answer}")
                _ds_hist.append({"role": "assistant", "content": _ds_answer})
            except ImportError:
                st.error(tt("❌ Chybí `openai`. Nainstaluj: `pip install openai`",
                            "❌ Missing `openai`. Install: `pip install openai`", _L))
                if _ds_hist and _ds_hist[-1]["role"] == "user":
                    _ds_hist.pop()
            except Exception as _ds_err:
                _s = str(_ds_err)
                if "401" in _s or "unauthorized" in _s.lower():
                    st.error(tt("❌ Neplatný API klíč (401).", "❌ Invalid API key (401).", _L))
                elif "402" in _s or "insufficient" in _s.lower():
                    st.error(tt("❌ Nedostatek kreditů (402). Doplň na platform.deepseek.com.",
                                "❌ Insufficient credits (402). Top up at platform.deepseek.com.", _L))
                elif "429" in _s:
                    st.error(tt("❌ Rate limit (429) — zkus za chvíli.",
                                "❌ Rate limit (429) — retry shortly.", _L))
                else:
                    st.error(f"❌ DeepSeek: {_s[:200]}")
                if _ds_hist and _ds_hist[-1]["role"] == "user":
                    _ds_hist.pop()
        st.session_state["ds_chat_history"] = _ds_hist
        st.rerun()

    elif _ds_send and _ds_user_input.strip() and st.session_state.get("ds_api_key"):
        # Stupeň 1: přidej zprávu + nastav čekací příznak + rerun
        _ds_hist.append({"role": "user", "content": _ds_user_input.strip()})
        st.session_state["ds_chat_history"] = _ds_hist
        st.session_state["_ds_waiting"] = True
        st.rerun()


# 4 – ČIŠTĚNÍ DAT
# ══════════════════════════════════════════════════════
with main_tabs[4]:
    keep_current_tab(4)   # ← v24.17
    _L = st.session_state.get("lang","cz")
    st.markdown(t("clean_title",_L))
    with st.expander(tt("ℹ️ Jak funguje čištění dat?",
                         "ℹ️ How does data cleaning work?", _L), expanded=False):
        st.markdown(tt(
            """
**Čištění dat** normalizuje hodnoty v CSV/Excel sloupci pomocí LLM — opravuje překlepy, sjednocuje formáty, odstraní artefakty.

**Postup:**
1. Nahraj CSV nebo Excel soubor
2. Vyber sloupec k normalizaci
3. Uprav nebo použij výchozí prompt
4. Klikni **▶️ Normalizovat sloupec**
5. Ve vizuální validaci **přijmi nebo odmítni** každou navrhovanou změnu
6. Klikni **✅ Aplikovat vybrané změny**

**Automatická detekce typů:**
- Rok (4 číslice, rozsah 1758–2025)
- GPS souřadnice
- Taxonomické jméno (začíná velkým písmenem)

**Detekce anomálií:** automaticky označí hodnoty mimo očekávaný rozsah

**Undo/Redo:** tlačítka ↩️ / ↪️ vrátí nebo zopakují poslední operaci
            """,
            """
**Data cleaning** normalizes values in a CSV/Excel column using the LLM — fixes typos, unifies formats, removes artifacts.

**How to:**
1. Upload a CSV or Excel file
2. Pick a column to normalize
3. Edit or use the default prompt
4. Click **▶️ Normalize column**
5. In visual validation, **accept or reject** each proposed change
6. Click **✅ Apply selected changes**

**Auto type detection:**
- Year (4 digits, range 1758–2025)
- GPS coordinates
- Taxon name (starts with a capital letter)

**Anomaly detection:** automatically flags values outside the expected range

**Undo/Redo:** ↩️ / ↪️ buttons revert or redo the last operation
            """, _L))

    # ── PROMPT – vždy viditelný, nahoře, před nahráním souboru ──
    st.markdown(tt("#### 🤖 LLM normalizace / čištění",
                   "#### 🤖 LLM normalization / cleaning", _L))
    clean_prompts_dict = st.session_state.get("clean_prompts", {})
    cp_names = list(clean_prompts_dict.keys())
    cp_cols = st.columns([3, 2, 1])
    _cp_custom = tt("— vlastní prompt —", "— custom prompt —", _L)
    with cp_cols[0]:
        cp_sel = st.selectbox(tt("📂 Uložené prompty", "📂 Saved prompts", _L),
                              [_cp_custom] + cp_names, key="cp_sel")
    with cp_cols[1]:
        cp_new_name = st.text_input(tt("Název pro uložení", "Name to save as", _L),
                                     key="cp_new_name",
                                    placeholder=tt("např. Taxony_v1", "e.g. Taxony_v1", _L))
    with cp_cols[2]:
        st.write(""); st.write("")
        save_cp_btn = st.button(tt("💾 Uložit", "💾 Save", _L), key="cp_save_btn")

    default_clean_prompt = (clean_prompts_dict[cp_sel]
                             if cp_sel != _cp_custom and cp_sel in clean_prompts_dict
                             else "Normalizuj taxonomické názvy: oprav překlepy, sjednoť formát 'Rod druh Autor rok'.")

    norm_prompt_c = st.text_area(
        tt("📝 Prompt – instrukce pro LLM (co a jak čistit / normalizovat)",
           "📝 Prompt – LLM instructions (what and how to clean / normalize)", _L),
        height=110, value=default_clean_prompt, key="norm_prompt_c")

    if save_cp_btn:
        name_to_save = cp_new_name.strip() or cp_sel
        if name_to_save and name_to_save != _cp_custom:
            st.session_state["clean_prompts"][name_to_save] = norm_prompt_c
            save_clean_prompts(st.session_state["clean_prompts"])
            st.success(tt(f"Prompt '{name_to_save}' uložen",
                          f"Prompt '{name_to_save}' saved", _L)); st.rerun()
        else:
            st.warning(tt("Zadej název pro uložení.",
                          "Enter a name to save as.", _L))

    if cp_names:
        with st.expander(tt("🗑️ Správa uložených promptů",
                             "🗑️ Manage saved prompts", _L)):
            cp_del_sel = st.selectbox(tt("Smazat prompt", "Delete prompt", _L),
                                       ["—"] + cp_names, key="cp_del_sel")
            if cp_del_sel != "—":
                if st.checkbox(tt(f"Opravdu smazat prompt '{cp_del_sel}'?",
                                  f"Really delete prompt '{cp_del_sel}'?", _L),
                               key="chk_del_cp"):
                    if st.button(tt("🗑️ Ano, smazat", "🗑️ Yes, delete", _L),
                                 key="cp_del_btn", type="secondary"):
                        st.session_state["clean_prompts"].pop(cp_del_sel, None)
                        save_clean_prompts(st.session_state["clean_prompts"])
                        st.success(tt(f"Prompt '{cp_del_sel}' smazán",
                                      f"Prompt '{cp_del_sel}' deleted", _L))
                        st.rerun()

    st.divider()

    # ── NAHRÁNÍ SOUBORU ──────────────────────────────
    data_file = st.file_uploader(t("clean_upload",_L), type=["csv","xlsx"], key="clean_file")
    if data_file:
        df_clean = (pd.read_csv(data_file) if data_file.name.endswith(".csv")
                    else pd.read_excel(data_file))

        # Undo/redo stack
        if "df_clean_history" not in st.session_state:
            st.session_state["df_clean_history"] = []
        if "df_clean_future"  not in st.session_state:
            st.session_state["df_clean_future"]  = []

        def _push_undo(df):
            st.session_state["df_clean_history"].append(df.copy())
            st.session_state["df_clean_future"] = []
            if len(st.session_state["df_clean_history"]) > 20:
                st.session_state["df_clean_history"].pop(0)

        undo_col, redo_col, _ = st.columns([1,1,6])
        with undo_col:
            if st.button(tt("↩️ Zpět", "↩️ Undo", _L), key="cl_undo",
                         disabled=not st.session_state["df_clean_history"]):
                st.session_state["df_clean_future"].append(df_clean.copy())
                df_clean = st.session_state["df_clean_history"].pop()
                st.session_state["df_clean_result"] = df_clean
                st.rerun()
        with redo_col:
            if st.button(tt("↪️ Vpřed", "↪️ Redo", _L), key="cl_redo",
                         disabled=not st.session_state["df_clean_future"]):
                st.session_state["df_clean_history"].append(df_clean.copy())
                df_clean = st.session_state["df_clean_future"].pop()
                st.session_state["df_clean_result"] = df_clean
                st.rerun()

        st.caption(t("clean_loaded",_L, r=len(df_clean), c=len(df_clean.columns)))
        st.dataframe(df_clean.head(20), width='stretch')

        # ── Detekce anomálií ─────────────────────────
        with st.expander(tt("🔍 Detekce anomálií", "🔍 Anomaly detection", _L)):
            anomalies = []
            for col in df_clean.columns:
                col_l = col.lower()
                ser = df_clean[col].dropna()
                # Rok: 4 číslice mimo rozsah 1758–2025
                if any(k in col_l for k in ("year","rok","datum","date")):
                    nums = pd.to_numeric(ser, errors="coerce").dropna()
                    bad = nums[(nums < 1758) | (nums > 2025)]
                    if not bad.empty:
                        anomalies.append(tt(
                            f"**{col}**: {len(bad)} hodnot mimo 1758–2025 "
                            f"(např. {list(bad.head(3).astype(int))})",
                            f"**{col}**: {len(bad)} values outside 1758–2025 "
                            f"(e.g. {list(bad.head(3).astype(int))})", _L))
                # GPS: lat mimo ±90, lon mimo ±180
                if any(k in col_l for k in ("lat","latitude","zeměpisná")):
                    nums = pd.to_numeric(ser, errors="coerce").dropna()
                    bad = nums[nums.abs() > 90]
                    if not bad.empty:
                        anomalies.append(tt(
                            f"**{col}**: {len(bad)} hodnot mimo ±90°",
                            f"**{col}**: {len(bad)} values outside ±90°", _L))
                if any(k in col_l for k in ("lon","lng","longitude")):
                    nums = pd.to_numeric(ser, errors="coerce").dropna()
                    bad = nums[nums.abs() > 180]
                    if not bad.empty:
                        anomalies.append(tt(
                            f"**{col}**: {len(bad)} hodnot mimo ±180°",
                            f"**{col}**: {len(bad)} values outside ±180°", _L))
                # Příliš dlouhé jméno (>100 znaků)
                if ser.dtype == object:
                    long_vals = ser[ser.str.len() > 100]
                    if not long_vals.empty:
                        anomalies.append(tt(
                            f"**{col}**: {len(long_vals)} hodnot delších než 100 znaků",
                            f"**{col}**: {len(long_vals)} values longer than 100 chars", _L))
            if anomalies:
                st.warning(tt(f"Nalezeno {len(anomalies)} typů anomálií:",
                              f"Found {len(anomalies)} types of anomalies:", _L))
                for a in anomalies:
                    st.markdown(f"- {a}")
            else:
                st.success(tt("✅ Žádné zjevné anomálie nenalezeny",
                              "✅ No obvious anomalies found", _L))

        # ── Automatická detekce typu sloupce ─────────
        with st.expander(tt("🤖 Automatická detekce typu sloupce a doporučení promptu",
                             "🤖 Auto-detect column type and recommend prompt", _L)):
            if df_clean is not None and not df_clean.empty:
                col_types = {}
                for col in df_clean.columns:
                    ser = df_clean[col].dropna()
                    cl  = col.lower()
                    if pd.to_numeric(ser.head(20), errors="coerce").notna().mean() > 0.8:
                        nums = pd.to_numeric(ser, errors="coerce").dropna()
                        if nums.between(1758,2025).mean() > 0.7:
                            col_types[col] = ("rok 📅", "Normalizuj jako 4místný rok (YYYY).")
                        elif nums.between(-90,90).mean() > 0.9:
                            col_types[col] = ("GPS lat 🌍", "Normalizuj jako desetinné číslo v rozsahu ±90.")
                        else:
                            col_types[col] = ("číslo 🔢", "Normalizuj jako číslo.")
                    elif ser.str.match(r'^[A-Z][a-z]+ [a-z]+', na=False).mean() > 0.5:
                        col_types[col] = ("taxonomické jméno 🦪",
                                          "Normalizuj jako 'Rod druh Autor, rok'. Oprav překlepy.")
                    else:
                        col_types[col] = ("text 📝", "Normalizuj a sjednoť formát.")
                for col, (typ, prompt_hint) in list(col_types.items())[:10]:
                    c1, c2 = st.columns([2,3])
                    c1.caption(f"**{col}**: {typ}")
                    if c2.button(tt("Použít prompt", "Use prompt", _L),
                                 key=f"use_col_prompt_{col}"):
                        st.session_state["_suggested_clean_prompt"] = prompt_hint
                        st.rerun()
                if st.session_state.get("_suggested_clean_prompt"):
                    st.info(tt(f"💡 Doporučený prompt: *{st.session_state['_suggested_clean_prompt']}*",
                               f"💡 Suggested prompt: *{st.session_state['_suggested_clean_prompt']}*", _L))

        ops1, ops2, ops3 = st.columns(3)
        with ops1:
            if st.button(t("clean_dedup",_L), key="cl_dd"):
                _push_undo(df_clean)
                before = len(df_clean); df_clean = df_clean.drop_duplicates()
                st.success(t("clean_dedup_done",_L, n=before-len(df_clean)))
                st.session_state["df_clean_result"] = df_clean
        with ops2:
            if st.button(t("clean_empty",_L), key="cl_em"):
                _push_undo(df_clean)
                before = len(df_clean); df_clean = df_clean.dropna(how="all")
                st.success(t("clean_empty_done",_L, n=before-len(df_clean)))
                st.session_state["df_clean_result"] = df_clean
        with ops3:
            if st.button(t("clean_trim",_L), key="cl_tr"):
                for col in df_clean.select_dtypes(include="object").columns:
                    df_clean[col] = df_clean[col].str.strip()
                st.success(t("clean_trim_done",_L))
                st.session_state["df_clean_result"] = df_clean

        # ── LLM normalizace sloupce ──────────────────
        cl_col1, cl_col2 = st.columns([3, 1])
        with cl_col1:
            norm_col = st.selectbox(t("clean_col",_L), df_clean.columns, key="norm_col")
        with cl_col2:
            clean_iter = st.slider(tt("🔁 Iterace", "🔁 Iterations", _L), 1, 5, 1, 1, key="cl_iter",
                                   help=tt("Počet průchodů LLM (výsledky se sloučí).",
                                           "Number of LLM passes (results are merged).", _L))

        if clean_iter > 1:
            _cl_merge_options = {
                "🤖 LLM Judge": "llm_judge",
                "📊 Consensus (voting)": "consensus",
                "🔗 Union + dedup": "union",
            }
            st.selectbox(
                "Metoda slučování iterací",
                list(_cl_merge_options.keys()),
                key="cl_merge_method",
                help=(
                    "**LLM Judge** — model vybere nejlepší normalizaci ze všech iterací\n\n"
                    "**Consensus** — pro každou hodnotu vybere normalizaci z majority iterací\n\n"
                    "**Union + dedup** — sloučí vše, odstraní duplicity (původní chování)"
                )
            )

        if st.button("▶️ Normalizovat sloupec", type="primary",
                     disabled=not selected_model, key="cl_norm_run"):
            sample = df_clean[norm_col].dropna().unique()[:50].tolist()
            msg_n  = norm_prompt_c + "\n\nHodnoty:\n" + "\n".join(str(s) for s in sample)
            _cl_merge_label  = st.session_state.get("cl_merge_method", "🔗 Union + dedup")
            _cl_merge_method = _cl_merge_options.get(_cl_merge_label, "union") if clean_iter > 1 else "union"

            with st.spinner("Normalizuji…"):
                iter_outputs_cl = []
                for _ci in range(clean_iter):
                    resp_ci = chat_completion_queued(base_url, selected_model,
                        [{"role":"system","content":"Vrať JSON [{original,normalized}]."},
                         {"role":"user","content":msg_n}],
                        temp=0.05 + _ci*0.02, max_tokens=4000)
                    iter_outputs_cl.append(resp_ci)

                resp_n = iter_outputs_cl[0]
                if clean_iter > 1:
                    if _cl_merge_method == "llm_judge":
                        # LLM Judge pro normalizaci
                        _cl_judge_sys = (
                            "Dostaneš několik variant normalizace stejných hodnot jako JSON [{original,normalized}]. "
                            "Pro každou původní hodnotu (original) vyber nejlepší normalizaci. "
                            "Výstup: POUZE JSON [{original,normalized}] bez komentářů."
                        )
                        _cl_variants = "\n\n".join(
                            f"=== Iterace {i+1} ===\n{o}" for i,o in enumerate(iter_outputs_cl))
                        try:
                            _cl_judged = chat_completion_queued(
                                base_url, selected_model,
                                [{"role":"system","content":_cl_judge_sys},
                                 {"role":"user","content":_cl_variants}],
                                temp=0.05, max_tokens=4000)
                            json.loads(_JSON_FENCE_RE.sub("", _cl_judged).strip())
                            resp_n = _cl_judged
                        except Exception:
                            resp_n = iter_outputs_cl[0]

                    elif _cl_merge_method == "consensus":
                        # Consensus: pro každý original vyber nejčastější normalized
                        from collections import Counter as _Counter
                        all_maps = []
                        for it_out in iter_outputs_cl:
                            try:
                                parsed_it = json.loads(_JSON_FENCE_RE.sub("", it_out).strip())
                                m = {r.get("original",""): r.get("normalized","")
                                     for r in (parsed_it if isinstance(parsed_it,list) else [])}
                                all_maps.append(m)
                            except Exception:
                                pass
                        if all_maps:
                            all_keys = set(k for m in all_maps for k in m)
                            consensus = {}
                            for k in all_keys:
                                vals = [m[k] for m in all_maps if k in m and m[k]]
                                if vals:
                                    top, _ = _Counter(vals).most_common(1)[0]
                                    consensus[k] = top
                            resp_n = json.dumps(
                                [{"original":k,"normalized":v} for k,v in consensus.items()],
                                ensure_ascii=False, indent=2)

                    else:
                        # Union + dedup (původní chování)
                        try:
                            all_norm = {}
                            for it_out in iter_outputs_cl:
                                parsed_it = json.loads(_JSON_FENCE_RE.sub("", it_out).strip())
                                for rec in (parsed_it if isinstance(parsed_it, list) else []):
                                    k = rec.get("original","")
                                    if k and k not in all_norm:
                                        all_norm[k] = rec.get("normalized","")
                            resp_n = json.dumps(
                                [{"original":k,"normalized":v} for k,v in all_norm.items()],
                                ensure_ascii=False, indent=2)
                        except Exception:
                            resp_n = iter_outputs_cl[0]

                # Stažení iterací
                if clean_iter > 1:
                    with st.expander(tt(f"⬇️ Stáhnout jednotlivé iterace ({clean_iter})",
                                         f"⬇️ Download individual iterations ({clean_iter})", _L)):
                        _cl_icols = st.columns(clean_iter + 1)
                        for _ci2, _co in enumerate(iter_outputs_cl):
                            with _cl_icols[_ci2]:
                                st.download_button(
                                    f"Iter {_ci2+1}",
                                    _co.encode("utf-8"),
                                    f"cisteni_iter{_ci2+1:02d}.json",
                                    key=f"cl_iter_dl_{_ci2}",
                                    width='stretch'
                                )
                        with _cl_icols[-1]:
                            st.download_button(
                                "📦 ZIP",
                                _make_iterations_zip(iter_outputs_cl, prefix="cisteni"),
                                "cisteni_iterace.zip",
                                mime="application/zip",
                                key="cl_iter_zip",
                                width='stretch'
                            )

            # ── Vizuální validace per-řádek (přijmout/odmítnout) ──
            try:
                norm_parsed = json.loads(_JSON_FENCE_RE.sub("", resp_n).strip())
                if isinstance(norm_parsed, list) and norm_parsed:
                    st.markdown(tt("#### ✅ Vizuální validace — přijmout / odmítnout per-řádek",
                                   "#### ✅ Visual validation — accept / reject per row", _L))
                    st.caption(tt("Zaškrtni řádky které chceš přijmout, pak klikni 'Aplikovat vybrané'.",
                                  "Tick the rows you want to accept, then click 'Apply selected'.", _L))

                    if "cl_norm_accepted" not in st.session_state:
                        st.session_state["cl_norm_accepted"] = {
                            r.get("original",""): True for r in norm_parsed}

                    # Zobraz jako editovatelnou tabulku s checkboxy
                    _vv_lbl_accept = tt("✅ Přijmout",            "✅ Accept",        _L)
                    _vv_lbl_orig   = tt("Původní hodnota",        "Original value",  _L)
                    _vv_lbl_norm   = tt("Normalizovaná hodnota",  "Normalized value", _L)
                    _vv_lbl_chg    = tt("Změna",                   "Change",          _L)
                    vv_rows = []
                    for rec in norm_parsed:
                        orig = rec.get("original","")
                        norm = rec.get("normalized","")
                        changed = orig != norm
                        vv_rows.append({
                            _vv_lbl_accept: st.session_state["cl_norm_accepted"].get(orig, True),
                            _vv_lbl_orig:   orig,
                            _vv_lbl_norm:   norm,
                            _vv_lbl_chg:    "✏️" if changed else "—",
                        })

                    edited_vv = st.data_editor(
                        pd.DataFrame(vv_rows),
                        column_config={
                            _vv_lbl_accept: st.column_config.CheckboxColumn(width="small"),
                            _vv_lbl_orig:   st.column_config.TextColumn(width="medium"),
                            _vv_lbl_norm:   st.column_config.TextColumn(width="medium"),
                            _vv_lbl_chg:    st.column_config.TextColumn(width="small"),
                        },
                        width='stretch',
                        key="cl_vv_editor",
                        num_rows="fixed"
                    )

                    vv_col1, vv_col2, vv_col3 = st.columns([2, 2, 3])
                    accepted_count = int(edited_vv[_vv_lbl_accept].sum())
                    with vv_col1:
                        st.metric(tt("Přijato", "Accepted", _L), accepted_count)
                    with vv_col2:
                        st.metric(tt("Odmítnuto", "Rejected", _L), len(vv_rows) - accepted_count)
                    with vv_col3:
                        if st.button(tt(f"✅ Aplikovat {accepted_count} přijatých změn",
                                         f"✅ Apply {accepted_count} accepted changes", _L),
                                     key="cl_vv_apply", type="primary"):
                            # Sestav mapu pouze přijatých normalizací
                            norm_map = {}
                            for _, row in edited_vv.iterrows():
                                if row[_vv_lbl_accept]:
                                    norm_map[row[_vv_lbl_orig]] = row[_vv_lbl_norm]
                            _push_undo(df_clean)
                            df_clean[norm_col] = df_clean[norm_col].map(
                                lambda x: norm_map.get(str(x), x) if pd.notna(x) else x)
                            st.session_state["df_clean_result"] = df_clean
                            st.success(tt(f"✅ Aplikováno {len(norm_map)} normalizací na sloupec '{norm_col}'",
                                          f"✅ Applied {len(norm_map)} normalizations to column '{norm_col}'", _L))
                            st.rerun()
                else:
                    st.text_area(tt("Výsledek normalizace", "Normalization result", _L),
                                 preview_text(resp_n), height=200,
                                 key="cl_norm_result_raw")
            except Exception:
                st.text_area(tt("Výsledek normalizace", "Normalization result", _L),
                             preview_text(resp_n), height=200,
                             key="cl_norm_result_fallback")

        st.divider()
        st.write(df_clean.describe(include="all"))
        res_df = st.session_state.get("df_clean_result", df_clean)
        if res_df is not None:
            st.download_button(t("clean_dl",_L), res_df.to_csv(index=False).encode("utf-8"),
                               "cista_data.csv", "text/csv", width='stretch')


# ══════════════════════════════════════════════════════
# 5 – STYLISTICKÉ ČIŠTĚNÍ ANGLICKÉHO TEXTU
# ══════════════════════════════════════════════════════
with main_tabs[5]:
    keep_current_tab(5)   # ← v24.17
    _L = st.session_state.get("lang","cz")
    st.markdown(tt("### ✍️ Stylistické čištění anglického textu",
                   "### ✍️ English Text Style Polish", _L))
    st.caption(tt("Zlepšení stylu, gramatiky a srozumitelnosti anglických vědeckých textů.",
                  "Improve style, grammar and clarity of English scientific writing.", _L))
    with st.expander(tt("ℹ️ Jak funguje stylistické čištění?",
                         "ℹ️ How does style polish work?", _L), expanded=False):
        st.markdown(tt(
            """
**Stylistika** analyzuje a vylepšuje anglický vědecký text — kontroluje pasivum, délku vět, hedging slova a Flesch-Kincaid skóre.

**Metriky (zobrazeny před/po):**
- **Flesch RE** — čitelnost 0–100 (vyšší = snazší; vědecký text: 30–50)
- **Pasivum** — podíl vět s pasivní konstrukcí (doporučeno <40 %)
- **Hedging** — frekvence slov jako "perhaps", "might", "could"
- **Průměrná délka věty** — doporučeno <25 slov

**Cílové časopisy:**
- **Palaeontologia Electronica** — formální, technický styl
- **Journal of Paleontology** — konzervativní, jasný
- **ZooKeys/BDJ** — otevřenější styl, kratší věty
- **PLOS ONE** — přístupný, bez žargonu

**Generátor abstraktů:** ze zadané metodiky + výsledků automaticky vygeneruje strukturovaný abstract (Úvod / Metody / Výsledky / Závěry).
            """,
            """
**Style polish** analyzes and improves English scientific prose — checks passive voice, sentence length, hedging words and Flesch-Kincaid score.

**Metrics (shown before/after):**
- **Flesch RE** — readability 0–100 (higher = easier; scholarly text: 30–50)
- **Passive voice** — share of sentences with passive construction (recommended <40 %)
- **Hedging** — frequency of words like "perhaps", "might", "could"
- **Average sentence length** — recommended <25 words

**Target journals:**
- **Palaeontologia Electronica** — formal, technical style
- **Journal of Paleontology** — conservative, clear
- **ZooKeys/BDJ** — more open style, shorter sentences
- **PLOS ONE** — accessible, no jargon

**Abstract writer:** from given methods + results automatically generates a structured abstract (Background / Methods / Results / Conclusions).
            """, _L))

    # ── Vstup textu ──────────────────────────────────
    sp_input_mode = st.radio(
        tt("Zdroj textu", "Text source", _L),
        ["✏️ Ručně", "📄 Ze souboru", "🔗 Z překladu"],
        format_func=lambda k: {"✏️ Ručně":      tt("✏️ Ručně",      "✏️ Manual",           _L),
                                "📄 Ze souboru": tt("📄 Ze souboru", "📄 From file",        _L),
                                "🔗 Z překladu": tt("🔗 Z překladu", "🔗 From translation", _L)}[k],
        horizontal=True, key="sp_input_mode")

    sp_src_text = ""
    if sp_input_mode == "✏️ Ručně":
        sp_src_text = st.text_area(
            tt("Anglický text ke stylistickému čištění",
               "English text to polish", _L),
            height=220, key="sp_manual_text",
            placeholder="Paste your English text here…")
    elif sp_input_mode == "📄 Ze souboru":
        sp_file = st.file_uploader(tt("Soubor (PDF, DOCX, TXT)", "File (PDF, DOCX, TXT)", _L),
                                   type=["pdf","docx","txt"], key="sp_file")
        if sp_file:
            sp_src_text = read_uploaded_file(sp_file)
            st.caption(tt(f"{len(sp_src_text):,} znaků (~{estimate_tokens(sp_src_text):,} tokenů)",
                          f"{len(sp_src_text):,} chars (~{estimate_tokens(sp_src_text):,} tokens)", _L))
    else:
        sp_src_text = st.session_state.get("translation_result", "")
        if sp_src_text:
            st.text_area(tt("Náhled z překladu", "Translation preview", _L),
                         sp_src_text[:600] + ("…" if len(sp_src_text) > 600 else ""),
                         height=120, disabled=True, key="sp_tr_preview")
        else:
            st.info(tt("Nejprve proveď překlad.", "Run a translation first.", _L))

    st.divider()

    # ── Úroveň čištění ───────────────────────────────
    st.markdown("#### 🎚️ Úroveň čištění" if _L == "cz" else "#### 🎚️ Polish level")

    STYLE_LEVELS = {
        "1 – Minimální (překlepy & mezery)": {
            "cz": "**Minimální** – opraví pouze překlepy, chybné mezery a interpunkci. Žádné obsahové změny.",
            "en": "**Minimal** – fix typos, spacing and punctuation only. No content changes.",
            "prompt": (
                "You are a copy-editor. Fix ONLY: spelling errors, spacing mistakes, and punctuation. "
                "Do NOT change wording, sentence structure, or content. "
                "Return only the corrected text, no explanations."
            )
        },
        "2 – Lehká (gramatika)": {
            "cz": "**Lehké** – opraví gramatiku, člen (a/an/the), shodu podmět–přísudek, časy. Zachová strukturu.",
            "en": "**Light** – fix grammar, articles (a/an/the), subject–verb agreement, tenses. Preserve structure.",
            "prompt": (
                "You are a scientific copy-editor. Fix grammar: articles, subject-verb agreement, "
                "verb tenses, and sentence fragments. Preserve the author's sentence structure and vocabulary. "
                "Preserve all Latin taxonomic names, stratigraphic terms and proper nouns unchanged. "
                "Return only the corrected text."
            )
        },
        "3 – Standardní (srozumitelnost)": {
            "cz": "**Standardní** – gramatika + zlepšení srozumitelnosti, přirozenosti vět. Doporučeno pro vědecké texty.",
            "en": "**Standard** – grammar + improved clarity and naturalness. Recommended for scientific papers.",
            "prompt": (
                "You are an expert scientific editor for palaeontology journals. "
                "Improve the text for clarity, naturalness and flow while strictly preserving all factual content, "
                "taxonomic names, stratigraphic units, and technical terminology. "
                "Fix grammar, awkward phrasing, and unnatural constructions typical of non-native English. "
                "Do not add or remove information. Return only the improved text."
            )
        },
        "4 – Silná (přepracování)": {
            "cz": "**Silné** – kompletní redakce: věty zkrátit, redundance odstranit, strukturu zlepšit. Obsah zachovat.",
            "en": "**Strong** – full editorial overhaul: shorten sentences, remove redundancy, improve structure. Preserve content.",
            "prompt": (
                "You are a senior editor at a palaeontology journal (e.g. Palaeontologia Electronica). "
                "Perform a thorough stylistic revision: shorten overly long sentences, eliminate redundant phrases, "
                "improve paragraph transitions, unify terminology, and ensure the text reads as natural academic English. "
                "All factual content, taxonomic names, stratigraphic units and citations must remain unchanged. "
                "Return only the revised text."
            )
        },
        "5 – Academické přepsání": {
            "cz": "**Akademický přepis** – přepíše text do stylu prestižního vědeckého časopisu. Zachová obsah a fakta.",
            "en": "**Academic rewrite** – rewrite to match the style of a high-impact journal. Preserve all facts.",
            "prompt": (
                "You are a professional scientific writer specialising in systematic palaeontology. "
                "Rewrite the following text in the style of a high-quality international palaeontology journal "
                "(concise, precise, impersonal, past tense for observations). "
                "Preserve all factual content, taxonomic names, stratigraphic terms, measurements, and citations exactly. "
                "Do not invent or omit information. Return only the rewritten text."
            )
        },
    }

    sp_level = st.selectbox(
        tt("Úroveň", "Level", _L),
        list(STYLE_LEVELS.keys()), index=2, key="sp_level")
    st.info(STYLE_LEVELS[sp_level]["cz" if _L == "cz" else "en"])

    # ── Doplňkové volby ──────────────────────────────
    sp_col1, sp_col2, sp_col3 = st.columns(3)
    with sp_col1:
        sp_iters = st.slider(tt("🔁 Iterace", "🔁 Iterations", _L), 1, 3, 1, key="sp_iters",
                             help=tt("Více iterací = opakované čištění (vhodné pro úrovně 4–5).",
                                     "More iterations = repeated polishing (useful for levels 4–5).", _L))
    with sp_col2:
        sp_preserve = st.checkbox(tt("Zachovat latinská jména",
                                      "Preserve Latin names", _L),
                                  value=True, key="sp_preserve")
    with sp_col3:
        sp_show_diff = st.checkbox("Zobrazit diff", value=True, key="sp_diff")

    sp_extra = st.text_input(
        "Doplňující instrukce (volitelné)" if _L == "cz" else "Additional instructions (optional)",
        key="sp_extra",
        placeholder=("Doplňující instrukce, např. Použij britský pravopis / max délka věty 25 slov" if _L=="cz" else "e.g. Use British spelling / avoid passive voice / max sentence length 25 words")
    )

    # ── Spuštění ──────────────────────────────────────
    if st.button("✍️ Spustit čištění" if _L == "cz" else "✍️ Run polish",
                 type="primary", disabled=not (selected_model and sp_src_text.strip()),
                 key="btn_style_polish"):

        sys_sp = STYLE_LEVELS[sp_level]["prompt"]
        if sp_preserve:
            sys_sp += (" Preserve all Latin taxonomic names, geological unit names, "
                       "and locality names exactly as written.")
        if sp_extra.strip():
            sys_sp += f" Additional requirement: {sp_extra.strip()}"

        with st.status("Čistím text…" if _L == "cz" else "Polishing text…", expanded=True) as sp_status:
            sp_result = sp_src_text
            for sp_i in range(sp_iters):
                if sp_iters > 1:
                    sp_status.write(f"Iterace {sp_i+1}/{sp_iters}…")
                # Použij smart chunking pro dlouhé texty
                if len(sp_result) > 7000:
                    sp_chunks = chunk_text_smart(sp_result, 7000)
                    sp_parts = []
                    for sp_ch in sp_chunks:
                        sp_parts.append(chat_completion(
                            base_url, selected_model,
                            [{"role": "system", "content": sys_sp},
                             {"role": "user", "content": sp_ch}],
                            temp=0.1, max_tokens=_MAX_TOKENS_TRANSLATE))
                    sp_result = "\n\n".join(sp_parts)
                else:
                    sp_result = chat_completion(
                        base_url, selected_model,
                        [{"role": "system", "content": sys_sp},
                         {"role": "user", "content": sp_result}],
                        temp=0.1, max_tokens=_MAX_TOKENS_TRANSLATE)
            sp_status.update(label="✅ Hotovo" if _L == "cz" else "✅ Done", state="complete")

        # ── Metriky ──────────────────────────────────
        m1, m2, m3 = st.columns(3)
        m1.metric("Znaků původní" if _L == "cz" else "Original chars", f"{len(sp_src_text):,}")
        m2.metric("Znaků výsledek" if _L == "cz" else "Result chars", f"{len(sp_result):,}")
        change_pct = (len(sp_result) - len(sp_src_text)) / max(1, len(sp_src_text)) * 100
        m3.metric("Změna" if _L == "cz" else "Change", f"{change_pct:+.1f}%")

        # ── Výsledek ──────────────────────────────────
        st.text_area(tt("✅ Výsledek (náhled 4000 znaků)",
                        "✅ Result (preview 4000 chars)", _L),
                     sp_result[:4000] + ("…" if len(sp_result) > 4000 else ""),
                     height=320, key="sp_result_area")

        # ── Diff ──────────────────────────────────────
        if sp_show_diff:
            with st.expander(tt("↔️ Srovnání originál / výsledek",
                                 "↔️ Original vs. result comparison", _L)):
                st.markdown(simple_diff_html(sp_src_text, sp_result), unsafe_allow_html=True)

        # ── Export ────────────────────────────────────
        render_export_buttons(sp_result, "style_polish", {"level": sp_level})

        # Uložit do session state + _temp ihned
        st.session_state["style_polish_result"] = sp_result
        try:
            _temp_save_txt("style_polish", 1, sp_result)
        except Exception:
            pass
        save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
        save_to_history("style_polish", {"level": sp_level, "chars": len(sp_src_text)})

        if st.button(tt("📥 Použít jako aktivní překlad",
                         "📥 Set as active translation", _L), key="sp_accept_as_tr"):
            st.session_state["translation_result"] = sp_result
            st.session_state["translation_ready"] = True
            save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
            st.success(tt("Uloženo jako aktivní překlad.",
                          "Saved as active translation.", _L))

    st.divider()

    # ── Vědecký writing style checker ────────────────
    with st.expander(tt("📊 Kontrola textu (metriky před/po)",
                         "📊 Writing Style Checker (before/after metrics)", _L)):
        wsc_text = st.text_area(tt("Text k analýze", "Text to analyze", _L),
                                height=150, key="wsc_input",
                                placeholder=tt("Vlož anglický vědecký text…",
                                                "Paste English scientific text…", _L))
        if wsc_text.strip():
            sentences  = re.split(r'[.!?]+', wsc_text)
            sentences  = [s.strip() for s in sentences if len(s.strip()) > 10]
            words      = re.findall(r'\b\w+\b', wsc_text)
            avg_sent   = sum(len(re.findall(r'\b\w+\b',s)) for s in sentences) / max(1,len(sentences))
            passive_n  = len(re.findall(r'\b(?:was|were|been|is|are|be)\s+\w+ed\b', wsc_text, re.I))
            hedge_words = ["perhaps","might","possibly","seemingly","apparently","could","may","suggest"]
            hedge_n    = sum(wsc_text.lower().count(h) for h in hedge_words)

            # Flesch-Kincaid Reading Ease
            def count_syllables(word: str) -> int:
                word = word.lower()
                count = len(re.findall(r'[aeiouy]+', word))
                if word.endswith('e') and count > 1:
                    count -= 1
                return max(1, count)
            total_syllables = sum(count_syllables(w) for w in words)
            avg_syllables = total_syllables / max(1, len(words))
            # Flesch Reading Ease: 206.835 - 1.015*(words/sentences) - 84.6*(syllables/words)
            fre = 206.835 - 1.015 * avg_sent - 84.6 * avg_syllables
            fre = max(0, min(100, fre))
            fre_label = ("Very Easy" if fre >= 90 else "Easy" if fre >= 70
                         else "Standard" if fre >= 50 else "Difficult" if fre >= 30
                         else "Very Difficult")

            # Non-native Czech/German calques
            calques    = {
                "in the work of":      "by",
                "were found to be present": "were present",
                "in the framework of": "within",
                "due to the fact that": "because",
                "in case of":          "for",
                "is characterized by the presence of": "has",
                "it is possible to":   "can",
            }
            found_calques = [(orig,fix) for orig,fix in calques.items()
                             if orig in wsc_text.lower()]

            wc1,wc2,wc3,wc4,wc5 = st.columns(5)
            wc1.metric(tt("Průměr slov/věta", "Avg. words/sentence", _L), f"{avg_sent:.1f}",
                       delta=tt("↑ příliš dlouhé", "↑ too long", _L) if avg_sent > 25 else None,
                       delta_color="inverse")
            wc2.metric(tt("Passivum", "Passive voice", _L), passive_n)
            wc3.metric(tt("Hedging slov", "Hedging words", _L), hedge_n)
            wc4.metric(tt("Slov celkem", "Total words", _L), len(words))
            wc5.metric("Flesch RE", f"{fre:.0f}",
                       help=tt(f"{fre_label} — 0=nejtěžší, 100=nejsnazší\n"
                               "Vědecké texty obvykle 30–50.",
                               f"{fre_label} — 0=hardest, 100=easiest\n"
                               "Scientific texts usually 30–50.", _L),
                       delta=fre_label)

            if found_calques:
                st.warning(tt("**Nalezeny kalky z češtiny/němčiny:**",
                              "**Czech/German calques found:**", _L))
                for orig, fix in found_calques:
                    st.markdown(f"- ~~*{orig}*~~ → `{fix}`")
            else:
                st.success(tt("✅ Žádné typické non-native kalky nenalezeny",
                              "✅ No typical non-native calques found", _L))


    # ── Cílový časopis ────────────────────────────────
    with st.expander(tt("📰 Stylistika pro cílový časopis",
                         "📰 Style for target journal", _L)):
        JOURNAL_STYLES = {
            "Palaeontologia Electronica": {
                "desc": "Online journal, clear prose, detailed methods, extensive synonymy lists",
                "prompt_add": ("Write in the style of Palaeontologia Electronica: "
                               "clear, precise, detailed systematic descriptions with full synonymy. "
                               "Use past tense for observations. Spell out numbers below 10.")
            },
            "Journal of Paleontology": {
                "desc": "Classic US journal, concise, formal, active voice preferred",
                "prompt_add": ("Write in the style of the Journal of Paleontology: "
                               "concise and formal, prefer active voice, minimize hedging. "
                               "SI units throughout. Brief but complete descriptions.")
            },
            "ZooKeys / BDJ": {
                "desc": "Pensoft journal, structured sections, data-rich, open-access focus",
                "prompt_add": ("Write in the style of ZooKeys/Biodiversity Data Journal: "
                               "structured with clear section headers, data-focused, "
                               "include structured occurrence data, use GBIF/Darwin Core terminology.")
            },
            "PLOS ONE": {
                "desc": "Broad audience, clear methods, accessible language",
                "prompt_add": ("Write in the style of PLOS ONE: "
                               "accessible to a broad scientific audience, "
                               "explicit and reproducible methods section, "
                               "clear statement of contribution in introduction.")
            },
        }
        sel_journal = st.selectbox(tt("Cílový časopis", "Target journal", _L),
                                    list(JOURNAL_STYLES.keys()), key="sp_journal")
        st.caption(JOURNAL_STYLES[sel_journal]["desc"])

        journal_text = st.text_area(tt("Text k úpravě pro časopis",
                                        "Text to adapt for the journal", _L),
                                     height=150, key="sp_journal_text",
                                    placeholder=tt("Vlož text k úpravě…",
                                                    "Paste text to polish…", _L))
        if st.button(tt("🗞️ Upravit pro časopis", "🗞️ Adapt for journal", _L),
                     key="sp_journal_btn",
                     disabled=not (selected_model and journal_text.strip())):
            jnl_sys = JOURNAL_STYLES[sel_journal]["prompt_add"]
            with st.spinner(tt(f"Upravuji pro {sel_journal}…",
                               f"Adapting for {sel_journal}…", _L)):
                jnl_result = chat_completion_queued(
                    base_url, selected_model,
                    [{"role":"system","content": jnl_sys},
                     {"role":"user","content": journal_text}],
                    temp=0.15, max_tokens=_MAX_TOKENS_TRANSLATE)
            st.text_area(tt("Výsledek", "Result", _L), jnl_result[:3000], height=250, key="sp_jnl_result")
            render_export_buttons(jnl_result, f"journal_{sel_journal.split()[0]}", {})

    # ── Abstract writer ───────────────────────────────
    with st.expander(tt("📄 Abstract writer (ze zadané metodiky + výsledků)",
                         "📄 Abstract writer (from given methods + results)", _L)):
        st.caption(tt("Vygeneruje strukturovaný abstract (Background/Methods/Results/Conclusions) "
                      "ve formátu zvoleného časopisu.",
                      "Generates a structured abstract (Background/Methods/Results/Conclusions) "
                      "in the format of the chosen journal.", _L))
        abs_methods  = st.text_area(tt("Metody (výňatek)", "Methods (excerpt)", _L),
                                    height=100, key="abs_methods",
                                    placeholder=tt("Popis použitých metod…",
                                                    "Description of methods used…", _L))
        abs_results  = st.text_area(tt("Výsledky (výňatek)", "Results (excerpt)", _L),
                                    height=100, key="abs_results",
                                    placeholder=tt("Klíčové výsledky…", "Key results…", _L))
        abs_journal  = st.selectbox(tt("Styl časopisu", "Journal style", _L),
                                     list(JOURNAL_STYLES.keys()),
                                    key="abs_journal")
        abs_max_words = st.slider(tt("Max. počet slov abstractu",
                                      "Max. abstract word count", _L),
                                   150, 400, 250, 25, key="abs_words")

        if st.button(tt("📝 Generovat abstract", "📝 Generate abstract", _L),
                     key="abs_gen_btn",
                     disabled=not (selected_model and (abs_methods.strip() or abs_results.strip()))):
            abs_sys = (
                f"You are an expert scientific writer. "
                f"Write a structured abstract of maximum {abs_max_words} words "
                f"in the style of {abs_journal}. "
                "Structure it as: Background (1–2 sentences), Methods (2–3 sentences), "
                "Results (3–4 sentences), Conclusions (1–2 sentences). "
                "Be concise and precise. Preserve all taxonomic names and stratigraphic terms. "
                "Return only the abstract text, no headings."
            )
            abs_input = f"METHODS:\n{abs_methods}\n\nRESULTS:\n{abs_results}"
            with st.spinner(tt("Generuji abstract…", "Generating abstract…", _L)):
                abs_result = chat_completion_queued(
                    base_url, selected_model,
                    [{"role":"system","content":abs_sys},
                     {"role":"user","content":abs_input}],
                    temp=0.2, max_tokens=1000)
            st.text_area(tt("Vygenerovaný abstract", "Generated abstract", _L),
                         abs_result, height=200, key="abs_result_area")
            word_count = len(abs_result.split())
            st.caption(tt(f"Délka: {word_count} slov",
                          f"Length: {word_count} words", _L))
            st.download_button(tt("⬇️ Stáhnout TXT", "⬇️ Download TXT", _L),
                               abs_result.encode("utf-8"),
                               "abstract.txt", key="abs_dl")


# ══════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════


with main_tabs[6]:
    keep_current_tab(6)   # ← v24.17
    _L = st.session_state.get("lang","cz")
    st.markdown(t("wf_title",_L))

    st.subheader(tt("📦 Globální nastavení bloků pro workflow",
                    "📦 Global chunking settings for workflow", _L))
    wf_col1, wf_col2, wf_col3 = st.columns([3, 2, 2])
    with wf_col1:
        wf_chunk_target = st.slider(
            tt("Cílová velikost bloku (znaky)", "Target chunk size (chars)", _L),
            1200, 6000, 2200, 100, key="wf_chunk_target"
        )
    with wf_col2:
        wf_use_smart = st.checkbox(
            tt("Chytré rozdělení po blocích (konec věty/odstavce)",
               "Smart chunking (sentence/paragraph)", _L),
            value=True, key="wf_use_smart"
        )
    with wf_col3:
        wf_overlap = st.slider(
            tt("Překryv (znaky)", "Overlap (chars)", _L),
            200, 600, 350, 50, key="wf_overlap"
        )
    with st.expander("ℹ️ Jak funguje workflow?", expanded=False):
        st.markdown("""
**Workflow** spustí celý proces automaticky — extrakci, překlad, validaci taxonů a export — jedním kliknutím.

**Postup:**
1. Nahraj soubor (PDF/DOCX/TXT)
2. Přidej kroky v pořadí (např. Extrakce → Překlad → Validace)
3. Nastav parametry pro každý krok
4. Volitelně nastav **webhook URL** pro notifikaci po dokončení (Slack, Teams)
5. Klikni **▶️ Spustit celý workflow**

**Šablony:**
- **Rychlá extrakce CZ→EN** — extrakce + překlad do angličtiny
- **Plný proces s validací** — extrakce + překlad + validace taxonů
- **Hyolitha kompletní** — plný proces + export pro Hyolitha DB

**Podmíněné kroky:** validace se spustí jen pokud extrakce vrátila >0 taxonů; překlad jen pokud text není v angličtině.

**Paralelní větve:** spustí extrakci s více různými prompty najednou pro porovnání výsledků.
        """)
    st.subheader(tt("1. Vstupní soubor", "1. Input file", _L))
    wf_file      = st.file_uploader(tt("Zdrojový dokument", "Source document", _L),
                                     type=["pdf","docx","txt"],key="wf_file_main")
    wf_page_spec = st.text_input(tt("Rozsah stran PDF", "PDF page range", _L),
                                  "",key="wf_pages")

    st.subheader(tt("2. Kroky procesu", "2. Pipeline steps", _L))
    WORKFLOW_ACTIONS = {
        "🔍 Extrakce":                    "extrakce",
        "🌐 Překlad":                     "preklad",
        "🧬 Validace taxonů":             "validace",
        "🧹 Čištění textu (LLM)":         "cisteni",
        "🗺️ Stratigrafická normalizace":  "strat_norm",
        "🦪 Export Hyolitha DB":          "hyolitha_exp",
    }
    _WF_ACTION_LABELS = {
        "🔍 Extrakce":                    tt("🔍 Extrakce",                    "🔍 Extraction",                  _L),
        "🌐 Překlad":                     tt("🌐 Překlad",                     "🌐 Translation",                 _L),
        "🧬 Validace taxonů":             tt("🧬 Validace taxonů",             "🧬 Taxon validation",            _L),
        "🧹 Čištění textu (LLM)":         tt("🧹 Čištění textu (LLM)",         "🧹 Text cleaning (LLM)",         _L),
        "🗺️ Stratigrafická normalizace":  tt("🗺️ Stratigrafická normalizace",  "🗺️ Stratigraphic normalization", _L),
        "🦪 Export Hyolitha DB":          tt("🦪 Export Hyolitha DB",          "🦪 Hyolitha DB export",          _L),
    }

    # ── Šablony workflow ──────────────────────────────
    wf_templates = load_wf_templates()
    _wf_custom = tt("— vlastní —", "— custom —", _L)
    wt_col1, wt_col2, wt_col3 = st.columns([3,2,2])
    with wt_col1:
        wf_tpl_sel = st.selectbox(tt("📋 Šablona workflow", "📋 Workflow template", _L),
                                   [_wf_custom] + list(wf_templates.keys()),
                                   key="wf_tpl_sel")
    with wt_col2:
        if st.button(tt("📋 Načíst šablonu", "📋 Load template", _L), key="wf_tpl_load",
                     disabled=wf_tpl_sel == _wf_custom):
            tpl = wf_templates[wf_tpl_sel]
            st.session_state["workflow_steps"] = list(tpl.get("steps", []))
            for k, v in tpl.get("params", {}).items():
                st.session_state[k] = v
            st.toast(tt(f"✅ Šablona '{wf_tpl_sel}' načtena",
                        f"✅ Template '{wf_tpl_sel}' loaded", _L))
            st.rerun()
    with wt_col3:
        wf_tpl_new_name = st.text_input(tt("Uložit jako šablonu", "Save as template", _L),
                                         key="wf_tpl_new_name",
                                         placeholder=tt("Název šablony…", "Template name…", _L),
                                         label_visibility="collapsed")
        if st.button(tt("💾 Uložit šablonu", "💾 Save template", _L), key="wf_tpl_save",
                     disabled=not wf_tpl_new_name.strip()):
            save_wf_template(wf_tpl_new_name,
                             st.session_state.get("workflow_steps", []),
                             {"wf_lang": st.session_state.get("wf_lang","angličtina")})
            st.toast(tt(f"✅ Šablona '{wf_tpl_new_name}' uložena",
                        f"✅ Template '{wf_tpl_new_name}' saved", _L))

    new_step = st.selectbox(tt("Přidat krok", "Add step", _L),
                             [""] + list(WORKFLOW_ACTIONS.keys()),
                             format_func=lambda k: _WF_ACTION_LABELS.get(k, k),
                             key="wf_new")
    if st.button(tt("➕ Přidat", "➕ Add", _L), key="wf_add") and new_step:
        st.session_state["workflow_steps"].append(new_step)

    if st.session_state["workflow_steps"]:
        badges = " → ".join(f'<span class="step-badge">{_WF_ACTION_LABELS.get(s,s)}</span>'
                             for s in st.session_state["workflow_steps"])
        st.markdown(badges, unsafe_allow_html=True)

        # ── Podmíněné kroky ───────────────────────────
        with st.expander(tt("⚙️ Podmíněné kroky", "⚙️ Conditional steps", _L)):
            st.caption(tt("Nastav podmínky pro přeskočení kroků.",
                          "Set conditions for skipping steps.", _L))
            if "wf_conditions" not in st.session_state:
                st.session_state["wf_conditions"] = {}
            for i, step in enumerate(st.session_state["workflow_steps"]):
                action = WORKFLOW_ACTIONS.get(step, "")
                cond_key = f"cond_{action}"
                # Interní hodnoty CZ (používají se v logice níže), zobrazované přes format_func
                _wf_always = "Vždy spustit"
                cond_opts = [_wf_always]
                _cond_labels = {_wf_always: tt("Vždy spustit", "Always run", _L)}
                if action == "validace":
                    k1 = "Jen pokud extrakce vrátila >0 taxonů"
                    cond_opts += [k1]
                    _cond_labels[k1] = tt("Jen pokud extrakce vrátila >0 taxonů",
                                           "Only if extraction returned >0 taxa", _L)
                if action == "preklad":
                    k2 = "Jen pokud detekovaný jazyk není angličtina"
                    cond_opts += [k2]
                    _cond_labels[k2] = tt("Jen pokud detekovaný jazyk není angličtina",
                                           "Only if detected language is not English", _L)
                if action == "hyolitha_exp":
                    k3 = "Jen pokud validace nalezla >0 taxonů"
                    cond_opts += [k3]
                    _cond_labels[k3] = tt("Jen pokud validace nalezla >0 taxonů",
                                           "Only if validation found >0 taxa", _L)
                sel = st.selectbox(tt(f"Podmínka pro: **{step}**",
                                       f"Condition for: **{_WF_ACTION_LABELS.get(step,step)}**", _L),
                                   cond_opts,
                                   format_func=lambda k: _cond_labels.get(k, k),
                                   key=f"wf_cond_{i}")
                st.session_state["wf_conditions"][action] = sel

        for i, step in enumerate(st.session_state["workflow_steps"]):
            if st.button(tt(f"🗑️ Odebrat: {step}",
                             f"🗑️ Remove: {_WF_ACTION_LABELS.get(step,step)}", _L),
                         key=f"wf_del_{i}"):
                st.session_state["workflow_steps"].pop(i); st.rerun()
        if st.button(tt("🗑️ Smazat celou pipeline",
                         "🗑️ Clear whole pipeline", _L)):
            st.session_state["workflow_steps"]=[]; st.rerun()
    else:
        st.info(tt("Přidej alespoň jeden krok.", "Add at least one step.", _L))

    st.subheader(tt("3. Parametry", "3. Parameters", _L))
    wf_ext_prompt = st.text_area(tt("Prompt pro Extrakci", "Extraction prompt", _L),
        value="Extrahuj záznamy jako JSON: [{rod,druh,autor,rok,lokalita,stratigrafie}]",
        height=80,key="wf_ext_p")
    wf_tgt_lang  = st.selectbox(tt("Cílový jazyk Překladu", "Translation target language", _L),
                                 LANGUAGE_OPTIONS,index=1,key="wf_lang")
    wf_valid_dbs = st.multiselect(tt("Databáze pro Validaci", "Databases for Validation", _L),
                                  list(TAXONOMIC_DATABASES.keys()),
                                  default=["PaleoDB","GBIF","IRMNG"],key="wf_dbs")
    wf_strat_llm = st.checkbox(
        tt("Stratigrafická normalizace: použít LLM (pomalejší, přesnější)",
           "Stratigraphic normalization: use LLM (slower, more accurate)", _L),
        value=False, key="wf_strat_llm",
        help=tt("Pokud není zaškrtnuto, použije se rychlá lokální náhrada ze slovníku ICS_SYNONYMS.",
                "If unchecked, a fast local substitution from the ICS_SYNONYMS dictionary is used.", _L))

    # ── Webhook notifikace ────────────────────────────
    with st.expander(tt("🔔 Webhook notifikace po dokončení",
                         "🔔 Webhook notification on completion", _L)):
        wf_webhook = st.text_input(
            tt("Webhook URL (Slack, Teams, nebo generic JSON)",
               "Webhook URL (Slack, Teams, or generic JSON)", _L),
            value=st.session_state.get("webhook_url",""),
            key="wf_webhook_url",
            placeholder="https://hooks.slack.com/services/…")
        wf_wh_threshold = st.slider(
            tt("Notifikovat jen pokud workflow trvalo déle než (min)",
               "Notify only if workflow took longer than (min)", _L),
            0, 30, 2, key="wf_wh_threshold",
            help=tt("0 = vždy notifikovat po dokončení",
                    "0 = always notify on completion", _L))

    st.subheader(tt("4. Spuštění", "4. Run", _L))

    # ── Paralelní větve ───────────────────────────────
    with st.expander(tt("🔀 Paralelní větve — extrakce s více prompty najednou",
                         "🔀 Parallel branches — extraction with multiple prompts at once", _L)):
        st.caption(tt("Spustí extrakci paralelně s různými prompty a zobrazí výsledky vedle sebe "
                      "pro porovnání. Užitečné pro ladění promptů.",
                      "Runs extraction in parallel with different prompts and shows results side by side "
                      "for comparison. Useful for prompt tuning.", _L))
        pb_enabled = st.checkbox(tt("Zapnout paralelní větve",
                                     "Enable parallel branches", _L),
                                 key="wf_pb_enabled", value=False)
        if pb_enabled:
            pb_prompts_raw = st.text_area(
                tt("Prompty pro paralelní větve (jeden per řádek)",
                   "Prompts for parallel branches (one per line)", _L),
                height=100, key="wf_pb_prompts",
                value=("Extrahuj taxonomické záznamy jako JSON: [{rod,druh,autor,rok}]\n"
                       "Extrahuj POUZE druhová jména a autory: [{druh,autor,rok,stratigrafie}]"))
            pb_prompts = [p.strip() for p in pb_prompts_raw.split("\n") if p.strip()]
            st.caption(tt(f"Počet větví: {len(pb_prompts)}",
                          f"Number of branches: {len(pb_prompts)}", _L))

    run_wf = st.button(tt("▶️ Spustit celý workflow", "▶️ Run whole workflow", _L),
                       type="primary",
                       disabled=not(selected_model and wf_file and st.session_state["workflow_steps"]))

    # ── Paralelní větve: spuštění ─────────────────────
    if st.session_state.get("wf_pb_enabled") and wf_file:
        pb_prompts_for_run = [p.strip() for p in
                              st.session_state.get("wf_pb_prompts","").split("\n") if p.strip()]
        if st.button(tt("🔀 Spustit paralelní větve", "🔀 Run parallel branches", _L),
                     key="wf_pb_run",
                     disabled=not (selected_model and wf_file and len(pb_prompts_for_run) > 0)):
            wf_file.seek(0)
            pb_text = read_uploaded_file(wf_file, wf_page_spec or None)
            pb_sample = pb_text[:3000]
            pb_results = {}
            pb_sys = EXTRACTION_MODES.get("Standard", {}).get("system", "")
            with st.status(tt(f"🔀 Spouštím {len(pb_prompts_for_run)} paralelních větví…",
                               f"🔀 Running {len(pb_prompts_for_run)} parallel branches…", _L),
                           expanded=True) as pb_status:
                from concurrent.futures import ThreadPoolExecutor, as_completed as _as_completed
                futures_pb = {}
                with ThreadPoolExecutor(max_workers=min(4, len(pb_prompts_for_run))) as pool_pb:
                    for pi, pp in enumerate(pb_prompts_for_run):
                        fut_pb = pool_pb.submit(
                            chat_completion, base_url, selected_model,
                            [{"role":"system","content":pb_sys},
                             {"role":"user","content":f"{pp}\n\n{pb_sample}"}],
                            0.1, 4000)
                        futures_pb[fut_pb] = (pi, pp)
                    for fut_pb in _as_completed(futures_pb):
                        pi, pp = futures_pb[fut_pb]
                        try:
                            pb_results[pi] = {"prompt": pp, "result": fut_pb.result()}
                            pb_status.write(tt(f"✅ Větev {pi+1}: {pp[:60]}…",
                                               f"✅ Branch {pi+1}: {pp[:60]}…", _L))
                        except Exception as e:
                            pb_results[pi] = {"prompt": pp, "result": f"Chyba: {e}"}
                pb_status.update(label=tt("✅ Paralelní větve dokončeny",
                                          "✅ Parallel branches complete", _L), state="complete")

            pb_cols = st.columns(len(pb_results))
            for col_idx, (pi, pb_data) in enumerate(sorted(pb_results.items())):
                with pb_cols[col_idx % len(pb_cols)]:
                    st.markdown(tt(f"**Větev {pi+1}:** `{pb_data['prompt'][:50]}…`",
                                   f"**Branch {pi+1}:** `{pb_data['prompt'][:50]}…`", _L))
                    st.text_area(tt(f"Výsledek {pi+1}", f"Result {pi+1}", _L),
                                 pb_data["result"][:1000],
                                 height=300, key=f"pb_res_{pi}")
                    try:
                        clean_pb = _JSON_FENCE_RE.sub("", pb_data["result"]).strip()
                        recs_pb = json.loads(clean_pb)
                        if isinstance(recs_pb, list):
                            st.success(tt(f"✅ {len(recs_pb)} záznamů",
                                          f"✅ {len(recs_pb)} records", _L))
                    except Exception:
                        pass


    if run_wf:
        wf_log: List[str] = []
        wf_file.seek(0)
        current_text = read_uploaded_file(wf_file, wf_page_spec or None)
        wf_log.append(f"[START] {len(current_text):,} znaků z {wf_file.name}")
        wf_taxa: List[str] = []
        wf_val_results: List[Dict] = []
        _wf_start_time = _time.time()
        _wf_conditions = st.session_state.get("wf_conditions", {})

        gsb_start(
            f"⚙️ Workflow: {wf_file.name}" if _L=="cz"
            else f"⚙️ Workflow: {wf_file.name}",
            tab=t("tab_workflow", _L)
        )

        status = st.empty()
        prog   = st.progress(0)
        n      = len(st.session_state["workflow_steps"])
        # Udržuj _temp/ čistý před novým workflow
        _temp_prune("wf_state", keep_last=n + 2)
        _temp_prune("wf_final", keep_last=2)
        _temp_prune("workflow_preklad", keep_last=3)
        _temp_prune("extraction_partial", keep_last=3)

        for i,step in enumerate(st.session_state["workflow_steps"]):
            prog.progress(i/n, text=f"Krok {i+1}/{n}: {step}")
            gsb_update(detail=f"Krok {i+1}/{n}: {step}", progress=i/n)
            status.info(f"⚙️ {step}")
            action = WORKFLOW_ACTIONS.get(step,"")

            # ── Podmíněné přeskakování kroků ─────────
            cond = _wf_conditions.get(action, "Vždy spustit")
            skip = False
            if cond == "Jen pokud extrakce vrátila >0 taxonů" and not wf_taxa:
                wf_log.append(f"[SKIP] {step} — žádné taxony z extrakce")
                status.warning(f"⏭️ Přeskočeno: {step} (žádné taxony)")
                skip = True
            elif cond == "Jen pokud detekovaný jazyk není angličtina":
                det_lang = detect_language(current_text)
                if det_lang == "angličtina":
                    wf_log.append(f"[SKIP] {step} — text je již v angličtině")
                    status.warning(f"⏭️ Přeskočeno: {step} (text je EN)")
                    skip = True
            elif cond == "Jen pokud validace nalezla >0 taxonů":
                found_any = any(r["summary"]["found"] > 0 for r in wf_val_results)
                if not found_any:
                    wf_log.append(f"[SKIP] {step} — validace nenalezla žádné taxony")
                    skip = True
            if skip:
                continue

            try:
                if action == "extrakce":
                    blocks = chunk_text(current_text,
                                        suggest_chunk_size(len(current_text), selected_model))
                    outs   = [chat_completion_queued(base_url, selected_model,
                        [{"role":"system","content":EXTRACTION_MODES["Standard"]["system"]},
                         {"role":"user","content":f"{wf_ext_prompt}\n\n{b}"}],
                        temp=0.05, max_tokens=_MAX_TOKENS_EXTRACT,
                        stop=_STOP_JSON) for b in blocks]
                    current_text = _merge_json_chunks(outs)
                    st.session_state["last_extraction_text"]  = current_text
                    st.session_state["last_extracted_text"]   = current_text  # sync
                    wf_taxa = extract_taxa_from_text(current_text)
                    st.session_state["last_extraction_taxa"]  = wf_taxa
                    wf_log.append(f"[EXTRAKCE] {len(current_text):,} znaků, {len(wf_taxa)} taxonů, {len(blocks)} bloků")

                elif action == "preklad":
                    det = detect_language(current_text) or "neznámý"
                    sys_wf = build_translate_system(det, wf_tgt_lang, {}, preserve_terms=True)
                    _wf_chunk = suggest_chunk_size(len(current_text), selected_model)
                    _wf_par   = st.session_state.get("lms_max_concurrent", 1) > 1
                    _wf_mw    = st.session_state.get("lms_max_concurrent", 1)
                    current_text = do_translate(base_url, selected_model, current_text,
                                                sys_wf, temp=0.1,
                                                chunk_size=_wf_chunk,
                                                parallel=_wf_par, max_workers=_wf_mw,
                                                src_lang=det)
                    wf_log.append(f"[PŘEKLAD] {det} → {wf_tgt_lang}, chunk={_wf_chunk}")

                elif action == "validace":
                    tv = wf_taxa or extract_taxa_from_text(current_text)
                    if not tv:
                        wf_log.append("[VALIDACE] Žádné taxony")
                    else:
                        tv_unique = list(set(tv[:50]))
                        _wf_val_workers = max(4, st.session_state.get("lms_max_concurrent", 4))
                        _wf_val_futures = {}
                        with ThreadPoolExecutor(max_workers=_wf_val_workers) as _wf_pool:
                            for _wf_t in tv_unique:
                                _wf_val_futures[_wf_pool.submit(
                                    validate_taxon_name, _wf_t, wf_valid_dbs,
                                    st.session_state["validation_cache"]
                                )] = _wf_t
                            wf_val_results = [f.result() for f in as_completed(_wf_val_futures)]
                        st.session_state["last_validation_results"] = wf_val_results
                        # Autosave validace ihned po dokončení kroku
                        try:
                            _temp_save_json("wf_validace", 1, [
                                {"taxon": v["taxon"],
                                 "found": v["summary"]["found"],
                                 "results": v.get("results", {})}
                                for v in wf_val_results
                            ])
                        except Exception:
                            pass
                        found_n = sum(1 for v in wf_val_results if v["summary"]["found"]>0)
                        current_text = (f"Validováno {len(wf_val_results)} taxonů, potvrzeno {found_n}\n\n"
                                        + "\n".join(f"{v['taxon']}: {v['summary']['confidence']}"
                                                    for v in wf_val_results))
                        wf_log.append(f"[VALIDACE] {len(wf_val_results)} taxonů, {found_n} potvrzeno")

                elif action == "cisteni":
                    current_text = chat_completion_queued(base_url, selected_model,
                        [{"role":"system","content":"Vyčisti a normalizuj text. Zachovej obsah."},
                         {"role":"user","content":current_text}],
                        temp=0.05, max_tokens=_MAX_TOKENS_TRANSLATE)
                    wf_log.append("[ČIŠTĚNÍ] Hotovo")

                elif action == "hyolitha_exp":
                    try:
                        clean_h = _JSON_FENCE_RE.sub("", current_text).strip()
                        recs_h  = json.loads(clean_h)
                        if isinstance(recs_h, dict): recs_h = [recs_h]
                        st.session_state["hyolitha_export_records"] = recs_h
                        try:
                            _temp_save_json("wf_hyolitha", 1, recs_h)
                        except Exception:
                            pass
                        wf_log.append(f"[HYOLITHA] {len(recs_h)} záznamů připraveno k exportu")
                    except Exception as e_h:
                        wf_log.append(f"[HYOLITHA] Chyba parsování JSON: {e_h}")

                elif action == "strat_norm":
                    if wf_strat_llm:
                        ics_list = "\n".join(f"  {k}: {v}" for k,v in list(ICS_SYNONYMS.items())[:30])
                        norm_resp = chat_completion_queued(base_url, selected_model,
                            [{"role":"system","content":
                              f"Normalizuj stratigrafické termíny v textu dle ICS nomenklatury.\n"
                              f"Příklady záměn:\n{ics_list}\n"
                              f"Zachovej veškerý ostatní obsah. Nahraď pouze stratigrafické termíny."},
                             {"role":"user","content":current_text}],
                            temp=0.05, max_tokens=_MAX_TOKENS_TRANSLATE)
                        wf_log.append("[STRAT_NORM] LLM normalizace dokončena")
                        current_text = norm_resp
                    else:
                        current_text, changes = normalize_stratigraphy_local(current_text)
                        wf_log.append(f"[STRAT_NORM] Lokální: {len(changes)} záměn: "
                                      + "; ".join(changes[:5]) + ("…" if len(changes)>5 else ""))

            except Exception as e:
                wf_log.append(f"[CHYBA {step}] {e}")
                status.error(f"❌ {step}: {e}")

            # ── Průběžné uložení stavu workflow po každém kroku ──
            try:
                _wf_safe_step  = re.sub(r'[^\w]', '_', step)
                _wf_step_label = f"wf_{_wf_safe_step}_{i+1}of{n}"
                _temp_save_txt(_wf_step_label, i + 1,
                               f"=== Krok {i+1}/{n}: {step} ===\n{current_text}",
                               total=n)
                _temp_save_json(f"wf_state_{i+1}of{n}", i + 1, {
                    "step": step, "step_idx": i, "total_steps": n,
                    "current_text_preview": current_text[:500],
                    "log": wf_log[:],
                    "saved_at": datetime.now().isoformat(),
                }, total=n)
            except Exception:
                pass

        prog.progress(1.0,tt("✅ Workflow dokončeno", "✅ Workflow done", _L))
        status.success(tt("✅ Workflow dokončeno", "✅ Workflow done", _L))
        # ── Uložení finálního výsledku workflow ──────────────────────
        try:
            _wf_safe_name    = re.sub(r'[^\w]', '_', wf_file.name[:20])
            _wf_final_prefix = f"wf_final_{_wf_safe_name}"
            _temp_save_txt(_wf_final_prefix, 1, current_text)
            _temp_save_json(f"{_wf_final_prefix}_log", 1, {
                "file": wf_file.name, "steps": st.session_state["workflow_steps"],
                "log": wf_log, "saved_at": datetime.now().isoformat(),
                "text_length": len(current_text),
            })
        except Exception:
            pass
        # Synchronizuj výsledky workflow do autosave klíčů
        if current_text.strip():
            try:
                st.session_state["last_extracted_text"] = current_text
                st.session_state["last_extraction_text"] = current_text
            except Exception:
                pass
        save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
        gsb_done_stay_on_tab(tt(f"✅ Workflow dokončen: {wf_file.name}",
                    f"✅ Workflow done: {wf_file.name}", _L))

        st.markdown(tt("### 📊 Výsledek workflow", "### 📊 Workflow result", _L))
        st.text_area(tt("Výstup", "Output", _L), current_text, height=300)
        with st.expander(tt("📋 Log", "📋 Log", _L)):
            st.code("\n".join(wf_log))
        if wf_val_results:
            df_wv = validation_results_to_df(wf_val_results)
            st.markdown(tt("#### Validace taxonů", "#### Taxon validation", _L))
            st.dataframe(df_wv, width='stretch')
            st.download_button(tt("⬇️ Validace CSV", "⬇️ Validation CSV", _L),
                               df_wv.to_csv(index=False).encode("utf-8"),
                               "wf_validace.csv","text/csv")
        render_export_buttons(current_text,"workflow_result",
                              {"steps":st.session_state["workflow_steps"]})

        # Hyolitha export z workflow
        if st.session_state.get("hyolitha_export_records"):
            recs_wf = st.session_state["hyolitha_export_records"]
            st.markdown(tt(f"#### 🦪 Hyolitha export ({len(recs_wf)} záznamů)",
                           f"#### 🦪 Hyolitha export ({len(recs_wf)} records)", _L))
            col_wh1, col_wh2 = st.columns(2)
            with col_wh1:
                if XLSX_AVAILABLE:
                    st.download_button(tt("⬇️ XLSX (4 listy)", "⬇️ XLSX (4 sheets)", _L),
                        hyolitha_export_xlsx(recs_wf), "hyolitha_wf.xlsx", width='stretch')
            with col_wh2:
                st.download_button(tt("⬇️ CSV ZIP", "⬇️ CSV ZIP", _L),
                    hyolitha_export_csv_zip(recs_wf), "hyolitha_wf_csv.zip", width='stretch')

        save_to_history("workflow",{"file":wf_file.name,"steps":st.session_state["workflow_steps"]})

        # ── Webhook notifikace ────────────────────────
        _wf_elapsed = (_time.time() - _wf_start_time) / 60  # minuty
        _wh_url = st.session_state.get("webhook_url","") or st.session_state.get("wf_webhook_url","")
        _wh_thresh = st.session_state.get("wf_wh_threshold", 2)
        if _wh_url and _wf_elapsed >= _wh_thresh:
            ok_steps = sum(1 for l in wf_log if "[OK]" in l)
            msg = tt(f"Workflow dokončen za {_wf_elapsed:.1f} min | "
                     f"Soubor: {wf_file.name} | Kroky: {ok_steps}/{n} OK",
                     f"Workflow completed in {_wf_elapsed:.1f} min | "
                     f"File: {wf_file.name} | Steps: {ok_steps}/{n} OK", _L)
            if send_webhook_notification(_wh_url, msg, "SciNexus — Workflow"):
                st.toast(tt("🔔 Webhook notifikace odeslána",
                            "🔔 Webhook notification sent", _L))


# ══════════════════════════════════════════════════════
# 7 – HISTORIE  (rozšířená: replay, anotace, filtrování)
# ══════════════════════════════════════════════════════
with main_tabs[7]:
    keep_current_tab(7)   # ← v24.17
    _L = st.session_state.get("lang","cz")
    st.markdown(t("hist_title",_L))
    with st.expander("ℹ️ Jak funguje historie?", expanded=False):
        st.markdown("""
**Historie** zaznamenává každou provedenou operaci (extrakce, překlad, validace, chat) s časovým razítkem, parametry a modelem.

**Funkce:**
- **Filtrování** podle operace a data, fulltext vyhledávání
- **Anotace** — přidej poznámku ke každému záznamu ("špatná kvalita OCR")
- **Replay** — obnov parametry historické operace a spusť ji znovu s úpravami
- **Diff extrakcí** — porovnej dva snapshoty extrakce stejného dokumentu (přibylo/ubylo záznamů)
- **Export CSV** — stáhni celou historii jako tabulku
- **Automatická záloha** — každý den při prvním spuštění se uloží kopie do `backups/`
        """)
    hist = load_history()

    # ── FTS full-text search ──────────────────────────
    _fts_count = fts_record_count()
    with st.expander(tt(f"🔎 Full-text search ve výsledcích extrakce ({_fts_count:,} záznamů v indexu)",
                         f"🔎 Full-text search in extraction results ({_fts_count:,} records in index)", _L)):
        fts_q = st.text_input(tt("Hledaný výraz (taxon, lokalita, autor…)",
                                  "Search term (taxon, locality, author…)", _L),
                              key="fts_query",
                              placeholder=tt("např. Hyolithes, Cambrian, Barrande",
                                             "e.g. Hyolithes, Cambrian, Barrande", _L))
        _fts_col1, _fts_col2 = st.columns([3, 1])
        with _fts_col2:
            fts_limit = st.number_input(tt("Max výsledků", "Max results", _L),
                                        5, 100, 20, 5, key="fts_limit")
        with _fts_col1:
            fts_run = st.button(tt("🔎 Hledat", "🔎 Search", _L), key="fts_run", type="primary",
                                disabled=not fts_q.strip())
        if fts_run and fts_q.strip():
            with st.spinner(tt("Prohledávám index…", "Searching index…", _L)):
                fts_hits = fts_search(fts_q.strip(), limit=int(fts_limit))
            if fts_hits:
                st.success(tt(f"Nalezeno {len(fts_hits)} výsledků",
                              f"Found {len(fts_hits)} results", _L))
                for hit in fts_hits:
                    with st.expander(
                        f"📄 **{hit['source_file']}** — {hit['timestamp']}", expanded=False):
                        # Zobraz snippet s vyznačením
                        snippet = hit.get("snippet", "")
                        if snippet:
                            st.markdown(
                                snippet.replace("<mark>", "**`").replace("</mark>", "`**"),
                                unsafe_allow_html=False)
                        # Tlačítko pro stažení plného výsledku
                        if hit.get("result_json"):
                            st.download_button(
                                tt("⬇️ Stáhnout výsledek", "⬇️ Download result", _L),
                                hit["result_json"].encode("utf-8"),
                                f"fts_{hit['source_file']}.json",
                                key=f"fts_dl_{hit['id']}")
            else:
                st.info(tt(f"Žádné výsledky pro '{fts_q}'",
                           f"No results for '{fts_q}'", _L))
        if _fts_count == 0:
            st.caption(tt("💡 Index se plní automaticky při každé extrakci. "
                          "Spusť libovolnou extrakci pro první zaindexování.",
                          "💡 The index is populated automatically with each extraction. "
                          "Run any extraction for first indexing.", _L))

    if not hist:
        st.info(t("hist_empty",_L))
    else:
        # ── Filtrování a vyhledávání ──────────────────
        hf_col1, hf_col2, hf_col3 = st.columns([2,2,1])
        with hf_col1:
            hist_search = st.text_input("🔍 Hledat v historii" if _L=="cz" else "🔍 Search history",
                                         key="hist_search", placeholder="taxon, operace…" if _L=="cz" else "taxon, operation…")
        with hf_col2:
            all_ops = sorted(set(h["operation"] for h in hist))
            hist_op_filter = st.multiselect(
                "Filtr operace" if _L=="cz" else "Filter operation",
                all_ops, default=all_ops, key="hist_op_filter")
        with hf_col3:
            hist_sort = st.selectbox("Řazení" if _L=="cz" else "Sort",
                                     ["Nejnovější" if _L=="cz" else "Newest",
                                      "Nejstarší" if _L=="cz" else "Oldest"],
                                     key="hist_sort")

        # Aplikuj filtry
        filtered_hist = hist
        if hist_op_filter:
            filtered_hist = [h for h in filtered_hist if h["operation"] in hist_op_filter]
        if hist_search.strip():
            q = hist_search.strip().lower()
            filtered_hist = [h for h in filtered_hist
                             if q in str(h).lower()]
        if hist_sort in ("Nejstarší","Oldest"):
            filtered_hist = list(reversed(filtered_hist))

        st.caption(f"Zobrazeno {len(filtered_hist)} z {len(hist)} záznamů"
                   if _L=="cz" else f"Showing {len(filtered_hist)} of {len(hist)} records")

        # ── Anotace pro záznamy (in-memory per session) ─
        if "hist_annotations" not in st.session_state:
            st.session_state["hist_annotations"] = {}

        for hidx, h in enumerate(filtered_hist[:50]):
            ts    = h["timestamp"][:16].replace("T"," ")
            op    = h["operation"]
            meta  = h.get("metadata", {})
            ann   = st.session_state["hist_annotations"].get(ts, "")
            op_icons = {"extrakce":"🔍","překlad":"🌐","validace":"🧬",
                        "chat":"💬","clean":"🧹","workflow":"⚙️"}
            op_icon = op_icons.get(op, "📋")

            with st.expander(f"{op_icon} **{op}** — {ts}" + (f"  ✏️ *{ann[:40]}*" if ann else ""),
                             expanded=False):
                # Metadata
                mc = st.columns(3)
                for ci, (k,v) in enumerate(list(meta.items())[:6]):
                    mc[ci%3].caption(f"**{k}**: {v}")

                # Anotace
                new_ann = st.text_input(
                    tt("✏️ Poznámka", "✏️ Note", _L),
                    value=ann, key=f"hist_ann_{hidx}",
                    placeholder=tt("např. špatné OCR, zkusit jiný prompt…",
                                   "e.g. bad OCR, try different prompt…", _L))
                if new_ann != ann:
                    st.session_state["hist_annotations"][ts] = new_ann

                # Replay tlačítko
                if op == "validace" and meta.get("databases"):
                    if st.button("🔄 Znovu spustit validaci" if _L=="cz"
                                 else "🔄 Replay validation",
                                 key=f"hist_replay_{hidx}"):
                        # Přepneme na záložku validace a předvyplníme databáze
                        st.session_state["hist_replay_dbs"] = meta.get("databases",[])
                        st.toast("✅ Parametry obnoveny — přejdi na záložku Validace"
                                 if _L=="cz" else "✅ Parameters restored — go to Validation tab")

        st.divider()
        hcl1, hcl2 = st.columns([1,3])
        with hcl1:
            if st.button(t("hist_clear",_L), key="hist_clr", type="secondary"):
                if os.path.exists(HISTORY_FILE): os.remove(HISTORY_FILE)
                st.success(t("hist_cleared",_L)); st.rerun()
        with hcl2:
            df_h = pd.DataFrame([{
                t("hist_time",_L): h["timestamp"][:16],
                t("hist_op",_L):   h["operation"],
                t("hist_detail",_L): str(h.get("metadata",""))[:80],
                "Poznámka": st.session_state["hist_annotations"].get(h["timestamp"][:16],"")
            } for h in filtered_hist])
            st.download_button(
                "⬇️ Export CSV",
                df_h.to_csv(index=False).encode("utf-8"),
                "historie.csv", "text/csv", key="hist_export")

        # ── Diff mezi operacemi (extraction snapshots) ──
        st.divider()
        with st.expander(tt("🔀 Diff mezi uloženými extrakcemi",
                             "🔀 Diff between saved extractions", _L)):
            st.caption(tt("Porovnej výsledky dvou timestampovaných snapshotů extrakce.",
                          "Compare results of two timestamped extraction snapshots.", _L))
            ext_versions = list_extraction_versions()
            if len(ext_versions) >= 2:
                diff_col1, diff_col2 = st.columns(2)
                with diff_col1:
                    snap_a_label = st.selectbox(
                        tt("Snapshot A (starší)", "Snapshot A (older)", _L),
                        [f"{v['timestamp']} {v['label']}" for v in ext_versions],
                        key="hist_diff_a", index=min(1, len(ext_versions)-1))
                with diff_col2:
                    snap_b_label = st.selectbox(
                        tt("Snapshot B (novější)", "Snapshot B (newer)", _L),
                        [f"{v['timestamp']} {v['label']}" for v in ext_versions],
                        key="hist_diff_b", index=0)

                if st.button(tt("🔀 Porovnat snapshoty", "🔀 Compare snapshots", _L),
                             key="hist_diff_run"):
                    idx_a = [f"{v['timestamp']} {v['label']}" for v in ext_versions].index(snap_a_label)
                    idx_b = [f"{v['timestamp']} {v['label']}" for v in ext_versions].index(snap_b_label)
                    try:
                        with open(ext_versions[idx_a]["file"],"r",encoding="utf-8") as fa:
                            data_a = json.load(fa).get("data","")
                        with open(ext_versions[idx_b]["file"],"r",encoding="utf-8") as fb:
                            data_b = json.load(fb).get("data","")
                        diff = diff_extraction_records(
                            data_a if isinstance(data_a,str) else json.dumps(data_a),
                            data_b if isinstance(data_b,str) else json.dumps(data_b))
                        d1,d2,d3,d4 = st.columns(4)
                        d1.metric(tt("Záznamy A", "Records A", _L), diff["total_old"])
                        d2.metric(tt("Záznamy B", "Records B", _L), diff["total_new"])
                        d3.metric(tt("➕ Přibylo",  "➕ Added",    _L), len(diff["added"]))
                        d4.metric(tt("➖ Ubylo",   "➖ Removed",  _L), len(diff["removed"]))
                        if diff["added"]:
                            with st.expander(tt(f"➕ Přidané záznamy ({len(diff['added'])})",
                                                 f"➕ Added records ({len(diff['added'])})", _L)):
                                for r in diff["added"][:10]:
                                    st.code(json.dumps(r, ensure_ascii=False)[:200])
                        if diff["removed"]:
                            with st.expander(tt(f"➖ Odebrané záznamy ({len(diff['removed'])})",
                                                 f"➖ Removed records ({len(diff['removed'])})", _L)):
                                for r in diff["removed"][:10]:
                                    st.code(json.dumps(r, ensure_ascii=False)[:200])
                        if diff["changed"]:
                            with st.expander(tt(f"✏️ Změněné záznamy ({len(diff['changed'])})",
                                                 f"✏️ Changed records ({len(diff['changed'])})", _L)):
                                for ch in diff["changed"][:10]:
                                    st.markdown(f"- `{ch['key']}`: " +
                                               ", ".join(f"{k}: `{v[0]}→{v[1]}`"
                                                         for k,v in list(ch["diffs"].items())[:3]))
                    except Exception as e:
                        st.error(tt(f"Chyba diff: {e}", f"Diff error: {e}", _L))
            elif len(ext_versions) == 1:
                st.info(tt("Pro diff jsou potřeba alespoň 2 snapshoty. Ulož více verzí extrakce.",
                           "Diff needs at least 2 snapshots. Save more extraction versions.", _L))
            else:
                st.info(tt("Žádné snapshoty extrakcí. Extrahuj data a ulož snapshot v záložce Extrakce.",
                           "No extraction snapshots. Extract data and save a snapshot in the Extraction tab.", _L))


# ══════════════════════════════════════════════════════
# EXPORT CELÉ SESSION JAKO ZIP
# ══════════════════════════════════════════════════════
def export_session_zip() -> bytes:
    """Zabalí všechna data session do ZIP archivu."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # Session state (serializovatelné klíče)
        session_data = {}
        for k in SESSION_PERSIST_KEYS:
            v = st.session_state.get(k)
            if v is not None:
                try:
                    json.dumps(v)
                    session_data[k] = v
                except (TypeError, ValueError):
                    pass
        zf.writestr("session_state.json",
                    json.dumps({"exported_at": datetime.now().isoformat(),
                                "data": session_data},
                               ensure_ascii=False, indent=2))
        # Disk cache info
        if _DISK_CACHE_PATH.exists():
            zf.write(str(_DISK_CACHE_PATH), "validation_cache.db")
        # Processing history
        if os.path.exists(HISTORY_FILE):
            zf.write(HISTORY_FILE, "processing_history.json")
        # Extraction snapshots
        if EXTRACTION_HISTORY_DIR.exists():
            for snap in EXTRACTION_HISTORY_DIR.glob("extraction_*.json"):
                zf.write(str(snap), f"extraction_history/{snap.name}")
        # Prompt profiles
        if os.path.exists(PROMPT_PROFILE_FILE):
            zf.write(PROMPT_PROFILE_FILE, "prompt_profiles.json")
        # Chat conversations
        if os.path.exists("chat_conversations.json"):
            zf.write("chat_conversations.json", "chat_conversations.json")
        # README
        readme = (f"SciNexus v{VERSION} — Session Export\n"
                  f"Exported: {datetime.now():%Y-%m-%d %H:%M}\n\n"
                  "Files:\n"
                  "  session_state.json   — validation results, extractions, translations\n"
                  "  validation_cache.db  — SQLite validation cache\n"
                  "  processing_history.json — operation history\n"
                  "  extraction_history/  — timestamped extraction snapshots\n"
                  "  prompt_profiles.json — saved prompt profiles\n"
                  "  chat_conversations.json — named chat conversations\n")
        zf.writestr("README.txt", readme)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════
# ONBOARDING WIZARD (při prvním spuštění)
# ══════════════════════════════════════════════════════
_is_first_run = (
    not st.session_state.get("glossaries") and
    not load_history() and
    not st.session_state.get("last_extracted_text") and
    not st.session_state.get("_onboarding_dismissed")
)
if _is_first_run:
    _L_ob = st.session_state.get("lang","cz")
    with st.expander(tt("👋 Vítej! Jak začít? (průvodce prvním spuštěním)",
                         "👋 Welcome! How to start? (first-run guide)", _L_ob), expanded=True):
        st.markdown(tt(
            """
**Krok 1 — Připoj LM Studio**
- Spusť LM Studio na svém počítači
- Nastav API URL v sidebaru (výchozí: `http://127.0.0.1:1234/v1`)
- Klikni **Načíst modely** — zobrazí se dostupné modely

**Krok 2 — Nahraj první PDF**
- Přejdi na záložku **🔍 Extrakce**
- Nahraj PDF s paleontologickými daty
- Nastav prompt (nebo použij šablonu)

**Krok 3 — Spusť extrakci**
- Klikni **Spustit extrakci** — výsledky se zobrazí průběžně
- Výsledky se automaticky uloží jako snapshot

**Krok 4 — Validuj taxony**
- Přejdi na záložku **🧬 Validace**
- Taxony z extrakce se automaticky načtou
- Vyber databáze a klikni **Spustit validaci**
            """,
            """
**Step 1 — Connect LM Studio**
- Launch LM Studio on your computer
- Set the API URL in the sidebar (default: `http://127.0.0.1:1234/v1`)
- Click **Load models** — available models will appear

**Step 2 — Upload your first PDF**
- Go to the **🔍 Extraction** tab
- Upload a PDF with paleontological data
- Set the prompt (or use a template)

**Step 3 — Run extraction**
- Click **Run extraction** — results appear progressively
- Results are automatically saved as a snapshot

**Step 4 — Validate taxa**
- Go to the **🧬 Validation** tab
- Taxa from the extraction load automatically
- Select databases and click **Run validation**
            """,
            _L_ob))
        if st.button(tt("✅ Rozumím, zavřít průvodce",
                         "✅ Got it, close guide", _L_ob), key="onboarding_dismiss"):
            st.session_state["_onboarding_dismissed"] = True
            st.rerun()


# ══════════════════════════════════════════════════════
# PATIČKA  (s export ZIP tlačítkem)
# ══════════════════════════════════════════════════════
st.divider()
_L_ft = st.session_state.get("lang","cz")
ft_col1, ft_col2 = st.columns([4,1])
with ft_col1:
    st.caption(f"🧠 **{APP_TITLE} v{VERSION}** | " + t("footer", _L_ft))
with ft_col2:
    if st.button(tt("📦 Export session ZIP", "📦 Export session ZIP", _L_ft),
                 key="ft_export_zip",
                 help=tt("Stáhni celou session jako ZIP archiv",
                         "Download the entire session as a ZIP archive", _L_ft)):
        zip_bytes = export_session_zip()
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        st.download_button(
            f"⬇️ Stáhnout session_{ts}.zip",
            zip_bytes, f"session_{ts}.zip", "application/zip",
            key="ft_dl_zip")


# ══════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════
# 8 – A/B PROMPT TESTING
# ══════════════════════════════════════════════════════
with main_tabs[8]:
    keep_current_tab(8)   # ← v24.17
    _L = st.session_state.get("lang", "cz")
    st.markdown("## 🧪 A/B Prompt Testing")
    st.caption(tt("Porovnej výsledky různých promptů na stejném textu — s metrikami.",
                  "Compare results of different prompts on the same text — with metrics.", _L))

    with st.expander(tt("ℹ️ Jak funguje A/B testing?",
                         "ℹ️ How does A/B testing work?", _L), expanded=False):
        st.markdown(tt(
            """
**A/B testing** spustí extrakci stejného textu s různými prompty najednou (paralelně)
a výsledky zobrazí vedle sebe s metrikami:

- **Počet extrahovaných záznamů** — recall
- **Completeness score** — průměrné zaplnění polí (0–100 %)
- **Čas zpracování** — rychlost každého promptu
- **JSON validita** — zda výstup je parsovatelný JSON

**Doporučený postup:**
1. Nahraj testovací PDF nebo vlož text
2. Zadej 2–4 varianty promptu
3. Spusť test — výsledky se zobrazí vedle sebe
4. Ulož nejlepší prompt jako šablonu
            """,
            """
**A/B testing** runs extraction on the same text with different prompts in parallel,
showing results side by side with metrics:

- **Number of extracted records** — recall
- **Completeness score** — average field fill rate (0–100 %)
- **Processing time** — speed of each prompt
- **JSON validity** — whether output is parseable JSON

**Recommended workflow:**
1. Upload a test PDF or paste text
2. Enter 2–4 prompt variants
3. Run the test — results appear side by side
4. Save the best prompt as a template
            """,
            _L))

    # ── Vstupní text ─────────────────────────────────
    _ab_src_map = {
        "✏️ Ručně":      tt("✏️ Ručně", "✏️ Manual",   _L),
        "📄 Ze souboru": tt("📄 Ze souboru", "📄 From file", _L),
    }
    ab_src_mode = st.radio(tt("Zdroj textu", "Text source", _L),
                           list(_ab_src_map.keys()),
                           format_func=lambda k: _ab_src_map[k],
                           horizontal=True, key="ab_src_mode")
    ab_text = ""
    if ab_src_mode == "✏️ Ručně":
        ab_text = st.text_area(tt("Testovací text", "Test text", _L),
                               height=150, key="ab_text_manual",
                               placeholder=tt("Vlož výňatek z paleontologické publikace…",
                                              "Paste an excerpt from a paleontological paper…", _L))
    else:
        ab_file = st.file_uploader(tt("Soubor (PDF/DOCX/TXT)", "File (PDF/DOCX/TXT)", _L),
                                   type=["pdf","docx","txt"],
                                   key="ab_file")
        if ab_file:
            ab_text = read_uploaded_file(ab_file)[:8000]  # jen prvních 8k znaků pro test
            st.caption(tt(f"Načteno {len(ab_text):,} znaků (zkráceno na 8 000 pro test)",
                          f"Loaded {len(ab_text):,} chars (truncated to 8 000 for test)", _L))

    st.divider()

    # ── Prompty ───────────────────────────────────────
    st.markdown(tt("### Varianty promptů", "### Prompt variants", _L))
    ab_n = st.slider(tt("Počet variant", "Number of variants", _L), 2, 4, 2, 1, key="ab_n_variants")

    ab_prompts = []
    ab_labels  = []
    ab_cols    = st.columns(ab_n)
    for _abi in range(ab_n):
        with ab_cols[_abi]:
            lbl = st.text_input(tt(f"Název varianty {_abi+1}",
                                    f"Variant {_abi+1} name", _L),
                                value=tt(f"Varianta {_abi+1}",
                                         f"Variant {_abi+1}", _L),
                                key=f"ab_label_{_abi}")
            prm = st.text_area(
                f"Prompt {_abi+1}",
                height=150,
                key=f"ab_prompt_{_abi}",
                value=(
                    "Extrahuj taxonomické záznamy jako JSON: [{rod,druh,autor,rok,lokalita,stratigrafie}]"
                    if _abi == 0 else
                    "Extrahuj VŠECHNY taxonomické záznamy. Pro každý záznam uveď: rod, druh, autor, rok vydání, lokalita nálezu, stratigrafická jednotka. Výstup jako JSON pole objektů."
                    if _abi == 1 else ""
                )
            )
            ab_labels.append(lbl)
            ab_prompts.append(prm)

    st.divider()

    # ── Nastavení testu ───────────────────────────────
    ab_set_col1, ab_set_col2 = st.columns(2)
    with ab_set_col1:
        ab_sys_mode = st.selectbox(tt("Extrakční režim (system prompt)",
                                       "Extraction mode (system prompt)", _L),
                                   list(EXTRACTION_MODES.keys()),
                                   key="ab_sys_mode")
    with ab_set_col2:
        ab_temp = st.slider(tt("Teplota", "Temperature", _L), 0.0, 1.0, 0.1, 0.05, key="ab_temp")

    ab_run = st.button(tt("🧪 Spustit A/B test", "🧪 Run A/B test", _L), type="primary",
                       disabled=not (selected_model and ab_text.strip()
                                     and all(p.strip() for p in ab_prompts[:ab_n])),
                       width='stretch')

    if ab_run:
        _ab_sys = EXTRACTION_MODES[ab_sys_mode]["system"]
        _ab_results: Dict[int, dict] = {}

        with st.status(tt(f"🧪 Spouštím {ab_n} variant paralelně…",
                           f"🧪 Running {ab_n} variants in parallel…", _L),
                       expanded=True) as ab_status:
            ab_start_all = _time.time()

            def _run_ab_variant(idx_prompt):
                idx, prompt_text = idx_prompt
                t0 = _time.time()
                try:
                    resp = chat_completion(
                        base_url, selected_model,
                        [{"role": "system", "content": _ab_sys},
                         {"role": "user",   "content": f"{prompt_text}\n\n{ab_text}"}],
                        temp=ab_temp, max_tokens=6000)
                    elapsed = _time.time() - t0
                    # Metriky
                    records = []
                    is_json = False
                    try:
                        clean_ab = _JSON_FENCE_RE.sub("", resp).strip()
                        parsed_ab = json.loads(clean_ab)
                        if isinstance(parsed_ab, list):
                            records = parsed_ab
                        elif isinstance(parsed_ab, dict):
                            records = [parsed_ab]
                        is_json = True
                    except Exception:
                        pass
                    # Completeness: průměrné % neprázdných polí
                    completeness = 0.0
                    if records:
                        scores = []
                        for rec in records:
                            if isinstance(rec, dict):
                                vals = [v for v in rec.values()
                                        if v and str(v).strip().lower()
                                        not in ("null","none","n/a","")]
                                scores.append(len(vals) / max(1, len(rec)))
                        completeness = (sum(scores) / len(scores) * 100) if scores else 0.0
                    return idx, {
                        "label":        ab_labels[idx],
                        "prompt":       prompt_text,
                        "response":     resp,
                        "records":      records,
                        "n_records":    len(records),
                        "is_json":      is_json,
                        "completeness": completeness,
                        "elapsed":      elapsed,
                        "ok":           True,
                    }
                except Exception as e:
                    return idx, {"label": ab_labels[idx], "prompt": prompt_text,
                                 "response": str(e), "records": [], "n_records": 0,
                                 "is_json": False, "completeness": 0.0,
                                 "elapsed": _time.time() - t0, "ok": False}

            with ThreadPoolExecutor(max_workers=ab_n) as _ab_pool:
                _ab_futures = {
                    _ab_pool.submit(_run_ab_variant, (i, ab_prompts[i])): i
                    for i in range(ab_n)
                }
                for fut in as_completed(_ab_futures):
                    idx, res = fut.result()
                    _ab_results[idx] = res
                    icon = "✅" if res["ok"] else "❌"
                    ab_status.write(
                        f"{icon} **{res['label']}** — "
                        f"{res['n_records']} záznamů, "
                        f"completeness {res['completeness']:.0f} %, "
                        f"{res['elapsed']:.1f} s"
                    )

        st.session_state["ab_results"] = _ab_results
        save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
        desktop_notify("🧪 A/B test dokončen",
                       f"{ab_n} variant, {_time.time()-ab_start_all:.0f} s")

    # ── Výsledky A/B testu ────────────────────────────
    if st.session_state.get("ab_results"):
        _abr = st.session_state["ab_results"]
        st.divider()
        st.markdown(tt("### 📊 Výsledky", "### 📊 Results", _L))

        # Souhrnná tabulka metrik
        _ab_metric_rows = []
        _m_variant = tt("Varianta",      "Variant",     _L)
        _m_records = tt("Záznamy",       "Records",     _L)
        _m_compl   = tt("Completeness",  "Completeness", _L)
        _m_json    = tt("JSON validní",  "JSON valid",  _L)
        _m_time    = tt("Čas (s)",       "Time (s)",    _L)
        for idx in sorted(_abr.keys()):
            r = _abr[idx]
            _ab_metric_rows.append({
                _m_variant: r["label"],
                _m_records: r["n_records"],
                _m_compl:   f"{r['completeness']:.1f} %",
                _m_json:    "✅" if r["is_json"] else "❌",
                _m_time:    f"{r['elapsed']:.1f}",
            })
        st.dataframe(pd.DataFrame(_ab_metric_rows), width='stretch',
                     hide_index=True)

        # Nejlepší varianta podle počtu záznamů × completeness
        _best_idx = max(_abr.keys(),
                        key=lambda i: _abr[i]["n_records"] * (_abr[i]["completeness"] / 100 + 0.01))
        st.success(tt(f"🏆 Doporučená varianta: **{_abr[_best_idx]['label']}** "
                      f"({_abr[_best_idx]['n_records']} záznamů, "
                      f"{_abr[_best_idx]['completeness']:.0f} % completeness)",
                      f"🏆 Recommended variant: **{_abr[_best_idx]['label']}** "
                      f"({_abr[_best_idx]['n_records']} records, "
                      f"{_abr[_best_idx]['completeness']:.0f} % completeness)", _L))

        # Detail výsledků vedle sebe
        _ab_detail_cols = st.columns(len(_abr))
        for ci, idx in enumerate(sorted(_abr.keys())):
            r = _abr[idx]
            with _ab_detail_cols[ci]:
                crown = " 🏆" if idx == _best_idx else ""
                st.markdown(f"**{r['label']}{crown}**")
                # Preview prvních 3 záznamů
                if r["records"]:
                    for ri, rec in enumerate(r["records"][:3]):
                        with st.expander(
                            tt(f"Záznam {ri+1}: {rec.get('druh') or rec.get('rod','?')}",
                               f"Record {ri+1}: {rec.get('druh') or rec.get('rod','?')}", _L),
                            expanded=ri==0):
                            for k, v in rec.items():
                                if v and k != "_confidence":
                                    st.caption(f"**{k}**: {str(v)[:80]}")
                else:
                    st.text(r["response"][:300])

                # Export + uložení jako šablona
                st.download_button(
                    "⬇️ JSON",
                    r["response"].encode("utf-8"),
                    f"ab_{r['label']}.json",
                    key=f"ab_dl_{idx}",
                    width='stretch')

                _tpl_name_key = f"ab_tpl_name_{idx}"
                if _tpl_name_key not in st.session_state:
                    st.session_state[_tpl_name_key] = f"AB_{r['label']}"
                tpl_name = st.text_input(tt("Uložit jako šablonu",
                                             "Save as template", _L),
                                         value=st.session_state[_tpl_name_key],
                                         key=f"ab_tpl_input_{idx}",
                                         label_visibility="collapsed")
                if st.button(tt("💾 Uložit prompt jako šablonu",
                                 "💾 Save prompt as template", _L),
                             key=f"ab_save_tpl_{idx}",
                             width='stretch'):
                    _templates = st.session_state.get("prompt_templates", {})
                    _templates[tpl_name] = {
                        "desc":      f"A/B test — {r['label']}",
                        "lang_hint": "angličtina",
                        "prompt":    r["prompt"],
                    }
                    st.session_state["prompt_templates"] = _templates
                    save_templates(_templates)
                    st.toast(tt(f"✅ Šablona '{tpl_name}' uložena",
                                f"✅ Template '{tpl_name}' saved", _L))


# ══════════════════════════════════════════════════════
# 9 – NASTAVENÍ
# ══════════════════════════════════════════════════════
with main_tabs[9]:
    keep_current_tab(9)   # ← v24.17
    _L = st.session_state.get("lang", "cz")
    st.markdown(tt("## ⚙️ Nastavení", "## ⚙️ Settings", _L))

    _ns_tab1, _ns_tab2, _ns_tab3, _ns_tab4 = st.tabs([
        tt("🤖 Model manager",         "🤖 Model manager",        _L),
        tt("🧠 Prompty per model",     "🧠 Prompts per model",    _L),
        tt("🔔 Notifikace & Webhook",  "🔔 Notifications & Webhook", _L),
        tt("🗄️ Data & Cache",          "🗄️ Data & Cache",         _L),
    ])

    # ── Model manager ────────────────────────────────
    with _ns_tab1:
        st.markdown(tt("#### 🤖 Dostupné modely", "#### 🤖 Available models", _L))
        MODEL_USECASE = {
            "qwen2.5": ("Extrakce, překlad, čištění", "128k ctx"),
            "qwen3":   ("Reasoning, extrakce", "32k ctx"),
            "mistral": ("OCR korekce, překlad", "32k ctx"),
            "deepseek":("Reasoning, složitá extrakce", "64k ctx"),
            "llama":   ("Rychlé úlohy, chat", "8k ctx"),
        }
        if st.session_state.get("available_models"):
            _mm_rows = []
            _mm_model    = tt("Model",                "Model",               _L)
            _mm_use      = tt("Doporučené použití",   "Recommended use",     _L)
            _mm_ctx      = tt("Context",              "Context",             _L)
            _mm_active   = tt("Aktivní",              "Active",              _L)
            for m in st.session_state["available_models"]:
                ml = m.lower()
                use, ctx = next(
                    ((u,c) for k,(u,c) in MODEL_USECASE.items() if k in ml),
                    ("obecné použití", "?"))
                _mm_rows.append({
                    _mm_model:  m,
                    _mm_use:    use,
                    _mm_ctx:    ctx,
                    _mm_active: "⭐" if m == selected_model else "",
                })
            st.dataframe(pd.DataFrame(_mm_rows), width='stretch', hide_index=True)
        else:
            st.info(tt("Načti modely tlačítkem 🔄 Připojit v sidebaru.",
                       "Load models with 🔄 Connect in the sidebar.", _L))

        st.divider()
        st.markdown(tt("#### 🔄 LM Studio CLI", "#### 🔄 LM Studio CLI", _L))
        _ns_lms_ok = _lms_available()
        if _ns_lms_ok:
            st.success(tt("✅ lms CLI nalezeno v PATH", "✅ lms CLI found in PATH", _L))
            if st.button(tt("📋 Zobrazit stažené modely (lms ls)",
                             "📋 Show downloaded models (lms ls)", _L), key="ns_lms_ls"):
                _ns_models, _ns_raw = lms_ls()
                if _ns_models:
                    st.dataframe(pd.DataFrame(_ns_models)[["id","size_gb","architecture","context"]],
                                 width='stretch', hide_index=True)
                else:
                    st.code(_ns_raw)
        else:
            st.warning(tt("⚠️ lms CLI není dostupný — nainstaluj LM Studio CLI pro správu modelů.",
                          "⚠️ lms CLI not available — install LM Studio CLI for model management.", _L))

    # ── Prompty per model ─────────────────────────────
    with _ns_tab2:
        st.markdown(tt("#### 🧠 System prompt per model",
                       "#### 🧠 System prompt per model", _L))
        st.caption(tt("Každý model dostane tento prompt automaticky přidaný ke každému volání.",
                      "Each model gets this prompt automatically prepended to every call.", _L))
        _pmp2 = load_per_model_prompts()
        _pmp2_keys = list(_pmp2.keys())
        _add_new = tt("+ přidat nový", "+ add new", _L)
        _ns_pmp_sel = st.selectbox(tt("Model pattern", "Model pattern", _L),
                                   _pmp2_keys + [_add_new],
                                   key="ns_pmp_sel")
        if _ns_pmp_sel == _add_new:
            _ns_pmp_key = st.text_input(tt("Pattern (část názvu modelu)",
                                            "Pattern (part of model name)", _L),
                                        key="ns_pmp_new_key",
                                        placeholder=tt("např. qwen, llama, mistral",
                                                        "e.g. qwen, llama, mistral", _L))
            _ns_pmp_val = st.text_area(tt("Systémový prompt", "System prompt", _L),
                                        key="ns_pmp_new_val", height=120)
            if st.button(tt("💾 Uložit nový prompt", "💾 Save new prompt", _L),
                         key="ns_pmp_save_new"):
                if _ns_pmp_key.strip():
                    _pmp2[_ns_pmp_key.strip()] = _ns_pmp_val
                    save_per_model_prompts(_pmp2)
                    st.toast(tt(f"✅ Uloženo: {_ns_pmp_key}",
                                f"✅ Saved: {_ns_pmp_key}", _L))
                    st.rerun()
        else:
            _ns_pmp_val_edit = st.text_area(
                tt("Prompt", "Prompt", _L),
                value=_pmp2.get(_ns_pmp_sel, ""),
                key="ns_pmp_edit", height=150)
            _nsp1, _nsp2 = st.columns(2)
            with _nsp1:
                if st.button(tt("💾 Uložit", "💾 Save", _L), key="ns_pmp_save", width='stretch'):
                    _pmp2[_ns_pmp_sel] = _ns_pmp_val_edit
                    save_per_model_prompts(_pmp2)
                    st.toast(tt("✅ Uloženo", "✅ Saved", _L))
            with _nsp2:
                if st.checkbox(tt(f"Opravdu smazat '{_ns_pmp_sel}'?",
                               f"Really delete '{_ns_pmp_sel}'?", _L),
                               key="chk_del_pmp"):
                    if st.button(tt("🗑️ Ano, smazat", "🗑️ Yes, delete", _L),
                                 key="ns_pmp_del", type="secondary", use_container_width=True):
                        _pmp2.pop(_ns_pmp_sel, None)
                        save_per_model_prompts(_pmp2)
                        st.toast(tt(f"Smazáno: {_ns_pmp_sel}", f"Deleted: {_ns_pmp_sel}", _L))
                        st.rerun()
            if selected_model:
                _active_pmp2 = get_model_system_prompt(selected_model)
                if _active_pmp2:
                    st.caption(tt(f"✅ Aktivní prompt pro **{selected_model[:40]}**",
                                  f"✅ Active prompt for **{selected_model[:40]}**", _L))
                else:
                    st.caption(tt("⚠️ Žádný prompt pro aktuální model",
                                  "⚠️ No prompt for the current model", _L))

        # ── Per-model nastavení (teplota, max_tokens) ─────
        st.divider()
        st.markdown(tt("#### 🎛️ Nastavení per model (teplota / max tokens)",
                       "#### 🎛️ Per-model settings (temperature / max tokens)", _L))
        st.caption(tt(
            "Tato nastavení se automaticky aplikují při výběru modelu v sidebaru.",
            "These settings are automatically applied when a model is selected in the sidebar.", _L))
        _pms = load_per_model_settings()
        _pms_keys = list(_pms.keys())
        _add_new_pms = tt("+ přidat nový", "+ add new", _L)
        _ns_pms_sel = st.selectbox(
            tt("Model pattern", "Model pattern", _L),
            _pms_keys + [_add_new_pms],
            key="ns_pms_sel")
        if _ns_pms_sel == _add_new_pms:
            _pms_key = st.text_input(
                tt("Pattern (část názvu modelu)", "Pattern (part of model name)", _L),
                key="ns_pms_new_key",
                placeholder=tt("např. qwen, llama, mistral", "e.g. qwen, llama, mistral", _L))
            _c1, _c2 = st.columns(2)
            with _c1:
                _pms_temp = st.slider(
                    tt("Teplota", "Temperature", _L), 0.0, 1.0, 0.1, 0.05,
                    key="ns_pms_new_temp")
            with _c2:
                _pms_tok = st.number_input(
                    tt("Max tokens", "Max tokens", _L), 512, 32768, 4096, 256,
                    key="ns_pms_new_tok")
            if st.button(tt("💾 Uložit nastavení", "💾 Save settings", _L),
                         key="ns_pms_save_new"):
                if _pms_key.strip():
                    _pms[_pms_key.strip()] = {"temperature": _pms_temp, "max_tokens": _pms_tok}
                    save_per_model_settings(_pms)
                    st.session_state["per_model_settings"] = _pms
                    st.toast(tt(f"✅ Uloženo: {_pms_key}", f"✅ Saved: {_pms_key}", _L))
                    st.rerun()
        elif _pms_keys:
            _cur = _pms.get(_ns_pms_sel, {})
            _e1, _e2 = st.columns(2)
            with _e1:
                _pms_temp_e = st.slider(
                    tt("Teplota", "Temperature", _L), 0.0, 1.0,
                    float(_cur.get("temperature", 0.1)), 0.05,
                    key="ns_pms_edit_temp")
            with _e2:
                _pms_tok_e = st.number_input(
                    tt("Max tokens", "Max tokens", _L), 512, 32768,
                    int(_cur.get("max_tokens", 4096)), 256,
                    key="ns_pms_edit_tok")
            _ps1, _ps2 = st.columns(2)
            with _ps1:
                if st.button(tt("💾 Uložit", "💾 Save", _L),
                             key="ns_pms_save", use_container_width=True):
                    _pms[_ns_pms_sel] = {"temperature": _pms_temp_e, "max_tokens": _pms_tok_e}
                    save_per_model_settings(_pms)
                    st.session_state["per_model_settings"] = _pms
                    st.toast(tt("✅ Uloženo", "✅ Saved", _L))
            with _ps2:
                if st.checkbox(tt(f"Opravdu smazat '{_ns_pms_sel}'?",
                                  f"Really delete '{_ns_pms_sel}'?", _L),
                               key="chk_del_pms"):
                    if st.button(tt("🗑️ Smazat", "🗑️ Delete", _L),
                                 key="ns_pms_del", type="secondary",
                                 use_container_width=True):
                        _pms.pop(_ns_pms_sel, None)
                        save_per_model_settings(_pms)
                        st.session_state["per_model_settings"] = _pms
                        st.toast(tt(f"Smazáno: {_ns_pms_sel}",
                                    f"Deleted: {_ns_pms_sel}", _L))
                        st.rerun()

        # ── v24.27: Správa Quick Presets ─────────────────
        st.divider()
        st.markdown(tt("#### ⚡ Správa Rychlých nastavení",
                       "#### ⚡ Quick Presets Management", _L))
        st.caption(tt("Vytvoř vlastní Rychlá nastavení — zobrazí se v sidebaru pro rychlé použití.",
                      "Create custom presets — they appear in the sidebar for quick access.", _L))

        _qp_name = st.text_input(
            tt("Název nového Rychlého nastavení", "New preset name", _L),
            placeholder=tt("např. Hyolitha rychlá extrakce", "e.g. Hyolitha fast extraction", _L),
            key="ns_qp_name",
        )
        _qp_desc = st.text_area(
            tt("Popis (volitelný)", "Description (optional)", _L),
            height=60, key="ns_qp_desc",
        )
        _qp_c1, _qp_c2 = st.columns(2)
        with _qp_c1:
            _qp_temp  = st.slider(tt("Teplota", "Temperature", _L), 0.0, 1.0, 0.1, 0.05, key="ns_qp_temp")
            _qp_chunk = st.slider(tt("Velikost bloku", "Chunk size", _L), 1000, 6000, 3200, 100, key="ns_qp_chunk")
        with _qp_c2:
            _qp_maxtok  = st.number_input(tt("Max tokens", "Max tokens", _L), 512, 32768, 4096, 256, key="ns_qp_maxtok")
            _qp_large   = st.checkbox(tt("Large Document Mode", "Large Document Mode", _L), key="ns_qp_large")
            _qp_parallel = st.checkbox(tt("Paralelní zpracování", "Parallel processing", _L), key="ns_qp_parallel")
        _qp_prompt = st.text_area(
            tt("Výchozí prompt (volitelný)", "Default prompt (optional)", _L),
            height=80, key="ns_qp_prompt",
        )

        if st.button(tt("💾 Uložit jako nové Rychlé nastavení", "💾 Save as new Quick Preset", _L),
                     type="primary", key="ns_qp_save"):
            if _qp_name.strip():
                st.session_state["quick_presets"][_qp_name.strip()] = {
                    "description": _qp_desc,
                    "temperature": _qp_temp,
                    "chunk_size":  _qp_chunk,
                    "max_tokens":  _qp_maxtok,
                    "large_mode":  _qp_large,
                    "parallel":    _qp_parallel,
                    "prompt":      _qp_prompt,
                }
                save_quick_presets(st.session_state["quick_presets"])
                st.success(tt(f"✅ Rychlé nastavení '{_qp_name.strip()}' uloženo",
                              f"✅ Preset '{_qp_name.strip()}' saved", _L))
                st.rerun()
            else:
                st.warning(tt("Zadej název presetu.", "Enter a preset name.", _L))

        # ── Vestavěné Rychlá nastavení (scénáře) ────────────────────────────
        st.divider()
        st.markdown(tt("**Vestavěné scénáře:**", "**Built-in scenarios:**", _L))
        _builtin_scenarios = {
            "Rychlá extrakce + validace": {
                "desc": tt("Optimalizované nastavení pro extrakci taxonomických dat + automatická validace.",
                            "Optimized for taxonomic data extraction + automatic validation.", _L),
                "settings": {"wf_large_mode": False, "temp_t": 0.05, "lms_max_concurrent": 4},
            },
            "Hyolitha velký PDF (300+ stran)": {
                "desc": tt("Large Document Mode, velké bloky, paralelní zpracování pro PDF 300+ stran.",
                            "Large Document Mode, large blocks, parallel processing for 300+ page PDFs.", _L),
                "settings": {"wf_large_mode": True, "wf_chunk_target": 4200, "lms_max_concurrent": 4},
            },
            "Historický text překlad (CZ/RU/DE→EN)": {
                "desc": tt("Sekvenční překlad, nižší teplota, vhodné pro historické dokumenty.",
                            "Sequential translation, lower temperature, suitable for historical documents.", _L),
                "settings": {"wf_large_mode": False, "wf_chunk_target": 2200, "temp_t": 0.1},
            },
            "Plný Workflow – velký dokument": {
                "desc": tt("Extrakce + překlad + validace, Large Mode, větší bloky.",
                            "Extraction + translation + validation, Large Mode, larger blocks.", _L),
                "settings": {"wf_large_mode": True, "wf_chunk_target": 3500, "lms_max_concurrent": 4},
            },
        }
        for _sc_name, _sc_info in _builtin_scenarios.items():
            with st.expander(f"🔷 {_sc_name}"):
                st.caption(_sc_info["desc"])
                st.json(_sc_info["settings"])
                if st.button(tt("🚀 Použít tento scénář", "🚀 Apply this scenario", _L),
                             key=f"ns_apply_builtin_{_sc_name}", type="primary",
                             use_container_width=True):
                    _apply = dict(_sc_info["settings"])
                    _apply["quick_preset"] = "— vyber scénář —" if _L=="cz" else "— select scenario —"
                    defer_apply_settings(_apply)
                    st.session_state["_active_preset_name"] = _sc_name
                    st.toast(tt(f"✅ Scénář '{_sc_name}' aplikován",
                                f"✅ Scenario '{_sc_name}' applied", _L))
                    st.rerun()

        # ── Uživatelská Rychlá nastavení ─────────────────────────────────────
        st.divider()
        _user_qp = st.session_state.get("quick_presets", {})
        if _user_qp:
            st.markdown(tt("**Moje Rychlá nastavení:**", "**My Quick Presets:**", _L))
            for _qp_n, _qp_d in list(_user_qp.items()):
                with st.expander(f"📌 {_qp_n}"):
                    if _qp_d.get("description"):
                        st.caption(_qp_d["description"])
                    _qp_del1, _qp_del2, _qp_del3 = st.columns([1, 1, 2])
                    with _qp_del1:
                        if st.button(tt("🗑️ Smazat", "🗑️ Delete", _L),
                                     key=f"ns_del_preset_{_qp_n}", type="secondary",
                                     width="stretch"):
                            st.session_state["quick_presets"].pop(_qp_n, None)
                            save_quick_presets(st.session_state["quick_presets"])
                            st.rerun()
                    with _qp_del2:
                        if st.button(tt("🚀 Použít nyní", "🚀 Apply now", _L),
                                     key=f"ns_apply_preset_{_qp_n}", width="stretch"):
                            _dq = {}
                            if "temperature" in _qp_d: _dq.update({"ext_temp": _qp_d["temperature"], "tr_temp": _qp_d["temperature"], "temp_t": _qp_d["temperature"]})
                            if "chunk_size"  in _qp_d: _dq.update({"tr_chunk_target": _qp_d["chunk_size"], "ext_chunk_size": _qp_d["chunk_size"]})
                            if "max_tokens"  in _qp_d: _dq["ext_max_tokens"] = _qp_d["max_tokens"]
                            if "large_mode"  in _qp_d: _dq["wf_large_mode"]  = _qp_d["large_mode"]
                            if "parallel"    in _qp_d: _dq["tr_parallel"]    = _qp_d["parallel"]
                            if _qp_d.get("prompt"):    _dq["ext_prompt"]     = _qp_d["prompt"]
                            defer_apply_settings(_dq)
                            st.session_state["_active_preset_name"] = _qp_n
                            st.toast(tt(f"✅ Nastavení '{_qp_n}' aplikováno",
                                        f"✅ Preset '{_qp_n}' applied", _L))
                    with _qp_del3:
                        with st.expander(tt("🔍 Zobrazit parametry", "🔍 Show parameters", _L)):
                            st.json(_qp_d)
        else:
            st.caption(tt("Zatím žádná uživatelská Rychlá nastavení.",
                          "No custom presets yet.", _L))
                                  # ── EDITACE existujících Rychlých nastavení ─────────────────────────
        st.divider()
        st.markdown(tt("**Upravit existující Rychlé nastavení:**", 
                       "**Edit existing Quick Preset:**", _L))
        
        _user_qp = st.session_state.get("quick_presets", {})
        if _user_qp:
            edit_preset_name = st.selectbox(
                tt("Vyber preset k úpravě", "Select preset to edit", _L),
                options=list(_user_qp.keys()),
                key="ns_edit_preset_select"
            )

            if edit_preset_name:
                _ep = _user_qp[edit_preset_name].copy()

                # Zobrazení aktuálních hodnot pro kontrolu
                with st.expander(tt("🔍 Aktuální hodnoty presetu", "🔍 Current preset values", _L),
                                 expanded=False):
                    st.json(_ep)

                _en_name = st.text_input(
                    tt("Název presetu", "Preset name", _L),
                    value=edit_preset_name,
                    key="ns_edit_preset_newname"
                )
                _en_desc = st.text_area(
                    tt("Popis", "Description", _L),
                    value=_ep.get("description", ""),
                    height=60, key="ns_edit_preset_desc"
                )
                _enc1, _enc2 = st.columns(2)
                with _enc1:
                    _en_temp  = st.slider(
                        tt("Teplota", "Temperature", _L),
                        0.0, 1.0, float(_ep.get("temperature", 0.1)), 0.05,
                        key="ns_edit_preset_temp"
                    )
                    _en_chunk = st.slider(
                        tt("Velikost bloku", "Chunk size", _L),
                        1000, 6000, int(_ep.get("chunk_size", 3200)), 100,
                        key="ns_edit_preset_chunk"
                    )
                with _enc2:
                    _en_maxtok = st.number_input(
                        tt("Max tokens", "Max tokens", _L),
                        512, 32768, int(_ep.get("max_tokens", 4096)), 256,
                        key="ns_edit_preset_maxtok"
                    )
                    _en_large = st.checkbox(
                        tt("Large Document Mode", "Large Document Mode", _L),
                        value=bool(_ep.get("large_mode", False)),
                        key="ns_edit_preset_large"
                    )
                    _en_parallel = st.checkbox(
                        tt("Paralelní zpracování", "Parallel processing", _L),
                        value=bool(_ep.get("parallel", False)),
                        key="ns_edit_preset_parallel"
                    )
                _en_prompt = st.text_area(
                    tt("Výchozí prompt (volitelný)", "Default prompt (optional)", _L),
                    value=_ep.get("prompt", ""),
                    height=80, key="ns_edit_preset_prompt"
                )

                _esave_col, _edel_col = st.columns(2)
                with _esave_col:
                    if st.button(tt("💾 Uložit změny", "💾 Save changes", _L),
                                 key="ns_save_edited_preset", use_container_width=True,
                                 type="primary"):
                        if _en_name.strip():
                            if _en_name != edit_preset_name:
                                _user_qp.pop(edit_preset_name, None)
                            _user_qp[_en_name.strip()] = {
                                "description": _en_desc,
                                "temperature": _en_temp,
                                "chunk_size":  _en_chunk,
                                "max_tokens":  _en_maxtok,
                                "large_mode":  _en_large,
                                "parallel":    _en_parallel,
                                "prompt":      _en_prompt,
                            }
                            st.session_state["quick_presets"] = _user_qp
                            save_quick_presets(_user_qp)
                            st.success(tt("✅ Preset byl uložen", "✅ Preset saved", _L))
                            st.rerun()
                        else:
                            st.warning(tt("Název nesmí být prázdný.", "Name cannot be empty.", _L))

                with _edel_col:
                    if st.checkbox(tt(f"Opravdu smazat '{edit_preset_name}'?",
                                      f"Really delete '{edit_preset_name}'?", _L),
                                   key="chk_del_preset_edit"):
                        if st.button(tt("🗑️ Ano, smazat", "🗑️ Yes, delete", _L),
                                     type="secondary", key="ns_delete_edited_preset",
                                     use_container_width=True):
                            _user_qp.pop(edit_preset_name, None)
                            st.session_state["quick_presets"] = _user_qp
                            save_quick_presets(_user_qp)
                            st.success(tt("✅ Preset smazán", "✅ Preset deleted", _L))
                            st.rerun()
        else:
            st.caption(tt("Nejdříve vytvoř alespoň jedno Rychlé nastavení.",
                          "Create at least one Quick Preset first.", _L))

    # ── Notifikace & Webhook ──────────────────────────
    with _ns_tab3:
        st.markdown(tt("#### 🔔 Desktop notifikace",
                       "#### 🔔 Desktop notifications", _L))
        st.markdown(tt(
            "Browser notifikace — zobrazí se systémové oznámení po dokončení extrakce nebo překladu, "
            "i když máš jiný tab aktivní. Funguje v Chrome/Edge/Firefox.",
            "Browser notifications — system notice appears after extraction or translation finishes, "
            "even when another tab is active. Works in Chrome/Edge/Firefox.", _L))
        if st.button(tt("🔔 Test desktop notifikace", "🔔 Test desktop notification", _L),
                     key="ns_notif_test"):
            desktop_notify(tt("🧪 Test notifikace", "🧪 Test notification", _L),
                           tt("SciNexus funguje správně!",
                              "SciNexus is working correctly!", _L))
            st.toast(tt("Notifikace odeslána — zkontroluj rozeznívání prohlížeče",
                        "Notification sent — check your browser permission", _L))
        st.caption(tt("Pokud notifikace nefunguje: prohlížeč → Nastavení → Oznámení → povolit pro tuto stránku",
                      "If notifications don't work: browser → Settings → Notifications → allow for this page", _L))

        st.divider()
        st.markdown(tt("#### 🔗 Webhook notifikace (Slack/Teams)",
                       "#### 🔗 Webhook notifications (Slack/Teams)", _L))
        _ns_wh_url = st.text_input(
            tt("Webhook URL", "Webhook URL", _L),
            value=st.session_state.get("webhook_url", ""),
            key="ns_webhook_url",
            placeholder="https://hooks.slack.com/services/…")
        if _ns_wh_url != st.session_state.get("webhook_url", ""):
            st.session_state["webhook_url"] = _ns_wh_url
        _nswh1, _nswh2 = st.columns(2)
        with _nswh1:
            if st.button(tt("🔔 Test webhook", "🔔 Test webhook", _L), key="ns_wh_test",
                         disabled=not _ns_wh_url.strip(),
                         width='stretch'):
                ok = send_webhook_notification(
                    _ns_wh_url,
                    tt("Test notifikace z SciNexus ✅",
                       "Test notification from SciNexus ✅", _L))
                st.toast(tt("✅ Webhook doručen", "✅ Webhook delivered", _L) if ok
                         else tt("❌ Webhook selhal", "❌ Webhook failed", _L))

    # ── Data & Cache ──────────────────────────────────
    with _ns_tab4:
        st.markdown(tt("#### 🗄️ Cache a databáze",
                       "#### 🗄️ Cache and databases", _L))
        _nsd_col1, _nsd_col2 = st.columns(2)
        with _nsd_col1:
            st.metric(tt("Validační cache (paměť)", "Validation cache (memory)", _L),
                      len(st.session_state.get("validation_cache", {})))
            st.metric(tt("Validační cache (disk)", "Validation cache (disk)", _L),
                      _disk_cache_count())
            st.metric(tt("FTS index (extrakce)", "FTS index (extractions)", _L),
                      fts_record_count())
        with _nsd_col2:
            st.metric(tt("Checkpointy", "Checkpoints", _L), len(checkpoint_list()))
            _odc2 = offline_db_count()
            st.metric(tt("Offline DB taxonů", "Offline DB taxa", _L),
                      f"{_odc2:,}" if _odc2 else "0")
            st.metric(tt("Šablony promptů", "Prompt templates", _L),
                      len(st.session_state.get("prompt_templates", {})))

        st.divider()
        _nsd_c1, _nsd_c2, _nsd_c3 = st.columns(3)
        with _nsd_c1:
            pass
        with _nsd_c2:
            if st.button(tt("🗑️ Vymazat FTS index", "🗑️ Clear FTS index", _L),
                         key="ns_clear_fts", width='stretch'):
                try:
                    _fts_conn2 = _fts_conn()
                    _fts_conn2.execute("DELETE FROM ext_records")
                    _fts_conn2.execute("DELETE FROM ext_fts")
                    _fts_conn2.commit()
                    st.toast(tt("✅ FTS index vymazán", "✅ FTS index cleared", _L))
                    st.rerun()
                except Exception as e:
                    st.error(tt(f"Chyba: {e}", f"Error: {e}", _L))
        with _nsd_c3:
            if st.button(tt("🗑️ Smazat checkpointy", "🗑️ Delete checkpoints", _L),
                         key="ns_clear_checkpoints", width='stretch'):
                for ck in checkpoint_list():
                    checkpoint_delete(ck["name"])
                st.toast(tt("✅ Checkpointy smazány", "✅ Checkpoints deleted", _L))
                st.rerun()

        st.divider()
        st.markdown(tt("#### 📁 Cesty k souborům", "#### 📁 File paths", _L))
        _path_data = {
            tt("Validační cache", "Validation cache", _L): str(_DISK_CACHE_PATH.resolve()),
            tt("FTS index",       "FTS index",        _L): str(_FTS_DB_PATH.resolve()),
            tt("Checkpointy",     "Checkpoints",      _L): str(_CHECKPOINT_DIR.resolve()),
            tt("Historie",        "History",          _L): str(Path(HISTORY_FILE).resolve()),
            tt("Šablony",         "Templates",        _L): str(Path(TEMPLATES_FILE).resolve()),
        }
        for name, path in _path_data.items():
            st.caption(f"**{name}:** `{path}`")

        # ── v24.24: Diagnostics + Health Check ───────────────────
        st.divider()
        st.markdown(tt("#### 🧪 Diagnostics & Health Check", "#### 🧪 Diagnostics & Health Check", _L))
        show_diagnostics()
        _hc_base = st.session_state.get("lms_base_url", "http://localhost:1234")
        _hc_model = st.session_state.get("selected_model", "")
        if st.button(tt("🚀 Spustit Self-test", "🚀 Run Self-test", _L),
                     key="ns_run_healthcheck"):
            run_health_check(base_url=_hc_base, model=_hc_model)

        # ── v24.24: AI Prompt Optimizer ──────────────────────────
        st.divider()
        st.markdown(tt("#### 🤖 AI Prompt Optimizer", "#### 🤖 AI Prompt Optimizer", _L))
        _po_task = st.selectbox(
            tt("Typ úlohy", "Task type", _L),
            ["extraction", "translation", "validation", "chat"],
            key="ns_po_task",
        )
        _po_prompt = st.text_area(
            tt("Původní prompt", "Original prompt", _L),
            height=120, key="ns_po_prompt",
            value=st.session_state.get("ext_prompt", ""),
        )
        if st.button(tt("✨ Vylepšit prompt", "✨ Optimize prompt", _L),
                     key="ns_po_run", type="primary"):
            _hc_model2 = st.session_state.get("selected_model", "")
            if _hc_model2 and _po_prompt.strip():
                with st.spinner(tt("Optimalizuji…", "Optimizing…", _L)):
                    _optimized = optimize_prompt_with_llm(
                        _po_prompt, _hc_base, _hc_model2, _po_task)
                st.text_area(tt("✅ Vylepšený prompt", "✅ Optimized prompt", _L),
                             value=_optimized, height=180, key="ns_po_result")
                if st.button(tt("📋 Použít jako aktivní prompt extrakce",
                                "📋 Use as active extraction prompt", _L),
                             key="ns_po_apply"):
                    st.session_state["ext_prompt"] = _optimized
                    st.success(tt("✅ Prompt nastaven", "✅ Prompt applied", _L))
            else:
                st.warning(tt("Nejprve načti model a zadej prompt.",
                              "Load a model and enter a prompt first.", _L))

        # ── v24.24: Remote Access (info pokud je vypnuto) ─────────
        if REMOTE_AUTH_ENABLED:
            st.divider()
            st.markdown(tt("#### 🔒 Vzdálený přístup", "#### 🔒 Remote Access", _L))
            st.caption(tt("Ochrana heslem aktivní (env LMU_PASSWORD nastaven).",
                          "Password protection active (env LMU_PASSWORD is set).", _L))
            if st.button(tt("🔓 Odhlásit vzdálený přístup", "🔓 Log out remote access", _L),
                         key="ns_remote_logout"):
                st.session_state["remote_authenticated"] = False
                st.rerun()
        else:
            with st.expander(tt("🔒 Vzdálený přístup (volitelné)", "🔒 Remote Access (optional)", _L)):
                st.info(tt("Ochrana je vypnuta. Nastav env proměnnou `LMU_PASSWORD` pro aktivaci.",
                           "Protection is off. Set env variable `LMU_PASSWORD` to enable.", _L))

        # ── v24.25: Unified Settings Dashboard ───────────────────
        st.divider()
        render_unified_settings()

        # ── v24.25: Operation Timeline ────────────────────────────
        st.divider()
        render_operation_timeline()

        # ── v24.25: Keyboard Shortcuts Cheat Sheet ────────────────
        st.divider()
        show_keyboard_cheatsheet()

        # ── v24.25: Factory Reset ─────────────────────────────────
        st.divider()
        render_factory_reset()

        # ── v24.26: LLM-as-Judge & Agentic Workflow ──────────────
        st.divider()
        st.markdown(tt("#### 🤖 LLM-as-Judge & Agentic Workflow",
                       "#### 🤖 LLM-as-Judge & Agentic Workflow", _L))
        st.checkbox(
            tt("✅ Zapnout LLM-as-Judge (automatické hodnocení kvality extrakce)",
               "✅ Enable LLM-as-Judge (automatic extraction quality scoring)", _L),
            value=st.session_state.get("enable_llm_judge", False),
            key="enable_llm_judge",
            help=tt(
                "Po každé extrakci nechá LLM ohodnotit kvalitu výsledku (0–100) "
                "a případně navrhne opravu. Vyžaduje načtený model.",
                "After each extraction, asks the LLM to score result quality (0–100) "
                "and optionally suggest corrections. Requires a loaded model.", _L),
        )
        if RAG_AVAILABLE:
            st.caption(tt("✅ Semantic RAG Cache dostupná (sentence-transformers + ChromaDB)",
                          "✅ Semantic RAG Cache available (sentence-transformers + ChromaDB)", _L))
        else:
            st.caption(tt("ℹ️ RAG Cache nedostupná — nainstaluj: `pip install sentence-transformers chromadb`",
                          "ℹ️ RAG Cache unavailable — install: `pip install sentence-transformers chromadb`", _L))
        if st.button(
            tt("🧪 Spustit Agentic Workflow test", "🧪 Run Agentic Workflow test", _L),
            key="ns_agentic_test",
        ):
            _ag_text = st.session_state.get("last_extracted_text", "")
            if _ag_text:
                with st.spinner(tt("Agent pracuje…", "Agent working…", _L)):
                    _ag_result = run_agentic_workflow(
                        _ag_text[:6000],
                        base_url=st.session_state.get("lms_base_url", ""),
                        model=st.session_state.get("selected_model", ""),
                    )
                st.markdown(tt("**Log agenta:**", "**Agent log:**", _L))
                for _ag_line in _ag_result["log"]:
                    st.caption(_ag_line)
                st.json(_ag_result["quality"])
            else:
                st.warning(tt("Nejdříve proveď extrakci (záložka Extrakce).",
                              "Run an extraction first (Extraction tab).", _L))


# ══════════════════════════════════════════════════════
# 10 – LM STUDIO CHAT (lokální model)
# ══════════════════════════════════════════════════════
with main_tabs[10]:
    keep_current_tab(10)
    _L = st.session_state.get("lang", "cz")
    st.markdown(tt("## 💬 LM Studio Chat", "## 💬 LM Studio Chat", _L))
    st.caption(tt(
        "Chat s lokálním LLM modelem přes LM Studio API. "
        "Model musí být načten (sidebar → LMS CLI).",
        "Chat with local LLM via LM Studio API. "
        "Model must be loaded (sidebar → LMS CLI).",
        _L))

    # ── Uložené konverzace ───────────────────────────────────────────────────
    _CHAT_CONV_FILE = "chat_conversations.json"
    def _load_chat_convs() -> Dict:
        if os.path.exists(_CHAT_CONV_FILE):
            try:
                with open(_CHAT_CONV_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {}
    def _save_chat_convs(convs: Dict):
        try:
            _tmp = _CHAT_CONV_FILE + ".tmp"
            with open(_tmp, "w", encoding="utf-8") as f:
                json.dump(convs, f, ensure_ascii=False, indent=2)
            os.replace(_tmp, _CHAT_CONV_FILE)
        except Exception:
            pass

    _chat_convs = _load_chat_convs()
    _conv_names = list(_chat_convs.keys())
    _cur_placeholder = tt("— aktuální —", "— current —", _L)

    _lcc1, _lcc2, _lcc3 = st.columns([3, 2, 2])
    with _lcc1:
        _active_conv = st.selectbox(
            tt("💬 Konverzace", "💬 Conversation", _L),
            [_cur_placeholder] + _conv_names,
            key="chat_active_conv")
        if _active_conv != _cur_placeholder and _active_conv in _chat_convs:
            if st.session_state.get("_chat_loaded_conv") != _active_conv:
                st.session_state["chat_messages"] = _chat_convs[_active_conv].get("messages", [])
                st.session_state["_chat_loaded_conv"] = _active_conv
    with _lcc2:
        _new_conv_name = st.text_input(
            tt("Název", "Name", _L), key="chat_new_conv_name",
            label_visibility="collapsed",
            placeholder=tt("Název konverzace…", "Conversation name…", _L))
        if st.button(tt("💾 Uložit", "💾 Save", _L), key="chat_save_conv"):
            if _new_conv_name.strip() and st.session_state.get("chat_messages"):
                _chat_convs[_new_conv_name] = {
                    "messages": st.session_state["chat_messages"],
                    "saved_at": datetime.now().isoformat()}
                _save_chat_convs(_chat_convs)
                st.toast(tt(f"✅ Uloženo: {_new_conv_name}",
                            f"✅ Saved: {_new_conv_name}", _L))
    with _lcc3:
        if st.button(tt("🗑️ Smazat konv.", "🗑️ Delete conv.", _L),
                     key="chat_del_conv",
                     disabled=_active_conv == _cur_placeholder):
            _chat_convs.pop(_active_conv, None)
            _save_chat_convs(_chat_convs)
            st.rerun()

    _col_c1, _col_c2 = st.columns([3, 1])
    with _col_c1:
        _system_chat = st.text_input(
            t("chat_sys_prompt", _L),
            value=tt("Jsi paleontologický asistent specializovaný na hyolity.",
                     "You are a palaeontological assistant specialising in hyoliths.", _L),
            key="sys_chat")
    with _col_c2:
        _temp_chat = st.slider(t("chat_temp", _L), 0.0, 1.5, 0.7, 0.05, key="temp_chat")
        if st.button(t("chat_clear", _L), key="chat_clr_btn"):
            st.session_state["chat_messages"] = []
            st.rerun()

    # ── Injekce kontextu ─────────────────────────────────────────────────────
    _linj1, _linj2 = st.columns([1, 3])
    with _linj1:
        if st.button(tt("📎 Vložit extrakci", "📎 Inject extraction", _L),
                     key="chat_inject_ctx",
                     disabled=not st.session_state.get("last_extracted_text", "")):
            _ext_ctx_c = st.session_state.get("last_extracted_text", "")[:2000]
            st.session_state["chat_injected_context"] = (
                tt("Kontext z extrakce:", "Extraction context:", _L)
                + f"\n\n{_ext_ctx_c}\n\n"
                + tt("Odpovídej na otázky o datech.",
                     "Answer questions about this data.", _L))
            st.toast(tt("✅ Kontext vložen", "✅ Context injected", _L))
    with _linj2:
        _inj_ctx_c = st.session_state.get("chat_injected_context", "")
        if _inj_ctx_c:
            st.caption(tt(f"📌 {_inj_ctx_c[:80]}…", f"📌 {_inj_ctx_c[:80]}…", _L))
            if st.button(tt("✖️ Odebrat", "✖️ Remove", _L), key="chat_remove_ctx"):
                st.session_state["chat_injected_context"] = ""
                st.rerun()

    # ── Historie konverzace ──────────────────────────────────────────────────
    _chat_container_lms = st.container(height=400, border=True)
    with _chat_container_lms:
        if not st.session_state.get("chat_messages"):
            st.caption(tt("Konverzace se zobrazí zde…",
                          "Conversation will appear here…", _L))
        for _cmsg in st.session_state.get("chat_messages", []):
            with st.chat_message(_cmsg["role"],
                                 avatar="🧬" if _cmsg["role"] == "user" else "🤖"):
                st.markdown(_cmsg["content"])

    if st.session_state.get("chat_messages"):
        _chat_exp_txt = "\n\n".join(
            f"{'TY' if m['role']=='user' else 'ASISTENT'}:\n{m['content']}"
            for m in st.session_state["chat_messages"])
        st.download_button(
            t("chat_export", _L), to_txt_bytes(_chat_exp_txt),
            "lms_chat.txt", key="chat_exp_btn")

    # ── Vstup zprávy ─────────────────────────────────────────────────────────
    with st.form("chat_form", clear_on_submit=True):
        _user_msg_input = st.text_area(
            tt("✏️ Zpráva (Ctrl+Enter = Odeslat)",
               "✏️ Message (Ctrl+Enter = Send)", _L),
            height=100,
            placeholder=tt("Napište zprávu…", "Type a message…", _L),
            key="chat_form_text")
        _send_btn = st.form_submit_button(
            tt("📨 Odeslat", "📨 Send", _L),
            type="primary",
            disabled=not selected_model)

    # Dvoustupňový pattern — stejně jako DeepSeek chat
    if st.session_state.pop("_lms_waiting", False) and selected_model:
        # Stupeň 2: API volání
        _sys_lms = st.session_state.get("_lms_pending_sys", "")
        _ctx_lms = st.session_state.pop("_lms_pending_ctx", [])
        with st.spinner(tt("Model přemýšlí…", "Model is thinking…", _L)):
            try:
                _resp_lms = chat_completion(
                    base_url, selected_model, _ctx_lms,
                    temp=_temp_chat, max_tokens=4000)
            except Exception as _ce:
                _resp_lms = tt(f"Chyba: {_ce}", f"Error: {_ce}", _L)
        st.session_state["chat_messages"].append(
            {"role": "assistant", "content": _resp_lms})
        save_session_to_disk(SESSION_AUTOSAVE_FILE, SESSION_AUTOSAVE_KEYS)
        st.rerun()

    elif _send_btn and _user_msg_input.strip():
        if not selected_model:
            st.error(t("chat_no_model", _L))
        else:
            # Stupeň 1: přidej zprávu, ulož kontext, nastav příznak, rerun
            st.session_state.setdefault("chat_messages", [])
            st.session_state["chat_messages"].append(
                {"role": "user", "content": _user_msg_input.strip()})
            _full_sys_lms = _system_chat
            if st.session_state.get("chat_injected_context"):
                _full_sys_lms = _system_chat + "\n\n" + st.session_state["chat_injected_context"]
            _ctx_lms = [{"role": "system", "content": _full_sys_lms}] \
                       + st.session_state["chat_messages"][-20:]
            st.session_state["_lms_pending_ctx"] = _ctx_lms
            st.session_state["_lms_waiting"] = True
            st.rerun()


# ❓ ZÁLOŽKA NÁPOVĚDA
# ══════════════════════════════════════════════════════
with main_tabs[11]:
    keep_current_tab(11)   # ← v24.17 (opraveno: byl 10)
    _L = st.session_state.get("lang","cz")
    st.markdown(tt("## ❓ Nápověda — SciNexus",
                   "## ❓ Help — SciNexus", _L))
    st.caption(tt("Kompletní průvodce aplikací pro paleontologickou taxonomickou databázi.",
                  "Complete guide to the paleontological taxonomic database app.", _L))

    st.markdown("---")

    # ── Rychlý start ──────────────────────────────────
    with st.expander(tt("🚀 Rychlý start — první spuštění",
                         "🚀 Quick start — first run", _L), expanded=True):
        st.markdown("""
### Předpoklady
1. **LM Studio** běží na `http://localhost:1234` (výchozí port)
2. Je načten alespoň jeden model (doporučeno: **Gemma 4 31b it**)
3. V LM Studio je zapnuto: `Developer → Server → Start Server`

### První kroky
1. Zadej URL LM Studia do pole **LM Studio URL** v sidebaru (výchozí: `http://localhost:1234`)
2. Klikni **🔄 Načíst modely** — zobrazí se dostupné modely
3. Vyber model ze seznamu **Aktivní model**
4. Nastav **⚡ Max souběžných požadavků** na **4** (= výchozí limit LM Studio)
5. Přejdi do záložky **🔍 Extrakce** a nahraj první PDF

### Doporučené nastavení modelů (Hyolitha výzkum)
| Úloha | Doporučený model |
|---|---|
| Extrakce taxonomických dat | Qwen2.5-72B nebo Qwen3-32B |
| Překlad (CZ/RU/FR → EN) | Qwen2.5-72B nebo Mistral-24B |
| Chat / dotazy | Mistral-Small-24B nebo Llama-3 |
| Reasoning / složitá extrakce | DeepSeek-R1-32B |
        """)

    # ── Extrakce ──────────────────────────────────────
    with st.expander(tt("🔍 Extrakce — podrobný návod",
                         "🔍 Extraction — detailed guide", _L)):
        st.markdown("""
### Extrakce taxonomických dat z PDF

**Podporované formáty:** PDF (s textem nebo OCR), DOCX, TXT

**Extrakční režimy:**
- **Verbatim** — přesná extrakce, žádná interpretace. Používej pro záznamy s jasnou strukturou.
- **Standard** — vyvážený režim. Doporučeno pro většinu hyolith. publikací.
- **Inferential** — LLM doplní chybějící pole ze kontextu. Používej opatrně.
- **Taxonomic focus** — optimalizováno pro taxonomické popisy druhů.

**Konfidence per-pole:**
Každé pole dostane skóre 0.0–1.0 (vlastní hodnocení LLM):
- 🟢 ≥0.8 — vysoká jistota
- 🟡 ≥0.5 — střední jistota, ověř ručně
- 🔴 <0.5 — nízká jistota, pravděpodobně nesprávné

**Chunk size (velikost bloku):**
- Qwen2.5-72B: doporučeno 20 000–30 000 znaků
- Mistral-24B: 8 000–12 000 znaků
- Llama-3-8B: 4 000–6 000 znaků
- Nech **Auto** — aplikace odhadne optimální hodnotu podle modelu

**Two-pass extrakce:**
Druhý průchod dostane výstup prvního jako kontext a doplní/opraví chybějící pole.
Zvyšuje recall o ~15-25 %, ale zdvojnásobí čas zpracování.

**Prompt Lab:**
Před spuštěním celé extrakce otestuj prompt na prvních 500 znacích.
Ušetří čas při ladění promptu pro nový typ dokumentu.
        """)

    # ── Překlad ───────────────────────────────────────
    with st.expander(tt("🌐 Překlad — podrobný návod",
                         "🌐 Translation — detailed guide", _L)):
        st.markdown("""
### Překlad vědeckých textů

**Překlad velkých PDF (200 stran):**
1. Nahraj PDF v záložce Překlad → Jeden text
2. Nastav **Rozsah stran** (např. `1-50`) nebo nech prázdné pro celý dokument
3. Nastav **Dávky po N stránkách** = 10 (doporučeno pro 100+ stran)
4. Zaškrtni **⚡ Paralelně** pro 3–4× rychlejší zpracování
5. Klikni **🌐 Přeložit**
6. Stáhni jednotlivé dávky nebo ZIP se všemi dávkami

**Výběr chunk size:**
- Aplikace automaticky odhaduje chunk size podle délky textu a modelu
- Pro technické vědecké texty je menší chunk size (6 000–8 000) přesnější
- Pro plynulé prózy jsou větší chunky (12 000+) vhodnější

**Slovník:**
- Přidej doménový slovník (např. `Hyolithes → Hyolithes (zachovat)`)
- Slovník se sdílí přes celou dávku — konzistentní terminologie

**Iterativní překlad (1–3 iterace):**
- 1 iterace = rychlé, dostatečné pro technické texty
- 2–3 iterace = LLM přeloží vícekrát a výsledky inteligentně sloučí; lepší pro historické texty

**Doladění překladu:**
Po překladu jsou k dispozici tři typy doladění:
- **Styl** — upraví text pro cílový časopis (PE, JoP, BDJ, PLOS ONE)
- **Terminologie** — zkontroluje odborné pojmy v dané doméně (paleontologie, geologie…)
- **Zpětný překlad** — přeloží výsledek zpět do originálu a porovná — odhalí ztrátu informací
        """)

    # ── Validace ──────────────────────────────────────
    with st.expander(tt("🧬 Validace taxonů — podrobný návod",
                         "🧬 Taxon validation — detailed guide", _L)):
        st.markdown("""
### Validace taxonomických jmen

**Dostupné databáze:**
| Databáze | Vhodné pro |
|---|---|
| PaleoDB | Fosilní taxony — primární zdroj pro paleontologii |
| GBIF | Recentní i fosilní, rozsáhlá globální databáze |
| IRMNG | Rod-úrovňová validace, mořské i suchozemské |
| Fossilworks | Veřejné rozhraní PaleoDB |
| CoL | Catalogue of Life — taxony všech skupin |
| WoRMS | Mořské organismy |
| ZooBank | Registrace nových jmen (ICZN) |
| BioLib | Česká databáze, funguje pouze z IP Národního muzea |
| Plazi | Taxonová jména z vědeckých publikací (TreatmentBank) |
| IPNI | Rostliny (International Plant Names Index) |

**Výsledky — karta taxonu:**
- Taxonomická pozice (kingdom → genus → species) z nejlepší DB
- Tabulka všech DB se statusem a přímým odkazem na záznam
- Konflikty mezi DB (CoL: synonym × PaleoDB: valid)

**Offline DB:**
- Importuj taxony z validačních výsledků do lokální SQLite DB
- Při zaškrtnutém "Offline fallback" se nejdřív zkusí lokální DB, pak internet
- Vhodné pro opakovanou validaci stejných taxonů bez internetu

**BioLib API:**
Funguje pouze z IP adres registrovaných u BioLib.cz (Národní muzeum Praha má přístup).
Z jiných IP vrací 403 — zobrazí se fallback search odkaz.
        """)

    # ── Klávesové zkratky ─────────────────────────────
    with st.expander(tt("⌨️ Klávesové zkratky",
                         "⌨️ Keyboard shortcuts", _L)):
        st.markdown("""
| Zkratka | Akce |
|---|---|
| **Ctrl+Enter** | Odeslat zprávu v chatu |
| **Alt+E** | Přejít na záložku Extrakce |
| **Alt+T** | Přejít na záložku Překlad |
| **Alt+V** | Přejít na záložku Validace |
| **Alt+C** | Přejít na záložku Chat |
| **Alt+D** | Přejít na záložku Čištění dat |
| **Alt+S** | Přejít na záložku Stylistika |
| **Alt+W** | Přejít na záložku Workflow |
        """)

    # ── Quick Start — doporučená výchozí nastavení ───
    with st.expander(tt("🚀 Quick Start — doporučená nastavení pro batch",
                         "🚀 Quick Start — recommended batch defaults", _L), expanded=False):
        st.markdown(tt("""
### 🚀 Doporučená výchozí nastavení pro batch (když nevíte kde začít)

Tato nastavení fungují dobře pro 90 % úloh — taxonomická extrakce, překlad vědeckých textů, dávková validace.

#### ⚡ Concurrency (souběžnost)
| Zdroj nastavení | Hodnota |
|---|---|
| LM Studio → Developer → Max Concurrent Predictions | **4** (default) nebo **8** pro RTX 4090 |
| Sidebar → Max souběžných LLM požadavků | stejná hodnota jako výše |
| Validace (HTTP dotazy na DB) | automaticky max(4, nastavená hodnota) |

> Pro lokální LM Studio na jedné GPU: začni na **4**, zvyš na **8** pokud model nestačí využívat VRAM.

#### 🌡️ Temperature & max_tokens
| Úloha | Temperature | max_tokens |
|---|---|---|
| JSON extrakce taxonů | **0.05–0.1** | **2048–4096** (přísný strop) |
| Překlad vědeckého textu | **0.1–0.15** | adaptivní (auto) |
| Překlad historického textu | **0.15–0.2** | adaptivní (auto) |
| Chat / kreativní úlohy | **0.5–0.8** | bez omezení |

> Nižší temperature = konzistentnější JSON, méně halucinací u taxonomických jmen.

#### 📦 Chunkování
| Parametr | Doporučení |
|---|---|
| Chunk size (extrakce) | **auto** (aplikace odhadne dle modelu) |
| Chunk size (překlad) | **auto** (odhadne dle délky textu) |
| Překryv (extrakce) | **200–400 znaků** pro taxonomické záznamy |
| Minimální chunk | aplikace automaticky slučuje bloky < 2 000 znaků |

#### 🔁 Stop sekvence (automatické)
Aplikace automaticky přidává stop sekvence pro JSON extrakci (`\\`\\`\\`\\n`), takže model přestane generovat hned za koncem výstupu — šetří tokeny a čas.

#### 🔄 Retry & odolnost
- Všechna LLM volání mají **3 pokusy** s exponential backoff + jitter
- Selželvší chunky jdou do **DLQ** (dead letter queue) — batch pokračuje, chyby jsou viditelné po dokončení
- Po každém dokončeném souboru se automaticky ukládá **checkpoint** (obnova po pádu)

#### 📊 Typické časy (RTX 4090, Qwen2.5-32B Q4)
| Úloha | Velikost | Čas |
|---|---|---|
| Extrakce JSON, 1 PDF (20 str.) | ~1 chunk | ~15–30 s |
| Překlad PDF (100 str.), paralelně | ~10 chunků | ~3–5 min |
| Validace 100 taxonů | 100 HTTP dotazů | ~30–60 s |
        """, """
### 🚀 Recommended batch defaults (when you don't know where to start)

These settings work well for 90% of tasks — taxonomic extraction, scientific text translation, batch validation.

#### ⚡ Concurrency
| Setting location | Value |
|---|---|
| LM Studio → Developer → Max Concurrent Predictions | **4** (default) or **8** for RTX 4090 |
| Sidebar → Max concurrent LLM requests | same value as above |
| Validation (HTTP DB queries) | automatically max(4, configured value) |

> For local LM Studio on a single GPU: start at **4**, increase to **8** if VRAM is underutilized.

#### 🌡️ Temperature & max_tokens
| Task | Temperature | max_tokens |
|---|---|---|
| JSON taxon extraction | **0.05–0.1** | **2048–4096** (strict cap) |
| Scientific text translation | **0.1–0.15** | adaptive (auto) |
| Historical text translation | **0.15–0.2** | adaptive (auto) |
| Chat / creative tasks | **0.5–0.8** | uncapped |

> Lower temperature = more consistent JSON, fewer hallucinations in taxon names.

#### 📦 Chunking
| Parameter | Recommendation |
|---|---|
| Chunk size (extraction) | **auto** (app estimates from model context) |
| Chunk size (translation) | **auto** (estimates from text length) |
| Overlap (extraction) | **200–400 chars** for taxonomic records |
| Minimum chunk | app automatically merges blocks < 2,000 chars |

#### 🔁 Stop sequences (automatic)
The app automatically adds stop sequences for JSON extraction (`\\`\\`\\`\\n`), stopping generation right after the output ends — saves tokens and time.

#### 🔄 Retry & resilience
- All LLM calls have **3 attempts** with exponential backoff + jitter
- Failed chunks go to the **DLQ** (dead letter queue) — batch continues, errors shown after completion
- After every completed file, a **checkpoint** is auto-saved (recovery after crash)

#### 📊 Typical times (RTX 4090, Qwen2.5-32B Q4)
| Task | Size | Time |
|---|---|---|
| JSON extraction, 1 PDF (20 pp.) | ~1 chunk | ~15–30 s |
| PDF translation (100 pp.), parallel | ~10 chunks | ~3–5 min |
| Validation of 100 taxa | 100 HTTP queries | ~30–60 s |
        """, _L))

    # ── Výkon a optimalizace ──────────────────────────
    with st.expander(tt("⚡ Výkon a optimalizace",
                         "⚡ Performance and optimization", _L)):
        st.markdown("""
### Paralelní requesty (LM Studio 0.4+)

LM Studio podporuje **continuous batching** — více požadavků najednou:

1. V LM Studio: `Developer → Server → Max Concurrent Predictions` = **4**
2. V aplikaci (sidebar): **⚡ Max souběžných požadavků** = **4**

**Kde se paralelismus projeví:**
- **Extrakce** — chunky dokumentu se zpracovávají souběžně → ~4× rychlejší
- **Překlad** — bloky textu se překládají najednou → ~3–4× rychlejší
- **Validace** — HTTP dotazy na databáze běží paralelně (neomezeno LM Studio limitem)
- **Dávkový překlad** — odstavce s deduplikací se překládají souběžně

**Důležité:** Paralelní překlad nesdílí kontext mezi bloky.
Pro texty kde je plynulost klíčová (historické dokumenty), použij sekvenční režim.
Pro vědecké záznamy s krátkými odstavci (hyolith. databáze) je paralelní ideální.

### Doporučená hardwarová konfigurace
| Konfigurace | Vhodné pro |
|---|---|
| RTX 4090 (24 GB VRAM) | Qwen2.5-72B Q4, rychlá extrakce 200-stránkových PDF |
| RTX 3090/4080 (16–24 GB) | Qwen2.5-32B Q4, Mistral-24B |
| CPU (128 GB RAM) | Qwen2.5-72B Q8 — pomalé, ale vysoká kvalita |
        """)

    # ── Časté problémy ────────────────────────────────
    with st.expander(tt("🔧 Řešení častých problémů",
                         "🔧 Troubleshooting", _L)):
        st.markdown("""
### Časté chybové hlášky

**`Context size has been exceeded`**
- Text nebo kontext přesahuje context window modelu
- Řešení: zmenši chunk size, použij model s větším oknem (Qwen2.5: 128k), nebo rozdělí text na menší části

**`Connection refused / nelze se připojit`**
- LM Studio server neběží
- Řešení: v LM Studio klikni Developer → Start Server

**`No models available`**
- Model není načten v LM Studio
- Řešení: v LM Studio klikni na model → Load Model

**`403 Forbidden (BioLib)`**
- Volání přichází z IP adresy mimo whitelist BioLib.cz
- Řešení: funguje pouze z Národního muzea Praha; z domova není přístup

**`Validace vrací prázdné výsledky`**
- Starý cache před opravou bugu v16.3
- Řešení: v sidebaru klikni **🗑️ Vymazat cache validace** a spusť validaci znovu

**Překlad nebo extrakce trvá příliš dlouho**
- Nastav ⚡ Max souběžných požadavků na 4 v sidebaru
- Zaškrtni ⚡ Paralelně v záložce překladu
- Zkontroluj chunk size — příliš malé chunky = zbytečně mnoho požadavků
        """)

    # ── O aplikaci ────────────────────────────────────
    with st.expander("ℹ️ O aplikaci"):
        st.markdown("""
### SciNexus
**Verze:** 24.34  
**Autor:** Martin Valent, Národní muzeum Praha — Paleontologické oddělení  
**Specializace:** Hyolitha (Cambrian marine invertebrates)

Aplikace slouží k automatizaci taxonomické práce:
- Extrakce dat z vědeckých publikací pomocí LLM
- Překlad historické literatury (čínština, ruština, francouzština, němčina → angličtina)
- Validace taxonomických jmen ve 14 mezinárodních databázích
- Normalizace a čištění databázových záznamů
- Stylistické úpravy pro vědecké časopisy

**Technologie:** Streamlit, LM Studio (llama.cpp), Python  
**Lokální zpracování:** Všechna data zůstávají na tvém počítači — žádná data se neodesílají do cloudu.
        """)